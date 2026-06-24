# SecBoard/app_risk/report_views.py

import json
import logging
import os
import tempfile
from datetime import datetime, timedelta
from decimal import Decimal
import decimal
from io import BytesIO

from django.contrib.auth.decorators import login_required, user_passes_test
from django.core.exceptions import PermissionDenied
from django.db.models import Count, Q, Sum, Avg, Max, Min
from django.http import JsonResponse, HttpResponse, Http404
from django.shortcuts import render, get_object_or_404
from django.template.loader import render_to_string
from django.utils import timezone
from django.utils.translation import gettext as _, get_language
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods

# Third-party imports
try:
    from weasyprint import HTML, CSS  # noqa: F401  # pyright: ignore[reportMissingImports]
    from weasyprint.text.fonts import FontConfiguration  # noqa: F401  # pyright: ignore[reportMissingImports]
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import xlsxwriter  # noqa: F401  # pyright: ignore[reportUnusedImport, reportMissingImports, reportMissingImports]
    XLSXWRITER_AVAILABLE = True
except ImportError:
    XLSXWRITER_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BEAUTIFULSOUP_AVAILABLE = True
except ImportError:
    BEAUTIFULSOUP_AVAILABLE = False

# Local imports
from .models import (
    RiskTreatment, AssetVulnerability, Vulnerability, Threat, 
    RiskLevel, AccessRisk, RiskAssessmentAuditLog,
    ScheduledReport, ScheduledReportExecution, ScheduledReportAttachment, ReportProfile, RiskReportEmailConfig
)
from app_cabinet.models import CabinetUser
from app_asset.models import InformationAsset
from app_conf.models import Company
from .risk_assessment_views import calculate_risk_level, calculate_value_of_risk
from .access_utils import can_add_risk_report, can_edit_risk_report, can_delete_risk_report

logger = logging.getLogger(__name__)


def check_risk_assessment_access(user, action='view'):
    """Check if user has access to risk assessment functionality"""
    if user.is_superuser:
        return True
    
    # Check if user has access through groups
    user_groups = user.groups.all()
    access_records = AccessRisk.objects.filter(
        group__in=user_groups,
        has_access_assessment=True
    )
    
    return access_records.exists()

def check_risk_report_access(user, action='view'):
    """Check if user has access to risk report functionality"""
    if user.is_superuser:
        return True
    
    # Check if user has access through groups
    user_groups = user.groups.all()
    access_records = AccessRisk.objects.filter(
        group__in=user_groups,
        has_access_report=True
    )
    
    return access_records.exists()

@login_required
@user_passes_test(check_risk_report_access)
@require_http_methods(["POST"])
def preview_email_content(request):
    """Preview email content with processed tags"""
    try:
        # Get form data
        email_subject = request.POST.get('email_subject', '')
        email_body = request.POST.get('email_body', '')
        schedule_id = request.POST.get('schedule_id')
        schedule_name = request.POST.get('schedule_name', '')
        company_id = request.POST.get('company_id', '')
        
        # Get scheduled report if exists, otherwise create a temporary one for preview
        if schedule_id:
            try:
                schedule = ScheduledReport.objects.get(id=schedule_id)
            except ScheduledReport.DoesNotExist:
                return JsonResponse({
                    'status': 'error',
                    'message': _('Scheduled report not found')
                })
        else:
            # Create a temporary ScheduledReport object for preview
            # This allows preview to work even before saving the schedule
            schedule = ScheduledReport(
                name=schedule_name or _('Preview Report'),
                created_by=request.user
            )
            
            # Set company if provided
            if company_id:
                try:
                    schedule.company = Company.objects.get(id=company_id)
                except Company.DoesNotExist:
                    pass
        
        # Create a mock execution for preview
        # We need to create it without saving to database
        class MockExecution:
            def __init__(self, scheduled_report, started_at):
                self.scheduled_report = scheduled_report
                self.started_at = started_at
                self.status = 'completed'
                self.id = 'preview-mock-id'
            
            def get_snapshot_url(self):
                # Return a placeholder URL path for preview (without domain)
                # The domain will be added by _build_absolute_url in process_email_tags
                from django.urls import reverse
                try:
                    return reverse('view_scheduled_report_snapshot', args=['preview-mock-id'])
                except:
                    return '/app_risk/scheduled-reports/snapshot/preview-mock-id/'
        
        mock_execution = MockExecution(schedule, timezone.now())
        
        # Process tags
        processed_subject = schedule.process_email_tags(email_subject, mock_execution)
        processed_body = schedule.process_email_tags(email_body, mock_execution)
        
        return JsonResponse({
            'status': 'success',
            'data': {
                'subject': processed_subject,
                'body': processed_body
            }
        })
        
    except Exception as e:
        logger.error(f"Error previewing email content: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error previewing email content: {}').format(str(e))
        })


@login_required
@user_passes_test(check_risk_report_access)
def download_scheduled_report_attachment(request, attachment_id):
    """Download attachment file"""
    try:
        from .models import ScheduledReportAttachment
        attachment = get_object_or_404(ScheduledReportAttachment, id=attachment_id)
        
        # Check permissions
        user_companies = get_user_companies(request.user)
        if attachment.scheduled_report.created_by != request.user and \
           (attachment.scheduled_report.company and attachment.scheduled_report.company not in user_companies):
            raise PermissionDenied(_("You do not have permission to download this file"))
        
        # Check if file exists
        if not attachment.file:
            raise Http404(_("File not found"))
        
        # Serve file
        response = HttpResponse(attachment.file, content_type='application/octet-stream')
        response['Content-Disposition'] = f'attachment; filename="{attachment.original_filename}"'
        return response
        
    except Exception as e:
        logger.error(f"Error downloading attachment {attachment_id}: {str(e)}")
        raise Http404(_("File not found"))


@login_required
@user_passes_test(check_risk_report_access)
def view_scheduled_report_snapshot(request, execution_id):
    """Render a frozen snapshot of a scheduled report execution that was saved at send time."""
    try:
        execution = get_object_or_404(ScheduledReportExecution, id=execution_id)
        report = execution.scheduled_report
        
        # Check permissions: owner, same company, or email recipient
        user_companies = get_user_companies(request.user)
        user_has_access = False
        
        # Check if user is the owner
        if report.created_by == request.user:
            user_has_access = True
        # Check if user belongs to the same company
        elif report.company and report.company in user_companies:
            user_has_access = True
        # Check if user is in the email recipients list
        elif report.email_recipients.filter(user=request.user).exists():
            user_has_access = True
        
        if not user_has_access:
            raise PermissionDenied(_("You do not have permission to view this report"))

        html = execution.snapshot_html or ''
        if not html:
            html = '<div class="alert alert-warning"><i class="fas fa-info-circle me-2"></i>' + \
                  _('Snapshot is not available for this execution.') + '</div>'

        return render(request, 'app_risk/risk_report_preview_page.html', {
            'preview_html': html,
            'page_title': _('Risk Assessment Report'),
            'is_profile': True,
            'profile_id': str(execution.id),  # Use execution ID as profile_id for export
            'language': execution.snapshot_language or 'uk',
        })
    except Exception as e:
        logger.error(f"Error rendering scheduled report snapshot {execution_id}: {e}", exc_info=True)
        return render(request, 'app_risk/risk_report_preview_page.html', {
            'preview_html': f'<div class="alert alert-danger">{_("Error generating preview")}: {str(e)}</div>',
            'page_title': _('Risk Assessment Report')
        })
@require_http_methods(["GET", "POST"])
def generate_risk_report(request):
    """Generate comprehensive risk assessment reports with multilingual support"""
    
    # Check permissions
    if not check_risk_report_access(request.user, 'view'):
        raise PermissionDenied(_("You don't have permission to generate reports"))
    
    if request.method == 'GET':
        # Check if this is a direct report generation request with parameters
        report_type = request.GET.get('reportType', 'summary')
        report_format = request.GET.get('format', 'word')
        report_language = request.GET.get('language', get_language()[:2])
        
        if report_type and report_format:
            # Direct generation via GET parameters
            try:
                # Activate the selected language
                from django.utils.translation import activate
                activate(report_language)
                
                # Validate format availability
                if not is_format_available(report_format):
                    return JsonResponse({
                        'status': 'error',
                        'message': _('Selected format is not available. Please install required dependencies.')
                    })
                
                # Generate report data
                report_data = generate_report_data(request.user, {
                    'reportType': report_type,
                    'format': report_format,
                    'language': report_language
                })
                
                # Generate report file
                response = generate_report_file(report_data, report_type, report_format)
                
                return response
                
            except Exception as e:
                logger.error(f"Error generating risk report: {str(e)}", exc_info=True)
                return JsonResponse({
                    'status': 'error',
                    'message': _('Error generating report: {}').format(str(e))
                })
        else:
            # Return report configuration page
            return render(request, 'app_risk/reports/report_config.html', {
                'available_formats': get_available_formats(),
                'companies': get_user_companies(request.user),
            })
    
    elif request.method == 'POST':
        try:
            # Parse request data for POST - handle both JSON and form data
            if request.content_type == 'application/json':
                data = json.loads(request.body) if request.body else {}
            else:
                # Handle form data
                data = {
                    'reportType': request.POST.get('reportType', 'full'),
                    'format': request.POST.get('reportFormat', 'pdf'),
                    'language': request.POST.get('reportLanguage', get_language()[:2]),
                    'company_id': request.POST.get('reportCompany', ''),
                    'notes': request.POST.get('reportNotes', ''),
                }
            
            # Parse selected sections if provided
            selected_sections_json = request.POST.get('selectedSections', '{}')
            try:
                selected_sections = json.loads(selected_sections_json)
                data['selectedSections'] = selected_sections
            except json.JSONDecodeError:
                # If JSON parsing fails, use empty dict
                data['selectedSections'] = {}
            
            report_type = data.get('reportType', 'full')
            report_format = data.get('format', data.get('reportFormat', 'pdf'))
            report_language = data.get('language', data.get('reportLanguage', get_language()[:2]))
            
            # Activate the selected language
            from django.utils.translation import activate
            activate(report_language)
            
            # Validate format availability
            if not is_format_available(report_format):
                return JsonResponse({
                    'status': 'error',
                    'message': _('Selected format is not available. Please install required dependencies.')
                })
            
            # Generate report data
            report_data = generate_report_data(request.user, data)
            
            # Generate report file
            response = generate_report_file(report_data, report_type, report_format)
            
            return response
            
        except Exception as e:
            logger.error(f"Error generating risk report: {str(e)}", exc_info=True)
            return JsonResponse({
                'status': 'error',
                'message': _('Error generating report: {}').format(str(e))
            })

# Import data collection and generation functions from new modules
from .report_data import generate_report_data
from .report_generators_core import (
    generate_report_file, is_format_available, get_available_formats,
    get_localized_vulnerability_field, format_localized_date, get_report_translations
)


def get_user_companies(user):
    """Get companies accessible to user"""
    if user.is_superuser:
        return Company.objects.all()
    
    # Check if user is staff and has no specific access restrictions
    if user.is_staff:
        user_groups = user.groups.all()
        access_records = AccessRisk.objects.filter(
            group__in=user_groups,
            has_access_assessment=True
        )
        
        if not access_records.exists():
            # If staff user has no access restrictions, give access to all companies
            return Company.objects.all()
    
    # Get companies through access records
    user_groups = user.groups.all()
    access_records = AccessRisk.objects.filter(
        group__in=user_groups,
        has_access_assessment=True
    )
    
    companies = Company.objects.none()
    for access in access_records:
        companies = companies | access.companies.all()
    
    # If no companies found through access records, try to get companies where user is a cabinet user
    if not companies.exists():
        try:
            from app_cabinet.models import CabinetUser
            cabinet_user = CabinetUser.objects.filter(user=user).first()
            if cabinet_user and cabinet_user.company:
                companies = Company.objects.filter(id=cabinet_user.company.id)
        except:
            pass
    
    # Last fallback: if still no companies and user is authenticated, give access to all companies
    if not companies.exists() and user.is_authenticated:
        companies = Company.objects.all()
    
    return companies.distinct()


@login_required
@user_passes_test(check_risk_report_access)
@require_http_methods(["POST"])
def preview_risk_report(request):
    """Generate HTML preview of risk assessment report for modal display"""
    
    # Check permissions
    if not check_risk_report_access(request.user, 'view'):
        raise PermissionDenied(_("You don't have permission to preview reports"))
    
    try:
        # Parse request data
        report_type = request.POST.get('reportType', 'summary')
        report_format = request.POST.get('reportFormat', 'word')
        report_language = request.POST.get('reportLanguage', get_language()[:2])
        
        # Parse selected sections from the configuration matrix
        selected_sections_json = request.POST.get('selectedSections', '{}')
        try:
            selected_sections = json.loads(selected_sections_json) if selected_sections_json else {}
        except json.JSONDecodeError:
            logger.warning(f"Invalid JSON in selectedSections: {selected_sections_json}")
            selected_sections = {}
        
        # Activate the selected language
        from django.utils.translation import activate
        activate(report_language)
        
        # Generate report data with selected sections
        report_data = generate_report_data(request.user, {
            'reportType': report_type,
            'format': report_format,
            'language': report_language,
            'company_id': request.POST.get('reportCompany', ''),
            'selectedSections': selected_sections,  # Pass selected sections to report data generation
        })
        
        # Generate HTML preview with selected sections
        preview_html = generate_html_preview(report_data, report_type, report_language, selected_sections)
        
        return HttpResponse(preview_html)
        
    except Exception as e:
        logger.error(f"Error generating report preview: {str(e)}", exc_info=True)
        error_html = f'''
        <div class="alert alert-danger">
            <i class="fas fa-exclamation-triangle me-2"></i>
            {_("Error generating preview")}: {str(e)}
        </div>
        '''
        return HttpResponse(error_html)


@login_required
@user_passes_test(check_risk_report_access)
@require_http_methods(["GET"])
def preview_risk_report_page(request):
    """Render preview as a standalone page using query params."""
    try:
        report_type = request.GET.get('reportType', 'summary')
        report_format = request.GET.get('format', 'word')
        report_language = request.GET.get('language', get_language()[:2])
        selected_sections_json = request.GET.get('selectedSections', '{}')
        try:
            selected_sections = json.loads(selected_sections_json) if selected_sections_json else {}
        except json.JSONDecodeError:
            selected_sections = {}

        # Activate language
        from django.utils.translation import activate
        activate(report_language)

        report_data = generate_report_data(request.user, {
            'reportType': report_type,
            'format': report_format,
            'language': report_language,
            'company_id': request.GET.get('reportCompany', ''),
            'selectedSections': selected_sections,
        })

        html = generate_html_preview(report_data, report_type, report_language, selected_sections)
        return render(request, 'app_risk/risk_report_preview_page.html', {
            'preview_html': html,
            'page_title': _('Report Preview'),
            'is_profile': False,
            'report_params': {
                'reportType': report_type,
                'language': report_language,
                'format': report_format,
                'company_id': request.GET.get('reportCompany', ''),
                'selectedSections': selected_sections_json,
            }
        })
    except Exception as e:
        logger.error(f"Error generating standalone preview: {str(e)}", exc_info=True)
        return render(request, 'app_risk/risk_report_preview_page.html', {
            'preview_html': f'<div class="alert alert-danger">{_("Error generating preview")}: {str(e)}</div>',
            'page_title': _('Report Preview')
        })


@login_required
@user_passes_test(check_risk_report_access)
@require_http_methods(["GET"])
def preview_risk_report_from_profile_page(request):
    """Render profile-based preview as a standalone page."""
    try:
        profile_id = request.GET.get('profile_id')
        if not profile_id:
            raise ValueError(_('Profile ID is required'))

        profile = ReportProfile.objects.get(id=profile_id)
        if not profile.can_be_used_by(request.user):
            raise PermissionDenied(_("You don't have permission to use this profile"))

        override_language = request.GET.get('language') or profile.default_language
        from django.utils.translation import activate
        activate(override_language)

        sections_config = profile.get_sections_config()
        company_id = sections_config.get('_company_id', '')

        params = {
            'reportType': 'profile',
            'format': 'html',
            'language': override_language,
            'company_id': str(company_id) if company_id else '',
            'selectedSections': sections_config,
            'profile': profile
        }

        # Enrich profile header
        company_name = _('All Companies')
        if company_id:
            try:
                from app_cabinet.models import Company
                company = Company.objects.get(id=company_id)
                company_name = company.name
            except Exception:
                pass
        params['profile'] = {
            'name': profile.name,
            'description': profile.description,
            'created_by': profile.created_by.get_full_name(),
            'created_at': profile.created_at,
            'company': company_name,
        }

        report_data = generate_report_data(request.user, params)
        html = generate_html_preview(report_data, 'profile', override_language, sections_config)

        return render(request, 'app_risk/risk_report_preview_page.html', {
            'preview_html': html,
            'page_title': _('Report Preview'),
            'is_profile': True,
            'profile_id': str(profile.id),
            'language': override_language,
        })
    except Exception as e:
        logger.error(f"Error generating standalone profile preview: {str(e)}", exc_info=True)
        return render(request, 'app_risk/risk_report_preview_page.html', {
            'preview_html': f'<div class="alert alert-danger">{_("Error generating preview")}: {str(e)}</div>',
            'page_title': _('Report Preview')
        })

def generate_ai_conclusion(report_data, selected_sections, language='uk'):
    """
    Generate AI Conclusion based on report data using selected AI model.
    
    Args:
        report_data: Dictionary containing all report data
        selected_sections: Dictionary with selected sections configuration
        language: Language code for the conclusion
    
    Returns:
        str: Generated AI conclusion text or None if generation failed
    """
    try:
        # Check if conclusions section is enabled (handle both boolean and string values)
        conclusions_enabled = selected_sections.get('conclusions', False)
        if not (conclusions_enabled == True or (isinstance(conclusions_enabled, str) and conclusions_enabled.lower() in ('true', '1'))):
            logger.debug("AI Conclusion section is not enabled")
            return None
        
        # Get AI model ID from selected sections (handle both string and int)
        ai_model_id = selected_sections.get('_ai_conclusion_model')
        if not ai_model_id:
            logger.warning("AI Conclusion is enabled but no AI model is selected")
            return None
        
        # Convert to int if it's a string
        if isinstance(ai_model_id, str):
            try:
                ai_model_id = int(ai_model_id)
            except ValueError:
                logger.error(f"Invalid AI model ID format: {ai_model_id}")
                return None
        
        # Get the AI model
        from app_ai.models import ModelChoice
        try:
            ai_model = ModelChoice.objects.get(id=ai_model_id, is_active=True)
        except ModelChoice.DoesNotExist:
            logger.error(f"AI model with ID {ai_model_id} not found or inactive")
            return None
        
        # Prepare report summary for AI analysis
        report_summary = prepare_report_summary_for_ai(report_data, language)
        
        # Create prompt for AI
        if language == 'uk':
            system_prompt = "Ти експерт з аналізу ризиків та безпеки. Проаналізуй дані звіту про оцінку ризиків та надай професійний висновок українською мовою."
            user_prompt = f"""Проаналізуй наступні дані звіту про оцінку ризиків та надай детальний професійний висновок:

{report_summary}

Висновок має містити:
1. Загальну оцінку поточного стану безпеки
2. Основні виявлені ризики та їх критичність
3. Рекомендації щодо пріоритетних дій
4. Оцінку ефективності запроваджених заходів безпеки
5. Загальні висновки та рекомендації для керівництва

Висновок має бути професійним, структурованим та зрозумілим."""
        elif language == 'ru':
            system_prompt = "Ты эксперт по анализу рисков и безопасности. Проанализируй данные отчета об оценке рисков и предоставь профессиональный вывод на русском языке."
            user_prompt = f"""Проанализируй следующие данные отчета об оценке рисков и предоставь детальный профессиональный вывод:

{report_summary}

Вывод должен содержать:
1. Общую оценку текущего состояния безопасности
2. Основные выявленные риски и их критичность
3. Рекомендации по приоритетным действиям
4. Оценку эффективности внедренных мер безопасности
5. Общие выводы и рекомендации для руководства

Вывод должен быть профессиональным, структурированным и понятным."""
        else:  # English
            system_prompt = "You are an expert in risk analysis and security. Analyze the risk assessment report data and provide a professional conclusion in English."
            user_prompt = f"""Analyze the following risk assessment report data and provide a detailed professional conclusion:

{report_summary}

The conclusion should include:
1. Overall assessment of current security status
2. Main identified risks and their criticality
3. Recommendations for priority actions
4. Assessment of effectiveness of implemented security measures
5. General conclusions and recommendations for management

The conclusion should be professional, structured, and clear."""
        
        # Get AI response based on provider
        from app_ai import ai_utils
        
        conversation_history = []
        ai_response = None
        
        # Normalize provider to lowercase for comparison
        provider = ai_model.provider.lower() if ai_model.provider else ''
        
        if provider == 'ollama':
            ai_response, _ = ai_utils.get_ollama_response(user_prompt, conversation_history, system_prompt)
        elif provider == 'google':
            ai_response, _ = ai_utils.get_google_response(user_prompt, conversation_history, system_prompt)
        elif provider == 'claude':
            ai_response, _ = ai_utils.get_claude_response(user_prompt, conversation_history, system_prompt)
        elif provider == 'groq':
            ai_response, _ = ai_utils.get_groq_response(user_prompt, conversation_history, system_prompt)
        elif provider == 'deepseek':
            ai_response, _ = ai_utils.get_deepseek_response(user_prompt, conversation_history, system_prompt)
        else:
            logger.error(f"Unsupported AI provider: {ai_model.provider}")
            return None
        
        if ai_response and not ai_response.startswith("Error"):
            return ai_response
        else:
            logger.error(f"AI response error: {ai_response}")
            return None
            
    except Exception as e:
        logger.error(f"Error generating AI conclusion: {str(e)}", exc_info=True)
        return None


def prepare_report_summary_for_ai(report_data, language='uk'):
    """
    Prepare a text summary of report data for AI analysis.
    
    Args:
        report_data: Dictionary containing all report data
        language: Language code for the summary
    
    Returns:
        str: Formatted text summary
    """
    summary_parts = []
    
    # Statistics
    if 'statistics' in report_data:
        stats = report_data['statistics']
        if language == 'uk':
            summary_parts.append(f"Загальна статистика:")
            summary_parts.append(f"- Всього активів: {stats.get('total_assets', 0)}")
            summary_parts.append(f"- Всього вразливостей: {stats.get('total_vulnerabilities', 0)}")
            summary_parts.append(f"- Активів з високим ризиком: {stats.get('high_risk_count', 0)}")
            summary_parts.append(f"- Рівень завершеності: {stats.get('completion_rate', 0):.1f}%")
        elif language == 'ru':
            summary_parts.append(f"Общая статистика:")
            summary_parts.append(f"- Всего активов: {stats.get('total_assets', 0)}")
            summary_parts.append(f"- Всего уязвимостей: {stats.get('total_vulnerabilities', 0)}")
            summary_parts.append(f"- Активов с высоким риском: {stats.get('high_risk_count', 0)}")
            summary_parts.append(f"- Уровень завершенности: {stats.get('completion_rate', 0):.1f}%")
        else:
            summary_parts.append(f"Overall Statistics:")
            summary_parts.append(f"- Total Assets: {stats.get('total_assets', 0)}")
            summary_parts.append(f"- Total Vulnerabilities: {stats.get('total_vulnerabilities', 0)}")
            summary_parts.append(f"- High Risk Assets: {stats.get('high_risk_count', 0)}")
            summary_parts.append(f"- Completion Rate: {stats.get('completion_rate', 0):.1f}%")
    
    # Risk Distribution
    if 'risk_distribution' in report_data and 'by_level' in report_data['risk_distribution']:
        risk_levels = report_data['risk_distribution']['by_level']
        if language == 'uk':
            summary_parts.append(f"\nРозподіл ризиків за рівнями:")
            for level, count in risk_levels.items():
                summary_parts.append(f"- {level}: {count}")
        elif language == 'ru':
            summary_parts.append(f"\nРаспределение рисков по уровням:")
            for level, count in risk_levels.items():
                summary_parts.append(f"- {level}: {count}")
        else:
            summary_parts.append(f"\nRisk Distribution by Level:")
            for level, count in risk_levels.items():
                summary_parts.append(f"- {level}: {count}")
    
    # Top Risks
    if 'statistics' in report_data and 'top_risks' in report_data['statistics']:
        top_risks = report_data['statistics']['top_risks']
        if top_risks:
            if language == 'uk':
                summary_parts.append(f"\nТоп-10 найбільш критичних ризиків:")
            elif language == 'ru':
                summary_parts.append(f"\nТоп-10 наиболее критичных рисков:")
            else:
                summary_parts.append(f"\nTop-10 Most Critical Risks:")
            
            for i, risk in enumerate(top_risks[:10], 1):
                asset_name = risk.get('asset_name', 'Unknown')
                vuln_name = risk.get('vulnerability_name', 'Unknown')
                risk_level = risk.get('risk_level', 'Unknown')
                if language == 'uk':
                    summary_parts.append(f"{i}. {asset_name} - {vuln_name} (Рівень ризику: {risk_level})")
                elif language == 'ru':
                    summary_parts.append(f"{i}. {asset_name} - {vuln_name} (Уровень риска: {risk_level})")
                else:
                    summary_parts.append(f"{i}. {asset_name} - {vuln_name} (Risk Level: {risk_level})")
    
    # Compliance
    if 'compliance' in report_data:
        compliance = report_data['compliance']
        if language == 'uk':
            summary_parts.append(f"\nСтатус відповідності:")
            if 'iso27001' in compliance:
                iso_rate = compliance['iso27001'].get('compliance_rate', 0)
                summary_parts.append(f"- ISO 27001: {iso_rate:.1f}% відповідності")
            if 'pcidss' in compliance:
                pci_rate = compliance['pcidss'].get('compliance_rate', 0)
                summary_parts.append(f"- PCI DSS: {pci_rate:.1f}% відповідності")
        elif language == 'ru':
            summary_parts.append(f"\nСтатус соответствия:")
            if 'iso27001' in compliance:
                iso_rate = compliance['iso27001'].get('compliance_rate', 0)
                summary_parts.append(f"- ISO 27001: {iso_rate:.1f}% соответствия")
            if 'pcidss' in compliance:
                pci_rate = compliance['pcidss'].get('compliance_rate', 0)
                summary_parts.append(f"- PCI DSS: {pci_rate:.1f}% соответствия")
        else:
            summary_parts.append(f"\nCompliance Status:")
            if 'iso27001' in compliance:
                iso_rate = compliance['iso27001'].get('compliance_rate', 0)
                summary_parts.append(f"- ISO 27001: {iso_rate:.1f}% compliance")
            if 'pcidss' in compliance:
                pci_rate = compliance['pcidss'].get('compliance_rate', 0)
                summary_parts.append(f"- PCI DSS: {pci_rate:.1f}% compliance")
    
    # Treatment Details
    if 'statistics' in report_data and 'overdue_treatments' in report_data['statistics']:
        overdue = report_data['statistics']['overdue_treatments']
        if overdue:
            overdue_count = len(overdue)
            if language == 'uk':
                summary_parts.append(f"\nПрострочені заходи з ліквідації ризиків: {overdue_count}")
            elif language == 'ru':
                summary_parts.append(f"\nПросроченные меры по устранению рисков: {overdue_count}")
            else:
                summary_parts.append(f"\nOverdue Risk Treatments: {overdue_count}")
    
    return "\n".join(summary_parts)


def generate_unified_report_content(data, report_type, selected_sections, translations, language=None):
    """
    Generate unified report content structure that can be used by HTML preview, Word, and PDF generators.
    This ensures all formats have identical data and structure.
    """
    
    # Use provided language or fallback to data language or default
    if language is None:
        language = data.get('language', 'uk')
    # Activate the language for proper multilingual support
    from django.utils.translation import activate
    activate(language)
    
    content = {
        'header': {},
        'sections': []
    }
    
    # Header information (always included)
    # Determine title based on profile or report type
    if data.get('profile') and data['profile'].get('name'):
        # Use profile name for title
        title = f"{translations['report_title']} - {data['profile']['name']}"
    else:
        # Fallback to report type (for backward compatibility)
        title = f"{translations['report_title']} - {report_type.title()}"
    
    content['header'] = {
        'title': title,
        'generated_date': format_localized_date(data['generated_date'], data.get('language', 'uk')),
        'report_period': f"{data['date_range']['start']} - {data['date_range']['end']}",
        'generated_by': data.get('generated_by', 'System'),
        'company_info': None
    }
    
    # Company information
    if data.get('companies'):
        if data['companies'].count() == 1:
            content['header']['company_info'] = {
                'type': 'single',
                'name': data['companies'].first().name
            }
        elif data['companies'].count() > 1:
            content['header']['company_info'] = {
                'type': 'multiple',
                'count': data['companies'].count(),
                'text': f"{translations['multiple_companies']} ({data['companies'].count()})"
            }
    
    # Executive Summary Section
    if selected_sections.get('statistics') or selected_sections.get('risk_distribution') or selected_sections.get('compliance_overview') or selected_sections.get('top_risks'):
        exec_summary = {
            'type': 'executive_summary',
            'title': translations['executive_summary'],
            'subsections': []
        }
        
        # Key Statistics
        if selected_sections.get('statistics') and 'statistics' in data:
            stats = data['statistics']
            exec_summary['subsections'].append({
                'type': 'statistics',
                'title': translations.get('key_metrics', 'Key Metrics'),
                'data': [
                    {
                        'label': translations['total_assets'],
                        'value': str(stats.get('total_assets', 0)),
                        'type': 'primary'
                    },
                    {
                        'label': translations['total_vulnerabilities'],
                        'value': str(stats.get('total_vulnerabilities', 0)),
                        'type': 'warning'
                    },
                    {
                        'label': translations['high_risk_assets'],
                        'value': str(stats.get('high_risk_count', 0)),
                        'type': 'danger'
                    },
                    {
                        'label': translations['completion_rate'],
                        'value': f"{stats.get('completion_rate', 0):.1f}%",
                        'type': 'success'
                    }
                ]
            })
        
        # Risk Details (Combined By Risk Level and TOP-10 Risks)
        risk_details_section = None
        
        # Check if either risk_distribution or top_risks is selected
        if (selected_sections.get('risk_distribution') and 'risk_distribution' in data) or (selected_sections.get('top_risks') and 'statistics' in data):
            risk_details_section = {
                'type': 'risk_details',
                'title': translations.get('risk_details', 'Risk Details'),
                'subsections': []
            }
            
            # By Risk Level subsection
            if selected_sections.get('risk_distribution') and 'risk_distribution' in data:
                risk_dist = data['risk_distribution']
                if risk_dist.get('by_level'):
                    total_risks = sum(risk_dist['by_level'].values())
                    risk_level_data = []
                    
                    # Get risk levels with colors from statistics if available
                    risk_levels_with_colors = data.get('statistics', {}).get('risk_levels_with_colors', {})
                    
                    # Create list with all info for sorting
                    for level, count in risk_dist['by_level'].items():
                        percentage = (count / total_risks * 100) if total_risks > 0 else 0
                        color_info = risk_levels_with_colors.get(level, {})
                        risk_level_data.append({
                            'risk_level': level,
                            'count': count,
                            'percentage': f"{percentage:.1f}%",
                            'color': color_info.get('color', '#000000'),
                            'min_value': color_info.get('min_value', 0),
                            'max_value': color_info.get('max_value', 0)
                        })
                    
                    # Sort by max_value in descending order (highest risk first)
                    risk_level_data.sort(key=lambda x: x['max_value'], reverse=True)
                    
                    risk_details_section['subsections'].append({
                        'type': 'risk_distribution_level',
                        'title': translations.get('by_risk_level', 'By Risk Level'),
                        'headers': [translations['risk_level'], translations['count'], translations.get('percentage', 'Percentage')],
                        'data': risk_level_data
                    })
            
            # TOP-10 Risks subsection
            if selected_sections.get('top_risks') and 'statistics' in data:
                stats = data['statistics']
                if stats.get('high_risk_assets'):
                    top_risks_data = []
                    for asset in stats['high_risk_assets'][:10]:  # Limit to 10
                        top_risks_data.append({
                            'asset_name': asset.get('asset_name', translations.get('unknown', 'Unknown')),
                            'vulnerability_name': asset.get('vulnerability_name', translations.get('unknown', 'Unknown')),
                            'risk_level': asset.get('risk_level', translations.get('unknown', 'Unknown'))
                        })
                    
                    risk_details_section['subsections'].append({
                        'type': 'top_risks',
                        'title': translations.get('top_risks', 'TOP-10 Risks'),
                        'headers': [translations.get('asset_name', 'Asset Name'), translations['vulnerability'], translations['risk_level']],
                        'data': top_risks_data
                    })
        
        content['sections'].append(exec_summary)
        
        # Add the combined section to content
        if risk_details_section and risk_details_section.get('subsections'):
            content['sections'].append(risk_details_section)
    
    # Compliance Section
    if (selected_sections.get('compliance_overview') or selected_sections.get('pci_dss') or 
        selected_sections.get('iso27001') or selected_sections.get('compliance_gaps') or
        selected_sections.get('framework_company_requirements') or 
        selected_sections.get('company_requirements') or 
        selected_sections.get('internal_requirements')):
        compliance_section = {
            'type': 'compliance',
            'title': translations['compliance_summary'],
            'subsections': []
        }
        
        # PCI DSS Compliance
        if selected_sections.get('pci_dss') and 'pcidss_compliance' in data:
            pcidss = data['pcidss_compliance']
            compliance_section['subsections'].append({
                'type': 'pci_dss',
                'title': translations['pci_dss_compliance'],
                'data': [
                    {
                        'label': translations['overall_compliance'],
                        'value': f"{pcidss.get('overall_compliance', 0):.1f}%"
                    },
                    {
                        'label': translations['total_requirements'],
                        'value': str(pcidss.get('total_requirements', 0))
                    },
                    {
                        'label': translations['compliant_vulnerabilities'],
                        'value': str(pcidss.get('compliant_vulnerabilities', 0))
                    }
                ]
            })
        
        # ISO 27001 Compliance
        if selected_sections.get('iso27001') and 'iso27001_compliance' in data:
            iso27001 = data['iso27001_compliance']
            compliance_section['subsections'].append({
                'type': 'iso27001',
                'title': translations['iso27001_compliance'],
                'data': [
                    {
                        'label': translations['overall_compliance'],
                        'value': f"{iso27001.get('overall_compliance', 0):.1f}%"
                    },
                    {
                        'label': translations['total_controls'],
                        'value': str(iso27001.get('total_controls', 0))
                    },
                    {
                        'label': translations['compliant_vulnerabilities'],
                        'value': str(iso27001.get('compliant_vulnerabilities', 0))
                    }
                ]
            })
        
        # Compliance Gaps
        if selected_sections.get('compliance_gaps'):
            gaps_data = []
            
            # PCI DSS Gaps
            if 'pcidss_gaps' in data and data['pcidss_gaps']:
                for gap in data['pcidss_gaps'][:10]:  # Limit to 10
                    gaps_data.append({
                        'type': 'PCI DSS',
                        'requirement': gap.get('requirement', translations.get('unknown', 'Unknown')),
                        'description': gap.get('description', translations.get('no_description', 'No description'))
                    })
            
            # ISO 27001 Gaps
            if 'iso27001_gaps' in data and data['iso27001_gaps']:
                for gap in data['iso27001_gaps'][:10]:  # Limit to 10
                    gaps_data.append({
                        'type': 'ISO 27001',
                        'requirement': gap.get('control', translations.get('unknown', 'Unknown')),
                        'description': gap.get('description', translations.get('no_description', 'No description'))
                    })
            
            if gaps_data:
                compliance_section['subsections'].append({
                    'type': 'compliance_gaps',
                    'title': translations.get('compliance_gaps', 'Compliance Gaps'),
                    'headers': [translations.get('type', 'Type'), translations.get('requirement', 'Requirement'), translations['description']],
                    'data': gaps_data
                })
        
        # Framework Company Requirements
        if selected_sections.get('framework_company_requirements') and 'compliance' in data:
            compliance_data = data.get('compliance', {})
            framework_data = compliance_data.get('framework_company_requirements')
            # Check if data exists (even if empty list, we should show the section if enabled)
            if framework_data is not None:
                frameworks_list = []
                if framework_data.get('frameworks') and len(framework_data['frameworks']) > 0:
                    for framework in framework_data['frameworks']:
                        frameworks_list.append({
                            'name': framework.get('name', ''),
                            'framework_type': framework.get('framework_type', ''),
                            'version': framework.get('version', ''),
                            'company': framework.get('company', ''),
                            'status': framework.get('status', ''),
                            'controls_total': framework.get('controls_total', 0),
                            'controls_completed': framework.get('controls_completed', 0),
                            'completion_percentage': f"{framework.get('completion_percentage', 0):.1f}%"
                        })
                
                compliance_section['subsections'].append({
                    'type': 'framework_company_requirements',
                    'title': translations.get('framework_company_requirements', 'Framework Company Requirements'),
                    'headers': [
                        translations.get('name', 'Name'),
                        translations.get('framework_type', 'Framework Type'),
                        translations.get('version', 'Version'),
                        translations.get('company', 'Company'),
                        translations.get('status', 'Status'),
                        translations.get('controls_total', 'Total Controls'),
                        translations.get('controls_completed', 'Completed'),
                        translations.get('completion_percentage', 'Completion %')
                    ],
                    'data': frameworks_list if frameworks_list else [{'name': translations.get('no_data', 'No data available'), 'framework_type': '', 'version': '', 'company': '', 'status': '', 'controls_total': 0, 'controls_completed': 0, 'completion_percentage': '0%'}]
                })
        
        # Local Company Requirements
        if selected_sections.get('company_requirements') and 'compliance' in data:
            compliance_data = data.get('compliance', {})
            local_data = compliance_data.get('company_requirements')
            # Check if data exists (even if empty list, we should show the section if enabled)
            if local_data is not None:
                requirements_list = []
                if local_data.get('requirements') and len(local_data['requirements']) > 0:
                    for req in local_data['requirements']:
                        requirements_list.append({
                            'code': req.get('code', ''),
                            'name': req.get('name', ''),
                            'requirement_type': req.get('requirement_type', ''),
                            'company': req.get('company', ''),
                            'regulator': req.get('regulator', ''),
                            'status': req.get('status', ''),
                            'controls_total': req.get('controls_total', 0),
                            'controls_completed': req.get('controls_completed', 0),
                            'completion_percentage': f"{req.get('completion_percentage', 0):.1f}%"
                        })
                
                compliance_section['subsections'].append({
                    'type': 'company_requirements',
                    'title': translations.get('local_company_requirements', 'Local Company Requirements'),
                    'headers': [
                        translations.get('code', 'Code'),
                        translations.get('name', 'Name'),
                        translations.get('requirement_type', 'Type'),
                        translations.get('company', 'Company'),
                        translations.get('regulator', 'Regulator'),
                        translations.get('status', 'Status'),
                        translations.get('controls_total', 'Total Controls'),
                        translations.get('controls_completed', 'Completed'),
                        translations.get('completion_percentage', 'Completion %')
                    ],
                    'data': requirements_list if requirements_list else [{'code': '', 'name': translations.get('no_data', 'No data available'), 'requirement_type': '', 'company': '', 'regulator': '', 'status': '', 'controls_total': 0, 'controls_completed': 0, 'completion_percentage': '0%'}]
                })
        
        # Internal Company Requirements
        if selected_sections.get('internal_requirements') and 'compliance' in data:
            compliance_data = data.get('compliance', {})
            internal_data = compliance_data.get('internal_requirements')
            # Check if data exists (even if empty list, we should show the section if enabled)
            if internal_data is not None:
                requirements_list = []
                if internal_data.get('requirements') and len(internal_data['requirements']) > 0:
                    for req in internal_data['requirements']:
                        requirements_list.append({
                            'code': req.get('code', ''),
                            'name': req.get('name', ''),
                            'requirement_type': req.get('requirement_type', ''),
                            'company': req.get('company', ''),
                            'status': req.get('status', ''),
                            'controls_total': req.get('controls_total', 0),
                            'controls_completed': req.get('controls_completed', 0),
                            'completion_percentage': f"{req.get('completion_percentage', 0):.1f}%"
                        })
                
                compliance_section['subsections'].append({
                    'type': 'internal_requirements',
                    'title': translations.get('internal_company_requirements', 'Internal Company Requirements'),
                    'headers': [
                        translations.get('code', 'Code'),
                        translations.get('name', 'Name'),
                        translations.get('requirement_type', 'Type'),
                        translations.get('company', 'Company'),
                        translations.get('status', 'Status'),
                        translations.get('controls_total', 'Total Controls'),
                        translations.get('controls_completed', 'Completed'),
                        translations.get('completion_percentage', 'Completion %')
                    ],
                    'data': requirements_list if requirements_list else [{'code': '', 'name': translations.get('no_data', 'No data available'), 'requirement_type': '', 'company': '', 'status': '', 'controls_total': 0, 'controls_completed': 0, 'completion_percentage': '0%'}]
                })
        
        if compliance_section['subsections']:
            content['sections'].append(compliance_section)
    
    # Incident Register Section
    if selected_sections.get('incident') and 'incident_data' in data:
        incident_data = data.get('incident_data')
        # Check if data exists (even if empty list, we should show the section if enabled)
        if incident_data is not None:
            incident_section = {
                'type': 'incident',
                'title': translations.get('incident_register', 'Incident Register'),
                'subsections': []
            }
            
            # Summary statistics
            if incident_data.get('total_incidents', 0) > 0:
                stats_data = [
                    {
                        'label': translations.get('total_incidents', 'Total Incidents'),
                        'value': str(incident_data.get('total_incidents', 0)),
                        'type': 'primary'
                    }
                ]
                
                # Add classification statistics
                if incident_data.get('statistics', {}).get('by_classification'):
                    classification_stats = incident_data['statistics']['by_classification']
                    for classification, count in sorted(classification_stats.items(), key=lambda x: x[1], reverse=True):
                        stats_data.append({
                            'label': f"{translations.get('by_classification', 'By Classification')}: {classification}",
                            'value': str(count),
                            'type': 'info'
                        })
                
                # Add state statistics
                if incident_data.get('statistics', {}).get('by_state'):
                    state_stats = incident_data['statistics']['by_state']
                    for state, count in sorted(state_stats.items(), key=lambda x: x[1], reverse=True):
                        stats_data.append({
                            'label': f"{translations.get('by_state', 'By State')}: {state}",
                            'value': str(count),
                            'type': 'warning'
                        })
                
                incident_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('incident_summary', 'Incident Summary'),
                    'data': stats_data
                })
            
            # Incident table
            incidents_list = []
            if incident_data.get('incidents') and len(incident_data['incidents']) > 0:
                for incident in incident_data['incidents']:
                    incidents_list.append({
                        'id': incident.get('id', ''),
                        'company': incident.get('company', ''),
                        'occurrence_datetime': incident.get('occurrence_datetime', ''),
                        'place': incident.get('place', ''),
                        'description': incident.get('description', '')[:200] + '...' if len(incident.get('description', '')) > 200 else incident.get('description', ''),
                        'classification': incident.get('classification', ''),
                        'incident_type': incident.get('incident_type', ''),
                        'responsible': incident.get('responsible', ''),
                        'reported_by': incident.get('reported_by', ''),
                        'current_state': incident.get('current_state', ''),
                        'impact': incident.get('impact', '')[:200] + '...' if len(incident.get('impact', '')) > 200 else incident.get('impact', ''),
                    })
            
            incident_section['subsections'].append({
                'type': 'incident_table',
                'title': translations.get('incidents', 'Incidents'),
                'headers': [
                    translations.get('id', 'ID'),
                    translations.get('company', 'Company'),
                    translations.get('occurrence_datetime', 'Occurrence Date'),
                    translations.get('place', 'Place'),
                    translations.get('incident_type', 'Type'),
                    translations.get('classification', 'Classification'),
                    translations.get('current_state', 'State'),
                    translations.get('responsible', 'Responsible'),
                    translations.get('description', 'Description')
                ],
                'data': incidents_list if incidents_list else [{'id': '', 'company': translations.get('no_data', 'No data available'), 'occurrence_datetime': '', 'place': '', 'incident_type': '', 'classification': '', 'current_state': '', 'responsible': '', 'description': ''}]
            })
            
            if incident_section['subsections']:
                content['sections'].append(incident_section)
    
    # Mandatory Processes Section
    if selected_sections.get('mandatory_processes') and 'mandatory_processes_data' in data:
        mandatory_processes_data = data.get('mandatory_processes_data')
        # Check if data exists (even if empty list, we should show the section if enabled)
        if mandatory_processes_data is not None:
            mandatory_processes_section = {
                'type': 'mandatory_processes',
                'title': translations.get('mandatory_processes', 'Mandatory Processes'),
                'subsections': []
            }
            
            # Summary statistics
            if mandatory_processes_data.get('total_processes', 0) > 0:
                stats_data = [
                    {
                        'label': translations.get('total_processes', 'Total Processes'),
                        'value': str(mandatory_processes_data.get('total_processes', 0)),
                        'type': 'primary'
                    }
                ]
                
                # Add status statistics
                if mandatory_processes_data.get('statistics', {}).get('by_status'):
                    status_stats = mandatory_processes_data['statistics']['by_status']
                    if status_stats.get('overdue', 0) > 0:
                        stats_data.append({
                            'label': translations.get('overdue', 'Overdue'),
                            'value': str(status_stats.get('overdue', 0)),
                            'type': 'danger'
                        })
                    if status_stats.get('upcoming', 0) > 0:
                        stats_data.append({
                            'label': translations.get('upcoming', 'Upcoming'),
                            'value': str(status_stats.get('upcoming', 0)),
                            'type': 'warning'
                        })
                    if status_stats.get('completed', 0) > 0:
                        stats_data.append({
                            'label': translations.get('completed', 'Completed'),
                            'value': str(status_stats.get('completed', 0)),
                            'type': 'success'
                        })
                
                # Add priority statistics
                if mandatory_processes_data.get('statistics', {}).get('by_priority'):
                    priority_stats = mandatory_processes_data['statistics']['by_priority']
                    for priority, count in sorted(priority_stats.items(), key=lambda x: x[1], reverse=True):
                        priority_display = priority.capitalize()
                        stats_data.append({
                            'label': f"{translations.get('by_priority', 'By Priority')}: {priority_display}",
                            'value': str(count),
                            'type': 'info'
                        })
                
                mandatory_processes_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('mandatory_processes_summary', 'Mandatory Processes Summary'),
                    'data': stats_data,
                    'display_inline': True  # Flag to display all blocks in one row
                })
            
            # Processes table - show only last 5 processes sorted by Next Due Date
            processes_list = []
            if mandatory_processes_data.get('processes') and len(mandatory_processes_data['processes']) > 0:
                # Sort processes by next_due_date (None values go to the end)
                sorted_processes = sorted(
                    mandatory_processes_data['processes'],
                    key=lambda x: (
                        x.get('next_due_date') if x.get('next_due_date') else '9999-12-31',
                        x.get('process_name', '')
                    )
                )
                # Take only first 5 processes (with nearest due dates)
                for process in sorted_processes[:5]:
                    processes_list.append({
                        'process_name': process.get('process_name', ''),
                        'company': process.get('company', ''),
                        'description': process.get('description', '')[:200] + '...' if len(process.get('description', '')) > 200 else process.get('description', ''),
                        'priority': process.get('priority', ''),
                        'frequency': process.get('frequency', ''),
                        'next_due_date': process.get('next_due_date', ''),
                        'last_completed_date': process.get('last_completed_date', ''),
                        'responsible_person': process.get('responsible_person', ''),
                        'status': process.get('status', ''),
                        'source_document': process.get('source_document', ''),
                    })
            
            mandatory_processes_section['subsections'].append({
                'type': 'mandatory_processes_table',
                'title': translations.get('last_5_mandatory_processes', 'Last 5 Mandatory Processes'),
                'headers': [
                    translations.get('process_name', 'Process Name'),
                    translations.get('company', 'Company'),
                    translations.get('priority', 'Priority'),
                    translations.get('frequency', 'Frequency'),
                    translations.get('next_due_date', 'Next Due Date'),
                    translations.get('last_completed_date', 'Last Completed Date'),
                    translations.get('responsible_person', 'Responsible Person'),
                    translations.get('status', 'Status'),
                    translations.get('description', 'Description')
                ],
                'data': processes_list if processes_list else [{'process_name': translations.get('no_data', 'No data available'), 'company': '', 'priority': '', 'frequency': '', 'next_due_date': '', 'last_completed_date': '', 'responsible_person': '', 'status': '', 'description': ''}]
            })
            
            if mandatory_processes_section['subsections']:
                content['sections'].append(mandatory_processes_section)
    
    # Certificate & Key Management Section
    if selected_sections.get('certificate_key_management') and 'certificate_key_management_data' in data:
        certificate_key_management_data = data.get('certificate_key_management_data')
        # Check if data exists (even if empty list, we should show the section if enabled)
        if certificate_key_management_data is not None:
            certificate_key_management_section = {
                'type': 'certificate_key_management',
                'title': translations.get('certificate_key_management', 'Certificate & Key Management'),
                'subsections': []
            }
            
            # Summary statistics
            if certificate_key_management_data.get('total_certificates', 0) > 0:
                stats_data = [
                    {
                        'label': translations.get('total_certificates', 'Total Certificates'),
                        'value': str(certificate_key_management_data.get('total_certificates', 0)),
                        'type': 'primary'
                    }
                ]
                
                # Add expiry status statistics
                if certificate_key_management_data.get('statistics', {}).get('by_expiry_status'):
                    expiry_stats = certificate_key_management_data['statistics']['by_expiry_status']
                    if expiry_stats.get('expired', 0) > 0:
                        stats_data.append({
                            'label': translations.get('expired', 'Expired'),
                            'value': str(expiry_stats.get('expired', 0)),
                            'type': 'danger'
                        })
                    if expiry_stats.get('expiring_soon', 0) > 0:
                        stats_data.append({
                            'label': translations.get('expiring_soon', 'Expiring Soon (≤30 days)'),
                            'value': str(expiry_stats.get('expiring_soon', 0)),
                            'type': 'warning'
                        })
                    if expiry_stats.get('valid', 0) > 0:
                        stats_data.append({
                            'label': translations.get('valid', 'Valid'),
                            'value': str(expiry_stats.get('valid', 0)),
                            'type': 'success'
                        })
                
                # Add type statistics
                if certificate_key_management_data.get('statistics', {}).get('by_type'):
                    type_stats = certificate_key_management_data['statistics']['by_type']
                    for cert_type, count in sorted(type_stats.items(), key=lambda x: x[1], reverse=True)[:5]:
                        stats_data.append({
                            'label': f"{translations.get('by_type', 'By Type')}: {cert_type}",
                            'value': str(count),
                            'type': 'info'
                        })
                
                certificate_key_management_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('certificate_key_management_summary', 'Certificate & Key Management Summary'),
                    'data': stats_data,
                    'display_inline': True  # Flag to display all blocks in one row
                })
            
            # Certificates table - show only last 5 certificates sorted by Expiry Date
            certificates_list = []
            if certificate_key_management_data.get('certificates') and len(certificate_key_management_data['certificates']) > 0:
                # Sort certificates by expiry_date (expiring soon first, then expired, then valid)
                sorted_certificates = sorted(
                    certificate_key_management_data['certificates'],
                    key=lambda x: (
                        x.get('expiry_status') == 'expired' and 0 or (x.get('expiry_status') == 'expiring_soon' and 1 or 2),
                        x.get('expiry_date', '9999-12-31'),
                        x.get('key_cert_num', '')
                    )
                )
                # Take only first 5 certificates (expiring soon/expired first)
                for cert in sorted_certificates[:5]:
                    certificates_list.append({
                        'key_cert_num': cert.get('key_cert_num', ''),
                        'company': cert.get('company', ''),
                        'type': cert.get('type', ''),
                        'purpose': cert.get('purpose', '')[:200] + '...' if len(cert.get('purpose', '')) > 200 else cert.get('purpose', ''),
                        'expiry_date': cert.get('expiry_date', ''),
                        'expiry_status': cert.get('expiry_status', ''),
                        'days_until_expiry': cert.get('days_until_expiry'),
                        'revocation_status': cert.get('revocation_status', ''),
                        'owner_name': cert.get('owner_name', ''),
                        'location': cert.get('location', ''),
                    })
            
            certificate_key_management_section['subsections'].append({
                'type': 'certificate_key_management_table',
                'title': translations.get('last_5_certificates', 'Last 5 Certificates (by Expiry Date)'),
                'headers': [
                    translations.get('key_cert_num', 'Key/Cert Number'),
                    translations.get('company', 'Company'),
                    translations.get('type', 'Type'),
                    translations.get('expiry_date', 'Expiry Date'),
                    translations.get('expiry_status', 'Status'),
                    translations.get('days_until_expiry', 'Days Until Expiry'),
                    translations.get('revocation_status', 'Revocation Status'),
                    translations.get('owner', 'Owner'),
                    translations.get('location', 'Location'),
                    translations.get('purpose', 'Purpose')
                ],
                'data': certificates_list
            })
            
            content['sections'].append(certificate_key_management_section)
    
    # Quiz Results Section
    if selected_sections.get('quiz_results') and 'quiz_results_data' in data:
        quiz_results_data = data.get('quiz_results_data')
        # Check if data exists (even if empty, we should show the section if enabled)
        if quiz_results_data is not None:
            quiz_results_section = {
                'type': 'quiz_results',
                'title': translations.get('quiz_results', 'Security Training & Quiz Results'),
                'subsections': []
            }
            
            # Summary statistics
            if quiz_results_data.get('total_attempts', 0) > 0:
                stats_data = [
                    {
                        'label': translations.get('total_attempts', 'Total Attempts'),
                        'value': str(quiz_results_data.get('total_attempts', 0)),
                        'type': 'primary'
                    },
                    {
                        'label': translations.get('success_rate', 'Success Rate'),
                        'value': f"{quiz_results_data.get('success_rate', 0):.1f}%",
                        'type': 'success' if quiz_results_data.get('success_rate', 0) >= 80 else 'warning'
                    },
                    {
                        'label': translations.get('average_score', 'Average Score'),
                        'value': f"{quiz_results_data.get('average_score', 0):.1f}",
                        'type': 'info'
                    },
                    {
                        'label': translations.get('failed_attempts', 'Failed Attempts'),
                        'value': str(quiz_results_data.get('failed_attempts', 0)),
                        'type': 'danger' if quiz_results_data.get('failed_attempts', 0) > 0 else 'success'
                    }
                ]
                
                quiz_results_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('training_summary', 'Training Summary'),
                    'data': stats_data,
                    'display_inline': True  # Display all blocks in one row
                })
            
            # Quiz Statistics Table
            if quiz_results_data.get('quiz_statistics') and len(quiz_results_data['quiz_statistics']) > 0:
                quiz_stats_list = []
                for quiz_stat in quiz_results_data['quiz_statistics']:
                    quiz_stats_list.append({
                        'quiz_title': quiz_stat.get('quiz_title', ''),
                        'total_attempts': str(quiz_stat.get('total_attempts', 0)),
                        'successful_attempts': str(quiz_stat.get('successful_attempts', 0)),
                        'failed_attempts': str(quiz_stat.get('failed_attempts', 0)),
                        'success_rate': f"{quiz_stat.get('success_rate', 0):.1f}%",
                        'average_score': f"{quiz_stat.get('average_score', 0):.1f}",
                        'passing_score': str(quiz_stat.get('passing_score', 0))
                    })
                
                quiz_results_section['subsections'].append({
                    'type': 'quiz_statistics_table',
                    'title': translations.get('quiz_statistics', 'Quiz Statistics'),
                    'headers': [
                        translations.get('quiz_title', 'Quiz Title'),
                        translations.get('total_attempts', 'Total Attempts'),
                        translations.get('successful_attempts', 'Successful'),
                        translations.get('failed_attempts', 'Failed'),
                        translations.get('success_rate', 'Success Rate'),
                        translations.get('average_score', 'Average Score'),
                        translations.get('passing_score', 'Passing Score')
                    ],
                    'data': quiz_stats_list if quiz_stats_list else [{'quiz_title': translations.get('no_data', 'No data available'), 'total_attempts': '', 'successful_attempts': '', 'failed_attempts': '', 'success_rate': '', 'average_score': '', 'passing_score': ''}]
                })
            
            # Users at Risk Table
            if quiz_results_data.get('users_at_risk') and len(quiz_results_data['users_at_risk']) > 0:
                users_at_risk_list = []
                for user_risk in quiz_results_data['users_at_risk']:
                    users_at_risk_list.append({
                        'user': user_risk.get('user', ''),
                        'company': user_risk.get('company', ''),
                        'attempts_count': str(user_risk.get('attempts_count', 0))
                    })
                
                quiz_results_section['subsections'].append({
                    'type': 'users_at_risk_table',
                    'title': translations.get('users_requiring_training', 'Users Requiring Additional Training'),
                    'headers': [
                        translations.get('user', 'User'),
                        translations.get('company', 'Company'),
                        translations.get('attempts_count', 'Attempts Count')
                    ],
                    'data': users_at_risk_list
                })
            
            if quiz_results_section['subsections']:
                content['sections'].append(quiz_results_section)
    
    # Third Party Risk Section
    if selected_sections.get('third_party_risk') and 'third_party_risk_data' in data:
        third_party_risk_data = data.get('third_party_risk_data')
        # Check if data exists (even if empty, we should show the section if enabled)
        if third_party_risk_data is not None:
            third_party_risk_section = {
                'type': 'third_party_risk',
                'title': translations.get('third_party_risk', 'Third Party Risk'),
                'subsections': []
            }
            
            # Vendor Risk Summary - Statistics
            if third_party_risk_data.get('total_vendors', 0) >= 0:
                stats_data = [
                    {
                        'label': translations.get('total_vendors', 'Total Vendors'),
                        'value': str(third_party_risk_data.get('total_vendors', 0)),
                        'type': 'primary'
                    },
                    {
                        'label': translations.get('high_risk_vendors', 'High Risk Vendors'),
                        'value': str(third_party_risk_data.get('high_risk_vendors_count', 0)),
                        'type': 'danger' if third_party_risk_data.get('high_risk_vendors_count', 0) > 0 else 'success'
                    }
                ]
                
                # Add risk distribution
                risk_dist = third_party_risk_data.get('risk_distribution', {})
                if risk_dist:
                    for risk_level in ['critical', 'high', 'medium', 'low']:
                        count = risk_dist.get(risk_level, 0)
                        if count > 0:
                            risk_level_display = risk_level.capitalize()
                            stats_data.append({
                                'label': f"{translations.get('risk_level', 'Risk Level')}: {risk_level_display}",
                                'value': str(count),
                                'type': 'warning' if risk_level in ['high', 'critical'] else 'info'
                            })
                
                third_party_risk_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('vendor_risk_summary', 'Vendor Risk Summary'),
                    'data': stats_data,
                    'display_inline': True
                })
            
            # Risk Distribution Table
            risk_dist = third_party_risk_data.get('risk_distribution', {})
            if risk_dist and sum(risk_dist.values()) > 0:
                risk_dist_list = []
                total = sum(risk_dist.values())
                for risk_level in ['low', 'medium', 'high', 'critical']:
                    count = risk_dist.get(risk_level, 0)
                    if count > 0 or total == 0:
                        percentage = (count / total * 100) if total > 0 else 0
                        risk_dist_list.append({
                            'risk_level': risk_level.capitalize(),
                            'count': str(count),
                            'percentage': f"{percentage:.1f}%"
                        })
                
                third_party_risk_section['subsections'].append({
                    'type': 'risk_distribution_table',
                    'title': translations.get('risk_distribution_by_level', 'Risk Distribution by Level'),
                    'headers': [
                        translations.get('risk_level', 'Risk Level'),
                        translations.get('count', 'Count'),
                        translations.get('percentage', 'Percentage')
                    ],
                    'data': risk_dist_list if risk_dist_list else [{'risk_level': translations.get('no_data', 'No data available'), 'count': '', 'percentage': ''}]
                })
            
            # Assessment Statuses Table
            assessment_statuses = third_party_risk_data.get('assessment_statuses', {})
            if assessment_statuses:
                status_list = []
                for status_key in ['pending', 'in_progress', 'completed', 'overdue']:
                    count = assessment_statuses.get(status_key, 0)
                    status_display = {
                        'pending': translations.get('pending', 'Pending'),
                        'in_progress': translations.get('in_progress', 'In Progress'),
                        'completed': translations.get('completed', 'Completed'),
                        'overdue': translations.get('overdue', 'Overdue')
                    }.get(status_key, status_key.capitalize())
                    
                    status_list.append({
                        'status': status_display,
                        'count': str(count)
                    })
                
                third_party_risk_section['subsections'].append({
                    'type': 'assessment_status_table',
                    'title': translations.get('assessment_status', 'Assessment Status'),
                    'headers': [
                        translations.get('status', 'Status'),
                        translations.get('count', 'Count')
                    ],
                    'data': status_list
                })
            
            # Overdue Assessments Table
            overdue_assessments = third_party_risk_data.get('overdue_assessments', [])
            if overdue_assessments:
                overdue_list = []
                for assessment in overdue_assessments:
                    overdue_list.append({
                        'vendor_name': assessment.get('vendor_name', ''),
                        'assessment_date': assessment.get('assessment_date', ''),
                        'next_review_date': assessment.get('next_review_date', ''),
                        'status': assessment.get('status', ''),
                        'days_overdue': str(assessment.get('days_overdue', 0))
                    })
                
                third_party_risk_section['subsections'].append({
                    'type': 'overdue_assessments_table',
                    'title': translations.get('overdue_assessments', 'Overdue Assessments'),
                    'headers': [
                        translations.get('vendor_name', 'Vendor Name'),
                        translations.get('assessment_date', 'Assessment Date'),
                        translations.get('next_review_date', 'Next Review Date'),
                        translations.get('status', 'Status'),
                        translations.get('days_overdue', 'Days Overdue')
                    ],
                    'data': overdue_list
                })
            
            # Incomplete Questionnaires Table
            incomplete_questionnaires = third_party_risk_data.get('incomplete_questionnaires', [])
            if incomplete_questionnaires:
                incomplete_list = []
                for questionnaire in incomplete_questionnaires:
                    incomplete_list.append({
                        'vendor_name': questionnaire.get('vendor_name', ''),
                        'template_name': questionnaire.get('template_name', ''),
                        'status': questionnaire.get('status', ''),
                        'started_date': questionnaire.get('started_date', ''),
                        'percentage_score': f"{questionnaire.get('percentage_score', 0):.1f}%"
                    })
                
                third_party_risk_section['subsections'].append({
                    'type': 'incomplete_questionnaires_table',
                    'title': translations.get('incomplete_questionnaires', 'Incomplete Questionnaires'),
                    'headers': [
                        translations.get('vendor_name', 'Vendor Name'),
                        translations.get('template_name', 'Template Name'),
                        translations.get('status', 'Status'),
                        translations.get('started_date', 'Started Date'),
                        translations.get('percentage_score', 'Score')
                    ],
                    'data': incomplete_list
                })
            
            # Risk Level Trends
            risk_level_trends = third_party_risk_data.get('risk_level_trends', [])
            if risk_level_trends:
                trends_list = []
                for trend in risk_level_trends:
                    trends_list.append({
                        'risk_level': trend.get('risk_level', '').capitalize(),
                        'count': str(trend.get('count', 0)),
                        'trend': trend.get('trend', 'stable').capitalize()
                    })
                
                third_party_risk_section['subsections'].append({
                    'type': 'risk_level_trends_table',
                    'title': translations.get('risk_level_trends', 'Risk Level Trends'),
                    'headers': [
                        translations.get('risk_level', 'Risk Level'),
                        translations.get('count', 'Count'),
                        translations.get('trend', 'Trend')
                    ],
                    'data': trends_list
                })
            
            if third_party_risk_section['subsections']:
                content['sections'].append(third_party_risk_section)
    
    # GDPR Compliance Section
    if selected_sections.get('gdpr_compliance') and 'gdpr_compliance_data' in data:
        gdpr_compliance_data = data.get('gdpr_compliance_data')
        # Check if data exists (even if empty, we should show the section if enabled)
        if gdpr_compliance_data is not None:
            gdpr_compliance_section = {
                'type': 'gdpr_compliance',
                'title': translations.get('gdpr_compliance', 'GDPR Compliance'),
                'subsections': []
            }
            
            # ===== DATA BREACH INCIDENTS =====
            # Breach Summary Statistics
            if gdpr_compliance_data.get('total_breaches', 0) >= 0:
                breach_stats_data = [
                    {
                        'label': translations.get('total_breaches', 'Total Breaches'),
                        'value': str(gdpr_compliance_data.get('total_breaches', 0)),
                        'type': 'primary'
                    },
                    {
                        'label': translations.get('affected_subjects', 'Affected Subjects'),
                        'value': str(gdpr_compliance_data.get('total_affected_subjects', 0)),
                        'type': 'danger' if gdpr_compliance_data.get('total_affected_subjects', 0) > 0 else 'success'
                    }
                ]
                
                # Add severity distribution
                severity_dist = gdpr_compliance_data.get('severity_distribution', {})
                if severity_dist:
                    for severity in ['critical', 'high', 'medium', 'low']:
                        count = severity_dist.get(severity, 0)
                        if count > 0:
                            severity_display = severity.capitalize()
                            breach_stats_data.append({
                                'label': f"{translations.get('severity', 'Severity')}: {severity_display}",
                                'value': str(count),
                                'type': 'danger' if severity in ['high', 'critical'] else 'warning' if severity == 'medium' else 'info'
                            })
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('data_breach_incidents', 'Data Breach Incidents'),
                    'data': breach_stats_data,
                    'display_inline': True
                })
            
            # Severity Distribution Table
            severity_dist = gdpr_compliance_data.get('severity_distribution', {})
            if severity_dist and sum(severity_dist.values()) > 0:
                severity_dist_list = []
                total = sum(severity_dist.values())
                for severity in ['low', 'medium', 'high', 'critical']:
                    count = severity_dist.get(severity, 0)
                    if count > 0 or total == 0:
                        percentage = (count / total * 100) if total > 0 else 0
                        severity_dist_list.append({
                            'severity': severity.capitalize(),
                            'count': str(count),
                            'percentage': f"{percentage:.1f}%"
                        })
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'breach_severity_table',
                    'title': translations.get('breach_severity_distribution', 'Breach Severity Distribution'),
                    'headers': [
                        translations.get('severity', 'Severity'),
                        translations.get('count', 'Count'),
                        translations.get('percentage', 'Percentage')
                    ],
                    'data': severity_dist_list if severity_dist_list else [{'severity': translations.get('no_data', 'No data available'), 'count': '', 'percentage': ''}]
                })
            
            # Notification Statuses Table
            notification_statuses = gdpr_compliance_data.get('notification_statuses', {})
            if notification_statuses:
                notification_list = []
                for status_key in ['reported_on_time', 'reported_late', 'not_reported_overdue', 'not_reported_within_deadline']:
                    count = notification_statuses.get(status_key, 0)
                    status_display = {
                        'reported_on_time': translations.get('reported_on_time', 'Reported On Time'),
                        'reported_late': translations.get('reported_late', 'Reported Late'),
                        'not_reported_overdue': translations.get('not_reported_overdue', 'Not Reported (Overdue)'),
                        'not_reported_within_deadline': translations.get('not_reported_within_deadline', 'Not Reported (Within Deadline)')
                    }.get(status_key, status_key.replace('_', ' ').title())
                    
                    notification_list.append({
                        'status': status_display,
                        'count': str(count)
                    })
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'notification_status_table',
                    'title': translations.get('notification_statuses', 'Notification Statuses (GDPR 72-hour deadlines)'),
                    'headers': [
                        translations.get('status', 'Status'),
                        translations.get('count', 'Count')
                    ],
                    'data': notification_list
                })
            
            # Breach Notification Details (overdue/late)
            breach_notification_details = gdpr_compliance_data.get('breach_notification_details', [])
            if breach_notification_details:
                notification_details_list = []
                for breach in breach_notification_details:
                    notification_details_list.append({
                        'incident_number': breach.get('incident_number', ''),
                        'title': breach.get('title', ''),
                        'incident_date': breach.get('incident_date', ''),
                        'discovery_date': breach.get('discovery_date', ''),
                        'notification_deadline': breach.get('notification_deadline', ''),
                        'reported': translations.get('yes', 'Yes') if breach.get('reported_to_authority') else translations.get('no', 'No'),
                        'authority_report_date': breach.get('authority_report_date', ''),
                        'days_overdue': str(breach.get('days_overdue', 0))
                    })
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'breach_notification_details_table',
                    'title': translations.get('breach_notification_details', 'Breach Notification Details (Overdue/Late)'),
                    'headers': [
                        translations.get('incident_number', 'Incident Number'),
                        translations.get('title', 'Title'),
                        translations.get('incident_date', 'Incident Date'),
                        translations.get('discovery_date', 'Discovery Date'),
                        translations.get('notification_deadline', 'Notification Deadline'),
                        translations.get('reported', 'Reported'),
                        translations.get('authority_report_date', 'Authority Report Date'),
                        translations.get('days_overdue', 'Days Overdue')
                    ],
                    'data': notification_details_list
                })
            
            # ===== DATA PROCESSING ACTIVITIES =====
            # Activities Summary Statistics
            if gdpr_compliance_data.get('total_active_activities', 0) >= 0:
                activities_stats_data = [
                    {
                        'label': translations.get('active_activities', 'Active Activities'),
                        'value': str(gdpr_compliance_data.get('total_active_activities', 0)),
                        'type': 'primary'
                    },
                    {
                        'label': translations.get('international_transfers', 'International Transfers'),
                        'value': str(gdpr_compliance_data.get('international_transfers_count', 0)),
                        'type': 'warning' if gdpr_compliance_data.get('international_transfers_count', 0) > 0 else 'success'
                    },
                    {
                        'label': translations.get('outdated_retention', 'Outdated Retention'),
                        'value': str(len(gdpr_compliance_data.get('outdated_retention_activities', []))),
                        'type': 'danger' if len(gdpr_compliance_data.get('outdated_retention_activities', [])) > 0 else 'success'
                    },
                    {
                        'label': translations.get('non_compliant_retention', 'Non-Compliant Retention'),
                        'value': str(len(gdpr_compliance_data.get('non_compliant_retention_activities', []))),
                        'type': 'danger' if len(gdpr_compliance_data.get('non_compliant_retention_activities', [])) > 0 else 'success'
                    }
                ]
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('data_processing_activities', 'Data Processing Activities'),
                    'data': activities_stats_data,
                    'display_inline': True
                })
            
            # Data Categories Table
            data_categories = gdpr_compliance_data.get('data_categories', [])
            if data_categories:
                categories_list = []
                for category in data_categories:
                    categories_list.append({
                        'category': category.get('category', ''),
                        'count': str(category.get('count', 0))
                    })
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'data_categories_table',
                    'title': translations.get('data_categories', 'Data Categories'),
                    'headers': [
                        translations.get('category', 'Category'),
                        translations.get('count', 'Count')
                    ],
                    'data': categories_list
                })
            
            # International Transfers Table
            international_transfers_list = gdpr_compliance_data.get('international_transfers_list', [])
            if international_transfers_list:
                transfers_list = []
                for transfer in international_transfers_list:
                    transfers_list.append({
                        'activity_name': transfer.get('activity_name', ''),
                        'company': transfer.get('company', ''),
                        'transfer_safeguards': transfer.get('transfer_safeguards', '')
                    })
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'international_transfers_table',
                    'title': translations.get('international_data_transfers', 'International Data Transfers (Risk)'),
                    'headers': [
                        translations.get('activity_name', 'Activity Name'),
                        translations.get('company', 'Company'),
                        translations.get('transfer_safeguards', 'Transfer Safeguards')
                    ],
                    'data': transfers_list
                })
            
            # Outdated Retention Periods Table
            outdated_retention = gdpr_compliance_data.get('outdated_retention_activities', [])
            if outdated_retention:
                outdated_list = []
                for activity in outdated_retention:
                    outdated_list.append({
                        'activity_name': activity.get('activity_name', ''),
                        'company': activity.get('company', ''),
                        'retention_period_days': str(activity.get('retention_period_days', 0)),
                        'created_date': activity.get('created_date', ''),
                        'updated_date': activity.get('updated_date', ''),
                        'days_overdue': str(activity.get('days_overdue', 0))
                    })
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'outdated_retention_table',
                    'title': translations.get('outdated_retention_periods', 'Outdated Retention Periods'),
                    'headers': [
                        translations.get('activity_name', 'Activity Name'),
                        translations.get('company', 'Company'),
                        translations.get('retention_period_days', 'Retention Period (days)'),
                        translations.get('created_date', 'Created Date'),
                        translations.get('updated_date', 'Updated Date'),
                        translations.get('days_overdue', 'Days Overdue')
                    ],
                    'data': outdated_list
                })
            
            # Non-Compliant Retention Periods Table
            non_compliant_retention = gdpr_compliance_data.get('non_compliant_retention_activities', [])
            if non_compliant_retention:
                non_compliant_list = []
                for activity in non_compliant_retention:
                    issue = activity.get('issue', '')
                    issue_display = {
                        'excessive_retention': translations.get('excessive_retention', 'Excessive Retention (>10 years)'),
                        'insufficient_retention': translations.get('insufficient_retention', 'Insufficient Retention (<30 days)')
                    }.get(issue, issue.replace('_', ' ').title())
                    
                    non_compliant_list.append({
                        'activity_name': activity.get('activity_name', ''),
                        'company': activity.get('company', ''),
                        'retention_period_days': str(activity.get('retention_period_days', 0)),
                        'issue': issue_display
                    })
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'non_compliant_retention_table',
                    'title': translations.get('non_compliant_retention_periods', 'Non-Compliant Retention Periods'),
                    'headers': [
                        translations.get('activity_name', 'Activity Name'),
                        translations.get('company', 'Company'),
                        translations.get('retention_period_days', 'Retention Period (days)'),
                        translations.get('issue', 'Issue')
                    ],
                    'data': non_compliant_list
                })
            
            if gdpr_compliance_section['subsections']:
                content['sections'].append(gdpr_compliance_section)
    
    # Access Risk Summary Section
    if selected_sections.get('access_risk_summary') and 'access_risk_summary_data' in data:
        access_risk_summary_data = data.get('access_risk_summary_data')
        # Check if data exists (even if empty, we should show the section if enabled)
        if access_risk_summary_data is not None:
            access_risk_section = {
                'type': 'access_risk_summary',
                'title': translations.get('access_risk_summary', 'Access Risk Summary'),
                'subsections': []
            }
            
            # Summary Statistics
            if access_risk_summary_data.get('total_active_accesses', 0) >= 0:
                stats_data = [
                    {
                        'label': translations.get('total_active_accesses', 'Total Active Accesses'),
                        'value': str(access_risk_summary_data.get('total_active_accesses', 0)),
                        'type': 'primary'
                    },
                    {
                        'label': translations.get('overdue_reviews', 'Overdue Access Reviews'),
                        'value': str(access_risk_summary_data.get('overdue_reviews_count', 0)),
                        'type': 'danger' if access_risk_summary_data.get('overdue_reviews_count', 0) > 0 else 'success'
                    },
                    {
                        'label': translations.get('third_party_access', 'Third-Party Access'),
                        'value': str(access_risk_summary_data.get('third_party_access_count', 0)),
                        'type': 'warning' if access_risk_summary_data.get('third_party_access_count', 0) > 0 else 'info'
                    },
                    {
                        'label': translations.get('privileged_access', 'Privileged Access'),
                        'value': str(access_risk_summary_data.get('privileged_access_count', 0)),
                        'type': 'danger' if access_risk_summary_data.get('privileged_access_count', 0) > 0 else 'info'
                    }
                ]
                
                access_risk_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('access_summary_statistics', 'Access Summary Statistics'),
                    'data': stats_data,
                    'display_inline': True
                })
            
            # Overdue Access Reviews Table
            overdue_reviews = access_risk_summary_data.get('overdue_reviews', [])
            if overdue_reviews:
                overdue_list = []
                for review in overdue_reviews:
                    overdue_list.append({
                        'asset_name': review.get('asset_name', ''),
                        'access_description': review.get('access_description', ''),
                        'last_review': review.get('last_review', '') or translations.get('never', 'Never'),
                        'reviewed_by': review.get('reviewed_by', ''),
                        'days_overdue': str(review.get('days_overdue', 0))
                    })
                
                access_risk_section['subsections'].append({
                    'type': 'overdue_reviews_table',
                    'title': translations.get('overdue_access_reviews', 'Overdue Access Reviews'),
                    'headers': [
                        translations.get('asset_name', 'Asset Name'),
                        translations.get('access_description', 'Access Description'),
                        translations.get('last_review', 'Last Review'),
                        translations.get('reviewed_by', 'Reviewed By'),
                        translations.get('days_overdue', 'Days Overdue')
                    ],
                    'data': overdue_list
                })
            
            # Third-Party Access Table
            third_party_access = access_risk_summary_data.get('third_party_access', [])
            if third_party_access:
                third_party_list = []
                for access in third_party_access:
                    third_party_list.append({
                        'asset_name': access.get('asset_name', ''),
                        'access_description': access.get('access_description', ''),
                        'last_review': access.get('last_review', '') or translations.get('never', 'Never'),
                        'start_date': access.get('start_date', '')
                    })
                
                access_risk_section['subsections'].append({
                    'type': 'third_party_access_table',
                    'title': translations.get('third_party_access_higher_risk', 'Third-Party Access (Higher Risk)'),
                    'headers': [
                        translations.get('asset_name', 'Asset Name'),
                        translations.get('access_description', 'Access Description'),
                        translations.get('last_review', 'Last Review'),
                        translations.get('start_date', 'Start Date')
                    ],
                    'data': third_party_list
                })
            
            # Access Without Last Review Table
            access_without_review = access_risk_summary_data.get('access_without_review', [])
            if access_without_review:
                without_review_list = []
                for access in access_without_review:
                    without_review_list.append({
                        'asset_name': access.get('asset_name', ''),
                        'access_description': access.get('access_description', ''),
                        'start_date': access.get('start_date', ''),
                        'days_since_start': str(access.get('days_since_start', 0))
                    })
                
                access_risk_section['subsections'].append({
                    'type': 'access_without_review_table',
                    'title': translations.get('access_without_last_review', 'Access Without Last Review'),
                    'headers': [
                        translations.get('asset_name', 'Asset Name'),
                        translations.get('access_description', 'Access Description'),
                        translations.get('start_date', 'Start Date'),
                        translations.get('days_since_start', 'Days Since Start')
                    ],
                    'data': without_review_list
                })
            
            # Privileged Access Section
            privileged_access = access_risk_summary_data.get('privileged_access', [])
            if privileged_access:
                privileged_list = []
                for access in privileged_access:
                    privileged_list.append({
                        'asset_name': access.get('asset_name', ''),
                        'access_description': access.get('access_description', ''),
                        'access_right': access.get('access_right', ''),
                        'roles': access.get('roles', ''),
                        'last_review': access.get('last_review', '') or translations.get('never', 'Never')
                    })
                
                access_risk_section['subsections'].append({
                    'type': 'privileged_access_table',
                    'title': translations.get('privileged_access_high_privileges', 'Privileged Access - High Privileges'),
                    'headers': [
                        translations.get('asset_name', 'Asset Name'),
                        translations.get('access_description', 'Access Description'),
                        translations.get('access_right', 'Access Right'),
                        translations.get('roles', 'Roles'),
                        translations.get('last_review', 'Last Review')
                    ],
                    'data': privileged_list
                })
            
            # Privileged Access Without Regular Reviews Table
            privileged_without_review = access_risk_summary_data.get('privileged_without_review', [])
            if privileged_without_review:
                privileged_no_review_list = []
                for access in privileged_without_review:
                    privileged_no_review_list.append({
                        'asset_name': access.get('asset_name', ''),
                        'access_description': access.get('access_description', ''),
                        'access_right': access.get('access_right', ''),
                        'roles': access.get('roles', ''),
                        'last_review': access.get('last_review', '') or translations.get('never', 'Never'),
                        'days_overdue': str(access.get('days_overdue', 0))
                    })
                
                access_risk_section['subsections'].append({
                    'type': 'privileged_without_review_table',
                    'title': translations.get('privileged_access_without_regular_reviews', 'Privileged Access Without Regular Reviews'),
                    'headers': [
                        translations.get('asset_name', 'Asset Name'),
                        translations.get('access_description', 'Access Description'),
                        translations.get('access_right', 'Access Right'),
                        translations.get('roles', 'Roles'),
                        translations.get('last_review', 'Last Review'),
                        translations.get('days_overdue', 'Days Overdue')
                    ],
                    'data': privileged_no_review_list
                })
            
            if access_risk_section['subsections']:
                content['sections'].append(access_risk_section)
    
    # SIEM Section
    if selected_sections.get('siem') and 'siem_data' in data:
        siem_data = data.get('siem_data')
        # Check if data exists (even if empty, we should show the section if enabled)
        if siem_data is not None:
            siem_section = {
                'type': 'siem',
                'title': translations.get('siem', 'SIEM'),
                'subsections': []
            }
            
            # Summary Statistics
            if siem_data.get('total_active_agents', 0) >= 0:
                stats_data = [
                    {
                        'label': translations.get('total_active_agents_wazuh', 'Total Active Wazuh Agents'),
                        'value': str(siem_data.get('total_active_agents', 0)),
                        'type': 'primary'
                    },
                    {
                        'label': translations.get('agents_with_problems', 'Agents with Problems'),
                        'value': str(siem_data.get('agents_with_problems_count', 0)),
                        'type': 'danger' if siem_data.get('agents_with_problems_count', 0) > 0 else 'success'
                    },
                    {
                        'label': translations.get('coverage_rate', 'Coverage Rate'),
                        'value': f"{siem_data.get('coverage_rate', 0):.1f}%",
                        'type': 'info'
                    },
                    {
                        'label': translations.get('monitored_assets', 'Monitored Assets'),
                        'value': f"{siem_data.get('monitored_assets', 0)}/{siem_data.get('total_assets', 0)}",
                        'type': 'info'
                    }
                ]
                
                siem_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('wazuh_agents_summary', 'Wazuh Agents Summary'),
                    'data': stats_data,
                    'display_inline': True
                })
            
            # Agents with Problems Table
            agents_with_problems = siem_data.get('agents_with_problems', [])
            if agents_with_problems:
                problems_list = []
                for agent in agents_with_problems:
                    problems_list.append({
                        'agent_name': agent.get('agent_name', ''),
                        'agent_id': agent.get('agent_id', ''),
                        'agent_ip': agent.get('agent_ip', ''),
                        'status': agent.get('status', ''),
                        'agent_version': agent.get('agent_version', '') or translations.get('not_specified', 'Not specified'),
                        'last_seen': agent.get('last_seen', '') or translations.get('never', 'Never'),
                        'problems': agent.get('problems', ''),
                        'days_since_seen': str(agent.get('days_since_seen', 0))
                    })
                
                siem_section['subsections'].append({
                    'type': 'agents_with_problems_table',
                    'title': translations.get('agents_with_problems_inactive_outdated', 'Agents with Problems (inactive, outdated)'),
                    'headers': [
                        translations.get('agent_name', 'Agent Name'),
                        translations.get('agent_id', 'Agent ID'),
                        translations.get('agent_ip', 'Agent IP'),
                        translations.get('status', 'Status'),
                        translations.get('agent_version', 'Version'),
                        translations.get('last_seen', 'Last Seen'),
                        translations.get('problems', 'Problems'),
                        translations.get('days_since_seen', 'Days Since Seen')
                    ],
                    'data': problems_list
                })
            
            # Coverage Rate Information
            coverage_rate = siem_data.get('coverage_rate', 0)
            monitored_assets = siem_data.get('monitored_assets', 0)
            total_assets = siem_data.get('total_assets', 0)
            
            if total_assets > 0:
                coverage_info = [
                    {
                        'label': translations.get('total_assets', 'Total Assets'),
                        'value': str(total_assets),
                        'type': 'info'
                    },
                    {
                        'label': translations.get('monitored_assets', 'Monitored Assets'),
                        'value': str(monitored_assets),
                        'type': 'success' if monitored_assets > 0 else 'warning'
                    },
                    {
                        'label': translations.get('coverage_rate', 'Coverage Rate'),
                        'value': f"{coverage_rate:.1f}%",
                        'type': 'primary' if coverage_rate >= 80 else 'warning' if coverage_rate >= 50 else 'danger'
                    }
                ]
                
                siem_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('coverage_rate_monitoring', 'Coverage Rate (Assets Under Monitoring)'),
                    'data': coverage_info,
                    'display_inline': True
                })
            
            # Threat Detection Trends
            threat_detection_trends = siem_data.get('threat_detection_trends', [])
            if threat_detection_trends:
                trends_list = []
                for trend in threat_detection_trends:
                    trends_list.append({
                        'date': trend.get('date', ''),
                        'alert_count': str(trend.get('alert_count', 0)),
                        'trend': trend.get('trend', 'stable').capitalize()
                    })
                
                siem_section['subsections'].append({
                    'type': 'threat_detection_trends_table',
                    'title': translations.get('threat_detection_trends', 'Threat Detection Trends'),
                    'headers': [
                        translations.get('date', 'Date'),
                        translations.get('alert_count', 'Alert Count'),
                        translations.get('trend', 'Trend')
                    ],
                    'data': trends_list
                })
            
            if siem_section['subsections']:
                content['sections'].append(siem_section)
    
    # Asset Details Section - Combined Assets by Criticality and TOP-10 Asset
    if selected_sections.get('asset_details') or selected_sections.get('asset_tables'):
        asset_details_section = {
            'type': 'asset_details_combined',
            'title': translations.get('asset_details', 'Asset Details'),
            'subsections': []
        }
        
        # Assets by Criticality
        if selected_sections.get('asset_details') and 'statistics' in data:
            stats = data['statistics']
            if stats.get('assets_by_criticality'):
                criticality_data = []
                for item in stats['assets_by_criticality']:
                    criticality_name = item.get('criticality__name_uk', translations.get('unknown', 'Unknown'))
                    criticality_data.append({
                        'criticality': criticality_name,
                        'count': str(item.get('count', 0)),
                        'percentage': str(item.get('percentage', 0)),
                        'color': item.get('color', '#000000'),
                        'cost': item.get('cost', 0)
                    })
                
                # Sort by cost in descending order (highest criticality first)
                criticality_data.sort(key=lambda x: x['cost'], reverse=True)
                
                asset_details_section['subsections'].append({
                    'type': 'assets_by_criticality',
                    'title': translations.get('assets_by_criticality', 'Assets by Criticality'),
                    'headers': [translations['criticality'], translations['count'], translations.get('percentage', 'Percentage')],
                    'data': criticality_data
                })
        
        # TOP-10 Asset
        if selected_sections.get('asset_tables') and 'assets' in data:
            asset_table_data = []
            # Get vulnerability data for assets
            assets_vuln_data = {}
            if 'statistics' in data and 'assets_vulnerability_data' in data['statistics']:
                for vuln_info in data['statistics']['assets_vulnerability_data']:
                    assets_vuln_data[vuln_info['asset_id']] = vuln_info
            
            for asset in data['assets']:
                # Handle both dictionary and model object cases
                if hasattr(asset, 'get'):  # Dictionary
                    criticality_cost = asset.get('criticality_cost', 0)
                    asset_id = asset.get('id', asset.get('asset_id', ''))
                    # Get vulnerability count info
                    vuln_info = assets_vuln_data.get(asset_id, {})
                    vulnerabilities_display = 'Undefined' if vuln_info.get('is_undefined', True) else str(vuln_info.get('vulnerabilities_count', 0))
                    
                    asset_table_data.append({
                        'asset_id': asset.get('asset_id', ''),
                        'asset_name': asset.get('name', ''),
                        'criticality': asset.get('criticality_name', ''),
                        'criticality_cost': criticality_cost,
                        'vulnerabilities': vulnerabilities_display
                    })
                else:  # Model object
                    # Get criticality using the proper method
                    criticality_info = asset.get_criticality()
                    criticality_name = criticality_info.get('name', '') if criticality_info else ''
                    criticality_color = criticality_info.get('color', '#000000') if criticality_info else '#000000'
                    criticality_cost = criticality_info.get('cost', 0) if criticality_info else 0
                    
                    # Get vulnerability count info
                    asset_id = getattr(asset, 'id', None)
                    vuln_info = assets_vuln_data.get(asset_id, {}) if asset_id else {}
                    vulnerabilities_display = 'Undefined' if vuln_info.get('is_undefined', True) else str(vuln_info.get('vulnerabilities_count', 0))
                    
                    asset_table_data.append({
                        'asset_id': getattr(asset, 'asset_id', ''),
                        'asset_name': getattr(asset, 'name', ''),
                        'criticality': criticality_name,
                        'criticality_color': criticality_color,
                        'criticality_cost': criticality_cost,
                        'vulnerabilities': vulnerabilities_display
                    })
            
            # Sort by criticality cost (highest first) and limit to 10
            asset_table_data.sort(key=lambda x: x['criticality_cost'], reverse=True)
            asset_table_data = asset_table_data[:10]
            
            asset_details_section['subsections'].append({
                'type': 'top_10_asset',
                'title': translations.get('top_10_asset', 'TOP-10 Asset'),
                'headers': [
                    translations.get('asset_id', 'Asset ID'),
                    translations.get('asset_name', 'Asset Name'),
                    translations['criticality'],
                    translations.get('vulnerabilities', 'Vulnerabilities'),
                    translations.get('company', 'Company')
                ],
                'data': asset_table_data
            })
        
        # Add the combined section to content
        if asset_details_section['subsections']:
            content['sections'].append(asset_details_section)
    
    # Vulnerability Details Section
    if selected_sections.get('vulnerability_details') or selected_sections.get('vulnerability_tables'):
        vuln_section = {
            'type': 'vulnerability_details_combined',
            'title': translations.get('vulnerability_details', 'Vulnerability Details'),
            'subsections': []
        }
        
        # Initialize vuln_data outside the condition to avoid "referenced before assignment" error
        vuln_data = []
        
        if selected_sections.get('vulnerability_details') and 'vulnerabilities' in data:
            
            # Process vulnerabilities and calculate additional data
            processed_vulns = []
            for vuln in data['vulnerabilities']:
                # Handle both dictionary and model object cases
                if hasattr(vuln, 'get'):  # Dictionary
                    asset_name = vuln.get('asset_name', '')
                    vulnerability_name = vuln.get('vulnerability_name', '')
                    status = vuln.get('status', '')
                    # For dictionary case, we need to get the actual objects
                    asset_obj = None
                    vuln_obj = None
                else:  # Model object
                    asset_obj = vuln.asset if hasattr(vuln, 'asset') else None
                    vuln_obj = vuln.vulnerability if hasattr(vuln, 'vulnerability') else None
                    asset_name = getattr(vuln, 'asset_name', '') if hasattr(vuln, 'asset_name') else (getattr(asset_obj, 'name', '') if asset_obj else '')
                    vulnerability_name = getattr(vuln, 'vulnerability_name', '') if hasattr(vuln, 'vulnerability_name') else (get_localized_vulnerability_field(vuln_obj, 'vulnerability', language) if vuln_obj else '')
                    status = getattr(vuln, 'status', '')
                
                # Calculate additional data
                threat_info = ''
                probability_impact = ''
                impact_value = ''
                risk_value = ''
                risk_level = ''
                risk_mitigation_controls = ''
                asset_criticality_cost = 0
                
                if asset_obj and vuln_obj:
                    # Get asset criticality for sorting
                    try:
                        criticality = asset_obj.get_criticality()
                        asset_criticality_cost = criticality.get('cost', 0)
                    except:
                        asset_criticality_cost = 0
                    
                    # Get risk mitigation controls and limit to 400 characters
                    risk_mitigation_controls = get_localized_vulnerability_field(vuln_obj, 'risk_mitigation_controls', language)
                    if risk_mitigation_controls and len(risk_mitigation_controls) > 400:
                        # Truncate to 400 characters and ensure it ends with a complete sentence
                        truncated = risk_mitigation_controls[:400]
                        # Find the last sentence ending (period, exclamation, question mark)
                        last_sentence_end = max(
                            truncated.rfind('.'), 
                            truncated.rfind('!'), 
                            truncated.rfind('?')
                        )
                        if last_sentence_end > 350:  # If we have a sentence ending in the last 50 chars
                            risk_mitigation_controls = truncated[:last_sentence_end + 1]
                        else:
                            # If no sentence ending found, just truncate and add ellipsis
                            risk_mitigation_controls = truncated + '...'
                    
                    # Calculate threat and risk information for all vulnerabilities (regardless of status)
                    try:
                        risk_value_calc = calculate_value_of_risk(asset_obj, vuln_obj)
                        risk_value = f"{risk_value_calc:.2f}"
                        
                        # Debug logging for risk calculation
                        logger.debug(f"Risk calculation for {vulnerability_name}: value={risk_value_calc}, asset={asset_name}, status={status}")
                        
                        risk_level_obj = calculate_risk_level(risk_value_calc)
                        if risk_level_obj:
                            risk_level = risk_level_obj.get_name_by_language(language) or risk_level_obj.get_name()
                            logger.debug(f"Found risk level: {risk_level} for value {risk_value_calc}")
                        else:
                            # If no risk level found, try to find the highest level for high risk values
                            if risk_value_calc > 50:  # High risk threshold
                                try:
                                    # First try to find "Extraordinary" level specifically
                                    extraordinary_level = RiskLevel.objects.filter(
                                        Q(name__icontains='Extraordinary') | 
                                        Q(translations__name_local__icontains='Extraordinary')
                                    ).first()
                                    
                                    if extraordinary_level:
                                        risk_level = extraordinary_level.get_name_by_language(language) or extraordinary_level.get_name()
                                        logger.debug(f"Using Extraordinary risk level: {risk_level} for high risk value {risk_value_calc}")
                                    else:
                                        # Fallback to highest level
                                        highest_level = RiskLevel.objects.order_by('-max_value').first()
                                        if highest_level:
                                            risk_level = highest_level.get_name_by_language(language) or highest_level.get_name()
                                            logger.debug(f"Using highest risk level: {risk_level} for high risk value {risk_value_calc}")
                                        else:
                                            risk_level = translations.get('unknown', 'Unknown')
                                            logger.warning(f"No risk levels found in database for high risk value {risk_value_calc}")
                                except Exception as e:
                                    risk_level = translations.get('unknown', 'Unknown')
                                    logger.error(f"Error finding highest risk level: {str(e)}")
                            else:
                                risk_level = translations.get('unknown', 'Unknown')
                                logger.debug(f"No risk level found for low risk value {risk_value_calc}")
                    except Exception as e:
                        logger.warning(f"Error calculating risk for vulnerability {vulnerability_name}: {str(e)}")
                        risk_value = '0.00'
                        risk_level = translations.get('unknown', 'Unknown')
                    
                    # Get threat information for all vulnerabilities
                    threats = vuln_obj.threats.all()
                    if threats.exists():
                        threat_names = []
                        for threat in threats:
                            threat_name = threat.get_name() or getattr(threat, f'name_{language}', '')
                            threat_names.append(threat_name)
                            
                            # Calculate probability/impact for this threat
                            try:
                                probability = threat.calculate_probability()
                                impact = threat.calculate_overall_impact() if hasattr(threat, 'calculate_overall_impact') else threat.impact / 100
                                prob_impact = probability * impact * 100
                                probability_impact = f"{prob_impact:.2f}%"
                                impact_value = f"{impact * 100:.2f}%"
                            except:
                                probability_impact = '0.00%'
                                impact_value = '0.00%'
                        
                        threat_info = ', '.join(threat_names)
                    else:
                        threat_info = translations.get('no_threats', 'No threats')
                        probability_impact = '0.00%'
                        impact_value = '0.00%'
                
                # Calculate risk level priority for sorting (higher number = higher priority)
                risk_level_priority = 0
                if risk_level:
                    if 'Надзвичайний' in risk_level or 'Extraordinary' in risk_level:
                        risk_level_priority = 6
                    elif 'Критичний' in risk_level or 'Critical' in risk_level:
                        risk_level_priority = 5
                    elif 'Високий' in risk_level or 'High' in risk_level:
                        risk_level_priority = 4
                    elif 'Середній' in risk_level or 'Medium' in risk_level:
                        risk_level_priority = 3
                    elif 'Низький' in risk_level or 'Low' in risk_level:
                        risk_level_priority = 2
                    elif 'Невизначено' in risk_level or 'Undefined' in risk_level:
                        risk_level_priority = 1
                
                processed_vulns.append({
                    'asset_name': asset_name,
                        'vulnerability_name': vulnerability_name,
                    'status': status,
                    'threat': threat_info,
                    'probability_impact': probability_impact,
                    'impact': impact_value,
                    'risk_value': risk_value,
                    'risk_level': risk_level,
                    'risk_mitigation_controls': risk_mitigation_controls,
                    'asset_criticality_cost': asset_criticality_cost,
                    'status_priority': 1 if status == 'Yes' else (2 if status == 'Undefined' else 3),
                    'risk_level_priority': risk_level_priority
                })
            
            # Sort vulnerabilities: first by status (Yes first), then by asset criticality (highest first), then by risk level (highest first)
            processed_vulns.sort(key=lambda x: (x['status_priority'], -x['asset_criticality_cost'], -x['risk_level_priority']))
            
            # Take top 10
            for vuln in processed_vulns[:10]:
                vuln_data.append({
                    'asset_name': vuln['asset_name'],
                    'vulnerability_name': vuln['vulnerability_name'],
                    'status': vuln['status'],
                    'threat': vuln['threat'],
                    'probability_impact': vuln['probability_impact'],
                    'impact': vuln['impact'],
                    'risk_value': vuln['risk_value'],
                    'risk_level': vuln['risk_level'],
                    'risk_mitigation_controls': vuln['risk_mitigation_controls']
                    })
            
        # Add Assets by Vulnerability subsection first (left side)
        if selected_sections.get('vulnerability_details') and 'statistics' in data:
            stats = data['statistics']
            if stats.get('assets_by_vulnerability'):
                vuln_by_asset_data = []
                for item in stats['assets_by_vulnerability']:
                    vuln_by_asset_data.append({
                        'criticality': item.get('criticality', ''),
                        'count': item.get('count', 0),
                        'percentage': item.get('percentage', 0),
                        'color': item.get('color', '#000000')
                    })
                
                vuln_section['subsections'].append({
                    'type': 'assets_by_vulnerability',
                    'title': translations.get('assets_by_vulnerability', 'Assets by Vulnerability'),
                    'headers': [
                        translations.get('criticality', 'Criticality'),
                        translations.get('count', 'Count'),
                        translations.get('percentage', 'Percentage')
                    ],
                    'data': vuln_by_asset_data
                })
        
        # Add TOP-10 Vulnerability subsection second (right side)
        vuln_section['subsections'].append({
            'type': 'top_10_vulnerability',
                'title': translations.get('top_10_vulnerability', 'TOP-10 Vulnerability'),
                'headers': [
                    translations.get('asset', 'Asset'),
                    translations['vulnerability'],
                    translations['status'],
                    translations['threat'],
                    translations['probability_impact'],
                    translations['impact'],
                    translations['risk_value'],
                    translations['risk_level'],
                    translations['risk_mitigation_controls']
                ],
                'data': vuln_data
            })
        
        if vuln_section['subsections']:
            content['sections'].append(vuln_section)
    
    # Treatment Details Section
    if selected_sections.get('treatment_details') or selected_sections.get('treatment_tables'):
        treatment_section = {
            'type': 'treatment_details',
            'title': translations.get('treatment_details', 'Treatment Details'),
            'subsections': []
        }
        
        if selected_sections.get('treatment_details') and 'statistics' in data:
            stats = data['statistics']
            if stats.get('overdue_treatments'):
                treatment_data = []
                for treatment in stats['overdue_treatments'][:10]:  # Limit to 10
                    treatment_data.append({
                        'asset_name': treatment.get('asset_name', translations.get('unknown', 'Unknown')),
                        'treatment': treatment.get('vulnerability_name', translations.get('unknown', 'Unknown')),
                        'deadline': treatment.get('due_date', translations.get('unknown', 'Unknown'))
                    })
                
                treatment_section['subsections'].append({
                    'type': 'overdue_treatments',
                    'title': translations.get('overdue_treatments', 'Overdue Treatments'),
                    'headers': [translations['asset'], translations.get('treatment', 'Treatment'), translations.get('deadline', 'Deadline')],
                    'data': treatment_data
                })
        
        # Add Treatment Details subsection if requested
        if selected_sections.get('treatment_details') and 'risk_treatments' in data:
            treatment_details_data = []
            for treatment in data['risk_treatments']:  # Process all treatments for sorting
                # Get vulnerability status and risk level if available
                vulnerability_status = translations.get('unknown', 'Unknown')
                risk_level = translations.get('unknown', 'Unknown')
                
                # Try to find the corresponding AssetVulnerability for this treatment
                if 'asset_vulnerabilities' in data:
                    for vuln in data['asset_vulnerabilities']:
                        # Handle both dictionary and model object cases
                        if hasattr(vuln, 'get'):  # Dictionary
                            vuln_asset_name = vuln.get('asset_name', '')
                            vuln_name = vuln.get('vulnerability_name', '')
                        else:  # Model object
                            vuln_asset_name = getattr(vuln.asset, 'name', '') if hasattr(vuln, 'asset') and vuln.asset else ''
                            vuln_name = get_localized_vulnerability_field(vuln.vulnerability, 'vulnerability', language) if hasattr(vuln, 'vulnerability') and vuln.vulnerability else ''
                        
                        # Match by asset name and vulnerability name
                        if (vuln_asset_name == treatment.get('asset_name', '') and 
                            vuln_name == treatment.get('vulnerability_name', '')):
                            vulnerability_status = getattr(vuln, 'status', translations.get('unknown', 'Unknown'))
                            
                            # Calculate risk level for this vulnerability
                            try:
                                if hasattr(vuln, 'asset') and hasattr(vuln, 'vulnerability'):
                                    # For model objects
                                    risk_value = calculate_value_of_risk(vuln.asset, vuln.vulnerability)
                                    risk_level_obj = calculate_risk_level(risk_value)
                                    if risk_level_obj:
                                        risk_level = risk_level_obj.get_name_by_language(language) or risk_level_obj.get_name()
                                    else:
                                        risk_level = translations.get('unknown', 'Unknown')
                                elif hasattr(vuln, 'get'):
                                    # For dictionary objects, try to get risk level if available
                                    risk_level = vuln.get('risk_level', translations.get('unknown', 'Unknown'))
                                else:
                                    risk_level = translations.get('unknown', 'Unknown')
                            except Exception as e:
                                logger.warning(f"Error calculating risk level: {e}")
                                risk_level = translations.get('unknown', 'Unknown')
                            break
                
                # Calculate priority for sorting
                status_priority = 0
                if vulnerability_status == 'Yes':
                    status_priority = 1
                
                risk_level_priority = 0
                if risk_level:
                    if 'Надзвичайний' in risk_level or 'Extraordinary' in risk_level:
                        risk_level_priority = 6
                    elif 'Критичний' in risk_level or 'Critical' in risk_level:
                        risk_level_priority = 5
                    elif 'Високий' in risk_level or 'High' in risk_level:
                        risk_level_priority = 4
                    elif 'Середній' in risk_level or 'Medium' in risk_level:
                        risk_level_priority = 3
                    elif 'Низький' in risk_level or 'Low' in risk_level:
                        risk_level_priority = 2
                    elif 'Невизначено' in risk_level or 'Undefined' in risk_level:
                        risk_level_priority = 1
                
                treatment_details_data.append({
                    'asset_name': treatment.get('asset_name', translations.get('unknown', 'Unknown')),
                    'vulnerability_name': treatment.get('vulnerability_name', translations.get('unknown', 'Unknown')),
                    'vulnerability_status': vulnerability_status,
                    'risk_level': risk_level,
                    'treatment_type': treatment.get('treatment_type', translations.get('unknown', 'Unknown')),
                    'treatment_status': treatment.get('status', translations.get('unknown', 'Unknown')),
                    'description_responsible': treatment.get('description', '') + ' | ' + treatment.get('assigned_to', translations.get('unknown', 'Unknown')),
                    'due_date': treatment.get('due_date', translations.get('unknown', 'Unknown')),
                    'status_priority': status_priority,
                    'risk_level_priority': risk_level_priority
                })
            
            # Sort by status priority (Yes first) and then by risk level priority (highest first)
            treatment_details_data.sort(key=lambda x: (x['status_priority'], x['risk_level_priority']), reverse=True)
            
            # Limit to top 10
            treatment_details_data = treatment_details_data[:10]
            
            # Remove priority fields from final data
            for item in treatment_details_data:
                item.pop('status_priority', None)
                item.pop('risk_level_priority', None)
            
            treatment_section['subsections'].append({
                'type': 'treatment_details',
                'title': translations.get('top_10_treatment', 'TOP-10 Treatment'),
                'headers': [translations['asset'], translations.get('vulnerability', 'Vulnerability'), translations.get('vulnerability_status', 'Vulnerability Status'), translations.get('risk_level', 'Risk Level'), translations.get('treatment_type', 'Treatment Type'), translations.get('treatment_status', 'Treatment Status'), translations.get('description_responsible', 'Description & Responsible'), translations.get('due_date', 'Due Date')],
                'data': treatment_details_data
            })
        
        if treatment_section['subsections']:
            content['sections'].append(treatment_section)
    
    # Recommendations Section
    if selected_sections.get('recommendations'):
        recommendations_section = {
            'type': 'recommendations',
            'title': translations['recommendations'],
            'subsections': [{
                'type': 'priority_actions',
                'title': translations.get('priority_actions', 'Priority Actions'),
                'data': [
                    translations.get('address_high_risk_vulnerabilities', 'Address high-risk vulnerabilities'),
                    translations.get('improve_compliance_rates', 'Improve compliance rates'),
                    translations.get('implement_regular_assessments', 'Implement regular assessments'),
                    translations.get('update_security_policies', 'Update security policies')
                ]
            }]
        }
        content['sections'].append(recommendations_section)
    
    # Financial Analysis Section
    if selected_sections.get('financial_analysis') or selected_sections.get('cost_benefit_analysis') or selected_sections.get('roi_assessment') or selected_sections.get('budget_allocation'):
        financial_section = {
            'type': 'financial_analysis',
            'title': translations.get('financial_analysis', 'Financial Analysis'),
            'subsections': []
        }
        
        if selected_sections.get('financial_analysis') and 'statistics' in data:
            stats = data['statistics']
            
            # Ensure stats is a dictionary
            if isinstance(stats, dict):
                financial_section['subsections'].append({
                    'type': 'financial_overview',
                    'title': translations.get('financial_overview', 'Financial Overview'),
                    'data': [
                        {
                            'label': translations.get('total_cost', 'Total Cost'),
                            'value': f"${stats.get('total_cost', 0):,.2f}"
                        },
                        {
                            'label': translations.get('potential_savings', 'Potential Savings'),
                            'value': f"${stats.get('potential_savings', 0):,.2f}"
                        },
                        {
                            'label': translations.get('roi_percentage', 'ROI'),
                            'value': f"{stats.get('roi_percentage', 0):.1f}%"
                        },
                        {
                            'label': translations.get('budget_allocated', 'Budget Allocated'),
                            'value': f"${stats.get('budget_allocated', 0):,.2f}"
                        }
                    ]
                })
            else:
                # Fallback if stats is not a dictionary
                financial_section['subsections'].append({
                    'type': 'financial_overview',
                    'title': translations.get('financial_overview', 'Financial Overview'),
                    'data': [
                        {
                            'label': translations.get('total_cost', 'Total Cost'),
                            'value': '$0.00'
                        },
                        {
                            'label': translations.get('potential_savings', 'Potential Savings'),
                            'value': '$0.00'
                        },
                        {
                            'label': translations.get('roi_percentage', 'ROI'),
                            'value': '0.0%'
                        },
                        {
                            'label': translations.get('budget_allocated', 'Budget Allocated'),
                            'value': '$0.00'
                        }
                    ]
                })
        
        content['sections'].append(financial_section)
    
    # Residual Risk & Priority Section
    if selected_sections.get('residual_risk_analysis') or selected_sections.get('risk_appetite') or selected_sections.get('priority_matrix') or selected_sections.get('resource_allocation'):
        residual_section = {
            'type': 'residual_risk_priority',
            'title': translations.get('residual_risk_priority', 'Residual Risk & Priority'),
            'subsections': []
        }
        
        if selected_sections.get('residual_risk_analysis') and 'statistics' in data:
            stats = data['statistics']
            
            # Ensure stats is a dictionary
            if isinstance(stats, dict):
                residual_section['subsections'].append({
                    'type': 'residual_overview',
                    'title': translations.get('residual_overview', 'Residual Risk Overview'),
                    'data': [
                        {
                            'label': translations.get('residual_risk_level', 'Residual Risk Level'),
                            'value': stats.get('residual_risk_level', translations.get('unknown', 'Unknown'))
                        },
                        {
                            'label': translations.get('appetite_threshold', 'Risk Appetite Threshold'),
                            'value': f"{stats.get('appetite_threshold', 0):.1f}%"
                        },
                        {
                            'label': translations.get('priority_level', 'Priority Level'),
                            'value': stats.get('priority_level', translations.get('unknown', 'Unknown'))
                        }
                    ]
                })
            else:
                # Fallback if stats is not a dictionary
                residual_section['subsections'].append({
                    'type': 'residual_overview',
                    'title': translations.get('residual_overview', 'Residual Risk Overview'),
                    'data': [
                        {
                            'label': translations.get('residual_risk_level', 'Residual Risk Level'),
                            'value': translations.get('unknown', 'Unknown')
                        },
                        {
                            'label': translations.get('appetite_threshold', 'Risk Appetite Threshold'),
                            'value': '0.0%'
                        },
                        {
                            'label': translations.get('priority_level', 'Priority Level'),
                            'value': translations.get('unknown', 'Unknown')
                        }
                    ]
                })
        
        content['sections'].append(residual_section)
    
    # Acceptable Risk Section
    if selected_sections.get('acceptable_risk_analysis') and 'acceptable_risk_data' in data:
        acceptable_risk_section = {
            'type': 'acceptable_risk_analysis',
            'title': translations.get('acceptable_risk_analysis', 'Acceptable Risk Analysis'),
            'subsections': []
        }
        
        # Summary statistics
        if data['acceptable_risk_data']['summary']:
            summary = data['acceptable_risk_data']['summary']
            acceptable_risk_section['subsections'].append({
                'type': 'statistics',
                'title': translations.get('acceptable_risk_summary', 'Acceptable Risk Summary'),
                'data': [
                    {
                        'label': translations.get('total_assets', 'Total Assets'),
                        'value': str(summary['total_assets']),
                        'type': 'primary'
                    },
                    {
                        'label': translations.get('assets_with_acceptable_risk', 'Assets with Acceptable Risk'),
                        'value': str(summary['assets_with_acceptable_risk']),
                        'type': 'success'
                    },
                    {
                        'label': translations.get('assets_without_acceptable_risk', 'Assets without Acceptable Risk'),
                        'value': str(summary['assets_without_acceptable_risk']),
                        'type': 'warning'
                    },
                    {
                        'label': translations.get('exceeding_acceptable_risk', 'Exceeding Acceptable Risk'),
                        'value': str(summary['exceeding_acceptable_risk']),
                        'type': 'danger'
                    }
                ]
            })
        
        # Assets with acceptable risk configuration
        if data['acceptable_risk_data']['assets_with_acceptable_risk']:
            acceptable_risk_section['subsections'].append({
                'type': 'table',
                'title': translations.get('top_10_assets_with_acceptable_risk_config', 'TOP-10 Assets with Acceptable Risk Configuration'),
                'headers': [
                    translations.get('asset_id', 'Asset ID'),
                    translations.get('asset_name', 'Asset Name'),
                    translations.get('current_criticality', 'Current Criticality'),
                    translations.get('acceptable_risk_level', 'Acceptable Risk Level'),
                    translations.get('exceeds_acceptable', 'Exceeds Acceptable')
                ],
                'data': [
                    {
                        'asset_id': asset['asset_id'],
                        'asset_name': asset['asset_name'],
                        'current_criticality': asset['current_criticality'],
                        'acceptable_risk_level': asset['acceptable_risk_level'],
                        'exceeds_acceptable': 'Yes' if asset['exceeds_acceptable'] else 'No'
                    }
                    for asset in data['acceptable_risk_data']['assets_with_acceptable_risk']
                ]
            })
        
        # Assets without acceptable risk configuration
        if data['acceptable_risk_data']['assets_without_acceptable_risk']:
            acceptable_risk_section['subsections'].append({
                'type': 'table',
                'title': translations.get('top_10_assets_without_acceptable_risk_config', 'TOP-10 Assets without Acceptable Risk Configuration'),
                'headers': [
                    translations.get('asset_id', 'Asset ID'),
                    translations.get('asset_name', 'Asset Name'),
                    translations.get('current_criticality', 'Current Criticality')
                ],
                'data': [
                    {
                        'asset_id': asset['asset_id'],
                        'asset_name': asset['asset_name'],
                        'current_criticality': asset['current_criticality']
                    }
                    for asset in data['acceptable_risk_data']['assets_without_acceptable_risk']
                ]
            })
        
        content['sections'].append(acceptable_risk_section)
    
    # AI Conclusion Section
    # Check if conclusions is enabled (handle both boolean and string values)
    conclusions_enabled = selected_sections.get('conclusions', False)
    logger.info(f"AI Conclusion check: conclusions={conclusions_enabled}, type={type(conclusions_enabled)}, model_id={selected_sections.get('_ai_conclusion_model')}")
    if conclusions_enabled == True or (isinstance(conclusions_enabled, str) and conclusions_enabled.lower() in ('true', '1')):
        logger.info(f"Generating AI Conclusion. Model ID: {selected_sections.get('_ai_conclusion_model')}")
        ai_conclusion_text = generate_ai_conclusion(data, selected_sections, language)
        logger.info(f"AI Conclusion generated: {ai_conclusion_text[:100] if ai_conclusion_text else 'None'}...")
        if ai_conclusion_text:
            conclusion_section = {
                'type': 'ai_conclusion',
                'title': translations.get('conclusions', 'Conclusions'),
                'content': ai_conclusion_text,
                'model_name': None
            }
            
            # Get AI model name if available
            ai_model_id = selected_sections.get('_ai_conclusion_model')
            if ai_model_id:
                # Convert to int if it's a string
                if isinstance(ai_model_id, str):
                    try:
                        ai_model_id = int(ai_model_id)
                    except ValueError:
                        logger.error(f"Invalid AI model ID format: {ai_model_id}")
                        ai_model_id = None
                
                if ai_model_id:
                    try:
                        from app_ai.models import ModelChoice
                        ai_model = ModelChoice.objects.get(id=ai_model_id, is_active=True)
                        conclusion_section['model_name'] = ai_model.model_name
                        conclusion_section['provider'] = ai_model.provider
                    except ModelChoice.DoesNotExist:
                        logger.error(f"AI model with ID {ai_model_id} not found")
                        pass
            
            content['sections'].append(conclusion_section)
    
    # Add translations to content for use in HTML rendering
    content['translations'] = translations
    
    return content


def markdown_to_html(text):
    """
    Convert basic markdown formatting to HTML.
    Supports: **bold**, *italic*, lists, paragraphs.
    """
    if not text:
        return ''
    
    import re
    from html import escape
    from django.utils.safestring import mark_safe
    
    # Split into lines for processing
    lines = text.split('\n')
    html_parts = []
    in_list = False
    list_type = None  # 'ul' or 'ol'
    list_items = []
    
    def flush_list():
        nonlocal in_list, list_type, list_items, html_parts
        if list_items:
            tag = 'ol' if list_type == 'ol' else 'ul'
            list_html = ''.join(list_items)
            html_parts.append(f'<{tag}>{list_html}</{tag}>')
            list_items = []
            in_list = False
            list_type = None
    
    def process_inline_formatting(text_content):
        """Process bold and italic formatting in text content"""
        # Process bold text **text** or __text__
        text_content = re.sub(r'\*\*([^*]+?)\*\*', r'<strong>\1</strong>', text_content)
        text_content = re.sub(r'__([^_]+?)__', r'<strong>\1</strong>', text_content)
        # Process italic text *text* (but not **text**)
        text_content = re.sub(r'(?<!\*)\*([^*\n]+?)\*(?!\*)', r'<em>\1</em>', text_content)
        text_content = re.sub(r'(?<!_)_([^_\n\s][^_\n]*?[^_\n\s])_(?!_)', r'<em>\1</em>', text_content)
        return text_content
    
    for line in lines:
        trimmed = line.strip()
        
        # Check for unordered list items (* or -)
        if re.match(r'^[\*\-]\s+', trimmed):
            content = re.sub(r'^[\*\-]\s+', '', trimmed)
            # Process markdown formatting
            content = process_inline_formatting(content)
            # Escape HTML but preserve formatting tags
            content = escape(content)
            content = content.replace('&lt;strong&gt;', '<strong>').replace('&lt;/strong&gt;', '</strong>')
            content = content.replace('&lt;em&gt;', '<em>').replace('&lt;/em&gt;', '</em>')
            
            if not in_list or list_type != 'ul':
                flush_list()
                in_list = True
                list_type = 'ul'
            list_items.append(f'<li>{content}</li>')
            continue
        
        # Check for ordered list items (1. 2. etc.)
        if re.match(r'^\d+\.\s+', trimmed):
            content = re.sub(r'^\d+\.\s+', '', trimmed)
            # Process markdown formatting
            content = process_inline_formatting(content)
            # Escape HTML but preserve formatting tags
            content = escape(content)
            content = content.replace('&lt;strong&gt;', '<strong>').replace('&lt;/strong&gt;', '</strong>')
            content = content.replace('&lt;em&gt;', '<em>').replace('&lt;/em&gt;', '</em>')
            
            if not in_list or list_type != 'ol':
                flush_list()
                in_list = True
                list_type = 'ol'
            list_items.append(f'<li>{content}</li>')
            continue
        
        # If not a list item, flush current list
        if in_list:
            flush_list()
        
        # Store regular line (will be processed as paragraph later)
        if trimmed:
            html_parts.append(('text', trimmed))
        else:
            html_parts.append(('blank', ''))
    
    # Flush any remaining list
    flush_list()
    
    # Now process paragraphs from stored text lines
    result_parts = []
    current_para_lines = []
    
    def finalize_paragraph():
        """Finalize current paragraph and add to result"""
        nonlocal current_para_lines, result_parts
        if current_para_lines:
            para_text = ' '.join(current_para_lines)
            para_text = process_inline_formatting(para_text)
            # Escape HTML but preserve formatting tags
            para_text = escape(para_text)
            para_text = para_text.replace('&lt;strong&gt;', '<strong>').replace('&lt;/strong&gt;', '</strong>')
            para_text = para_text.replace('&lt;em&gt;', '<em>').replace('&lt;/em&gt;', '</em>')
            result_parts.append(f'<p>{para_text}</p>')
            current_para_lines = []
    
    for item in html_parts:
        if isinstance(item, tuple):
            item_type, content = item
            if item_type == 'blank':
                # Blank line - end of paragraph
                finalize_paragraph()
            elif item_type == 'text':
                current_para_lines.append(content)
        else:
            # This is an HTML element (list) - finalize paragraph and add list
            finalize_paragraph()
            result_parts.append(item)
    
    # Handle remaining paragraph
    finalize_paragraph()
    
    return mark_safe(''.join(result_parts))


def render_unified_content_as_html(unified_content):
    """Convert unified content structure to HTML for preview"""
    
    # Get translations for confidentiality classification
    from django.utils.translation import gettext as _
    confidentiality_classification = _("Confidentiality Classification")
    confidential_information = _("Confidential Information")
    
    html = f'''
    <div class="report-preview">
        <div class="report-header text-center mb-4">
            <h2 class="text-primary">{unified_content['header']['title']}</h2>
    '''
    
    # Profile info (if available)
    if unified_content.get('profile_info'):
        profile_info = unified_content['profile_info']
        html += f'''
        <div class="profile-info mb-3">
            <h4 class="text-info">Profile Information</h4>
            <p><strong>Profile Name:</strong> {profile_info.get('profile_name', 'N/A')}</p>
            <p><strong>Report Type:</strong> {profile_info.get('report_type', 'N/A')}</p>
            <p><strong>Language:</strong> {profile_info.get('language', 'N/A')}</p>
            <p><strong>Date Range:</strong> {profile_info.get('date_range', 'N/A')}</p>
        </div>
    '''
    
    # Company info
    if unified_content['header']['company_info']:
        company_info = unified_content['header']['company_info']
        if company_info['type'] == 'single':
            html += f'<p class="text-muted mb-1"><i class="fas fa-building me-2"></i><strong>Company:</strong> {company_info["name"]}</p>'
        else:
            html += f'<p class="text-muted mb-1"><i class="fas fa-building me-2"></i><strong>Company:</strong> {company_info["text"]}</p>'
    
    # Add confidentiality classification
    html += f'''
            <p class="text-muted">{unified_content['header']['generated_date']}</p>
            <div class="confidentiality-classification mt-3 mb-2">
                <p class="text-danger fw-bold"><i class="fas fa-shield-alt me-2"></i><strong>{confidentiality_classification}:</strong> {confidential_information}</p>
            </div>
            <hr>
        </div>
    '''
    
    # Render sections
    logger.info(f"Rendering {len(unified_content['sections'])} sections")
    for section in unified_content['sections']:
        logger.debug(f"Rendering section: type={section.get('type')}, title={section.get('title')}")
        html += f'<div class="mb-4"><h3 class="text-secondary"><i class="fas fa-chart-line me-2"></i>{section["title"]}</h3>'
        
        # Special handling for Risk Details section - display subsections side by side
        if section['type'] == 'risk_details':
            html += '<div class="row">'
            
            for subsection in section['subsections']:
                if subsection['type'] == 'risk_distribution_level':
                    html += '<div class="col-md-2">'  # 15% width (closest to 15% is col-md-2 which is ~16.67%)
                else:
                    html += '<div class="col-md-10">'  # 85% width for TOP-10 Risks
                html += f'<h5>{subsection["title"]}</h5>'
                
                # Render subsection content
                if 'headers' in subsection and 'data' in subsection:
                    html += '<div class="table-responsive"><table class="table table-striped"><thead><tr>'
                    for header in subsection['headers']:
                        # Use translations for common field names
                        translated_header = header
                        if header.lower() in ['type', 'title', 'name', 'description', 'location', 'group', 'status', 'comment', 'notes']:
                            # Get translation from unified_content translations if available
                            if 'translations' in unified_content:
                                translation_key = header.lower()
                                translated_header = unified_content['translations'].get(translation_key, header)
                        html += f'<th>{translated_header}</th>'
                    html += '</tr></thead><tbody>'
                    
                    for row in subsection['data']:
                        html += '<tr>'
                        if subsection['type'] == 'risk_distribution_level':
                            color = row.get('color', '#000000')
                            html += f'<td><span class="badge" style="background-color: {color}; color: white;">{row["risk_level"]}</span></td>'
                            html += f'<td>{row["count"]}</td>'
                            html += f'<td>{row["percentage"]}</td>'
                        elif subsection['type'] == 'top_risks':
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td>{row["vulnerability_name"]}</td>'
                            html += f'<td>{row["risk_level"]}</td>'
                        html += '</tr>'
                    
                    html += '</tbody></table></div>'
                
                html += '</div>'
            
            html += '</div>'
        
        # Special handling for Asset Details section - display subsections side by side
        elif section['type'] == 'asset_details_combined':
            html += '<div class="row">'
            
            for subsection in section['subsections']:
                if subsection['type'] == 'assets_by_criticality':
                    html += '<div class="col-md-2">'  # 15% width (closest to 15% is col-md-2 which is ~16.67%)
                else:
                    html += '<div class="col-md-10">'  # 85% width for TOP-10 Asset
                html += f'<h5>{subsection["title"]}</h5>'
                
                # Render subsection content
                if 'headers' in subsection and 'data' in subsection:
                        html += '<div class="table-responsive"><table class="table table-striped"><thead><tr>'
                        for header in subsection['headers']:
                        # Use translations for common field names
                            translated_header = header
                            if header.lower() in ['type', 'title', 'name', 'description', 'location', 'group', 'status', 'comment', 'notes']:
                            # Get translation from unified_content translations if available
                                if 'translations' in unified_content:
                                    translation_key = header.lower()
                                    translated_header = unified_content['translations'].get(translation_key, header)
                            html += f'<th>{translated_header}</th>'
                        html += '</tr></thead><tbody>'
                        
                        for row in subsection['data']:
                            html += '<tr>'
                        if subsection['type'] == 'assets_by_criticality':
                            # Use color from data if available, otherwise use default
                            color = row.get('color', '#000000')
                            html += f'<td><span class="badge" style="background-color: {color}; color: white;">{row["criticality"]}</span></td>'
                            html += f'<td>{row["count"]}</td>'
                            html += f'<td>{row["percentage"]}%</td>'
                        elif subsection['type'] == 'top_10_asset':
                            html += f'<td>{row["asset_id"]}</td>'
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td><span style="color:{row.get("criticality_color", "#000000")}">{row["criticality"]}</span></td>'
                            html += '</tr>'
                    
                        html += '</tbody></table></div>'
                
                html += '</div>'
            
            html += '</div>'
        
        # Special handling for Vulnerability Details section - display subsections side by side with 15%/85% width distribution
        elif section['type'] == 'vulnerability_details_combined':
            html += '<div class="row">'
            
            for subsection in section['subsections']:
                # Set width based on subsection type: 15% for Assets by Vulnerability, 85% for TOP-10 Vulnerability
                if subsection['type'] == 'assets_by_vulnerability':
                    html += '<div class="col-md-2">'  # 15% width (closest to 15% is col-md-2 which is ~16.67%)
                else:
                    html += '<div class="col-md-10">'  # 85% width for TOP-10 Vulnerability
                
                html += f'<h5>{subsection["title"]}</h5>'
                
                # Render subsection content
                if 'headers' in subsection and 'data' in subsection:
                    html += '<div class="table-responsive"><table class="table table-striped"><thead><tr>'
                    for header in subsection['headers']:
                        # Use translations for common field names
                        translated_header = header
                        if header.lower() in ['type', 'title', 'name', 'description', 'location', 'group', 'status', 'comment', 'notes']:
                            # Get translation from unified_content translations if available
                            if 'translations' in unified_content:
                                translation_key = header.lower()
                                translated_header = unified_content['translations'].get(translation_key, header)
                        html += f'<th>{translated_header}</th>'
                    html += '</tr></thead><tbody>'
                    
                    for row in subsection['data']:
                        html += '<tr>'
                        if subsection['type'] == 'top_10_vulnerability':
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td>{row["vulnerability_name"]}</td>'
                            html += f'<td>{row["status"]}</td>'
                            html += f'<td>{row["threat"]}</td>'
                            html += f'<td>{row["probability_impact"]}</td>'
                            html += f'<td>{row["impact"]}</td>'
                            html += f'<td>{row["risk_value"]}</td>'
                            html += f'<td>{row["risk_level"]}</td>'
                            html += f'<td>{row["risk_mitigation_controls"]}</td>'
                        elif subsection['type'] == 'assets_by_vulnerability':
                            # Use color from data if available, otherwise use default
                            color = row.get('color', '#000000')
                            html += f'<td><span class="badge" style="background-color: {color}; color: white;">{row["criticality"]}</span></td>'
                            html += f'<td>{row["count"]}</td>'
                            html += f'<td>{row["percentage"]}%</td>'
                        html += '</tr>'
                    
                    html += '</tbody></table></div>'
                
                html += '</div>'
            
            html += '</div>'
        
        # Special handling for AI Conclusion section
        elif section['type'] == 'ai_conclusion':
            # Convert markdown formatting to HTML
            conclusion_content = section.get("content", "")
            formatted_content = markdown_to_html(conclusion_content)
            
            # Escape title but keep formatted_content as SafeString
            from django.utils.html import escape
            from django.utils.safestring import mark_safe
            
            escaped_title = escape(section["title"])
            # Use string concatenation to preserve SafeString
            html_part = '<div class="ai-conclusion-section mb-4"><div class="card border-info"><div class="card-header bg-info text-white"><h4 class="mb-0"><i class="fas fa-lightbulb me-2"></i>' + escaped_title + '</h4></div><div class="card-body"><div class="ai-conclusion-content" style="line-height: 1.6;">'
            html += mark_safe(html_part) + formatted_content + mark_safe('</div>')
            if section.get('model_name'):
                provider = section.get("provider", "")
                model_name = section.get("model_name", "")
                html += f'''
                        <hr>
                        <small class="text-muted">
                            <i class="fas fa-robot me-1"></i>
                            {_("Generated by")}: {escape(provider.upper())} - {escape(model_name)}
                        </small>
                '''
            html += '''
                    </div>
                </div>
            </div>
            '''
            html += '</div>'  # Close section div
        
        else:
            # Normal rendering for other sections
            for subsection in section['subsections']:
                html += f'<h5>{subsection["title"]}</h5>'
                
                if subsection['type'] == 'statistics':
                    # Render statistics as cards
                    # Check if display_inline flag is set to display all blocks in one row
                    if subsection.get('display_inline'):
                        html += '<div class="row flex-nowrap">'
                        for item in subsection['data']:
                            color_class = f"bg-{item['type']}"
                            html += f'''
                            <div class="col">
                                <div class="card {color_class} text-white">
                                    <div class="card-body text-center">
                                        <h4>{item['value']}</h4>
                                        <p class="mb-0">{item['label']}</p>
                                    </div>
                                </div>
                            </div>
                            '''
                        html += '</div>'
                    else:
                        # Default rendering for other statistics
                        html += '<div class="row">'
                        for item in subsection['data']:
                            color_class = f"bg-{item['type']}"
                            html += f'''
                            <div class="col-md-3">
                                <div class="card {color_class} text-white">
                                    <div class="card-body text-center">
                                        <h4>{item['value']}</h4>
                                        <p class="mb-0">{item['label']}</p>
                                    </div>
                                </div>
                            </div>
                            '''
                        html += '</div>'
                
                elif 'headers' in subsection and 'data' in subsection:
                    # Normal table rendering for all subsection types
                    html += '<div class="table-responsive"><table class="table table-striped"><thead><tr>'
                    for header in subsection['headers']:
                        # Use translations for common field names
                        translated_header = header
                        if header.lower() in ['type', 'title', 'name', 'description', 'location', 'group', 'status', 'comment', 'notes']:
                            # Get translation from unified_content translations if available
                            if 'translations' in unified_content:
                                translation_key = header.lower()
                                translated_header = unified_content['translations'].get(translation_key, header)
                        html += f'<th>{translated_header}</th>'
                    html += '</tr></thead><tbody>'
                    
                    for row in subsection['data']:
                        html += '<tr>'
                        if subsection['type'] == 'risk_distribution_level':
                            color = row.get('color', '#000000')
                            html += f'<td><span class="badge" style="background-color: {color}; color: white;">{row["risk_level"]}</span></td>'
                            html += f'<td>{row["count"]}</td>'
                            html += f'<td>{row["percentage"]}</td>'
                        elif subsection['type'] == 'top_risks':
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td>{row["vulnerability_name"]}</td>'
                            html += f'<td>{row["risk_level"]}</td>'
                        elif subsection['type'] == 'compliance_gaps':
                            html += f'<td>{row["type"]}</td>'
                            html += f'<td>{row["requirement"]}</td>'
                            html += f'<td>{row["description"]}</td>'
                        elif subsection['type'] == 'framework_company_requirements':
                            html += f'<td>{row.get("name", "")}</td>'
                            html += f'<td>{row.get("framework_type", "")}</td>'
                            html += f'<td>{row.get("version", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("status", "")}</td>'
                            html += f'<td>{row.get("controls_total", 0)}</td>'
                            html += f'<td>{row.get("controls_completed", 0)}</td>'
                            html += f'<td>{row.get("completion_percentage", "0%")}</td>'
                        elif subsection['type'] == 'company_requirements':
                            html += f'<td>{row.get("code", "")}</td>'
                            html += f'<td>{row.get("name", "")}</td>'
                            html += f'<td>{row.get("requirement_type", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("regulator", "")}</td>'
                            html += f'<td>{row.get("status", "")}</td>'
                            html += f'<td>{row.get("controls_total", 0)}</td>'
                            html += f'<td>{row.get("controls_completed", 0)}</td>'
                            html += f'<td>{row.get("completion_percentage", "0%")}</td>'
                        elif subsection['type'] == 'internal_requirements':
                            html += f'<td>{row.get("code", "")}</td>'
                            html += f'<td>{row.get("name", "")}</td>'
                            html += f'<td>{row.get("requirement_type", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("status", "")}</td>'
                            html += f'<td>{row.get("controls_total", 0)}</td>'
                            html += f'<td>{row.get("controls_completed", 0)}</td>'
                            html += f'<td>{row.get("completion_percentage", "0%")}</td>'
                        elif subsection['type'] == 'incident_table':
                            html += f'<td>{row.get("id", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("occurrence_datetime", "")}</td>'
                            html += f'<td>{row.get("place", "")}</td>'
                            html += f'<td>{row.get("incident_type", "")}</td>'
                            html += f'<td>{row.get("classification", "")}</td>'
                            html += f'<td>{row.get("current_state", "")}</td>'
                            html += f'<td>{row.get("responsible", "")}</td>'
                            html += f'<td>{row.get("description", "")}</td>'
                        elif subsection['type'] == 'mandatory_processes_table':
                            html += f'<td>{row.get("process_name", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("priority", "")}</td>'
                            html += f'<td>{row.get("frequency", "")}</td>'
                            html += f'<td>{row.get("next_due_date", "")}</td>'
                            html += f'<td>{row.get("last_completed_date", "")}</td>'
                            html += f'<td>{row.get("responsible_person", "")}</td>'
                            # Add status badge with color
                            status = row.get("status", "")
                            status_color = "secondary"
                            if status == "overdue":
                                status_color = "danger"
                            elif status == "upcoming":
                                status_color = "warning"
                            elif status == "completed":
                                status_color = "success"
                            html += f'<td><span class="badge bg-{status_color}">{status.capitalize() if status else ""}</span></td>'
                            html += f'<td>{row.get("description", "")}</td>'
                        elif subsection['type'] == 'certificate_key_management_table':
                            html += f'<td>{row.get("key_cert_num", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("type", "")}</td>'
                            html += f'<td>{row.get("expiry_date", "")}</td>'
                            # Add expiry status badge with color
                            expiry_status = row.get("expiry_status", "")
                            status_color = "secondary"
                            if expiry_status == "expired":
                                status_color = "danger"
                            elif expiry_status == "expiring_soon":
                                status_color = "warning"
                            elif expiry_status == "valid":
                                status_color = "success"
                            html += f'<td><span class="badge bg-{status_color}">{expiry_status.replace("_", " ").title() if expiry_status else ""}</span></td>'
                            days_until_expiry = row.get("days_until_expiry")
                            if days_until_expiry is not None:
                                html += f'<td>{days_until_expiry}</td>'
                            else:
                                html += '<td>-</td>'
                            html += f'<td>{row.get("revocation_status", "")}</td>'
                            html += f'<td>{row.get("owner_name", "")}</td>'
                            html += f'<td>{row.get("location", "")}</td>'
                            html += f'<td>{row.get("purpose", "")}</td>'
                        elif subsection['type'] == 'quiz_statistics_table':
                            html += f'<td>{row.get("quiz_title", "")}</td>'
                            html += f'<td>{row.get("total_attempts", "")}</td>'
                            html += f'<td>{row.get("successful_attempts", "")}</td>'
                            html += f'<td>{row.get("failed_attempts", "")}</td>'
                            html += f'<td>{row.get("success_rate", "")}</td>'
                            html += f'<td>{row.get("average_score", "")}</td>'
                            html += f'<td>{row.get("passing_score", "")}</td>'
                        elif subsection['type'] == 'users_at_risk_table':
                            html += f'<td>{row.get("user", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("attempts_count", "")}</td>'
                            if status == "overdue":
                                status_color = "danger"
                            elif status == "upcoming":
                                status_color = "warning"
                            elif status == "completed":
                                status_color = "success"
                            html += f'<td><span class="badge bg-{status_color}">{status.capitalize() if status else ""}</span></td>'
                            html += f'<td>{row.get("description", "")}</td>'
                        elif subsection['type'] == 'risk_distribution_table':
                            risk_level = row.get("risk_level", "")
                            risk_level_lower = risk_level.lower()
                            color = "#28a745"  # green for low
                            if risk_level_lower == "medium":
                                color = "#ffc107"  # yellow
                            elif risk_level_lower == "high":
                                color = "#fd7e14"  # orange
                            elif risk_level_lower == "critical":
                                color = "#dc3545"  # red
                            html += f'<td><span class="badge" style="background-color: {color}; color: white;">{risk_level}</span></td>'
                            html += f'<td>{row.get("count", "")}</td>'
                            html += f'<td>{row.get("percentage", "")}</td>'
                        elif subsection['type'] == 'assessment_status_table':
                            html += f'<td>{row.get("status", "")}</td>'
                            html += f'<td>{row.get("count", "")}</td>'
                        elif subsection['type'] == 'overdue_assessments_table':
                            html += f'<td>{row.get("vendor_name", "")}</td>'
                            html += f'<td>{row.get("assessment_date", "")}</td>'
                            html += f'<td>{row.get("next_review_date", "")}</td>'
                            html += f'<td>{row.get("status", "")}</td>'
                            days_overdue = row.get("days_overdue", "0")
                            html += f'<td><span class="badge bg-danger">{days_overdue}</span></td>'
                        elif subsection['type'] == 'incomplete_questionnaires_table':
                            html += f'<td>{row.get("vendor_name", "")}</td>'
                            html += f'<td>{row.get("template_name", "")}</td>'
                            html += f'<td>{row.get("status", "")}</td>'
                            html += f'<td>{row.get("started_date", "")}</td>'
                            html += f'<td>{row.get("percentage_score", "")}</td>'
                        elif subsection['type'] == 'risk_level_trends_table':
                            html += f'<td>{row.get("risk_level", "")}</td>'
                            html += f'<td>{row.get("count", "")}</td>'
                            html += f'<td>{row.get("trend", "")}</td>'
                        elif subsection['type'] == 'breach_severity_table':
                            severity = row.get("severity", "")
                            severity_lower = severity.lower()
                            color = "#28a745"  # green for low
                            if severity_lower == "medium":
                                color = "#ffc107"  # yellow
                            elif severity_lower == "high":
                                color = "#fd7e14"  # orange
                            elif severity_lower == "critical":
                                color = "#dc3545"  # red
                            html += f'<td><span class="badge" style="background-color: {color}; color: white;">{severity}</span></td>'
                            html += f'<td>{row.get("count", "")}</td>'
                            html += f'<td>{row.get("percentage", "")}</td>'
                        elif subsection['type'] == 'notification_status_table':
                            html += f'<td>{row.get("status", "")}</td>'
                            html += f'<td>{row.get("count", "")}</td>'
                        elif subsection['type'] == 'breach_notification_details_table':
                            html += f'<td>{row.get("incident_number", "")}</td>'
                            html += f'<td>{row.get("title", "")}</td>'
                            html += f'<td>{row.get("incident_date", "")}</td>'
                            html += f'<td>{row.get("discovery_date", "")}</td>'
                            html += f'<td>{row.get("notification_deadline", "")}</td>'
                            reported = row.get("reported", "")
                            reported_color = "success" if reported.lower() == "yes" or reported.lower() == "так" else "danger"
                            html += f'<td><span class="badge bg-{reported_color}">{reported}</span></td>'
                            html += f'<td>{row.get("authority_report_date", "")}</td>'
                            days_overdue = row.get("days_overdue", "0")
                            try:
                                if int(days_overdue) > 0:
                                    html += f'<td><span class="badge bg-danger">{days_overdue}</span></td>'
                                else:
                                    html += f'<td>{days_overdue}</td>'
                            except (ValueError, TypeError):
                                html += f'<td>{days_overdue}</td>'
                        elif subsection['type'] == 'data_categories_table':
                            html += f'<td>{row.get("category", "")}</td>'
                            html += f'<td>{row.get("count", "")}</td>'
                        elif subsection['type'] == 'international_transfers_table':
                            html += f'<td>{row.get("activity_name", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("transfer_safeguards", "")}</td>'
                        elif subsection['type'] == 'outdated_retention_table':
                            html += f'<td>{row.get("activity_name", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("retention_period_days", "")}</td>'
                            html += f'<td>{row.get("created_date", "")}</td>'
                            html += f'<td>{row.get("updated_date", "")}</td>'
                            days_overdue = row.get("days_overdue", "0")
                            try:
                                if int(days_overdue) > 0:
                                    html += f'<td><span class="badge bg-danger">{days_overdue}</span></td>'
                                else:
                                    html += f'<td>{days_overdue}</td>'
                            except (ValueError, TypeError):
                                html += f'<td>{days_overdue}</td>'
                        elif subsection['type'] == 'non_compliant_retention_table':
                            html += f'<td>{row.get("activity_name", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("retention_period_days", "")}</td>'
                            issue = row.get("issue", "")
                            issue_lower = issue.lower()
                            issue_color = "danger"
                            if "excessive" in issue_lower or "надмірний" in issue_lower:
                                issue_color = "warning"
                            elif "insufficient" in issue_lower or "недостатній" in issue_lower:
                                issue_color = "info"
                            html += f'<td><span class="badge bg-{issue_color}">{issue}</span></td>'
                        elif subsection['type'] == 'assets_by_criticality':
                            # Use color from data if available, otherwise use default
                            color = row.get('color', '#000000')
                            html += f'<td><span class="badge" style="background-color: {color}; color: white;">{row["criticality"]}</span></td>'
                            html += f'<td>{row["count"]}</td>'
                            html += f'<td>{row["percentage"]}%</td>'
                        elif subsection['type'] == 'assets_by_vulnerability':
                            # Use color from data if available, otherwise use default
                            color = row.get('color', '#000000')
                            html += f'<td><span class="badge" style="background-color: {color}; color: white;">{row["criticality"]}</span></td>'
                            html += f'<td>{row["count"]}</td>'
                            html += f'<td>{row["percentage"]}%</td>'
                        elif subsection['type'] == 'asset_table':
                            html += f'<td>{row["asset_id"]}</td>'
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td><span style="color:{row["criticality_color"]}">{row["criticality"]}</span></td>'
                            html += f'<td>{row.get("company", "")}</td>'
                        elif subsection['type'] == 'top_10_asset':
                            html += f'<td>{row["asset_id"]}</td>'
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td><span style="color:{row.get("criticality_color", "#000000")}">{row["criticality"]}</span></td>'
                            # Display vulnerabilities with proper logic: "Undefined" if no assessment, "0" if assessed but no vulnerabilities, or count
                            vulnerabilities_display = row.get("vulnerabilities", "Undefined")
                            if vulnerabilities_display == "Undefined":
                                html += f'<td><span class="badge bg-secondary">Undefined</span></td>'
                            elif vulnerabilities_display == "0":
                                html += f'<td><span class="badge bg-secondary">0</span></td>'
                            else:
                                html += f'<td><span class="badge bg-danger">{vulnerabilities_display}</span></td>'
                            html += f'<td>{row.get("company", "")}</td>'
                        elif subsection['type'] == 'vulnerability_table':
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td>{row["vulnerability_name"]}</td>'
                            html += f'<td>{row["status"]}</td>'
                            html += f'<td>{row["threat"]}</td>'
                            html += f'<td>{row["probability_impact"]}</td>'
                            html += f'<td>{row["impact"]}</td>'
                            html += f'<td>{row["risk_value"]}</td>'
                            html += f'<td>{row["risk_level"]}</td>'
                            html += f'<td>{row["risk_mitigation_controls"]}</td>'
                        elif subsection['type'] == 'overdue_treatments':
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td>{row["treatment"]}</td>'
                            html += f'<td>{row["deadline"]}</td>'
                        elif subsection['type'] == 'treatment_details':
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td>{row["vulnerability_name"]}</td>'
                            html += f'<td>{row["vulnerability_status"]}</td>'
                            html += f'<td>{row["risk_level"]}</td>'
                            html += f'<td>{row["treatment_type"]}</td>'
                            html += f'<td>{row["treatment_status"]}</td>'
                            html += f'<td>{row["description_responsible"]}</td>'
                            html += f'<td>{row["due_date"]}</td>'
                        elif subsection['type'] == 'table':
                            # Generic table rendering for acceptable risk tables
                            for key, value in row.items():
                                html += f'<td>{value}</td>'
                    html += '</tr>'
                
                    html += '</tbody></table></div>'
                
                elif subsection['type'] == 'pci_dss' or subsection['type'] == 'iso27001':
                    # Render compliance data as list
                    for item in subsection['data']:
                        html += f'<p><strong>{item["label"]}:</strong> {item["value"]}</p>'
                
                elif subsection['type'] == 'financial_overview':
                    # Render financial overview as cards
                    html += '<div class="row">'
                    for item in subsection['data']:
                        html += f'''
                        <div class="col-md-3">
                            <div class="card bg-info text-white">
                                <div class="card-body text-center">
                                    <h4>{item["value"]}</h4>
                                    <p class="mb-0">{item["label"]}</p>
                                </div>
                            </div>
                        </div>
                        '''
                    html += '</div>'
                
                elif subsection['type'] == 'residual_overview':
                    # Render residual overview as cards
                    html += '<div class="row">'
                    for item in subsection['data']:
                        html += f'''
                        <div class="col-md-4">
                            <div class="card bg-info text-white">
                                <div class="card-body text-center">
                                    <h4>{item["value"]}</h4>
                                    <p class="mb-0">{item["label"]}</p>
                                </div>
                            </div>
                        </div>
                        '''
                    html += '</div>'
                
                elif subsection['type'] == 'governance_overview':
                    # Render governance overview as cards
                    html += '<div class="row">'
                    for item in subsection['data']:
                        html += f'''
                        <div class="col-md-3">
                            <div class="card bg-primary text-white">
                                <div class="card-body text-center">
                                    <h4>{item["value"]}</h4>
                                    <p class="mb-0">{item["label"]}</p>
                                </div>
                            </div>
                        </div>
                        '''
                    html += '</div>'
                
                elif subsection['type'] == 'priority_actions':
                    # Render recommendations as list
                    html += '<ul>'
                    for item in subsection['data']:
                        html += f'<li>{item}</li>'
                    html += '</ul>'
        
            html += '</div>'
    
    html += '</div>'
    # Mark the entire HTML as safe since it contains SafeString parts
    from django.utils.safestring import mark_safe
    return mark_safe(html)


def generate_html_preview(data, report_type, language, selected_sections):
    """Generate HTML preview of the report content using unified content structure"""
    
    translations = get_report_translations(language)
    
    # Generate unified content structure
    unified_content = generate_unified_report_content(data, report_type, selected_sections, translations, language)
    
    # Convert unified content to HTML
    return render_unified_content_as_html(unified_content)


def generate_full_report_preview(data, translations, selected_sections):
    """Generate full report preview HTML"""
    
    # Get company information
    company_name = None
    if data.get('companies'):
        if data['companies'].count() == 1:
            company_name = data['companies'].first().name
        elif data['companies'].count() > 1:
            company_name = f"{translations.get('multiple_companies', 'Multiple Companies')} ({data['companies'].count()})"
    
    # Build company display text
    company_display = ""
    if company_name:
        company_display = f'<p class="text-muted mb-1"><i class="fas fa-building me-2"></i><strong>{translations.get("company", "Company")}:</strong> {company_name}</p>'
    
    # Get translations for confidentiality classification
    from django.utils.translation import gettext as _
    confidentiality_classification = _("Confidentiality Classification")
    confidential_information = _("Confidential Information")
    
    html = f'''
    <div class="report-preview">
        <div class="report-header text-center mb-4">
            <h2 class="text-primary">{translations['full_report_title']}</h2>
            {company_display}
            <p class="text-muted">{translations['generated_on']}: {format_localized_date(data['generation_date'], data['language'])}</p>
            <div class="confidentiality-classification mt-3 mb-2">
                <p class="text-danger fw-bold"><i class="fas fa-shield-alt me-2"></i><strong>{confidentiality_classification}:</strong> {confidential_information}</p>
            </div>
            <hr>
        </div>
    '''
    
    # Executive Summary Section
    if selected_sections.get('statistics') or selected_sections.get('top_risks'):
        html += f'''
        <div class="executive-summary mb-4">
            <h3 class="text-secondary"><i class="fas fa-chart-line me-2"></i>{translations['executive_summary']}</h3>
        '''
        
        # Key Statistics
        if selected_sections.get('statistics'):
            html += f'''
            <div class="row">
                <div class="col-md-3">
                    <div class="card bg-primary text-white">
                        <div class="card-body text-center">
                            <h4>{data['statistics']['total_assets']}</h4>
                            <p class="mb-0">{translations['total_assets']}</p>
                        </div>
                    </div>
                </div>
                <div class="col-md-3">
                    <div class="card bg-warning text-white">
                        <div class="card-body text-center">
                            <h4>{data['statistics']['total_vulnerabilities']}</h4>
                            <p class="mb-0">{translations['total_vulnerabilities']}</p>
                        </div>
                    </div>
                </div>
                <div class="col-md-3">
                    <div class="card bg-danger text-white">
                        <div class="card-body text-center">
                            <h4>{data['statistics']['high_risk_count']}</h4>
                            <p class="mb-0">{translations['high_risk_assets']}</p>
                        </div>
                    </div>
                </div>
                <div class="col-md-3">
                    <div class="card bg-success text-white">
                        <div class="card-body text-center">
                            <h4>{data['statistics']['completion_rate']:.1f}%</h4>
                            <p class="mb-0">{translations['completion_rate']}</p>
                        </div>
                    </div>
                </div>
            </div>
            '''
        
        html += '</div>'
    
    # Risk Details Section (Combined By Risk Level and TOP-10 Risks)
    if selected_sections.get('risk_distribution') or selected_sections.get('top_risks'):
        html += f'''
        <div class="risk-details mb-4">
            <h3 class="text-secondary"><i class="fas fa-chart-pie me-2"></i>{translations.get('risk_details', 'Risk Details')}</h3>
            <div class="row">
        '''
        
        # By Risk Level (left column)
        if selected_sections.get('risk_distribution'):
            html += f'''
                <div class="col-md-6">
                    <h5>{translations['by_risk_level']}</h5>
                    <div class="table-responsive">
                        <table class="table table-striped">
                            <thead>
                                <tr>
                                    <th>{translations['risk_level']}</th>
                                    <th>{translations['count']}</th>
                                    <th>{translations['percentage']}</th>
                                </tr>
                            </thead>
                            <tbody>
    '''
        
        # Add risk level distribution with sorting and colors
        total_risks = sum(data['risk_distribution']['by_level'].values())
        risk_levels_with_colors = data.get('statistics', {}).get('risk_levels_with_colors', {})
        
        # Create list for sorting
        risk_level_items = []
        for level, count in data['risk_distribution']['by_level'].items():
            percentage = (count / total_risks * 100) if total_risks > 0 else 0
            color_info = risk_levels_with_colors.get(level, {})
            risk_level_items.append({
                'level': level,
                'count': count,
                'percentage': percentage,
                'color': color_info.get('color', '#000000'),
                'max_value': color_info.get('max_value', 0)
            })
        
        # Sort by max_value in descending order (highest risk first)
        risk_level_items.sort(key=lambda x: x['max_value'], reverse=True)
        
        for item in risk_level_items:
            html += f'''
                                <tr>
                                    <td><span class="badge" style="background-color: {item['color']}; color: white;">{item['level']}</span></td>
                                    <td>{item['count']}</td>
                                    <td>{item['percentage']:.1f}%</td>
                                </tr>
            '''
        
            html += '''
                            </tbody>
                        </table>
                    </div>
                </div>
            '''
        
        # TOP-10 Risks (right column)
        if selected_sections.get('top_risks'):
            html += f'''
                <div class="col-md-6">
                    <h5>{translations.get('top_risks', 'TOP-10 Risks')}</h5>
                    <div class="table-responsive">
                        <table class="table table-striped">
                            <thead>
                                <tr>
                                    <th>{translations.get('asset_name', 'Asset Name')}</th>
                                    <th>{translations['vulnerability']}</th>
                                    <th>{translations['risk_level']}</th>
                                </tr>
                            </thead>
                            <tbody>
    '''
        
            # Add TOP-10 risks data
            if 'statistics' in data and data['statistics'].get('high_risk_assets'):
                for asset in data['statistics']['high_risk_assets'][:10]:  # Limit to 10
                    html += f'''
                                <tr>
                                        <td>{asset.get('asset_name', translations.get('unknown', 'Unknown'))}</td>
                                        <td>{asset.get('vulnerability_name', translations.get('unknown', 'Unknown'))}</td>
                                        <td>{asset.get('risk_level', translations.get('unknown', 'Unknown'))}</td>
                                </tr>
                    '''
        
        html += '''
                            </tbody>
                        </table>
                    </div>
                </div>
            '''
        
        html += '''
            </div>
        </div>
        '''
    
    # Compliance Overview Section
    if selected_sections.get('compliance_overview') or selected_sections.get('pci_dss') or selected_sections.get('iso27001'):
        html += f'''
        <div class="compliance-overview mb-4">
            <h3 class="text-secondary"><i class="fas fa-shield-alt me-2"></i>{translations['compliance_overview']}</h3>
            <div class="row">
        '''
        
        if selected_sections.get('pci_dss'):
            html += f'''
                <div class="col-md-6">
                    <div class="card">
                        <div class="card-header">
                            <h5 class="mb-0">PCI DSS</h5>
                        </div>
                        <div class="card-body">
                            <p><strong>{translations['compliant_vulnerabilities']}:</strong> {data['compliance']['pcidss']['compliant_count']}</p>
                            <p><strong>{translations['non_compliant_vulnerabilities']}:</strong> {data['compliance']['pcidss']['non_compliant_count']}</p>
                            <p><strong>{translations['compliance_rate']}:</strong> {data['compliance']['pcidss']['compliance_rate']:.1f}%</p>
                        </div>
                    </div>
                </div>
            '''
        
        if selected_sections.get('iso27001'):
            html += f'''
                <div class="col-md-6">
                    <div class="card">
                        <div class="card-header">
                            <h5 class="mb-0">ISO 27001</h5>
                        </div>
                        <div class="card-body">
                            <p><strong>{translations['compliant_vulnerabilities']}:</strong> {data['compliance']['iso27001']['compliant_count']}</p>
                            <p><strong>{translations['non_compliant_vulnerabilities']}:</strong> {data['compliance']['iso27001']['non_compliant_count']}</p>
                            <p><strong>{translations['compliance_rate']}:</strong> {data['compliance']['iso27001']['compliance_rate']:.1f}%</p>
                        </div>
                    </div>
                </div>
            '''
        
        html += '''
            </div>
        </div>
        '''
    
    # Asset Details Section - Combined Assets by Criticality and TOP-10 Asset
    if selected_sections.get('asset_details') or selected_sections.get('asset_tables'):
        html += f'''
        <div class="asset-details mb-4">
            <h3 class="text-secondary"><i class="fas fa-server me-2"></i>{translations.get('asset_details', 'Asset Details')}</h3>
            <div class="row">
        '''
        
        # Assets by Criticality
        if selected_sections.get('asset_details') and 'statistics' in data:
            stats = data['statistics']
            if stats.get('assets_by_criticality'):
                html += '''
                <div class="col-md-6">
                    <h5>Assets by Criticality</h5>
            <div class="table-responsive">
                <table class="table table-striped">
                    <thead>
                        <tr>
                                    <th>Criticality</th>
                                    <th>Count</th>
                        </tr>
                    </thead>
                    <tbody>
    '''
        
                for item in stats['assets_by_criticality']:
                    criticality_name = item.get('criticality__name_uk', translations.get('unknown', 'Unknown'))
                    color = item.get('color', '#000000')
                    count = item.get('count', 0)
                    
                    html += f'''
                                <tr>
                                    <td><span class="badge" style="background-color: {color}; color: white;">{criticality_name}</span></td>
                                    <td>{count}</td>
                                </tr>
                    '''
                
                html += '''
                            </tbody>
                        </table>
                    </div>
                </div>
                '''
        
        # TOP-10 Asset
        if selected_sections.get('asset_tables') and 'assets' in data:
            html += '''
            <div class="col-md-6">
                <h5>{translations.get('top_10_asset', 'TOP-10 Asset')}</h5>
                <div class="table-responsive">
                    <table class="table table-striped">
                        <thead>
                            <tr>
                                <th>Asset ID</th>
                                <th>Asset Name</th>
                                <th>Criticality</th>
                            </tr>
                        </thead>
                        <tbody>
            '''
            
            # Get top 10 assets by criticality cost
            asset_table_data = []
            for asset in data['assets']:
                if hasattr(asset, 'get_criticality'):
                    criticality_info = asset.get_criticality()
                    criticality_name = criticality_info.get('name', '') if criticality_info else ''
                    criticality_color = criticality_info.get('color', '#000000') if criticality_info else '#000000'
                    criticality_cost = criticality_info.get('cost', 0) if criticality_info else 0
                    
                    asset_table_data.append({
                        'asset_id': getattr(asset, 'asset_id', ''),
                        'asset_name': getattr(asset, 'name', ''),
                        'criticality': criticality_name,
                        'criticality_color': criticality_color,
                        'criticality_cost': criticality_cost
                    })
            
            # Sort by criticality cost (highest first) and limit to 10
            asset_table_data.sort(key=lambda x: x['criticality_cost'], reverse=True)
            asset_table_data = asset_table_data[:10]
            
            for asset_data in asset_table_data:
                html += f'''
                        <tr>
                                <td>{asset_data["asset_id"]}</td>
                                <td>{asset_data["asset_name"]}</td>
                                <td><span style="color:{asset_data["criticality_color"]}">{asset_data["criticality"]}</span></td>
                        </tr>
                '''
        
        html += '''
                    </tbody>
                </table>
            </div>
        </div>
        '''
    
        html += '''
            </div>
        </div>
        '''
    
    # Treatment Details Section
    if selected_sections.get('treatment_details') or selected_sections.get('treatment_tables'):
        html += f'''
        <div class="treatment-details mb-4">
            <h3 class="text-secondary"><i class="fas fa-tasks me-2"></i>{translations.get('treatment_details', 'Treatment Details')}</h3>
            <div class="table-responsive">
                <table class="table table-striped">
                    <thead>
                        <tr>
                            <th>{translations['asset']}</th>
                            <th>{translations['vulnerability']}</th>
                            <th>{translations.get('due_date', 'Due Date')}</th>
                            <th>{translations['status']}</th>
                        </tr>
                    </thead>
                    <tbody>
    '''
        
        # Add treatment details (first 10 for preview)
        for treatment in data.get('overdue_treatments', [])[:10]:
            html += f'''
                        <tr>
                            <td>{treatment['asset_name']}</td>
                            <td>{treatment['vulnerability_name']}</td>
                            <td>{format_localized_date(treatment['due_date'], data['language'])}</td>
                            <td><span class="badge bg-danger">{treatment['status']}</span></td>
                        </tr>
            '''
        
        html += '''
                    </tbody>
                </table>
            </div>
        </div>
        '''
        
        # Add Treatment Details subsection if requested
        if selected_sections.get('treatment_details') and 'risk_treatments' in data:
            html += f'''
        <div class="treatment-details mb-4">
            <h4 class="text-info"><i class="fas fa-clipboard-list me-2"></i>{translations.get('top_10_treatment', 'TOP-10 Treatment')}</h4>
            <div class="table-responsive">
                <table class="table table-striped">
                    <thead>
                        <tr>
                            <th>{translations['asset']}</th>
                            <th>{translations.get('vulnerability', 'Vulnerability')}</th>
                            <th>{translations.get('vulnerability_status', 'Vulnerability Status')}</th>
                            <th>{translations.get('risk_level', 'Risk Level')}</th>
                            <th>{translations.get('treatment_type', 'Treatment Type')}</th>
                            <th>{translations.get('treatment_status', 'Treatment Status')}</th>
                            <th>{translations.get('description_responsible', 'Description & Responsible')}</th>
                            <th>{translations.get('due_date', 'Due Date')}</th>
                        </tr>
                    </thead>
                    <tbody>
            '''
            
            # Add treatment details (first 10 for preview with sorting)
            treatments_for_preview = []
            for treatment in data.get('risk_treatments', []):
                # Get vulnerability status and risk level for this treatment
                vulnerability_status = translations.get('unknown', 'Unknown')
                risk_level = translations.get('unknown', 'Unknown')
                if 'vulnerabilities' in data:
                    for vuln in data['vulnerabilities']:
                        # Handle both dictionary and model object cases
                        if hasattr(vuln, 'get'):  # Dictionary
                            vuln_asset_name = vuln.get('asset_name', '')
                            vuln_name = vuln.get('vulnerability_name', '')
                        else:  # Model object
                            vuln_asset_name = getattr(vuln, 'asset_name', '') if hasattr(vuln, 'asset_name') else (getattr(vuln.asset, 'name', '') if hasattr(vuln, 'asset') and vuln.asset else '')
                            vuln_name = getattr(vuln, 'vulnerability_name', '') if hasattr(vuln, 'vulnerability_name') else (get_localized_vulnerability_field(vuln.vulnerability, 'vulnerability', data.get('language', 'uk')) if hasattr(vuln, 'vulnerability') and vuln.vulnerability else '')
                        
                        # Match by asset name and vulnerability name
                        if (vuln_asset_name == treatment.get('asset_name', '') and 
                            vuln_name == treatment.get('vulnerability_name', '')):
                            vulnerability_status = getattr(vuln, 'status', translations.get('unknown', 'Unknown'))
                            
                            # Calculate risk level for this vulnerability
                            try:
                                if hasattr(vuln, 'asset') and hasattr(vuln, 'vulnerability'):
                                    # For model objects
                                    risk_value = calculate_value_of_risk(vuln.asset, vuln.vulnerability)
                                    risk_level_obj = calculate_risk_level(risk_value)
                                    if risk_level_obj:
                                        risk_level = risk_level_obj.get_name_by_language(data.get("language", "uk")) or risk_level_obj.get_name()
                                    else:
                                        risk_level = translations.get('unknown', 'Unknown')
                                elif hasattr(vuln, 'get'):
                                    # For dictionary objects, try to get risk level if available
                                    risk_level = vuln.get('risk_level', translations.get('unknown', 'Unknown'))
                                else:
                                    risk_level = translations.get('unknown', 'Unknown')
                            except Exception as e:
                                logger.warning(f"Error calculating risk level: {e}")
                                risk_level = translations.get('unknown', 'Unknown')
                            break
                
                # Calculate priority for sorting
                status_priority = 0
                if vulnerability_status == 'Yes':
                    status_priority = 1
                
                risk_level_priority = 0
                if risk_level:
                    if 'Надзвичайний' in risk_level or 'Extraordinary' in risk_level:
                        risk_level_priority = 6
                    elif 'Критичний' in risk_level or 'Critical' in risk_level:
                        risk_level_priority = 5
                    elif 'Високий' in risk_level or 'High' in risk_level:
                        risk_level_priority = 4
                    elif 'Середній' in risk_level or 'Medium' in risk_level:
                        risk_level_priority = 3
                    elif 'Низький' in risk_level or 'Low' in risk_level:
                        risk_level_priority = 2
                    elif 'Невизначено' in risk_level or 'Undefined' in risk_level:
                        risk_level_priority = 1
                
                treatments_for_preview.append({
                    'asset_name': treatment.get('asset_name', translations.get('unknown', 'Unknown')),
                    'vulnerability_name': treatment.get('vulnerability_name', translations.get('unknown', 'Unknown')),
                    'vulnerability_status': vulnerability_status,
                    'risk_level': risk_level,
                    'treatment_type': treatment.get('treatment_type', translations.get('unknown', 'Unknown')),
                    'treatment_status': treatment.get('status', translations.get('unknown', 'Unknown')),
                    'description_responsible': treatment.get('description', '') + ' | ' + treatment.get('assigned_to', translations.get('unknown', 'Unknown')),
                    'due_date': format_localized_date(treatment.get('due_date'), data['language']) if treatment.get('due_date') else translations.get('unknown', 'Unknown'),
                    'status_priority': status_priority,
                    'risk_level_priority': risk_level_priority
                })
            
            # Sort by status priority (Yes first) and then by risk level priority (highest first)
            treatments_for_preview.sort(key=lambda x: (x['status_priority'], x['risk_level_priority']), reverse=True)
            
            # Limit to top 10
            treatments_for_preview = treatments_for_preview[:10]
            
            # Generate HTML for sorted treatments
            for treatment in treatments_for_preview:
                html += f'''
                        <tr>
                            <td>{treatment["asset_name"]}</td>
                            <td>{treatment["vulnerability_name"]}</td>
                            <td><span class="badge bg-info">{treatment["vulnerability_status"]}</span></td>
                            <td><span class="badge bg-warning">{treatment["risk_level"]}</span></td>
                            <td>{treatment["treatment_type"]}</td>
                            <td><span class="badge bg-primary">{treatment["treatment_status"]}</span></td>
                            <td>{treatment["description_responsible"]}</td>
                            <td>{treatment["due_date"]}</td>
                        </tr>
                '''
            
            html += '''
                    </tbody>
                </table>
            </div>
        </div>
        '''
    
    # Charts & Graphs Section
    if selected_sections.get('charts') or selected_sections.get('graphs') or selected_sections.get('metrics'):
        html += f'''
        <div class="charts-section mb-4">
            <h3 class="text-secondary"><i class="fas fa-chart-bar me-2"></i>{translations.get('charts_graphs', 'Charts & Graphs')}</h3>
            <div class="alert alert-info">
                <i class="fas fa-info-circle me-2"></i>
                {translations.get('charts_preview_note', 'Charts and graphs will be included in the final report.')}
            </div>
        </div>
        '''
    
    # Recommendations Section
    if selected_sections.get('recommendations'):
        html += f'''
        <div class="recommendations mb-4">
            <h3 class="text-secondary"><i class="fas fa-lightbulb me-2"></i>{translations.get('recommendations', 'Recommendations')}</h3>
            <div class="alert alert-warning">
                <h6>{translations.get('priority_recommendations', 'Priority Recommendations')}:</h6>
                <ul>
                    <li>{translations.get('recommendation_1', 'Address high-risk vulnerabilities immediately')}</li>
                    <li>{translations.get('recommendation_2', 'Implement regular security assessments')}</li>
                    <li>{translations.get('recommendation_3', 'Update compliance documentation')}</li>
                </ul>
            </div>
        </div>
        '''
    
    # Report Footer
    html += f'''
        <div class="report-footer text-center mt-4">
            <p class="text-muted">{translations['report_generated_by']}: {data['generated_by']}</p>
            <p class="text-muted">{translations.get('preview_note', 'This is a preview. The final report will contain complete data.')}</p>
        </div>
    </div>
    '''
    
    return html


def generate_summary_report_preview(data, translations, selected_sections):
    """Generate summary report preview HTML"""
    
    # Get company information
    company_name = None
    if data.get('companies'):
        if data['companies'].count() == 1:
            company_name = data['companies'].first().name
        elif data['companies'].count() > 1:
            company_name = f"{translations.get('multiple_companies', 'Multiple Companies')} ({data['companies'].count()})"
    
    # Build company display text
    company_display = ""
    if company_name:
        company_display = f'<p class="text-muted mb-1"><i class="fas fa-building me-2"></i><strong>{translations.get("company", "Company")}:</strong> {company_name}</p>'
    
    # Get translations for confidentiality classification
    from django.utils.translation import gettext as _
    confidentiality_classification = _("Confidentiality Classification")
    confidential_information = _("Confidential Information")
    
    html = f'''
    <div class="report-preview">
        <div class="report-header text-center mb-4">
            <h2 class="text-primary">{translations['executive_summary']}</h2>
            {company_display}
            <p class="text-muted">{translations['generated_on']}: {format_localized_date(data['generation_date'], data['language'])}</p>
            <div class="confidentiality-classification mt-3 mb-2">
                <p class="text-danger fw-bold"><i class="fas fa-shield-alt me-2"></i><strong>{confidentiality_classification}:</strong> {confidential_information}</p>
            </div>
            <hr>
        </div>
    '''
    
    # Key Statistics Section
    if selected_sections.get('statistics'):
        html += f'''
        <div class="key-metrics mb-4">
            <h3 class="text-secondary"><i class="fas fa-tachometer-alt me-2"></i>{translations['key_metrics']}</h3>
            <div class="row">
                <div class="col-md-4">
                    <div class="card bg-primary text-white">
                        <div class="card-body text-center">
                            <h3>{data['statistics']['total_assets']}</h3>
                            <p class="mb-0">{translations['total_assets']}</p>
                        </div>
                    </div>
                </div>
                <div class="col-md-4">
                    <div class="card bg-warning text-white">
                        <div class="card-body text-center">
                            <h3>{data['statistics']['total_vulnerabilities']}</h3>
                            <p class="mb-0">{translations['total_vulnerabilities']}</p>
                        </div>
                    </div>
                </div>
                <div class="col-md-4">
                    <div class="card bg-danger text-white">
                        <div class="card-body text-center">
                            <h3>{data['statistics']['high_risk_count']}</h3>
                            <p class="mb-0">{translations['high_risk_assets']}</p>
                        </div>
                    </div>
                </div>
            </div>
        </div>
        '''
    
    # Note: TOP-10 Risks is now combined with Risk Details section above
    
    # Compliance Overview Section
    if selected_sections.get('compliance_overview'):
        html += f'''
        <div class="compliance-summary mb-4">
            <h3 class="text-secondary"><i class="fas fa-shield-alt me-2"></i>{translations['compliance_overview']}</h3>
            <div class="row">
                <div class="col-md-6">
                    <div class="card">
                        <div class="card-header">
                            <h6 class="mb-0">PCI DSS</h6>
                        </div>
                        <div class="card-body">
                            <div class="d-flex justify-content-between">
                                <span>{translations.get('compliance_rate', 'Compliance Rate')}:</span>
                                <strong>{data['compliance']['pcidss']['compliance_rate']:.1f}%</strong>
                            </div>
                        </div>
                    </div>
                </div>
                <div class="col-md-6">
                    <div class="card">
                        <div class="card-header">
                            <h6 class="mb-0">ISO 27001</h6>
                        </div>
                        <div class="card-body">
                            <div class="d-flex justify-content-between">
                                <span>{translations.get('compliance_rate', 'Compliance Rate')}:</span>
                                <strong>{data['compliance']['iso27001']['compliance_rate']:.1f}%</strong>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
        </div>
        '''
    
    # Recommendations Section
    if selected_sections.get('recommendations'):
        html += f'''
        <div class="recommendations mb-4">
            <h3 class="text-secondary"><i class="fas fa-lightbulb me-2"></i>{translations.get('recommendations', 'Recommendations')}</h3>
            <div class="alert alert-info">
                <h6>{translations.get('key_recommendations', 'Key Recommendations')}:</h6>
                <ol>
                    <li>{translations.get('recommendation_1', 'Address high-risk vulnerabilities immediately')}</li>
                    <li>{translations.get('recommendation_2', 'Implement regular security assessments')}</li>
                    <li>{translations.get('recommendation_3', 'Update compliance documentation')}</li>
                </ol>
            </div>
        </div>
        '''
    
    # Charts & Metrics Section
    if selected_sections.get('charts') or selected_sections.get('metrics'):
        html += f'''
        <div class="charts-section mb-4">
            <h3 class="text-secondary"><i class="fas fa-chart-line me-2"></i>{translations.get('performance_metrics', 'Performance Metrics')}</h3>
            <div class="alert alert-info">
                <i class="fas fa-info-circle me-2"></i>
                {translations.get('metrics_preview_note', 'Performance metrics and charts will be included in the final report.')}
            </div>
        </div>
        '''
    
    # Report Footer
    html += f'''
        <div class="report-footer text-center mt-4">
            <p class="text-muted">{translations['report_generated_by']}: {data['generated_by']}</p>
            <p class="text-muted">{translations.get('preview_note', 'This is a preview. The final report will contain complete data.')}</p>
        </div>
    </div>
    '''
    
    return html


def generate_compliance_report_preview(data, translations, selected_sections):
    """Generate compliance report preview HTML"""
    
    # Get company information
    company_name = None
    if data.get('companies'):
        if data['companies'].count() == 1:
            company_name = data['companies'].first().name
        elif data['companies'].count() > 1:
            company_name = f"{translations.get('multiple_companies', 'Multiple Companies')} ({data['companies'].count()})"
    
    # Build company display text
    company_display = ""
    if company_name:
        company_display = f'<p class="text-muted mb-1"><i class="fas fa-building me-2"></i><strong>{translations.get("company", "Company")}:</strong> {company_name}</p>'
    
    # Get translations for confidentiality classification
    from django.utils.translation import gettext as _
    confidentiality_classification = _("Confidentiality Classification")
    confidential_information = _("Confidential Information")
    
    html = f'''
    <div class="report-preview">
        <div class="report-header text-center mb-4">
            <h2 class="text-primary">{translations.get('compliance_report', 'Compliance Report')}</h2>
            {company_display}
            <p class="text-muted">{translations['generated_on']}: {format_localized_date(data['generation_date'], data['language'])}</p>
            <div class="confidentiality-classification mt-3 mb-2">
                <p class="text-danger fw-bold"><i class="fas fa-shield-alt me-2"></i><strong>{confidentiality_classification}:</strong> {confidential_information}</p>
            </div>
            <hr>
        </div>
    '''
    
    # Compliance Overview Section
    if selected_sections.get('compliance_overview') or selected_sections.get('pci_dss') or selected_sections.get('iso27001'):
        html += f'''
        <div class="compliance-overview mb-4">
            <h3 class="text-secondary"><i class="fas fa-shield-alt me-2"></i>{translations['compliance_overview']}</h3>
            <div class="row">
        '''
        
        if selected_sections.get('pci_dss'):
            html += f'''
                <div class="col-md-6">
                    <div class="card">
                        <div class="card-header bg-primary text-white">
                            <h5 class="mb-0">PCI DSS {translations.get('compliance', 'Compliance')}</h5>
                        </div>
                        <div class="card-body">
                            <div class="row">
                                <div class="col-6">
                                    <p><strong>{translations.get('compliant', 'Compliant')}:</strong></p>
                                    <h4 class="text-success">{data['compliance']['pcidss']['compliant_count']}</h4>
                                </div>
                                <div class="col-6">
                                    <p><strong>{translations.get('non_compliant', 'Non-Compliant')}:</strong></p>
                                    <h4 class="text-danger">{data['compliance']['pcidss']['non_compliant_count']}</h4>
                                </div>
                            </div>
                            <div class="progress mt-3">
                                <div class="progress-bar bg-success" style="width: {data['compliance']['pcidss']['compliance_rate']:.1f}%">
                                    {data['compliance']['pcidss']['compliance_rate']:.1f}%
                                </div>
                            </div>
                        </div>
                    </div>
                </div>
            '''
        
        if selected_sections.get('iso27001'):
            html += f'''
                <div class="col-md-6">
                    <div class="card">
                        <div class="card-header bg-info text-white">
                            <h5 class="mb-0">ISO 27001 {translations.get('compliance', 'Compliance')}</h5>
                        </div>
                        <div class="card-body">
                            <div class="row">
                                <div class="col-6">
                                    <p><strong>{translations.get('compliant', 'Compliant')}:</strong></p>
                                    <h4 class="text-success">{data['compliance']['iso27001']['compliant_count']}</h4>
                                </div>
                                <div class="col-6">
                                    <p><strong>{translations.get('non_compliant', 'Non-Compliant')}:</strong></p>
                                    <h4 class="text-danger">{data['compliance']['iso27001']['non_compliant_count']}</h4>
                                </div>
                            </div>
                            <div class="progress mt-3">
                                <div class="progress-bar bg-success" style="width: {data['compliance']['iso27001']['compliance_rate']:.1f}%">
                                    {data['compliance']['iso27001']['compliance_rate']:.1f}%
                                </div>
                            </div>
                        </div>
                    </div>
                </div>
            '''
        
        html += '''
            </div>
        </div>
        '''
    
    # Compliance Gaps Section
    if selected_sections.get('compliance_gaps'):
        html += f'''
        <div class="compliance-gaps mb-4">
            <h3 class="text-secondary"><i class="fas fa-exclamation-circle me-2"></i>{translations.get('compliance_gaps', 'Compliance Gaps')}</h3>
            <div class="row">
        '''
        
        if selected_sections.get('pci_dss'):
            html += f'''
                <div class="col-md-6">
                    <h5>PCI DSS {translations.get('gaps', 'Gaps')}</h5>
                    <div class="table-responsive">
                        <table class="table table-striped">
                            <thead>
                                <tr>
                                    <th>{translations.get('vulnerability', 'Vulnerability')}</th>
                                    <th>{translations.get('status', 'Status')}</th>
                                </tr>
                            </thead>
                            <tbody>
            '''
            
            # Add PCI DSS gaps
            for gap in data['compliance']['pcidss'].get('gaps', [])[:5]:  # Show first 5
                html += f'''
                                <tr>
                                    <td>{gap.get('vulnerability_name', 'N/A')}</td>
                                    <td><span class="badge bg-danger">{translations.get('non_compliant', 'Non-Compliant')}</span></td>
                                </tr>
                '''
            
            html += '''
                            </tbody>
                        </table>
                    </div>
                </div>
            '''
        
        if selected_sections.get('iso27001'):
            html += f'''
                <div class="col-md-6">
                    <h5>ISO 27001 {translations.get('gaps', 'Gaps')}</h5>
                    <div class="table-responsive">
                        <table class="table table-striped">
                            <thead>
                                <tr>
                                    <th>{translations.get('vulnerability', 'Vulnerability')}</th>
                                    <th>{translations.get('status', 'Status')}</th>
                                </tr>
                            </thead>
                            <tbody>
            '''
            
            # Add ISO 27001 gaps
            for gap in data['compliance']['iso27001'].get('gaps', [])[:5]:  # Show first 5
                html += f'''
                                <tr>
                                    <td>{gap.get('vulnerability_name', 'N/A')}</td>
                                    <td><span class="badge bg-danger">{translations.get('non_compliant', 'Non-Compliant')}</span></td>
                                </tr>
                '''
            
            html += '''
                            </tbody>
                        </table>
                    </div>
                </div>
            '''
        
        html += '''
            </div>
        </div>
        '''
    
    # Recommendations Section
    if selected_sections.get('recommendations'):
        html += f'''
        <div class="recommendations mb-4">
            <h3 class="text-secondary"><i class="fas fa-lightbulb me-2"></i>{translations.get('recommendations', 'Recommendations')}</h3>
            <div class="alert alert-warning">
                <h6>{translations.get('compliance_recommendations', 'Compliance Recommendations')}:</h6>
                <ol>
                    <li>{translations.get('recommendation_compliance_1', 'Review and address all non-compliant vulnerabilities')}</li>
                    <li>{translations.get('recommendation_compliance_2', 'Implement required security controls for PCI DSS and ISO 27001')}</li>
                    <li>{translations.get('recommendation_compliance_3', 'Establish regular compliance monitoring and reporting')}</li>
                    <li>{translations.get('recommendation_compliance_4', 'Update security policies and procedures')}</li>
                </ol>
            </div>
        </div>
        '''
    
    # Charts Section
    if selected_sections.get('charts'):
        html += f'''
        <div class="charts-section mb-4">
            <h3 class="text-secondary"><i class="fas fa-chart-pie me-2"></i>{translations.get('compliance_charts', 'Compliance Charts')}</h3>
            <div class="alert alert-info">
                <i class="fas fa-info-circle me-2"></i>
                {translations.get('compliance_charts_note', 'Detailed compliance charts and visualizations will be included in the final report.')}
            </div>
        </div>
        '''
    
    # Report Footer
    html += f'''
        <div class="report-footer text-center mt-4">
            <p class="text-muted">{translations['report_generated_by']}: {data['generated_by']}</p>
            <p class="text-muted">{translations.get('preview_note', 'This is a preview. The final report will contain complete data.')}</p>
        </div>
    </div>
    '''
    
    return html


def get_risk_level_color(risk_level):
    """Get Bootstrap color class for risk level"""
    risk_colors = {
        'Low': 'success',
        'Medium': 'warning',
        'High': 'danger',
        'Critical': 'dark',
        'Низький': 'success',
        'Середній': 'warning',
        'Високий': 'danger',
        'Критичний': 'dark'
    }
    return risk_colors.get(risk_level, 'secondary') 


@login_required
@user_passes_test(check_risk_report_access)
def get_scheduled_reports(request):
    """Get list of scheduled reports for the current user"""
    try:
        # Check if user has access to risk reports
        if not check_risk_report_access(request.user, 'view'):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to access scheduled reports')
            }, status=403)
        user_companies = get_user_companies(request.user)
        
        # Get scheduled reports accessible to the user (owner, same company, or email recipient)
        scheduled_reports = ScheduledReport.objects.filter(
            Q(created_by=request.user) | 
            Q(company__in=user_companies) |
            Q(email_recipients__user=request.user)
        ).select_related('company', 'created_by', 'report_profile', 'mail_server', 'mail_account').prefetch_related('email_recipients').distinct()
        
        # Format data for frontend
        reports_data = []
        for report in scheduled_reports:
            # Calculate next run if not set
            if not report.next_run:
                report.next_run = report.calculate_next_run()
                report.save()
            
            company_name = report.company.name if report.company else _('All Companies')
            logger.info(f"Scheduled report {report.name}: company={report.company}, company_name={company_name}")
            
            # Get email recipients summary
            recipients = report.email_recipients.all()
            recipients_count = recipients.count()
            recipients_names = []
            for recipient in recipients[:3]:  # Show first 3 names
                user = recipient.user
                name = user.get_full_name() or user.username
                recipients_names.append(name)
            
            recipients_summary = ', '.join(recipients_names)
            if recipients_count > 3:
                recipients_summary += f' +{recipients_count - 3} more'
            
            # Determine email settings display
            email_settings_display = ""
            if report.send_email:
                if report.use_default_email_settings:
                    email_settings_display = f'<div class="text-center"><span class="badge bg-success"><i class="fas fa-cog me-1"></i>{_("Default")}</span><br><small class="text-muted">{_("System settings")}</small></div>'
                else:
                    if report.mail_server and report.mail_account:
                        email_settings_display = f'''<div class="text-center">
                            <div><span class="badge bg-info"><i class="fas fa-server me-1"></i>{report.mail_server.name}</span></div>
                            <small class="text-muted d-block mt-1"><i class="fas fa-user me-1"></i>{report.mail_account.username}</small>
                            <small class="text-muted"><i class="fas fa-envelope me-1"></i>{report.mail_server.smtp_host}:{report.mail_server.smtp_port}</small>
                        </div>'''
                    elif report.mail_server:
                        email_settings_display = f'''<div class="text-center">
                            <span class="badge bg-warning"><i class="fas fa-server me-1"></i>{report.mail_server.name}</span>
                            <br><small class="text-warning"><i class="fas fa-exclamation-triangle me-1"></i>{_("No account")}</small>
                            <br><small class="text-muted">{report.mail_server.smtp_host}:{report.mail_server.smtp_port}</small>
                        </div>'''
                    else:
                        email_settings_display = f'<div class="text-center"><span class="badge bg-warning"><i class="fas fa-exclamation-triangle me-1"></i>{_("Incomplete")}</span><br><small class="text-warning">{_("No server configured")}</small></div>'
            else:
                email_settings_display = f'<div class="text-center"><span class="badge bg-secondary"><i class="fas fa-envelope-slash me-1"></i>{_("Disabled")}</span><br><small class="text-muted">{_("Email sending disabled")}</small></div>'
            
            # Check for active execution (running or pending, or completed but very recent)
            active_execution = None
            try:
                from .models import ScheduledReportExecution
                from datetime import timedelta
                # Check for running or pending executions
                active_execution = ScheduledReportExecution.objects.filter(
                    scheduled_report=report,
                    status__in=['pending', 'running']
                ).order_by('-started_at').first()
                
                # If no active execution, check for very recent completed ones (within last 30 seconds)
                # This helps catch executions that completed very quickly
                if not active_execution:
                    recent_threshold = timezone.now() - timedelta(seconds=30)
                    recent_execution = ScheduledReportExecution.objects.filter(
                        scheduled_report=report,
                        status='completed',
                        started_at__gte=recent_threshold
                    ).order_by('-started_at').first()
                    if recent_execution:
                        # Show as completing if very recent
                        active_execution = recent_execution
            except Exception as e:
                logger.warning(f"Error checking active execution for report {report.id}: {e}")
            
            execution_status = None
            execution_progress = None
            execution_started_at = None
            if active_execution:
                execution_status = active_execution.status
                execution_started_at = active_execution.started_at.astimezone().strftime('%Y-%m-%d %H:%M:%S')
                # Determine progress based on execution state
                if active_execution.status == 'pending':
                    execution_progress = 'pending'
                elif active_execution.status == 'running':
                    # Check if report generation is complete (snapshot_html exists)
                    if active_execution.snapshot_html:
                        # Check if email is already sent
                        if active_execution.email_sent:
                            # Email sent, execution should be completed soon
                            execution_progress = 'completing'
                        else:
                            # Report generated, now sending email
                            execution_progress = 'sending_email'
                    else:
                        # Report generation in progress
                        execution_progress = 'generating_report'
                elif active_execution.status == 'completed':
                    # Show as completing if very recent (within last 30 seconds)
                    execution_progress = 'completing'
            elif report.next_run and report.next_run <= timezone.now() and report.status == 'active':
                # Next run time has passed but no execution created yet - should start soon
                # This can happen if Celery task hasn't run yet or is delayed
                # Check if execution was created very recently (within last 5 minutes)
                # If not, it's truly pending
                from datetime import timedelta
                recent_threshold = timezone.now() - timedelta(minutes=5)
                recent_execution = ScheduledReportExecution.objects.filter(
                    scheduled_report=report,
                    started_at__gte=recent_threshold
                ).order_by('-started_at').first()
                
                if not recent_execution:
                    # No recent execution - truly pending
                    execution_status = 'pending'
                    execution_progress = 'pending'
                    execution_started_at = None
                else:
                    # Recent execution exists - use its status
                    execution_status = recent_execution.status
                    execution_started_at = recent_execution.started_at.astimezone().strftime('%Y-%m-%d %H:%M:%S')
                    if recent_execution.status == 'running':
                        if recent_execution.snapshot_html:
                            if recent_execution.email_sent:
                                execution_progress = 'completing'
                            else:
                                execution_progress = 'sending_email'
                        else:
                            execution_progress = 'generating_report'
                    elif recent_execution.status == 'completed':
                        execution_progress = 'completing'
                    else:
                        execution_progress = 'pending'
            # Don't show execution status for completed reports
            elif report.status == 'completed':
                execution_status = None
                execution_progress = None
                execution_started_at = None
            
            # Calculate time until next_run
            time_until_next_run = None
            if report.next_run and report.status == 'active':
                time_diff = report.next_run - timezone.now()
                if time_diff.total_seconds() > 0:
                    total_seconds = int(time_diff.total_seconds())
                    days = total_seconds // 86400
                    hours = (total_seconds % 86400) // 3600
                    minutes = (total_seconds % 3600) // 60
                    
                    time_parts = []
                    if days > 0:
                        time_parts.append(f"{days} {_('day') if days == 1 else _('days')}")
                    if hours > 0:
                        time_parts.append(f"{hours} {_('hour') if hours == 1 else _('hours')}")
                    if minutes > 0 or len(time_parts) == 0:
                        time_parts.append(f"{minutes} {_('minute') if minutes == 1 else _('minutes')}")
                    
                    time_until_next_run = ', '.join(time_parts)
                else:
                    # Time has passed
                    time_until_next_run = _('Overdue')
            
            reports_data.append({
                'id': str(report.id),
                'name': report.name,
                'description': report.description,
                # 'report_type': report.get_report_type_display(),  # removed per new design
                'report_profile_id': str(report.report_profile.id) if report.report_profile else None,
                'report_profile_name': report.report_profile.name if report.report_profile else None,
                # 'report_format': report.report_format.upper(),  # removed in link mode
                'frequency': report.get_frequency_display(),
                'status': report.get_status_display(),
                'company': company_name,
                'created_by': report.created_by.get_full_name() or report.created_by.username,
                'created_at': report.created_at.astimezone().strftime('%Y-%m-%d %H:%M'),
                'last_run': report.last_run.astimezone().strftime('%Y-%m-%d %H:%M') if report.last_run else _('Never'),
                'next_run': report.next_run.astimezone().strftime('%Y-%m-%d %H:%M') if report.next_run else _('Not scheduled'),
                'next_run_timestamp': report.next_run.timestamp() if report.next_run else None,  # For client-side calculation
                'time_until_next_run': time_until_next_run,
                'run_count': report.run_count,
                'email_recipients_count': recipients_count,
                'email_recipients_summary': recipients_summary,
                'email_settings_display': email_settings_display,
                'execution_time': report.execution_time.strftime('%H:%M'),
                'weekdays': ', '.join(report.get_weekdays()) if report.frequency == 'weekly' else '',
                'day_of_month': report.day_of_month if report.frequency == 'monthly' else None,
                'execution_status': execution_status,
                'execution_progress': execution_progress,
                'execution_started_at': execution_started_at,
            })
        
        return JsonResponse({
            'status': 'success',
            'data': reports_data
        })
        
    except Exception as e:
        logger.error(f"Error getting scheduled reports: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error loading scheduled reports: {}').format(str(e))
        })


@login_required
@user_passes_test(can_add_risk_report)
@require_http_methods(["POST"])
def create_scheduled_report(request):
    """Create a new scheduled report"""
    try:
        # Check if user has permission to add reports
        if not can_add_risk_report(request.user):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to create scheduled reports')
            }, status=403)
        # Handle both JSON and form data
        if request.content_type == 'application/json':
            data = json.loads(request.body)
        else:
            # Handle form data
            data = {}
            for key, value in request.POST.items():
                if key == 'csrfmiddlewaretoken':
                    continue
                if key == 'email_recipients':
                    # Handle multiple values for email_recipients
                    data[key] = request.POST.getlist(key)
                elif value.lower() == 'true':
                    data[key] = True
                elif value.lower() == 'false':
                    data[key] = False
                else:
                    data[key] = value
        
        # Validate required fields
        required_fields = ['name', 'report_profile_id', 'report_format', 'frequency', 'start_date', 'execution_time']
        for field in required_fields:
            if not data.get(field):
                return JsonResponse({
                    'status': 'error',
                    'message': _('Field {} is required').format(field)
                })
        
        # Get report profile
        try:
            report_profile = ReportProfile.objects.get(id=data['report_profile_id'])
        except ReportProfile.DoesNotExist:
            return JsonResponse({
                'status': 'error',
                'message': _('Selected report profile not found')
                })
        
        # Create scheduled report
        scheduled_report = ScheduledReport(
            name=data['name'],
            description=data.get('description', ''),
            report_profile=report_profile,
            report_format=data['report_format'],
            report_language=data.get('report_language', 'uk'),
            frequency=data['frequency'],
            start_date=datetime.strptime(data['start_date'], '%Y-%m-%d').date(),
            execution_time=datetime.strptime(data['execution_time'], '%H:%M').time(),
            email_subject=data.get('email_subject', ''),
            email_body=data.get('email_body', ''),
            send_email=data.get('send_email', True),
            use_default_email_settings=data.get('use_default_email_settings', True),
            created_by=request.user
        )
        
        # Validate and set company (required field)
        company_id = data.get('company_id')
        logger.info(f"Creating scheduled report with company_id: {company_id} (type: {type(company_id)})")
        
        if not company_id:
            logger.error("Company selection is required for scheduled reports")
            return JsonResponse({
                'status': 'error',
                'message': _('Company selection is required for scheduled reports')
            })
        
        try:
            company = Company.objects.get(id=company_id)
            scheduled_report.company = company
            logger.info(f"Successfully set company: {company.name} (ID: {company.id})")
        except Company.DoesNotExist:
            logger.error(f"Company with ID {company_id} not found")
            return JsonResponse({
                'status': 'error',
                'message': _('Selected company not found')
            })
        
        # Set mail server and account if specified
        if data.get('mail_server') and not data.get('use_default_email_settings', True):
            try:
                from app_conf.models import MailServer
                mail_server = MailServer.objects.get(id=data['mail_server'])
                scheduled_report.mail_server = mail_server
            except MailServer.DoesNotExist:
                return JsonResponse({
                    'status': 'error',
                    'message': _('Selected mail server not found')
                })
        
        if data.get('mail_account') and not data.get('use_default_email_settings', True):
            try:
                from app_conf.models import MailAccount
                mail_account = MailAccount.objects.get(id=data['mail_account'])
                scheduled_report.mail_account = mail_account
            except MailAccount.DoesNotExist:
                return JsonResponse({
                    'status': 'error',
                    'message': _('Selected mail account not found')
                })
        
        # Set end date if specified
        if data.get('end_date'):
            scheduled_report.end_date = datetime.strptime(data['end_date'], '%Y-%m-%d').date()
        
        # Set frequency-specific fields
        if data['frequency'] == 'weekly':
            scheduled_report.monday = data.get('monday', False)
            scheduled_report.tuesday = data.get('tuesday', False)
            scheduled_report.wednesday = data.get('wednesday', False)
            scheduled_report.thursday = data.get('thursday', False)
            scheduled_report.friday = data.get('friday', False)
            scheduled_report.saturday = data.get('saturday', False)
            scheduled_report.sunday = data.get('sunday', False)
        elif data['frequency'] == 'monthly':
            scheduled_report.day_of_month = data.get('day_of_month', 1)
        
        # Calculate next run
        scheduled_report.next_run = scheduled_report.calculate_next_run()
        
        scheduled_report.save()
        
        # Set email recipients
        if data.get('email_recipients'):
            recipient_ids = data['email_recipients']
            # Get CabinetUser objects from IDs
            cabinet_users = CabinetUser.objects.filter(id__in=recipient_ids)
            scheduled_report.email_recipients.set(cabinet_users)
        
        # Handle file attachments
        if request.FILES.getlist('attachments'):
            for uploaded_file in request.FILES.getlist('attachments'):
                # Create attachment record
                attachment = ScheduledReportAttachment(
                    scheduled_report=scheduled_report,
                    file=uploaded_file,
                    original_filename=uploaded_file.name,
                    file_size=uploaded_file.size,
                    uploaded_by=request.user
                )
                attachment.save()
        
        return JsonResponse({
            'status': 'success',
            'message': _('Scheduled report created successfully'),
            'id': str(scheduled_report.id)
        })
        
    except Exception as e:
        logger.error(f"Error creating scheduled report: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error creating scheduled report: {}').format(str(e))
        })


@login_required
@user_passes_test(can_edit_risk_report)
@require_http_methods(["POST"])
def update_scheduled_report(request, report_id):
    """Update an existing scheduled report"""
    try:
        # Check if user has permission to edit reports
        if not can_edit_risk_report(request.user):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to edit scheduled reports')
            }, status=403)
        # Get the scheduled report
        scheduled_report = ScheduledReport.objects.get(id=report_id)
        
        # Check permissions
        user_companies = get_user_companies(request.user)
        if (scheduled_report.created_by != request.user and 
            scheduled_report.company not in user_companies):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to edit this scheduled report')
            })
        
        # Handle both JSON and form data
        if request.content_type == 'application/json':
            data = json.loads(request.body)
        else:
            # Handle form data
            data = {}
            for key, value in request.POST.items():
                if key == 'csrfmiddlewaretoken':
                    continue
                if key == 'email_recipients':
                    # Handle multiple values for email_recipients
                    data[key] = request.POST.getlist(key)
                elif value.lower() == 'true':
                    data[key] = True
                elif value.lower() == 'false':
                    data[key] = False
                else:
                    data[key] = value
        
        # Update fields
        if 'name' in data:
            scheduled_report.name = data['name']
        if 'description' in data:
            scheduled_report.description = data['description']
        if 'report_profile_id' in data:
            try:
                report_profile = ReportProfile.objects.get(id=data['report_profile_id'])
                scheduled_report.report_profile = report_profile
            except ReportProfile.DoesNotExist:
                return JsonResponse({
                    'status': 'error',
                    'message': _('Selected report profile not found')
                })
        if 'report_format' in data:
            scheduled_report.report_format = data['report_format']
        if 'report_language' in data:
            scheduled_report.report_language = data['report_language']
        if 'frequency' in data:
            scheduled_report.frequency = data['frequency']
        if 'start_date' in data:
            scheduled_report.start_date = datetime.strptime(data['start_date'], '%Y-%m-%d').date()
        if 'end_date' in data:
            scheduled_report.end_date = datetime.strptime(data['end_date'], '%Y-%m-%d').date() if data['end_date'] else None
        if 'execution_time' in data:
            scheduled_report.execution_time = datetime.strptime(data['execution_time'], '%H:%M').time()
        if 'email_subject' in data:
            scheduled_report.email_subject = data['email_subject']
        if 'email_body' in data:
            scheduled_report.email_body = data['email_body']
        if 'send_email' in data:
            scheduled_report.send_email = data['send_email']
        if 'use_default_email_settings' in data:
            scheduled_report.use_default_email_settings = data['use_default_email_settings']
        if 'status' in data:
            scheduled_report.status = data['status']
        
        # Validate and set company (required field)
        if 'company_id' in data:
            if not data['company_id']:
                return JsonResponse({
                    'status': 'error',
                    'message': _('Company selection is required for scheduled reports')
                })
            try:
                company = Company.objects.get(id=data['company_id'])
                scheduled_report.company = company
            except Company.DoesNotExist:
                return JsonResponse({
                    'status': 'error',
                    'message': _('Selected company not found')
                })
        
        # Update mail server and account if specified
        if 'mail_server' in data:
            if data['mail_server'] and not data.get('use_default_email_settings', True):
                try:
                    from app_conf.models import MailServer
                    mail_server = MailServer.objects.get(id=data['mail_server'])
                    scheduled_report.mail_server = mail_server
                except MailServer.DoesNotExist:
                    return JsonResponse({
                        'status': 'error',
                        'message': _('Selected mail server not found')
                    })
            else:
                scheduled_report.mail_server = None
        
        if 'mail_account' in data:
            if data['mail_account'] and not data.get('use_default_email_settings', True):
                try:
                    from app_conf.models import MailAccount
                    mail_account = MailAccount.objects.get(id=data['mail_account'])
                    scheduled_report.mail_account = mail_account
                except MailAccount.DoesNotExist:
                    return JsonResponse({
                        'status': 'error',
                        'message': _('Selected mail account not found')
                    })
            else:
                scheduled_report.mail_account = None
        
        # Update frequency-specific fields
        if scheduled_report.frequency == 'weekly':
            scheduled_report.monday = data.get('monday', False)
            scheduled_report.tuesday = data.get('tuesday', False)
            scheduled_report.wednesday = data.get('wednesday', False)
            scheduled_report.thursday = data.get('thursday', False)
            scheduled_report.friday = data.get('friday', False)
            scheduled_report.saturday = data.get('saturday', False)
            scheduled_report.sunday = data.get('sunday', False)
        elif scheduled_report.frequency == 'monthly':
            scheduled_report.day_of_month = data.get('day_of_month', 1)
        
        # Recalculate next run
        scheduled_report.next_run = scheduled_report.calculate_next_run()
        
        scheduled_report.save()
        
        # Update email recipients
        if 'email_recipients' in data:
            recipient_ids = data['email_recipients']
            # Get CabinetUser objects from IDs
            cabinet_users = CabinetUser.objects.filter(id__in=recipient_ids)
            scheduled_report.email_recipients.set(cabinet_users)
        
        # Handle file attachments
        if request.FILES.getlist('attachments'):
            for uploaded_file in request.FILES.getlist('attachments'):
                # Create attachment record
                attachment = ScheduledReportAttachment(
                    scheduled_report=scheduled_report,
                    file=uploaded_file,
                    original_filename=uploaded_file.name,
                    file_size=uploaded_file.size,
                    uploaded_by=request.user
                )
                attachment.save()
        
        # Handle existing attachments (keep only those in the list)
        if 'existing_attachments' in data:
            try:
                existing_attachment_ids = json.loads(data['existing_attachments'])
                # Delete attachments not in the list
                scheduled_report.attachments.exclude(id__in=existing_attachment_ids).delete()
            except (json.JSONDecodeError, KeyError):
                # If no existing_attachments provided, keep all existing
                pass
        
        return JsonResponse({
            'status': 'success',
            'message': _('Scheduled report updated successfully')
        })
        
    except ScheduledReport.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': _('Scheduled report not found')
        })
    except Exception as e:
        logger.error(f"Error updating scheduled report: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error updating scheduled report: {}').format(str(e))
        })


@login_required
@user_passes_test(can_edit_risk_report)
@require_http_methods(["POST"])
def send_scheduled_report_now(request, report_id):
    """Send a scheduled report immediately"""
    try:
        # Check if user has permission to edit reports
        if not can_edit_risk_report(request.user):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to send scheduled reports')
            }, status=403)
        
        # Get the scheduled report
        scheduled_report = ScheduledReport.objects.select_related(
            'report_profile', 'company', 'created_by', 'mail_server', 'mail_account'
        ).get(id=report_id)
        
        # Check permissions
        user_companies = get_user_companies(request.user)
        if (scheduled_report.created_by != request.user and 
            scheduled_report.company not in user_companies):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to send this scheduled report')
            }, status=403)
        
        # Allow manual send regardless of current status (failed/paused/completed)
        # Manual execution should always be permitted for authorized users
        
        request_time = timezone.now()
        
        # Create execution record for manual sending
        execution = ScheduledReportExecution.objects.create(
            scheduled_report=scheduled_report,
            status='running'
        )
        
        logger.info(f"Manual execution started for scheduled report {scheduled_report.name} (ID: {scheduled_report.id}) by user {request.user.username}")
        
        # Trigger the report execution task
        from .tasks import execute_scheduled_report
        task = execute_scheduled_report.delay(str(scheduled_report.id), manual_execution=True)
        
        # For production, we don't wait for the result synchronously
        # The task will execute asynchronously and update the execution record
        logger.info(f"Manual execution task started for scheduled report {scheduled_report.name} (task_id: {task.id})")
        
        return JsonResponse({
            'status': 'success',
            'message': _('Report execution started successfully. You will receive a notification when it completes.'),
            'task_id': str(task.id)
        })
        
    except ScheduledReport.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': _('Scheduled report not found')
        }, status=404)
    except Exception as e:
        logger.error(f"Error sending scheduled report: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error sending scheduled report: {}').format(str(e))
        }, status=500)





@login_required
@user_passes_test(can_delete_risk_report)
@require_http_methods(["POST"])
def delete_scheduled_report(request, report_id):
    """Delete a scheduled report"""
    try:
        # Check if user has permission to delete reports
        if not can_delete_risk_report(request.user):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to delete scheduled reports')
            }, status=403)
        # Get the scheduled report
        scheduled_report = ScheduledReport.objects.get(id=report_id)
        
        # Check permissions
        user_companies = get_user_companies(request.user)
        if (scheduled_report.created_by != request.user and 
            scheduled_report.company not in user_companies):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to delete this scheduled report')
            })
        
        report_name = scheduled_report.name
        scheduled_report.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': _('Scheduled report "{}" deleted successfully').format(report_name)
        })
        
    except ScheduledReport.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': _('Scheduled report not found')
        })
    except Exception as e:
        logger.error(f"Error deleting scheduled report: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error deleting scheduled report: {}').format(str(e))
        })


@login_required
@user_passes_test(check_risk_report_access)
def get_scheduled_report_details(request, report_id):
    """Get details of a specific scheduled report"""
    try:
        # Check if user has access to risk reports
        if not check_risk_report_access(request.user, 'view'):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to access scheduled reports')
            }, status=403)
        # Get the scheduled report
        scheduled_report = ScheduledReport.objects.select_related(
            'report_profile', 'company', 'created_by', 'mail_server', 'mail_account'
        ).get(id=report_id)
        
        # Check permissions: owner, same company, or email recipient
        user_companies = get_user_companies(request.user)
        user_has_access = False
        
        # Check if user is the owner
        if scheduled_report.created_by == request.user:
            user_has_access = True
        # Check if user belongs to the same company
        elif scheduled_report.company and scheduled_report.company in user_companies:
            user_has_access = True
        # Check if user is in the email recipients list
        elif scheduled_report.email_recipients.filter(user=request.user).exists():
            user_has_access = True
        
        if not user_has_access:
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to view this scheduled report')
            })
        
        # Get email recipients with department and position info
        recipients = []
        for recipient in scheduled_report.email_recipients.all():
            # Check if recipient is CabinetUser or User
            if hasattr(recipient, 'user'):
                # recipient is CabinetUser
                cabinet_user = recipient
                user = recipient.user
                recipient_data = {
                    'id': cabinet_user.id,  # Use CabinetUser ID for JavaScript
                    'name': user.get_full_name() or user.username,
                    'email': user.email
                }
                
                # Add department and position info
                recipient_data['department'] = cabinet_user.department.get_name() if cabinet_user.department else ''
                recipient_data['position'] = cabinet_user.position.get_name() if cabinet_user.position else ''
                recipient_data['company'] = cabinet_user.company.name if cabinet_user.company else ''
            else:
                # recipient is User, try to get associated CabinetUser
                user = recipient
                cabinet_user = None
                try:
                    cabinet_user = user.cabinet
                except:
                    pass
                    
                recipient_data = {
                    'id': cabinet_user.id if cabinet_user else user.id,  # Use CabinetUser ID for JavaScript
                    'name': user.get_full_name() or user.username,
                    'email': user.email
                }
                
                # Add department and position if available
                if cabinet_user:
                    recipient_data['department'] = cabinet_user.department.get_name() if cabinet_user.department else ''
                    recipient_data['position'] = cabinet_user.position.get_name() if cabinet_user.position else ''
                    recipient_data['company'] = cabinet_user.company.name if cabinet_user.company else ''
            
            recipients.append(recipient_data)
        
        # Get recent executions
        recent_executions = []
        for execution in scheduled_report.executions.all()[:10]:  # Last 10 executions
            recent_executions.append({
                'id': str(execution.id),
                'status': execution.get_status_display(),
                'started_at': execution.started_at.astimezone().strftime('%Y-%m-%d %H:%M'),
                'completed_at': execution.completed_at.astimezone().strftime('%Y-%m-%d %H:%M') if execution.completed_at else '',
                'file_size': execution.file_size,
                'email_sent': execution.email_sent,
                'email_recipients_count': execution.email_recipients_count,
                'error_message': execution.error_message
            })
        
        # Get attachments
        attachments = []
        for attachment in scheduled_report.attachments.all():
            attachments.append({
                'id': str(attachment.id),
                'original_filename': attachment.original_filename,
                'file_size': attachment.file_size,
                'file_url': attachment.get_file_url(),
                'uploaded_at': attachment.uploaded_at.astimezone().strftime('%Y-%m-%d %H:%M'),
                'uploaded_by': attachment.uploaded_by.get_full_name() if attachment.uploaded_by else ''
            })
        
        data = {
            'id': str(scheduled_report.id),
            'name': scheduled_report.name,
            'description': scheduled_report.description,
            'report_type': scheduled_report.report_type,
            'report_profile_id': str(scheduled_report.report_profile.id) if scheduled_report.report_profile else None,
            'report_profile_name': scheduled_report.report_profile.name if scheduled_report.report_profile else None,
            'report_format': scheduled_report.report_format,
            'report_language': scheduled_report.report_language,
            'frequency': scheduled_report.frequency,
            'start_date': scheduled_report.start_date.strftime('%Y-%m-%d'),
            'end_date': scheduled_report.end_date.strftime('%Y-%m-%d') if scheduled_report.end_date else '',
            'execution_time': scheduled_report.execution_time.strftime('%H:%M'),
            'status': scheduled_report.status,
            'company_id': scheduled_report.company.id if scheduled_report.company else '',
            'company_name': scheduled_report.company.name if scheduled_report.company else '',
            'email_subject': scheduled_report.email_subject,
            'email_body': scheduled_report.email_body,
            'send_email': scheduled_report.send_email,
            'use_default_email_settings': scheduled_report.use_default_email_settings,
            'mail_server': scheduled_report.mail_server.id if scheduled_report.mail_server else '',
            'mail_server_name': scheduled_report.mail_server.name if scheduled_report.mail_server else '',
            'mail_account': scheduled_report.mail_account.id if scheduled_report.mail_account else '',
            'mail_account_name': f"{scheduled_report.mail_account.username} ({scheduled_report.mail_account.server.name})" if scheduled_report.mail_account else '',
            'email_recipients': recipients,
            'last_run': scheduled_report.last_run.astimezone().strftime('%Y-%m-%d %H:%M') if scheduled_report.last_run else '',
            'next_run': scheduled_report.next_run.astimezone().strftime('%Y-%m-%d %H:%M') if scheduled_report.next_run else '',
            'run_count': scheduled_report.run_count,
            'recent_executions': recent_executions,
            # Weekly fields
            'monday': scheduled_report.monday,
            'tuesday': scheduled_report.tuesday,
            'wednesday': scheduled_report.wednesday,
            'thursday': scheduled_report.thursday,
            'friday': scheduled_report.friday,
            'saturday': scheduled_report.saturday,
            'sunday': scheduled_report.sunday,
            # Monthly fields
            'day_of_month': scheduled_report.day_of_month,
            # Attachments
            'attachments': attachments,
        }
        
        return JsonResponse({
            'status': 'success',
            'data': data
        })
        
    except ScheduledReport.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': _('Scheduled report not found')
        })
    except Exception as e:
        logger.error(f"Error getting scheduled report details: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error loading scheduled report details: {}').format(str(e))
        })


@login_required
@user_passes_test(check_risk_report_access)
def get_scheduled_report_history(request, report_id):
    """Get execution history for a specific scheduled report"""
    try:
        # Get the scheduled report
        scheduled_report = ScheduledReport.objects.get(id=report_id)
        
        # Check permissions: owner, same company, or email recipient
        user_companies = get_user_companies(request.user)
        user_has_access = False
        
        # Check if user is the owner
        if scheduled_report.created_by == request.user:
            user_has_access = True
        # Check if user belongs to the same company
        elif scheduled_report.company and scheduled_report.company in user_companies:
            user_has_access = True
        # Check if user is in the email recipients list
        elif scheduled_report.email_recipients.filter(user=request.user).exists():
            user_has_access = True
        
        if not user_has_access:
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to view this scheduled report')
            })
        
        # Get pagination parameters
        page = int(request.GET.get('page', 1))
        per_page = int(request.GET.get('per_page', 20))
        
        # Get all executions for this report
        all_executions = scheduled_report.executions.all()
        total_count = all_executions.count()
        
        # Calculate pagination
        start_index = (page - 1) * per_page
        end_index = start_index + per_page
        executions = all_executions[start_index:end_index]
        
        # Format execution data
        execution_history = []
        for execution in executions:
            import os
            
            # Get file info
            file_exists = bool(execution.file_path and os.path.exists(execution.file_path))
            file_name = os.path.basename(execution.file_path) if execution.file_path else ''
            file_size_mb = round(execution.file_size / (1024 * 1024), 2) if execution.file_size else 0
            
            # Calculate duration
            duration = ''
            if execution.completed_at and execution.started_at:
                duration_seconds = (execution.completed_at - execution.started_at).total_seconds()
                if duration_seconds < 60:
                    duration = f"{int(duration_seconds)}s"
                elif duration_seconds < 3600:
                    duration = f"{int(duration_seconds // 60)}m {int(duration_seconds % 60)}s"
                else:
                    hours = int(duration_seconds // 3600)
                    minutes = int((duration_seconds % 3600) // 60)
                    duration = f"{hours}h {minutes}m"
            
            # Build snapshot link if snapshot_html exists
            try:
                from django.urls import reverse
                snapshot_path = reverse('view_scheduled_report_snapshot', args=[str(execution.id)])
            except Exception:
                snapshot_path = ''

            execution_data = {
                'id': str(execution.id),
                'status': execution.status,
                'status_display': execution.get_status_display(),
                'started_at': execution.started_at.astimezone().strftime('%Y-%m-%d %H:%M:%S'),
                'completed_at': execution.completed_at.astimezone().strftime('%Y-%m-%d %H:%M:%S') if execution.completed_at else '',
                'duration': duration,
                'file_path': execution.file_path,
                'file_name': file_name,
                'file_size': execution.file_size,
                'file_size_mb': file_size_mb,
                'file_exists': file_exists,
                'snapshot_available': bool(execution.snapshot_html),
                'snapshot_path': snapshot_path,
                'email_sent': execution.email_sent,
                'email_recipients_count': execution.email_recipients_count,
                'email_error': execution.email_error,
                'error_message': execution.error_message,
                'report_type': scheduled_report.report_type,
                'report_format': scheduled_report.report_format,
                'report_language': scheduled_report.report_language,
                'company_name': scheduled_report.company.name if scheduled_report.company else _('All Companies')
            }
            
            execution_history.append(execution_data)
        
        # Calculate pagination info
        total_pages = (total_count + per_page - 1) // per_page
        has_next = page < total_pages
        has_previous = page > 1
        
        data = {
            'schedule_id': str(scheduled_report.id),
            'schedule_name': scheduled_report.name,
            'executions': execution_history,
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total_count': total_count,
                'total_pages': total_pages,
                'has_next': has_next,
                'has_previous': has_previous,
                'start_index': start_index + 1,
                'end_index': min(end_index, total_count)
            }
        }
        
        return JsonResponse({
            'status': 'success',
            'data': data
        })
        
    except ScheduledReport.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': _('Scheduled report not found')
        })
    except Exception as e:
        logger.error(f"Error getting scheduled report history: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error loading scheduled report history: {}').format(str(e))
        })


@login_required
@user_passes_test(check_risk_report_access)
def download_scheduled_report_file(request, execution_id):
    """Download a specific scheduled report file"""
    from django.utils.translation import gettext as _
    try:
        # Get the execution
        execution = ScheduledReportExecution.objects.get(id=execution_id)
        report = execution.scheduled_report
        
        # Check permissions: owner, same company, or email recipient
        user_companies = get_user_companies(request.user)
        user_has_access = False
        
        # Check if user is the owner
        if report.created_by == request.user:
            user_has_access = True
        # Check if user belongs to the same company
        elif report.company and report.company in user_companies:
            user_has_access = True
        # Check if user is in the email recipients list
        elif report.email_recipients.filter(user=request.user).exists():
            user_has_access = True
        
        if not user_has_access:
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to download this report')
            })
        
        # Check if file exists
        if not execution.file_path or not os.path.exists(execution.file_path):
            return JsonResponse({
                'status': 'error',
                'message': _('Report file not found or has been deleted')
            })
        
        # Serve the file
        import mimetypes
        from django.http import FileResponse
        
        file_path = execution.file_path
        file_name = os.path.basename(file_path)
        
        # Determine content type
        content_type, _ = mimetypes.guess_type(file_path)
        if not content_type:
            content_type = 'application/octet-stream'
        
        # Create response
        response = FileResponse(
            open(file_path, 'rb'),
            content_type=content_type,
            as_attachment=True,
            filename=file_name
        )
        
        return response
        
    except ScheduledReportExecution.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': _('Report execution not found')
        })
    except Exception as e:
        logger.error(f"Error downloading scheduled report file: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error downloading report file: {}').format(str(e))
        })


@login_required
@user_passes_test(check_risk_report_access)
def get_cabinet_users(request):
    """Get list of cabinet users for email recipients selection"""
    try:
        # Get company filter from request
        company_id = request.GET.get('company_id')
        
        # Get users from the same companies as the current user
        user_companies = get_user_companies(request.user)
        
        if company_id:
            # Filter by specific company if provided
            try:
                selected_company = Company.objects.get(id=company_id)
                if selected_company in user_companies or request.user.is_superuser:
                    cabinet_users = CabinetUser.objects.filter(
                        company=selected_company,
                        user__is_active=True
                    ).select_related('company', 'user', 'department', 'position')
                else:
                    # User doesn't have access to this company
                    cabinet_users = CabinetUser.objects.none()
            except Company.DoesNotExist:
                cabinet_users = CabinetUser.objects.none()
        else:
            # No specific company selected, show users from all accessible companies
            if user_companies:
                cabinet_users = CabinetUser.objects.filter(
                    company__in=user_companies,
                    user__is_active=True
                ).select_related('company', 'user', 'department', 'position')
            else:
                # If no companies, get all active users (admin case)
                cabinet_users = CabinetUser.objects.filter(user__is_active=True).select_related('company', 'user', 'department', 'position')
        
        users_data = []
        for cabinet_user in cabinet_users:
            user = cabinet_user.user
            
            # Get department info
            department_name = ''
            if cabinet_user.department:
                department_name = cabinet_user.department.get_name()
            
            # Get position info  
            position_name = ''
            if cabinet_user.position:
                position_name = cabinet_user.position.get_name()
            
            users_data.append({
                'id': cabinet_user.id,
                'name': user.get_full_name() or user.username,
                'email': user.email,
                'company': cabinet_user.company.name if cabinet_user.company else _('No Company'),
                'company_id': cabinet_user.company.id if cabinet_user.company else None,
                'department': department_name,
                'position': position_name
            })
        
        return JsonResponse({
            'status': 'success',
            'data': users_data
        })
        
    except Exception as e:
        logger.error(f"Error getting cabinet users: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error getting cabinet users: {}').format(str(e))
        })


@login_required
@user_passes_test(check_risk_report_access)
@require_http_methods(["POST"])
def send_profile_link(request):
    """Send Risk Assessment Report link for a profile to a specific email using global email settings."""
    try:
        profile_id = request.POST.get('profile_id')
        recipient_email = request.POST.get('email')
        if not profile_id or not recipient_email:
            return JsonResponse({'status': 'error', 'message': _('Profile ID and email are required')}, status=400)

        # Validate profile
        try:
            profile = ReportProfile.objects.get(id=profile_id)
        except ReportProfile.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': _('Report profile not found')}, status=404)

        if not profile.can_be_used_by(request.user):
            raise PermissionDenied(_("You don't have permission to use this profile"))

        # Build preview HTML (fresh) using existing helpers
        sections_config = profile.get_sections_config()
        company_id = sections_config.get('_company_id', '')

        # Prepare params and generate unified content
        from .report_views import generate_report_data, generate_html_preview
        params = {
            'reportType': 'profile',
            'format': 'html',
            'language': profile.default_language,
            'company_id': str(company_id) if company_id else '',
            'selectedSections': sections_config,
            'profile': {
                'name': profile.name,
                'description': profile.description,
                'created_by': profile.created_by.get_full_name() or profile.created_by.username,
                'created_at': profile.created_at,
                'company': _('All Companies'),
            }
        }
        # If company_id exists, resolve name
        if company_id:
            try:
                from app_conf.models import Company
                company = Company.objects.get(id=company_id)
                params['profile']['company'] = company.name
            except Exception:
                pass

        report_data = generate_report_data(request.user, params)
        preview_html = generate_html_preview(report_data, 'profile', profile.default_language, sections_config)

        # Save a transient ScheduledReportExecution-like snapshot for linking
        execution = ScheduledReportExecution.objects.create(
            scheduled_report=ScheduledReport.objects.create(
                name=f"Ad-hoc: {profile.name}",
                description=_('Ad-hoc email from profile'),
                report_profile=profile,
                report_type='full',
                report_format='html',
                report_language=profile.default_language,
                frequency='once',
                start_date=timezone.now().date(),
                execution_time=timezone.now().time(),
                company=None,
                created_by=request.user,
                status='active'
            ),
            status='completed',
            snapshot_html=preview_html,
            snapshot_language=profile.default_language,
            snapshot_created_at=timezone.now(),
            email_sent=False
        )

        # Build absolute URL to snapshot
        def _build_absolute_url(path: str) -> str:
            # Try to get Site Domain from SiteSettings (primary source)
            try:
                from app_conf.models import SiteSettings
                site_settings = SiteSettings.get_settings()
                if site_settings and site_settings.site_domain:
                    # Use get_site_url() method which combines protocol + domain
                    base = site_settings.get_site_url().rstrip('/')
                    logger.info(f"[send_profile_link] Using Site Domain from SiteSettings: {base}")
                    return f"{base}{path}"
            except Exception as e:
                logger.warning(f"[send_profile_link] Could not load SiteSettings: {e}")
            
            # Fallback to PUBLIC_BASE_URL
            base = getattr(__import__('django.conf').conf.settings, 'PUBLIC_BASE_URL', '').rstrip('/')
            if base:
                logger.info(f"[send_profile_link] Using PUBLIC_BASE_URL: {base}")
                return f"{base}{path}"
            
            # Fallback to ALLOWED_HOSTS
            scheme = 'https' if getattr(__import__('django.conf').conf.settings, 'PRODUCTION', False) else 'http'
            host = 'localhost:8000'
            from django.conf import settings as dj_settings
            try:
                if dj_settings.ALLOWED_HOSTS:
                    host = next((h for h in dj_settings.ALLOWED_HOSTS if h not in ('testserver',)), dj_settings.ALLOWED_HOSTS[0])
                    if ':' not in host and host in ('localhost', '127.0.0.1'):
                        host = f"{host}:8000"
            except Exception:
                pass
            fallback_url = f"{scheme}://{host}"
            logger.info(f"[send_profile_link] Using ALLOWED_HOSTS fallback: {fallback_url}")
            return f"{fallback_url}{path}"

        from django.urls import reverse
        snapshot_path = reverse('view_scheduled_report_snapshot', args=[str(execution.id)])
        link_url = _build_absolute_url(snapshot_path)

        # Send email using global settings (direct SMTP if configured)
        cfg = RiskReportEmailConfig.objects.first()
        subject = f"{cfg.default_subject if cfg and cfg.default_subject else 'Risk Assessment Report'}"
        body = (cfg.default_body if cfg and cfg.default_body else _('Please view your report at the link below.')) + f"\n\n{link_url}"

        try:
            sent = False
            if cfg and cfg.mail_account:
                import smtplib, ssl
                from email.mime.multipart import MIMEMultipart
                from email.mime.text import MIMEText
                smtp_host = cfg.mail_account.server.smtp_host
                smtp_port = cfg.mail_account.server.smtp_port
                use_ssl = getattr(cfg.mail_account.server, 'use_ssl', False)
                use_tls = getattr(cfg.mail_account.server, 'use_tls', False)
                username = cfg.mail_account.username
                password = cfg.mail_account.password

                msg = MIMEMultipart()
                msg['From'] = username
                msg['To'] = recipient_email
                msg['Subject'] = subject
                msg.attach(MIMEText(body, 'plain'))

                if use_ssl:
                    context = ssl.create_default_context()
                    context.check_hostname = False
                    context.verify_mode = ssl.CERT_NONE
                    server = smtplib.SMTP_SSL(host=smtp_host, port=smtp_port, context=context)
                else:
                    server = smtplib.SMTP(host=smtp_host, port=smtp_port)
                    if use_tls:
                        server.starttls()
                server.login(username, password)
                server.send_message(msg)
                server.quit()
                sent = True
            else:
                from django.core.mail import EmailMessage
                from_email = cfg.mail_account.username if (cfg and cfg.mail_account) else None
                EmailMessage(subject=subject, body=body, from_email=from_email, to=[recipient_email]).send(fail_silently=False)
                sent = True

            if sent:
                return JsonResponse({'status': 'success', 'message': _('Email sent successfully')})
            return JsonResponse({'status': 'error', 'message': _('Email was not sent')}, status=500)

        except Exception as e:
            logger.error(f"Error sending profile link: {e}", exc_info=True)
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

    except Exception as e:
        logger.error(f"send_profile_link error: {e}", exc_info=True)
        return JsonResponse({'status': 'error', 'message': _('Unexpected error: {}').format(str(e))}, status=500)

@login_required
@user_passes_test(check_risk_report_access)
@require_http_methods(["POST"])
def generate_risk_report_from_scheduled_execution(request):
    """Generate risk report from a scheduled report execution snapshot"""
    
    # Check permissions
    if not check_risk_report_access(request.user, 'view'):
        raise PermissionDenied(_("You don't have permission to generate reports"))
    
    try:
        execution_id = request.POST.get('execution_id')
        logger.info(f"generate_risk_report_from_scheduled_execution: received execution_id={execution_id}, POST data={dict(request.POST)}")
        if not execution_id:
            return JsonResponse({
                'status': 'error',
                'message': _('Execution ID is required')
            })
        
        # Get the execution
        try:
            logger.info(f"generate_risk_report_from_scheduled_execution: Looking for execution with ID: {execution_id}")
            execution = ScheduledReportExecution.objects.get(id=execution_id)
            logger.info(f"generate_risk_report_from_scheduled_execution: Found execution: {execution.id}")
        except ScheduledReportExecution.DoesNotExist:
            logger.error(f"generate_risk_report_from_scheduled_execution: Execution not found with ID: {execution_id}")
            return JsonResponse({
                'status': 'error',
                'message': _('Scheduled report execution not found')
            })
        
        # Check permissions: owner, same company, or email recipient
        report = execution.scheduled_report
        user_companies = get_user_companies(request.user)
        user_has_access = False
        
        # Check if user is the owner
        if report.created_by == request.user:
            user_has_access = True
        # Check if user belongs to the same company
        elif report.company and report.company in user_companies:
            user_has_access = True
        # Check if user is in the email recipients list
        elif report.email_recipients.filter(user=request.user).exists():
            user_has_access = True
        
        if not user_has_access:
            raise PermissionDenied(_("You don't have permission to use this report"))
        
        # Optional overrides from request
        override_format = request.POST.get('format') or 'pdf'
        override_language = request.POST.get('language') or execution.snapshot_language or 'uk'
        
        # Activate the language
        from django.utils.translation import activate
        activate(override_language)
        
        # Set language for the current request
        request.LANGUAGE_CODE = override_language
        
        # Validate format availability
        if not is_format_available(override_format):
            return JsonResponse({
                'status': 'error',
                'message': _('Selected format is not available. Please install required dependencies.')
            })
        
        # Get the snapshot HTML
        snapshot_html = execution.snapshot_html
        if not snapshot_html:
            return JsonResponse({
                'status': 'error',
                'message': _('Report snapshot is not available')
            })
        
        # Generate report file using the same approach as Report Profiles
        logger.info(f"generate_risk_report_from_scheduled_execution: Format requested: {override_format}")
        
        # Get the original scheduled report configuration
        scheduled_report = execution.scheduled_report
        
        # Prepare data from scheduled report (similar to profile approach)
        # Get sections_config from the associated report profile
        sections_config = {}
        company_id = ''
        if scheduled_report.report_profile:
            sections_config = scheduled_report.report_profile.get_sections_config()
            company_id = sections_config.get('_company_id', '')
        else:
            # Fallback: use basic configuration if no profile is associated
            sections_config = {
                'statistics': True,
                'risk_distribution': True,
                'compliance_overview': True,
                'top_risks': True,
                'asset_details': True,
                'vulnerability_details': True,
                'treatment_details': True,
                'recommendations': True
            }
        
        # Prepare parameters in the format expected by generate_report_data
        # Use company from scheduled report if available, otherwise from profile
        report_company_id = ''
        if scheduled_report.company:
            report_company_id = str(scheduled_report.company.id)
        elif company_id:
            report_company_id = str(company_id)
        
        params = {
            'reportType': 'scheduled',  # Use scheduled type
            'format': override_format,
            'language': override_language,
            'company_id': report_company_id,
            'selectedSections': sections_config,
            'scheduled_report': scheduled_report
        }
        
        # Add scheduled report information to params
        company_name = _('All Companies')
        if scheduled_report.company:
            company_name = scheduled_report.company.name
        elif company_id:
            try:
                from app_cabinet.models import Company
                company = Company.objects.get(id=company_id)
                company_name = company.name
            except:
                pass
        
        params['profile'] = {
            'name': scheduled_report.name,
            'description': scheduled_report.description or '',
            'created_by': scheduled_report.created_by.get_full_name(),
            'created_at': scheduled_report.created_at,
            'company': company_name,
        }
        
        # Generate report data using the same method as profiles
        report_data = generate_report_data(request.user, params)
        
        # Debug logging
        logger.info(f"Scheduled report info in params: {params.get('profile', {})}")
        logger.info(f"Scheduled report info in report_data: {report_data.get('profile', {})}")
        
        # Generate report file using the same method as profiles
        response = generate_report_file(report_data, 'scheduled', override_format)
        
        return response
        
    except Exception as e:
        logger.error(f"Error generating report from scheduled execution: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error generating report: {}').format(str(e))
        }, status=500)


@login_required
@user_passes_test(check_risk_report_access)
@require_http_methods(["POST"])
def generate_risk_report_from_profile(request):
    """Generate risk report using a saved profile"""
    
    # Check permissions
    if not check_risk_report_access(request.user, 'view'):
        raise PermissionDenied(_("You don't have permission to generate reports"))
    
    try:
        profile_id = request.POST.get('profile_id')
        logger.info(f"generate_risk_report_from_profile: received profile_id={profile_id}, POST data={dict(request.POST)}")
        if not profile_id:
            return JsonResponse({
                'status': 'error',
                'message': _('Profile ID is required')
            })
        
        # Get the profile
        try:
            profile = ReportProfile.objects.get(id=profile_id)
        except ReportProfile.DoesNotExist:
            return JsonResponse({
                'status': 'error',
                'message': _('Report profile not found')
            })
        
        # Check if user has access to this profile
        if not profile.can_be_used_by(request.user):
            raise PermissionDenied(_("You don't have permission to use this profile"))
        
        # Optional overrides from request
        override_format = request.POST.get('format') or profile.default_format
        override_language = request.POST.get('language') or profile.default_language
        
        # Activate the profile's default language for proper multilingual support
        from django.utils.translation import activate
        activate(override_language)
        
        # Set language for the current request
        request.LANGUAGE_CODE = override_language
        
        # Validate format availability
        if not is_format_available(override_format):
            return JsonResponse({
                'status': 'error',
                'message': _('Selected format is not available. Please install required dependencies.')
            })
        
        # Prepare data from profile
        # Get company_id from sections_config if available, otherwise empty
        sections_config = profile.get_sections_config()
        company_id = sections_config.get('_company_id', '')
        
        # Prepare parameters in the format expected by generate_report_data
        params = {
            'reportType': 'profile',  # Use profile type
            'format': override_format,
            'language': override_language,
            'company_id': str(company_id) if company_id else '',
            'selectedSections': sections_config,
            'profile': profile
        }
        
        # Add profile information to params before generating report data
        # Get company name from sections_config if available
        company_name = _('All Companies')
        if company_id:
            try:
                from app_cabinet.models import Company
                company = Company.objects.get(id=company_id)
                company_name = company.name
            except:
                pass
        
        params['profile'] = {
            'name': profile.name,
            'description': profile.description,
            'created_by': profile.created_by.get_full_name(),
            'created_at': profile.created_at,
            'company': company_name,
        }
        
        # Generate report data
        report_data = generate_report_data(request.user, params)
        
        # Debug logging
        logger.info(f"Profile info in params: {params.get('profile', {})}")
        logger.info(f"Profile info in report_data: {report_data.get('profile', {})}")
        
        # Generate report file
        response = generate_report_file(report_data, 'profile', override_format)
        
        # Update profile usage statistics
        profile.usage_count += 1
        profile.last_used_at = timezone.now()
        profile.save()
        
        return response
        
    except Exception as e:
        logger.error(f"Error generating report from profile: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error generating report: {}').format(str(e))
        })


@login_required
@user_passes_test(check_risk_report_access)
@require_http_methods(["POST"])
def preview_risk_report_from_profile(request):
    """Generate HTML preview for a report profile (always up-to-date)."""
    if not check_risk_report_access(request.user, 'view'):
        raise PermissionDenied(_("You don't have permission to preview reports"))
    try:
        profile_id = request.POST.get('profile_id')
        if not profile_id:
            return JsonResponse({'status': 'error', 'message': _('Profile ID is required')})

        try:
            profile = ReportProfile.objects.get(id=profile_id)
        except ReportProfile.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': _('Report profile not found')})

        if not profile.can_be_used_by(request.user):
            raise PermissionDenied(_("You don't have permission to use this profile"))

        # Activate language
        from django.utils.translation import activate
        activate(profile.default_language)
        request.LANGUAGE_CODE = profile.default_language

        # Prepare sections and company
        sections_config = profile.get_sections_config()
        company_id = sections_config.get('_company_id', '')

        # Build params for data generation
        params = {
            'reportType': 'profile',
            'format': 'html',
            'language': profile.default_language,
            'company_id': str(company_id) if company_id else '',
            'selectedSections': sections_config,
            'profile': profile
        }

        # Enrich profile info
        company_name = _('All Companies')
        if company_id:
            try:
                from app_cabinet.models import Company
                company = Company.objects.get(id=company_id)
                company_name = company.name
            except Exception:
                pass

        params['profile'] = {
            'name': profile.name,
            'description': profile.description,
            'created_by': profile.created_by.get_full_name(),
            'created_at': profile.created_at,
            'company': company_name,
        }

        # Generate fresh data and preview HTML
        report_data = generate_report_data(request.user, params)
        preview_html = generate_html_preview(report_data, 'profile', profile.default_language, sections_config)
        return HttpResponse(preview_html)

    except Exception as e:
        logger.error(f"Error previewing report from profile: {str(e)}", exc_info=True)
        error_html = f'<div class="alert alert-danger"><i class="fas fa-exclamation-triangle me-2"></i>{_("Error generating preview")}: {str(e)}</div>'
        return HttpResponse(error_html)
        timeline_section = {
            'type': 'timeline_trends',
            'title': translations.get('timeline_analysis', 'Timeline & Trends'),
            'subsections': []
        }
        
        if selected_sections.get('trend_analysis') and 'trend_data' in data:
            trend_data = []
            for trend in data['trend_data'][:10]:  # Limit to 10
                trend_data.append({
                    'period': trend.get('period', ''),
                    'metric': trend.get('metric', ''),
                    'value': trend.get('value', ''),
                    'direction': trend.get('direction', '')
                })
            
            if trend_data:
                timeline_section['subsections'].append({
                    'type': 'trend_overview',
                    'title': translations.get('trend_analysis', 'Trend Analysis'),
                    'headers': [
                        translations.get('time_period', 'Period'),
                        translations.get('metric', 'Metric'),
                        translations.get('value', 'Value'),
                        translations.get('trend_direction', 'Direction')
                    ],
                    'data': trend_data
                })
        
        if timeline_section['subsections']:
            content['sections'].append(timeline_section)
    
    # Risk Dependencies Section
    if selected_sections.get('dependency_analysis') or selected_sections.get('impact_assessment') or selected_sections.get('cascading_risks') or selected_sections.get('interdependency_matrix'):
        dependency_section = {
            'type': 'risk_dependencies',
            'title': translations.get('dependency_analysis', 'Risk Dependencies'),
            'subsections': []
        }
        
        if selected_sections.get('dependency_analysis') and 'dependency_data' in data:
            dependency_data = []
            for dep in data['dependency_data'][:10]:  # Limit to 10
                dependency_data.append({
                    'source': dep.get('source', ''),
                    'target': dep.get('target', ''),
                    'type': dep.get('type', ''),
                    'impact': dep.get('impact', '')
                })
            
            if dependency_data:
                dependency_section['subsections'].append({
                    'type': 'dependency_overview',
                    'title': translations.get('dependency_analysis', 'Dependency Analysis'),
                    'headers': [
                        translations.get('source', 'Source'),
                        translations.get('target', 'Target'),
                        translations.get('dependency_type', 'Type'),
                        translations.get('impact_level', 'Impact')
                    ],
                    'data': dependency_data
                })
        
        if dependency_section['subsections']:
            content['sections'].append(dependency_section)
    
    # Extended Asset Details Section
    if selected_sections.get('asset_criticality') or selected_sections.get('asset_location') or selected_sections.get('asset_owners') or selected_sections.get('asset_lifecycle'):
        extended_asset_section = {
            'type': 'extended_asset_details',
            'title': translations.get('asset_criticality', 'Extended Asset Details'),
            'subsections': []
        }
        
        if selected_sections.get('asset_criticality') and 'assets' in data:
            asset_criticality_data = []
            for asset in data['assets'][:15]:  # Limit to 15
                if hasattr(asset, 'get'):  # Dictionary
                    asset_criticality_data.append({
                        'asset_name': asset.get('name', ''),
                        'criticality': asset.get('criticality_name', ''),
                        'location': asset.get('location', ''),
                        'owner': asset.get('owner', '')
                    })
                else:  # Model object
                    # Get criticality using the proper method
                    criticality_info = asset.get_criticality()
                    criticality_name = criticality_info.get('name', '') if criticality_info else ''
                    criticality_color = criticality_info.get('color', '#000000') if criticality_info else '#000000'
                    
                    asset_criticality_data.append({
                        'asset_name': getattr(asset, 'name', ''),
                        'criticality': criticality_name,
                        'criticality_color': criticality_color,
                        'location': getattr(asset, 'location', ''),
                        'owner': getattr(asset, 'owner', '')
                    })
            
            if asset_criticality_data:
                extended_asset_section['subsections'].append({
                    'type': 'asset_criticality_overview',
                    'title': translations.get('asset_criticality', 'Asset Criticality'),
                    'headers': [
                        translations.get('asset_name', 'Asset Name'),
                        translations.get('criticality_level', 'Criticality'),
                        translations.get('physical_location', 'Location'),
                        translations.get('asset_owner', 'Owner')
                    ],
                    'data': asset_criticality_data
                })
        
        if extended_asset_section['subsections']:
            content['sections'].append(extended_asset_section)
    
    # Threat Analysis Section
    if selected_sections.get('threat_analysis') or selected_sections.get('threat_scenarios') or selected_sections.get('probability_analysis') or selected_sections.get('threat_landscape'):
        threat_section = {
            'type': 'threat_analysis',
            'title': translations.get('threat_analysis', 'Threat Analysis'),
            'subsections': []
        }
        
        if selected_sections.get('threat_analysis') and 'threats' in data:
            threat_data = []
            for threat in data['threats'][:10]:  # Limit to 10
                if hasattr(threat, 'get'):  # Dictionary
                    threat_data.append({
                        'threat_name': threat.get('name', ''),
                        'type': threat.get('type', ''),
                        'probability': threat.get('probability', ''),
                        'source': threat.get('source', '')
                    })
                else:  # Model object
                    threat_data.append({
                        'threat_name': getattr(threat, 'name', ''),
                        'type': getattr(threat, 'type', ''),
                        'probability': getattr(threat, 'probability', ''),
                        'source': getattr(threat, 'source', '')
                    })
            
            if threat_data:
                threat_section['subsections'].append({
                    'type': 'threat_overview',
                    'title': translations.get('threat_analysis', 'Threat Analysis'),
                    'headers': [
                        translations.get('threat_name', 'Threat Name'),
                        translations.get('threat_type', 'Type'),
                        translations.get('probability_level', 'Probability'),
                        translations.get('threat_source', 'Source')
                    ],
                    'data': threat_data
                })
        
        if threat_section['subsections']:
            content['sections'].append(threat_section)
    
    # Residual Risk & Priority Section
    if selected_sections.get('residual_risk_analysis') or selected_sections.get('risk_appetite') or selected_sections.get('priority_matrix') or selected_sections.get('resource_allocation'):
        residual_section = {
            'type': 'residual_risk_priority',
            'title': translations.get('residual_risk_analysis', 'Residual Risk & Priority'),
            'subsections': []
        }
        
        if selected_sections.get('residual_risk_analysis') and 'statistics' in data:
            stats = data['statistics']
            residual_section['subsections'].append({
                'type': 'residual_risk_overview',
                'title': translations.get('residual_risk_analysis', 'Residual Risk Analysis'),
                'data': [
                    {
                        'label': translations.get('residual_risk_level', 'Residual Risk Level'),
                        'value': stats.get('residual_risk_level', 'Medium'),
                        'type': 'warning'
                    },
                    {
                        'label': translations.get('appetite_threshold', 'Appetite Threshold'),
                        'value': f"{stats.get('appetite_threshold', 15):.1f}%",
                        'type': 'info'
                    },
                    {
                        'label': translations.get('priority_level', 'Priority Level'),
                        'value': stats.get('priority_level', 'High'),
                        'type': 'danger'
                    }
                ]
            })
        
        if residual_section['subsections']:
            content['sections'].append(residual_section)
    
    return content


def generate_audit_logs_data(assets, start_date, end_date):
    """Generate audit logs data for reporting"""
    try:
        # Get audit logs for the specified date range
        audit_logs = RiskAssessmentAuditLog.objects.filter(
            asset__in=assets,
            timestamp__date__range=[start_date, end_date]
        ).select_related('user', 'asset')[:50]  # Limit to 50 most recent
        
        return [{
            'timestamp': log.timestamp.strftime('%Y-%m-%d %H:%M:%S'),
            'user': log.user.get_full_name() or log.user.username if log.user else 'System',
            'action': log.action,
            'details': log.details or '',
            'asset': log.asset.name if log.asset else ''
        } for log in audit_logs]
    except Exception:
        return []


def generate_governance_data(assets, companies):
    """Generate governance data for reporting"""
    try:
        # Get users associated with the companies
        from app_cabinet.models import CabinetUser
        users = CabinetUser.objects.filter(company__in=companies)
        
        governance_data = []
        for user in users[:20]:  # Limit to 20
            governance_data.append({
                'role': user.position or 'User',
                'responsibility': f"Asset management for {user.company.name if user.company else 'N/A'}",
                'level': 'Operational',
                'status': 'Active' if user.is_active else 'Inactive'
            })
        
        return governance_data
    except Exception:
        return []


def generate_trend_data(assets, vulnerabilities, start_date, end_date):
    """Generate trend analysis data for reporting"""
    try:
        from datetime import timedelta
        import calendar
        
        # Generate monthly trend data
        trend_data = []
        current_date = start_date
        
        while current_date <= end_date:
            month_start = current_date.replace(day=1)
            month_end = (month_start + timedelta(days=32)).replace(day=1) - timedelta(days=1)
            
            # Count vulnerabilities for this month
            month_vulns = vulnerabilities.filter(
                modified_at__date__range=[month_start, min(month_end, end_date)]
            ).count()
            
            trend_data.append({
                'period': f"{calendar.month_name[current_date.month]} {current_date.year}",
                'metric': 'Vulnerabilities',
                'value': str(month_vulns),
                'direction': 'stable'  # Could be calculated based on previous month
            })
            
            # Move to next month
            if current_date.month == 12:
                current_date = current_date.replace(year=current_date.year + 1, month=1)
            else:
                current_date = current_date.replace(month=current_date.month + 1)
        
        return trend_data[:12]  # Limit to 12 months
    except Exception:
        return []


def generate_dependency_data(assets, vulnerabilities):
    """Generate dependency analysis data for reporting"""
    try:
        dependency_data = []
        
        # Simple dependency analysis based on asset relationships
        for vuln in vulnerabilities.select_related('asset', 'vulnerability')[:15]:
            if vuln.asset and vuln.vulnerability:
                dependency_data.append({
                    'source': vuln.asset.name,
                    'target': vuln.vulnerability.get_name() or vuln.vulnerability.get_name('en'),
                    'type': 'Vulnerability',
                    'impact': 'High' if vuln.status == 'Yes' else 'Low'
                })
        
        return dependency_data
    except Exception:
        return []


def generate_threat_data(vulnerabilities):
    """Generate threat analysis data for reporting"""
    try:
        from .models import Threat
        
        # Get unique threats from vulnerabilities
        threats = Threat.objects.filter(
            vulnerability__in=vulnerabilities.values_list('vulnerability', flat=True)
        ).distinct()[:15]
        
        threat_data = []
        for threat in threats:
            threat_data.append({
                'name': threat.threat_uk or threat.threat_en or 'Unknown',
                'type': 'Security Threat',
                'probability': 'Medium',  # Could be calculated
                'source': 'Internal/External'  # Could be determined from threat data
            })
        
        return threat_data
    except Exception:
        return []


def generate_pdf_from_html(html_content, execution):
    """Generate PDF from HTML content using ReportLab"""
    try:
        if not REPORTLAB_AVAILABLE:
            return JsonResponse({
                'status': 'error',
                'message': _('PDF generation is not available. Please install reportlab.')
            })
        
        # Log HTML content for debugging
        logger.info(f"generate_pdf_from_html: HTML content length: {len(html_content)}")
        logger.info(f"generate_pdf_from_html: First 500 chars: {html_content[:500]}")
        logger.info(f"generate_pdf_from_html: Contains 'black' or 'rectangle': {'black' in html_content.lower() or 'rectangle' in html_content.lower()}")
        
        # Create PDF using ReportLab
        from io import BytesIO
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.fonts import addMapping
        
        # Create buffer for PDF
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(A4))
        story = []
        
        # Register fonts for Cyrillic support
        try:
            # Try to register a font that supports Cyrillic
            pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
            cyrillic_font = 'DejaVuSans'
        except:
            try:
                # Fallback to Arial Unicode MS if available
                pdfmetrics.registerFont(TTFont('ArialUnicodeMS', 'ARIALUNI.TTF'))
                cyrillic_font = 'ArialUnicodeMS'
            except:
                # Use default font if no Cyrillic fonts available
                cyrillic_font = 'Helvetica'
        
        # Get styles and customize for Cyrillic
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=cyrillic_font,
            fontSize=16,
            spaceAfter=12
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=cyrillic_font,
            fontSize=10,
            spaceAfter=6
        )
        
        # Add title
        title = Paragraph(f"Risk Assessment Report - {execution.scheduled_report.name}", title_style)
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Add metadata
        story.append(Paragraph(f"Generated: {execution.started_at.strftime('%Y-%m-%d %H:%M')}", normal_style))
        story.append(Paragraph(f"Report: {execution.scheduled_report.name}", normal_style))
        story.append(Paragraph(f"Company: {execution.scheduled_report.company.name if execution.scheduled_report.company else 'All Companies'}", normal_style))
        story.append(Spacer(1, 12))
        
        # Extract and process HTML content
        if BEAUTIFULSOUP_AVAILABLE:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Clean up the HTML content
            html_content = str(soup)
            
            # Log processed content
            logger.info(f"generate_pdf_from_html: Processed HTML length: {len(html_content)}")
            
            # Process different HTML elements
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'span', 'table', 'tr', 'td', 'th']):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    # Headers
                    text = element.get_text().strip()
                    if text and not text.startswith('█'):  # Filter out black rectangles
                        if element.name == 'h1':
                            story.append(Paragraph(text, title_style))
                        elif element.name == 'h2':
                            header_style = ParagraphStyle(
                                'CustomHeading2',
                                parent=styles['Heading2'],
                                fontName=cyrillic_font,
                                fontSize=14,
                                spaceAfter=8
                            )
                            story.append(Paragraph(text, header_style))
                        else:
                            header_style = ParagraphStyle(
                                'CustomHeading3',
                                parent=styles['Heading3'],
                                fontName=cyrillic_font,
                                fontSize=12,
                                spaceAfter=6
                            )
                            story.append(Paragraph(text, header_style))
                        story.append(Spacer(1, 6))
                
                elif element.name == 'p':
                    # Paragraphs
                    text = element.get_text().strip()
                    if text and not text.startswith('█'):  # Filter out black rectangles
                        story.append(Paragraph(text, normal_style))
                        story.append(Spacer(1, 6))
                
                elif element.name == 'table':
                    # Tables - create proper table structure for PDF
                    try:
                        from reportlab.platypus import Table, TableStyle
                        from reportlab.lib import colors
                        
                        # Count rows and columns
                        rows = element.find_all('tr')
                        if rows:
                            # Get max columns
                            max_cols = 0
                            for row in rows:
                                cells = row.find_all(['td', 'th'])
                                max_cols = max(max_cols, len(cells))
                            
                            if max_cols > 0:
                                # Create table data
                                table_data = []
                                for row in rows:
                                    cells = row.find_all(['td', 'th'])
                                    row_data = []
                                    for cell in cells:
                                        cell_text = cell.get_text().strip()
                                        if cell_text and not cell_text.startswith('█'):
                                            row_data.append(cell_text)
                                        else:
                                            row_data.append('')
                                    # Pad row if needed
                                    while len(row_data) < max_cols:
                                        row_data.append('')
                                    table_data.append(row_data)
                                
                                if table_data:
                                    # Create PDF table
                                    pdf_table = Table(table_data)
                                    pdf_table.setStyle(TableStyle([
                                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                        ('FONTNAME', (0, 0), (-1, 0), cyrillic_font),
                                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                                    ]))
                                    story.append(pdf_table)
                                    story.append(Spacer(1, 12))
                        
                        # If table creation failed, fall back to text
                        else:
                            table_text = element.get_text().strip()
                            if table_text:
                                lines = table_text.split('\n')
                                for line in lines:
                                    if line.strip() and not line.strip().startswith('█'):
                                        story.append(Paragraph(line.strip(), normal_style))
                                story.append(Spacer(1, 6))
                    except:
                        # Fallback to text if table creation fails
                        table_text = element.get_text().strip()
                        if table_text:
                            lines = table_text.split('\n')
                            for line in lines:
                                if line.strip() and not line.strip().startswith('█'):
                                    story.append(Paragraph(line.strip(), normal_style))
                            story.append(Spacer(1, 6))
                
                elif element.name in ['div', 'span']:
                    # Divs and spans - extract text if it's meaningful
                    text = element.get_text().strip()
                    if text and len(text) > 10 and not text.startswith('█'):  # Only add if it's substantial text and not black rectangles
                        story.append(Paragraph(text, normal_style))
                        story.append(Spacer(1, 6))
            
            # If no structured content found, fall back to simple text extraction
            if not story:
                text_content = soup.get_text()
                paragraphs = text_content.split('\n\n')
                for para in paragraphs:
                    if para.strip() and not para.strip().startswith('█'):  # Filter out black rectangles
                        clean_text = para.strip().replace('\n', ' ')
                        story.append(Paragraph(clean_text, normal_style))
                        story.append(Spacer(1, 6))
        else:
            # Fallback: simple text extraction
            import re
            text_content = re.sub(r'<[^>]+>', '', html_content)
            paragraphs = text_content.split('\n\n')
            for para in paragraphs:
                if para.strip() and not para.strip().startswith('█'):  # Filter out black rectangles
                    clean_text = para.strip().replace('\n', ' ')
                    story.append(Paragraph(clean_text, normal_style))
                    story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        # Create response
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="risk_report_{execution.scheduled_report.name}_{execution.started_at.strftime("%Y%m%d")}.pdf"'
        
        return response
        
    except Exception as e:
        logger.error(f"Error generating PDF from HTML: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error generating PDF: {}').format(str(e))
        })


def generate_word_from_html(html_content, execution):
    """Generate Word document from HTML content"""
    logger.info(f"generate_word_from_html: Function called with execution_id: {execution.id}")
    try:
        if not DOCX_AVAILABLE:
            from django.utils.translation import gettext as _
            logger.error("generate_word_from_html: DOCX_AVAILABLE is False")
            return JsonResponse({
                'status': 'error',
                'message': _('Word generation is not available. Please install python-docx.')
            })
        
        logger.info(f"generate_word_from_html: DOCX_AVAILABLE is True, proceeding...")
        
        # Log HTML content for debugging
        logger.info(f"generate_word_from_html: HTML content length: {len(html_content)}")
        logger.info(f"generate_word_from_html: First 500 chars: {html_content[:500]}")
        logger.info(f"generate_word_from_html: Contains 'black' or 'rectangle': {'black' in html_content.lower() or 'rectangle' in html_content.lower()}")
        
        # Create Word document
        doc = Document()
        
        # Set landscape orientation for all sections
        try:
            from docx.shared import Inches
            for section in doc.sections:
                # Set landscape orientation (A4 landscape: 11.69" x 8.27")
                section.page_width = Inches(11.69)
                section.page_height = Inches(8.27)
        except Exception:
            pass  # Continue if orientation setting fails
        
        # Set document language for better Cyrillic support
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            # Set language to Ukrainian
            element = OxmlElement('w:lang')
            element.set(qn('w:val'), 'uk-UA')
            element.set(qn('w:eastAsia'), 'uk-UA')
            element.set(qn('w:bidi'), 'uk-UA')
            
            # Add to document settings
            settings = doc.settings._element
            settings.append(element)
        except:
            pass  # Continue if language setting fails
        
        # Add title
        title = doc.add_heading(f"Risk Assessment Report - {execution.scheduled_report.name}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Generated: {execution.started_at.strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Report: {execution.scheduled_report.name}")
        doc.add_paragraph(f"Company: {execution.scheduled_report.company.name if execution.scheduled_report.company else 'All Companies'}")
        
        # Add content from HTML with better parsing
        if BEAUTIFULSOUP_AVAILABLE:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Clean up the HTML content
            html_content = str(soup)
            
            # Log processed content
            logger.info(f"generate_word_from_html: Processed HTML length: {len(html_content)}")
            
            # Log all elements found
            all_elements = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'span', 'table', 'tr', 'td', 'th'])
            logger.info(f"generate_word_from_html: Found {len(all_elements)} elements to process")
            
            # Count elements by type
            element_counts = {}
            for elem in all_elements:
                element_counts[elem.name] = element_counts.get(elem.name, 0) + 1
            logger.info(f"generate_word_from_html: Element counts: {element_counts}")
            
            # Process different HTML elements
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'span', 'table', 'tr', 'td', 'th']):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    # Headers
                    text = element.get_text().strip()
                    if text and not text.startswith('█'):  # Filter out black rectangles
                        level = int(element.name[1])  # Extract number from h1, h2, etc.
                        heading = doc.add_heading(text, level=min(level, 6))
                        
                        # Add color to headings based on classes
                        classes = element.get('class', [])
                        if isinstance(classes, str):
                            classes = [classes]
                        
                        if 'text-primary' in classes:
                            for run in heading.runs:
                                run.font.color.rgb = RGBColor(0, 123, 255)  # Blue
                        elif 'text-secondary' in classes:
                            for run in heading.runs:
                                run.font.color.rgb = RGBColor(108, 117, 125)  # Gray
                
                elif element.name == 'p':
                    # Paragraphs
                    text = element.get_text().strip()
                    if text and not text.startswith('█'):  # Filter out black rectangles
                        p = doc.add_paragraph(text)
                        
                        # Add styling based on classes
                        classes = element.get('class', [])
                        if isinstance(classes, str):
                            classes = [classes]
                        
                        if 'text-muted' in classes:
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(108, 117, 125)  # Gray
                        elif 'text-center' in classes:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                elif element.name == 'table':
                    # Tables - create proper table structure
                    logger.info(f"generate_word_from_html: Found table element")
                    try:
                        from docx.shared import Inches
                        
                        # Count rows and columns
                        rows = element.find_all('tr')
                        logger.info(f"generate_word_from_html: Table has {len(rows)} rows")
                        
                        if rows:
                            # Get max columns
                            max_cols = 0
                            for row in rows:
                                cells = row.find_all(['td', 'th'])
                                max_cols = max(max_cols, len(cells))
                            
                            logger.info(f"generate_word_from_html: Table has {max_cols} columns")
                            
                            if max_cols > 0:
                                # Create table
                                table = doc.add_table(rows=len(rows), cols=max_cols)
                                table.style = 'Table Grid'
                                
                                # Fill table data
                                for i, row in enumerate(rows):
                                    cells = row.find_all(['td', 'th'])
                                    for j, cell in enumerate(cells):
                                        if j < max_cols:
                                            cell_text = cell.get_text().strip()
                                            if cell_text and not cell_text.startswith('█'):
                                                table_cell = table.cell(i, j)
                                                table_cell.text = cell_text
                                                
                                                # Style header row (first row)
                                                if i == 0:
                                                    for paragraph in table_cell.paragraphs:
                                                        for run in paragraph.runs:
                                                            run.bold = True
                                                            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                                
                                                logger.info(f"generate_word_from_html: Added cell [{i},{j}]: {cell_text[:50]}...")
                        
                        # If table creation failed, fall back to text
                        else:
                            logger.info(f"generate_word_from_html: No rows found in table, falling back to text")
                            table_text = element.get_text().strip()
                            if table_text:
                                lines = table_text.split('\n')
                                for line in lines:
                                    if line.strip() and not line.strip().startswith('█'):
                                        doc.add_paragraph(line.strip())
                    except Exception as e:
                        # Fallback to text if table creation fails
                        logger.error(f"generate_word_from_html: Error creating table: {str(e)}")
                        table_text = element.get_text().strip()
                        if table_text:
                            lines = table_text.split('\n')
                            for line in lines:
                                if line.strip() and not line.strip().startswith('█'):
                                    doc.add_paragraph(line.strip())
                
                elif element.name in ['div', 'span']:
                    # Divs and spans - extract text if it's meaningful
                    text = element.get_text().strip()
                    if text and len(text) > 10 and not text.startswith('█'):  # Only add if it's substantial text and not black rectangles
                        # Check if this might be a colored block (look for specific classes or styles)
                        classes = element.get('class', [])
                        if isinstance(classes, str):
                            classes = [classes]
                        
                        # Skip div elements that contain tables or other processed elements
                        if element.name == 'div':
                            if element.find('table'):
                                logger.info(f"generate_word_from_html: Skipping div with table inside")
                                continue
                            # Skip div elements that contain card elements (they will be processed separately)
                            if element.find(class_=lambda x: x and 'card' in x):
                                logger.info(f"generate_word_from_html: Skipping div with card inside")
                                continue
                            # Skip div elements that contain metric elements (they will be processed separately)
                            if element.find(class_=lambda x: x and ('metric' in x or 'block' in x)):
                                logger.info(f"generate_word_from_html: Skipping div with metric/block inside")
                                continue
                        
                        logger.info(f"generate_word_from_html: Processing {element.name} with classes: {classes}, text: {text[:100]}...")
                        
                        # Add colored background for certain elements (simulate colored blocks)
                        if any('card' in cls.lower() for cls in classes):
                            # Only process card elements, skip others to avoid duplication
                            # Only process if text is substantial (not just whitespace or very short)
                            if len(text.strip()) > 5:
                                p = doc.add_paragraph()
                                run = p.add_run(text)
                                run.bold = True
                            
                            # Add color based on card class
                            if 'bg-primary' in classes:
                                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                logger.info(f"generate_word_from_html: Made blue card bold with white text: {text[:50]}...")
                            elif 'bg-warning' in classes:
                                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                logger.info(f"generate_word_from_html: Made yellow card bold with white text: {text[:50]}...")
                            elif 'bg-danger' in classes:
                                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                logger.info(f"generate_word_from_html: Made red card bold with white text: {text[:50]}...")
                            elif 'bg-success' in classes:
                                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                logger.info(f"generate_word_from_html: Made green card bold with white text: {text[:50]}...")
                            else:
                                logger.info(f"generate_word_from_html: Made text bold for colored block: {text[:50]}...")
                        elif any('metric' in cls.lower() or 'block' in cls.lower() for cls in classes):
                            # Process other metric/block elements
                            # Only process if text is substantial
                            if len(text.strip()) > 5:
                                p = doc.add_paragraph()
                                run = p.add_run(text)
                                run.bold = True
                                logger.info(f"generate_word_from_html: Made text bold for metric/block: {text[:50]}...")
                        else:
                            # Only add if it's not a card or metric element
                            doc.add_paragraph(text)
            
            # If no structured content found, fall back to simple text extraction
            if len(doc.paragraphs) <= 4:  # Only title and metadata
                text_content = soup.get_text()
                paragraphs = text_content.split('\n\n')
                for para in paragraphs:
                    if para.strip() and not para.strip().startswith('█'):  # Filter out black rectangles
                        doc.add_paragraph(para.strip())
        else:
            # Fallback: simple text extraction
            import re
            text_content = re.sub(r'<[^>]+>', '', html_content)
            paragraphs = text_content.split('\n\n')
            for para in paragraphs:
                if para.strip() and not para.strip().startswith('█'):  # Filter out black rectangles
                    doc.add_paragraph(para.strip())
        
        # Save to BytesIO
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Create response
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="risk_report_{execution.scheduled_report.name}_{execution.started_at.strftime("%Y%m%d")}.docx"'
        
        return response
        
    except Exception as e:
        logger.error(f"Error generating Word from HTML: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error generating Word document: {}').format(str(e))
        })