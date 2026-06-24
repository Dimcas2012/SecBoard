# SecBoard/app_risk/report_generators_scheduled.py

import json
import logging
import os
import tempfile
from datetime import datetime, timedelta
from decimal import Decimal
import decimal
from io import BytesIO

from django.contrib.auth.decorators import login_required, user_passes_test
from django.core.exceptions import PermissionDenied
from django.db.models import Count, Q, Sum, Avg, Max, Min
from django.http import JsonResponse, HttpResponse, Http404
from django.shortcuts import render, get_object_or_404
from django.template.loader import render_to_string
from django.utils import timezone
from django.utils.translation import gettext as _, get_language
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods

# Third-party imports
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BEAUTIFULSOUP_AVAILABLE = True
except ImportError:
    BEAUTIFULSOUP_AVAILABLE = False

# Local imports
from .models import (
    RiskTreatment, AssetVulnerability, Vulnerability, Threat, 
    RiskLevel, AccessRisk, RiskAssessmentAuditLog,
    ScheduledReport, ScheduledReportExecution, ScheduledReportAttachment, ReportProfile, RiskReportEmailConfig
)
from app_cabinet.models import CabinetUser
from app_asset.models import InformationAsset
from app_conf.models import Company
from .risk_assessment_views import calculate_risk_level, calculate_value_of_risk
from .access_utils import can_add_risk_report, can_edit_risk_report, can_delete_risk_report
from .report_generators_core import (
    check_risk_assessment_access,
    check_risk_report_access,
    generate_report_file,
    generate_html_preview,
    is_format_available,
    logger,
)
from .report_data import generate_report_data
from .report_views import get_user_companies

def get_scheduled_reports(request):
    """Get list of scheduled reports for the current user"""
    try:
        # Check if user has access to risk reports
        if not check_risk_report_access(request.user, 'view'):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to access scheduled reports')
            }, status=403)
        user_companies = get_user_companies(request.user)
        
        # Get scheduled reports accessible to the user (owner, same company, or email recipient)
        scheduled_reports = ScheduledReport.objects.filter(
            Q(created_by=request.user) | 
            Q(company__in=user_companies) |
            Q(email_recipients__user=request.user)
        ).select_related('company', 'created_by', 'report_profile', 'mail_server', 'mail_account').prefetch_related('email_recipients').distinct()
        
        # Format data for frontend
        reports_data = []
        for report in scheduled_reports:
            # Calculate next run if not set
            if not report.next_run:
                report.next_run = report.calculate_next_run()
                report.save()
            
            company_name = report.company.name if report.company else _('All Companies')
            logger.info(f"Scheduled report {report.name}: company={report.company}, company_name={company_name}")
            
            # Get email recipients summary
            recipients = report.email_recipients.all()
            recipients_count = recipients.count()
            recipients_names = []
            for recipient in recipients[:3]:  # Show first 3 names
                user = recipient.user
                name = user.get_full_name() or user.username
                recipients_names.append(name)
            
            recipients_summary = ', '.join(recipients_names)
            if recipients_count > 3:
                recipients_summary += f' +{recipients_count - 3} more'
            
            # Determine email settings display
            email_settings_display = ""
            if report.send_email:
                if report.use_default_email_settings:
                    email_settings_display = f'<div class="text-center"><span class="badge bg-success"><i class="fas fa-cog me-1"></i>{_("Default")}</span><br><small class="text-muted">{_("System settings")}</small></div>'
                else:
                    if report.mail_server and report.mail_account:
                        email_settings_display = f'''<div class="text-center">
                            <div><span class="badge bg-info"><i class="fas fa-server me-1"></i>{report.mail_server.name}</span></div>
                            <small class="text-muted d-block mt-1"><i class="fas fa-user me-1"></i>{report.mail_account.username}</small>
                            <small class="text-muted"><i class="fas fa-envelope me-1"></i>{report.mail_server.smtp_host}:{report.mail_server.smtp_port}</small>
                        </div>'''
                    elif report.mail_server:
                        email_settings_display = f'''<div class="text-center">
                            <span class="badge bg-warning"><i class="fas fa-server me-1"></i>{report.mail_server.name}</span>
                            <br><small class="text-warning"><i class="fas fa-exclamation-triangle me-1"></i>{_("No account")}</small>
                            <br><small class="text-muted">{report.mail_server.smtp_host}:{report.mail_server.smtp_port}</small>
                        </div>'''
                    else:
                        email_settings_display = f'<div class="text-center"><span class="badge bg-warning"><i class="fas fa-exclamation-triangle me-1"></i>{_("Incomplete")}</span><br><small class="text-warning">{_("No server configured")}</small></div>'
            else:
                email_settings_display = f'<div class="text-center"><span class="badge bg-secondary"><i class="fas fa-envelope-slash me-1"></i>{_("Disabled")}</span><br><small class="text-muted">{_("Email sending disabled")}</small></div>'
            
            # Check for active execution (running or pending, or completed but very recent)
            active_execution = None
            try:
                from .models import ScheduledReportExecution
                from datetime import timedelta
                # Check for running or pending executions
                active_execution = ScheduledReportExecution.objects.filter(
                    scheduled_report=report,
                    status__in=['pending', 'running']
                ).order_by('-started_at').first()
                
                # If no active execution, check for very recent completed ones (within last 30 seconds)
                # This helps catch executions that completed very quickly
                if not active_execution:
                    recent_threshold = timezone.now() - timedelta(seconds=30)
                    recent_execution = ScheduledReportExecution.objects.filter(
                        scheduled_report=report,
                        status='completed',
                        started_at__gte=recent_threshold
                    ).order_by('-started_at').first()
                    if recent_execution:
                        # Show as completing if very recent
                        active_execution = recent_execution
            except Exception as e:
                logger.warning(f"Error checking active execution for report {report.id}: {e}")
            
            execution_status = None
            execution_progress = None
            execution_started_at = None
            if active_execution:
                execution_status = active_execution.status
                execution_started_at = active_execution.started_at.astimezone().strftime('%Y-%m-%d %H:%M:%S')
                # Determine progress based on execution state
                if active_execution.status == 'pending':
                    execution_progress = 'pending'
                elif active_execution.status == 'running':
                    # Check if report generation is complete (snapshot_html exists)
                    if active_execution.snapshot_html:
                        # Check if email is already sent
                        if active_execution.email_sent:
                            # Email sent, execution should be completed soon
                            execution_progress = 'completing'
                        else:
                            # Report generated, now sending email
                            execution_progress = 'sending_email'
                    else:
                        # Report generation in progress
                        execution_progress = 'generating_report'
                elif active_execution.status == 'completed':
                    # Show as completing if very recent (within last 30 seconds)
                    execution_progress = 'completing'
            elif report.next_run and report.next_run <= timezone.now() and report.status == 'active':
                # Next run time has passed but no execution created yet - should start soon
                # This can happen if Celery task hasn't run yet or is delayed
                # Check if execution was created very recently (within last 5 minutes)
                # If not, it's truly pending
                from datetime import timedelta
                recent_threshold = timezone.now() - timedelta(minutes=5)
                recent_execution = ScheduledReportExecution.objects.filter(
                    scheduled_report=report,
                    started_at__gte=recent_threshold
                ).order_by('-started_at').first()
                
                if not recent_execution:
                    # No recent execution - truly pending
                    execution_status = 'pending'
                    execution_progress = 'pending'
                    execution_started_at = None
                else:
                    # Recent execution exists - use its status
                    execution_status = recent_execution.status
                    execution_started_at = recent_execution.started_at.astimezone().strftime('%Y-%m-%d %H:%M:%S')
                    if recent_execution.status == 'running':
                        if recent_execution.snapshot_html:
                            if recent_execution.email_sent:
                                execution_progress = 'completing'
                            else:
                                execution_progress = 'sending_email'
                        else:
                            execution_progress = 'generating_report'
                    elif recent_execution.status == 'completed':
                        execution_progress = 'completing'
                    else:
                        execution_progress = 'pending'
            # Don't show execution status for completed reports
            elif report.status == 'completed':
                execution_status = None
                execution_progress = None
                execution_started_at = None
            
            # Calculate time until next_run
            time_until_next_run = None
            if report.next_run and report.status == 'active':
                time_diff = report.next_run - timezone.now()
                if time_diff.total_seconds() > 0:
                    total_seconds = int(time_diff.total_seconds())
                    days = total_seconds // 86400
                    hours = (total_seconds % 86400) // 3600
                    minutes = (total_seconds % 3600) // 60
                    
                    time_parts = []
                    if days > 0:
                        time_parts.append(f"{days} {_('day') if days == 1 else _('days')}")
                    if hours > 0:
                        time_parts.append(f"{hours} {_('hour') if hours == 1 else _('hours')}")
                    if minutes > 0 or len(time_parts) == 0:
                        time_parts.append(f"{minutes} {_('minute') if minutes == 1 else _('minutes')}")
                    
                    time_until_next_run = ', '.join(time_parts)
                else:
                    # Time has passed
                    time_until_next_run = _('Overdue')
            
            reports_data.append({
                'id': str(report.id),
                'name': report.name,
                'description': report.description,
                # 'report_type': report.get_report_type_display(),  # removed per new design
                'report_profile_id': str(report.report_profile.id) if report.report_profile else None,
                'report_profile_name': report.report_profile.name if report.report_profile else None,
                # 'report_format': report.report_format.upper(),  # removed in link mode
                'frequency': report.get_frequency_display(),
                'status': report.get_status_display(),
                'company': company_name,
                'created_by': report.created_by.get_full_name() or report.created_by.username,
                'created_at': report.created_at.astimezone().strftime('%Y-%m-%d %H:%M'),
                'last_run': report.last_run.astimezone().strftime('%Y-%m-%d %H:%M') if report.last_run else _('Never'),
                'next_run': report.next_run.astimezone().strftime('%Y-%m-%d %H:%M') if report.next_run else _('Not scheduled'),
                'next_run_timestamp': report.next_run.timestamp() if report.next_run else None,  # For client-side calculation
                'time_until_next_run': time_until_next_run,
                'run_count': report.run_count,
                'email_recipients_count': recipients_count,
                'email_recipients_summary': recipients_summary,
                'email_settings_display': email_settings_display,
                'execution_time': report.execution_time.strftime('%H:%M'),
                'weekdays': ', '.join(report.get_weekdays()) if report.frequency == 'weekly' else '',
                'day_of_month': report.day_of_month if report.frequency == 'monthly' else None,
                'execution_status': execution_status,
                'execution_progress': execution_progress,
                'execution_started_at': execution_started_at,
            })
        
        return JsonResponse({
            'status': 'success',
            'data': reports_data
        })
        
    except Exception as e:
        logger.error(f"Error getting scheduled reports: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error loading scheduled reports: {}').format(str(e))
        })


@login_required
@user_passes_test(can_add_risk_report)
@require_http_methods(["POST"])
def create_scheduled_report(request):
    """Create a new scheduled report"""
    try:
        # Check if user has permission to add reports
        if not can_add_risk_report(request.user):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to create scheduled reports')
            }, status=403)
        # Handle both JSON and form data
        if request.content_type == 'application/json':
            data = json.loads(request.body)
        else:
            # Handle form data
            data = {}
            for key, value in request.POST.items():
                if key == 'csrfmiddlewaretoken':
                    continue
                if key == 'email_recipients':
                    # Handle multiple values for email_recipients
                    data[key] = request.POST.getlist(key)
                elif value.lower() == 'true':
                    data[key] = True
                elif value.lower() == 'false':
                    data[key] = False
                else:
                    data[key] = value
        
        # Validate required fields
        required_fields = ['name', 'report_profile_id', 'report_format', 'frequency', 'start_date', 'execution_time']
        for field in required_fields:
            if not data.get(field):
                return JsonResponse({
                    'status': 'error',
                    'message': _('Field {} is required').format(field)
                })
        
        # Get report profile
        try:
            report_profile = ReportProfile.objects.get(id=data['report_profile_id'])
        except ReportProfile.DoesNotExist:
            return JsonResponse({
                'status': 'error',
                'message': _('Selected report profile not found')
                })
        
        # Create scheduled report
        scheduled_report = ScheduledReport(
            name=data['name'],
            description=data.get('description', ''),
            report_profile=report_profile,
            report_format=data['report_format'],
            report_language=data.get('report_language', 'uk'),
            frequency=data['frequency'],
            start_date=datetime.strptime(data['start_date'], '%Y-%m-%d').date(),
            execution_time=datetime.strptime(data['execution_time'], '%H:%M').time(),
            email_subject=data.get('email_subject', ''),
            email_body=data.get('email_body', ''),
            send_email=data.get('send_email', True),
            use_default_email_settings=data.get('use_default_email_settings', True),
            created_by=request.user
        )
        
        # Validate and set company (required field)
        company_id = data.get('company_id')
        logger.info(f"Creating scheduled report with company_id: {company_id} (type: {type(company_id)})")
        
        if not company_id:
            logger.error("Company selection is required for scheduled reports")
            return JsonResponse({
                'status': 'error',
                'message': _('Company selection is required for scheduled reports')
            })
        
        try:
            company = Company.objects.get(id=company_id)
            scheduled_report.company = company
            logger.info(f"Successfully set company: {company.name} (ID: {company.id})")
        except Company.DoesNotExist:
            logger.error(f"Company with ID {company_id} not found")
            return JsonResponse({
                'status': 'error',
                'message': _('Selected company not found')
            })
        
        # Set mail server and account if specified
        if data.get('mail_server') and not data.get('use_default_email_settings', True):
            try:
                from app_conf.models import MailServer
                mail_server = MailServer.objects.get(id=data['mail_server'])
                scheduled_report.mail_server = mail_server
            except MailServer.DoesNotExist:
                return JsonResponse({
                    'status': 'error',
                    'message': _('Selected mail server not found')
                })
        
        if data.get('mail_account') and not data.get('use_default_email_settings', True):
            try:
                from app_conf.models import MailAccount
                mail_account = MailAccount.objects.get(id=data['mail_account'])
                scheduled_report.mail_account = mail_account
            except MailAccount.DoesNotExist:
                return JsonResponse({
                    'status': 'error',
                    'message': _('Selected mail account not found')
                })
        
        # Set end date if specified
        if data.get('end_date'):
            scheduled_report.end_date = datetime.strptime(data['end_date'], '%Y-%m-%d').date()
        
        # Set frequency-specific fields
        if data['frequency'] == 'weekly':
            scheduled_report.monday = data.get('monday', False)
            scheduled_report.tuesday = data.get('tuesday', False)
            scheduled_report.wednesday = data.get('wednesday', False)
            scheduled_report.thursday = data.get('thursday', False)
            scheduled_report.friday = data.get('friday', False)
            scheduled_report.saturday = data.get('saturday', False)
            scheduled_report.sunday = data.get('sunday', False)
        elif data['frequency'] == 'monthly':
            scheduled_report.day_of_month = data.get('day_of_month', 1)
        
        # Calculate next run
        scheduled_report.next_run = scheduled_report.calculate_next_run()
        
        scheduled_report.save()
        
        # Set email recipients
        if data.get('email_recipients'):
            recipient_ids = data['email_recipients']
            # Get CabinetUser objects from IDs
            cabinet_users = CabinetUser.objects.filter(id__in=recipient_ids)
            scheduled_report.email_recipients.set(cabinet_users)
        
        # Handle file attachments
        if request.FILES.getlist('attachments'):
            for uploaded_file in request.FILES.getlist('attachments'):
                # Create attachment record
                attachment = ScheduledReportAttachment(
                    scheduled_report=scheduled_report,
                    file=uploaded_file,
                    original_filename=uploaded_file.name,
                    file_size=uploaded_file.size,
                    uploaded_by=request.user
                )
                attachment.save()
        
        return JsonResponse({
            'status': 'success',
            'message': _('Scheduled report created successfully'),
            'id': str(scheduled_report.id)
        })
        
    except Exception as e:
        logger.error(f"Error creating scheduled report: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error creating scheduled report: {}').format(str(e))
        })


@login_required
@user_passes_test(can_edit_risk_report)
@require_http_methods(["POST"])
def update_scheduled_report(request, report_id):
    """Update an existing scheduled report"""
    try:
        # Check if user has permission to edit reports
        if not can_edit_risk_report(request.user):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to edit scheduled reports')
            }, status=403)
        # Get the scheduled report
        scheduled_report = ScheduledReport.objects.get(id=report_id)
        
        # Check permissions
        user_companies = get_user_companies(request.user)
        if (scheduled_report.created_by != request.user and 
            scheduled_report.company not in user_companies):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to edit this scheduled report')
            })
        
        # Handle both JSON and form data
        if request.content_type == 'application/json':
            data = json.loads(request.body)
        else:
            # Handle form data
            data = {}
            for key, value in request.POST.items():
                if key == 'csrfmiddlewaretoken':
                    continue
                if key == 'email_recipients':
                    # Handle multiple values for email_recipients
                    data[key] = request.POST.getlist(key)
                elif value.lower() == 'true':
                    data[key] = True
                elif value.lower() == 'false':
                    data[key] = False
                else:
                    data[key] = value
        
        # Update fields
        if 'name' in data:
            scheduled_report.name = data['name']
        if 'description' in data:
            scheduled_report.description = data['description']
        if 'report_profile_id' in data:
            try:
                report_profile = ReportProfile.objects.get(id=data['report_profile_id'])
                scheduled_report.report_profile = report_profile
            except ReportProfile.DoesNotExist:
                return JsonResponse({
                    'status': 'error',
                    'message': _('Selected report profile not found')
                })
        if 'report_format' in data:
            scheduled_report.report_format = data['report_format']
        if 'report_language' in data:
            scheduled_report.report_language = data['report_language']
        if 'frequency' in data:
            scheduled_report.frequency = data['frequency']
        if 'start_date' in data:
            scheduled_report.start_date = datetime.strptime(data['start_date'], '%Y-%m-%d').date()
        if 'end_date' in data:
            scheduled_report.end_date = datetime.strptime(data['end_date'], '%Y-%m-%d').date() if data['end_date'] else None
        if 'execution_time' in data:
            scheduled_report.execution_time = datetime.strptime(data['execution_time'], '%H:%M').time()
        if 'email_subject' in data:
            scheduled_report.email_subject = data['email_subject']
        if 'email_body' in data:
            scheduled_report.email_body = data['email_body']
        if 'send_email' in data:
            scheduled_report.send_email = data['send_email']
        if 'use_default_email_settings' in data:
            scheduled_report.use_default_email_settings = data['use_default_email_settings']
        if 'status' in data:
            scheduled_report.status = data['status']
        
        # Validate and set company (required field)
        if 'company_id' in data:
            if not data['company_id']:
                return JsonResponse({
                    'status': 'error',
                    'message': _('Company selection is required for scheduled reports')
                })
            try:
                company = Company.objects.get(id=data['company_id'])
                scheduled_report.company = company
            except Company.DoesNotExist:
                return JsonResponse({
                    'status': 'error',
                    'message': _('Selected company not found')
                })
        
        # Update mail server and account if specified
        if 'mail_server' in data:
            if data['mail_server'] and not data.get('use_default_email_settings', True):
                try:
                    from app_conf.models import MailServer
                    mail_server = MailServer.objects.get(id=data['mail_server'])
                    scheduled_report.mail_server = mail_server
                except MailServer.DoesNotExist:
                    return JsonResponse({
                        'status': 'error',
                        'message': _('Selected mail server not found')
                    })
            else:
                scheduled_report.mail_server = None
        
        if 'mail_account' in data:
            if data['mail_account'] and not data.get('use_default_email_settings', True):
                try:
                    from app_conf.models import MailAccount
                    mail_account = MailAccount.objects.get(id=data['mail_account'])
                    scheduled_report.mail_account = mail_account
                except MailAccount.DoesNotExist:
                    return JsonResponse({
                        'status': 'error',
                        'message': _('Selected mail account not found')
                    })
            else:
                scheduled_report.mail_account = None
        
        # Update frequency-specific fields
        if scheduled_report.frequency == 'weekly':
            scheduled_report.monday = data.get('monday', False)
            scheduled_report.tuesday = data.get('tuesday', False)
            scheduled_report.wednesday = data.get('wednesday', False)
            scheduled_report.thursday = data.get('thursday', False)
            scheduled_report.friday = data.get('friday', False)
            scheduled_report.saturday = data.get('saturday', False)
            scheduled_report.sunday = data.get('sunday', False)
        elif scheduled_report.frequency == 'monthly':
            scheduled_report.day_of_month = data.get('day_of_month', 1)
        
        # Recalculate next run
        scheduled_report.next_run = scheduled_report.calculate_next_run()
        
        scheduled_report.save()
        
        # Update email recipients
        if 'email_recipients' in data:
            recipient_ids = data['email_recipients']
            # Get CabinetUser objects from IDs
            cabinet_users = CabinetUser.objects.filter(id__in=recipient_ids)
            scheduled_report.email_recipients.set(cabinet_users)
        
        # Handle file attachments
        if request.FILES.getlist('attachments'):
            for uploaded_file in request.FILES.getlist('attachments'):
                # Create attachment record
                attachment = ScheduledReportAttachment(
                    scheduled_report=scheduled_report,
                    file=uploaded_file,
                    original_filename=uploaded_file.name,
                    file_size=uploaded_file.size,
                    uploaded_by=request.user
                )
                attachment.save()
        
        # Handle existing attachments (keep only those in the list)
        if 'existing_attachments' in data:
            try:
                existing_attachment_ids = json.loads(data['existing_attachments'])
                # Delete attachments not in the list
                scheduled_report.attachments.exclude(id__in=existing_attachment_ids).delete()
            except (json.JSONDecodeError, KeyError):
                # If no existing_attachments provided, keep all existing
                pass
        
        return JsonResponse({
            'status': 'success',
            'message': _('Scheduled report updated successfully')
        })
        
    except ScheduledReport.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': _('Scheduled report not found')
        })
    except Exception as e:
        logger.error(f"Error updating scheduled report: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error updating scheduled report: {}').format(str(e))
        })


@login_required
@user_passes_test(can_edit_risk_report)
@require_http_methods(["POST"])
def send_scheduled_report_now(request, report_id):
    """Send a scheduled report immediately"""
    try:
        # Check if user has permission to edit reports
        if not can_edit_risk_report(request.user):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to send scheduled reports')
            }, status=403)
        
        # Get the scheduled report
        scheduled_report = ScheduledReport.objects.select_related(
            'report_profile', 'company', 'created_by', 'mail_server', 'mail_account'
        ).get(id=report_id)
        
        # Check permissions
        user_companies = get_user_companies(request.user)
        if (scheduled_report.created_by != request.user and 
            scheduled_report.company not in user_companies):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to send this scheduled report')
            }, status=403)
        
        # Allow manual send regardless of current status (failed/paused/completed)
        # Manual execution should always be permitted for authorized users
        
        request_time = timezone.now()
        
        # Create execution record for manual sending
        execution = ScheduledReportExecution.objects.create(
            scheduled_report=scheduled_report,
            status='running'
        )
        
        logger.info(f"Manual execution started for scheduled report {scheduled_report.name} (ID: {scheduled_report.id}) by user {request.user.username}")
        
        # Trigger the report execution task
        from .tasks import execute_scheduled_report
        task = execute_scheduled_report.delay(str(scheduled_report.id), manual_execution=True)
        
        # For production, we don't wait for the result synchronously
        # The task will execute asynchronously and update the execution record
        logger.info(f"Manual execution task started for scheduled report {scheduled_report.name} (task_id: {task.id})")
        
        return JsonResponse({
            'status': 'success',
            'message': _('Report execution started successfully. You will receive a notification when it completes.'),
            'task_id': str(task.id)
        })
        
    except ScheduledReport.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': _('Scheduled report not found')
        }, status=404)
    except Exception as e:
        logger.error(f"Error sending scheduled report: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error sending scheduled report: {}').format(str(e))
        }, status=500)





@login_required
@user_passes_test(can_delete_risk_report)
@require_http_methods(["POST"])
def delete_scheduled_report(request, report_id):
    """Delete a scheduled report"""
    try:
        # Check if user has permission to delete reports
        if not can_delete_risk_report(request.user):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to delete scheduled reports')
            }, status=403)
        # Get the scheduled report
        scheduled_report = ScheduledReport.objects.get(id=report_id)
        
        # Check permissions
        user_companies = get_user_companies(request.user)
        if (scheduled_report.created_by != request.user and 
            scheduled_report.company not in user_companies):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to delete this scheduled report')
            })
        
        report_name = scheduled_report.name
        scheduled_report.delete()
        
        return JsonResponse({
            'status': 'success',
            'message': _('Scheduled report "{}" deleted successfully').format(report_name)
        })
        
    except ScheduledReport.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': _('Scheduled report not found')
        })
    except Exception as e:
        logger.error(f"Error deleting scheduled report: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error deleting scheduled report: {}').format(str(e))
        })


@login_required
@user_passes_test(check_risk_report_access)
def get_scheduled_report_details(request, report_id):
    """Get details of a specific scheduled report"""
    try:
        # Check if user has access to risk reports
        if not check_risk_report_access(request.user, 'view'):
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to access scheduled reports')
            }, status=403)
        # Get the scheduled report
        scheduled_report = ScheduledReport.objects.select_related(
            'report_profile', 'company', 'created_by', 'mail_server', 'mail_account'
        ).get(id=report_id)
        
        # Check permissions: owner, same company, or email recipient
        user_companies = get_user_companies(request.user)
        user_has_access = False
        
        # Check if user is the owner
        if scheduled_report.created_by == request.user:
            user_has_access = True
        # Check if user belongs to the same company
        elif scheduled_report.company and scheduled_report.company in user_companies:
            user_has_access = True
        # Check if user is in the email recipients list
        elif scheduled_report.email_recipients.filter(user=request.user).exists():
            user_has_access = True
        
        if not user_has_access:
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to view this scheduled report')
            })
        
        # Get email recipients with department and position info
        recipients = []
        for recipient in scheduled_report.email_recipients.all():
            # Check if recipient is CabinetUser or User
            if hasattr(recipient, 'user'):
                # recipient is CabinetUser
                cabinet_user = recipient
                user = recipient.user
                recipient_data = {
                    'id': cabinet_user.id,  # Use CabinetUser ID for JavaScript
                    'name': user.get_full_name() or user.username,
                    'email': user.email
                }
                
                # Add department and position info
                recipient_data['department'] = cabinet_user.department.get_name() if cabinet_user.department else ''
                recipient_data['position'] = cabinet_user.position.get_name() if cabinet_user.position else ''
                recipient_data['company'] = cabinet_user.company.name if cabinet_user.company else ''
            else:
                # recipient is User, try to get associated CabinetUser
                user = recipient
                cabinet_user = None
                try:
                    cabinet_user = user.cabinet
                except:
                    pass
                    
                recipient_data = {
                    'id': cabinet_user.id if cabinet_user else user.id,  # Use CabinetUser ID for JavaScript
                    'name': user.get_full_name() or user.username,
                    'email': user.email
                }
                
                # Add department and position if available
                if cabinet_user:
                    recipient_data['department'] = cabinet_user.department.get_name() if cabinet_user.department else ''
                    recipient_data['position'] = cabinet_user.position.get_name() if cabinet_user.position else ''
                    recipient_data['company'] = cabinet_user.company.name if cabinet_user.company else ''
            
            recipients.append(recipient_data)
        
        # Get recent executions
        recent_executions = []
        for execution in scheduled_report.executions.all()[:10]:  # Last 10 executions
            recent_executions.append({
                'id': str(execution.id),
                'status': execution.get_status_display(),
                'started_at': execution.started_at.astimezone().strftime('%Y-%m-%d %H:%M'),
                'completed_at': execution.completed_at.astimezone().strftime('%Y-%m-%d %H:%M') if execution.completed_at else '',
                'file_size': execution.file_size,
                'email_sent': execution.email_sent,
                'email_recipients_count': execution.email_recipients_count,
                'error_message': execution.error_message
            })
        
        # Get attachments
        attachments = []
        for attachment in scheduled_report.attachments.all():
            attachments.append({
                'id': str(attachment.id),
                'original_filename': attachment.original_filename,
                'file_size': attachment.file_size,
                'file_url': attachment.get_file_url(),
                'uploaded_at': attachment.uploaded_at.astimezone().strftime('%Y-%m-%d %H:%M'),
                'uploaded_by': attachment.uploaded_by.get_full_name() if attachment.uploaded_by else ''
            })
        
        data = {
            'id': str(scheduled_report.id),
            'name': scheduled_report.name,
            'description': scheduled_report.description,
            'report_type': scheduled_report.report_type,
            'report_profile_id': str(scheduled_report.report_profile.id) if scheduled_report.report_profile else None,
            'report_profile_name': scheduled_report.report_profile.name if scheduled_report.report_profile else None,
            'report_format': scheduled_report.report_format,
            'report_language': scheduled_report.report_language,
            'frequency': scheduled_report.frequency,
            'start_date': scheduled_report.start_date.strftime('%Y-%m-%d'),
            'end_date': scheduled_report.end_date.strftime('%Y-%m-%d') if scheduled_report.end_date else '',
            'execution_time': scheduled_report.execution_time.strftime('%H:%M'),
            'status': scheduled_report.status,
            'company_id': scheduled_report.company.id if scheduled_report.company else '',
            'company_name': scheduled_report.company.name if scheduled_report.company else '',
            'email_subject': scheduled_report.email_subject,
            'email_body': scheduled_report.email_body,
            'send_email': scheduled_report.send_email,
            'use_default_email_settings': scheduled_report.use_default_email_settings,
            'mail_server': scheduled_report.mail_server.id if scheduled_report.mail_server else '',
            'mail_server_name': scheduled_report.mail_server.name if scheduled_report.mail_server else '',
            'mail_account': scheduled_report.mail_account.id if scheduled_report.mail_account else '',
            'mail_account_name': f"{scheduled_report.mail_account.username} ({scheduled_report.mail_account.server.name})" if scheduled_report.mail_account else '',
            'email_recipients': recipients,
            'last_run': scheduled_report.last_run.astimezone().strftime('%Y-%m-%d %H:%M') if scheduled_report.last_run else '',
            'next_run': scheduled_report.next_run.astimezone().strftime('%Y-%m-%d %H:%M') if scheduled_report.next_run else '',
            'run_count': scheduled_report.run_count,
            'recent_executions': recent_executions,
            # Weekly fields
            'monday': scheduled_report.monday,
            'tuesday': scheduled_report.tuesday,
            'wednesday': scheduled_report.wednesday,
            'thursday': scheduled_report.thursday,
            'friday': scheduled_report.friday,
            'saturday': scheduled_report.saturday,
            'sunday': scheduled_report.sunday,
            # Monthly fields
            'day_of_month': scheduled_report.day_of_month,
            # Attachments
            'attachments': attachments,
        }
        
        return JsonResponse({
            'status': 'success',
            'data': data
        })
        
    except ScheduledReport.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': _('Scheduled report not found')
        })
    except Exception as e:
        logger.error(f"Error getting scheduled report details: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error loading scheduled report details: {}').format(str(e))
        })


@login_required
@user_passes_test(check_risk_report_access)
def get_scheduled_report_history(request, report_id):
    """Get execution history for a specific scheduled report"""
    try:
        # Get the scheduled report
        scheduled_report = ScheduledReport.objects.get(id=report_id)
        
        # Check permissions: owner, same company, or email recipient
        user_companies = get_user_companies(request.user)
        user_has_access = False
        
        # Check if user is the owner
        if scheduled_report.created_by == request.user:
            user_has_access = True
        # Check if user belongs to the same company
        elif scheduled_report.company and scheduled_report.company in user_companies:
            user_has_access = True
        # Check if user is in the email recipients list
        elif scheduled_report.email_recipients.filter(user=request.user).exists():
            user_has_access = True
        
        if not user_has_access:
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to view this scheduled report')
            })
        
        # Get pagination parameters
        page = int(request.GET.get('page', 1))
        per_page = int(request.GET.get('per_page', 20))
        
        # Get all executions for this report
        all_executions = scheduled_report.executions.all()
        total_count = all_executions.count()
        
        # Calculate pagination
        start_index = (page - 1) * per_page
        end_index = start_index + per_page
        executions = all_executions[start_index:end_index]
        
        # Format execution data
        execution_history = []
        for execution in executions:
            import os
            
            # Get file info
            file_exists = bool(execution.file_path and os.path.exists(execution.file_path))
            file_name = os.path.basename(execution.file_path) if execution.file_path else ''
            file_size_mb = round(execution.file_size / (1024 * 1024), 2) if execution.file_size else 0
            
            # Calculate duration
            duration = ''
            if execution.completed_at and execution.started_at:
                duration_seconds = (execution.completed_at - execution.started_at).total_seconds()
                if duration_seconds < 60:
                    duration = f"{int(duration_seconds)}s"
                elif duration_seconds < 3600:
                    duration = f"{int(duration_seconds // 60)}m {int(duration_seconds % 60)}s"
                else:
                    hours = int(duration_seconds // 3600)
                    minutes = int((duration_seconds % 3600) // 60)
                    duration = f"{hours}h {minutes}m"
            
            # Build snapshot link if snapshot_html exists
            try:
                from django.urls import reverse
                snapshot_path = reverse('view_scheduled_report_snapshot', args=[str(execution.id)])
            except Exception:
                snapshot_path = ''

            execution_data = {
                'id': str(execution.id),
                'status': execution.status,
                'status_display': execution.get_status_display(),
                'started_at': execution.started_at.astimezone().strftime('%Y-%m-%d %H:%M:%S'),
                'completed_at': execution.completed_at.astimezone().strftime('%Y-%m-%d %H:%M:%S') if execution.completed_at else '',
                'duration': duration,
                'file_path': execution.file_path,
                'file_name': file_name,
                'file_size': execution.file_size,
                'file_size_mb': file_size_mb,
                'file_exists': file_exists,
                'snapshot_available': bool(execution.snapshot_html),
                'snapshot_path': snapshot_path,
                'email_sent': execution.email_sent,
                'email_recipients_count': execution.email_recipients_count,
                'email_error': execution.email_error,
                'error_message': execution.error_message,
                'report_type': scheduled_report.report_type,
                'report_format': scheduled_report.report_format,
                'report_language': scheduled_report.report_language,
                'company_name': scheduled_report.company.name if scheduled_report.company else _('All Companies')
            }
            
            execution_history.append(execution_data)
        
        # Calculate pagination info
        total_pages = (total_count + per_page - 1) // per_page
        has_next = page < total_pages
        has_previous = page > 1
        
        data = {
            'schedule_id': str(scheduled_report.id),
            'schedule_name': scheduled_report.name,
            'executions': execution_history,
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total_count': total_count,
                'total_pages': total_pages,
                'has_next': has_next,
                'has_previous': has_previous,
                'start_index': start_index + 1,
                'end_index': min(end_index, total_count)
            }
        }
        
        return JsonResponse({
            'status': 'success',
            'data': data
        })
        
    except ScheduledReport.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': _('Scheduled report not found')
        })
    except Exception as e:
        logger.error(f"Error getting scheduled report history: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error loading scheduled report history: {}').format(str(e))
        })


@login_required
@user_passes_test(check_risk_report_access)
def download_scheduled_report_file(request, execution_id):
    """Download a specific scheduled report file"""
    from django.utils.translation import gettext as _
    try:
        # Get the execution
        execution = ScheduledReportExecution.objects.get(id=execution_id)
        report = execution.scheduled_report
        
        # Check permissions: owner, same company, or email recipient
        user_companies = get_user_companies(request.user)
        user_has_access = False
        
        # Check if user is the owner
        if report.created_by == request.user:
            user_has_access = True
        # Check if user belongs to the same company
        elif report.company and report.company in user_companies:
            user_has_access = True
        # Check if user is in the email recipients list
        elif report.email_recipients.filter(user=request.user).exists():
            user_has_access = True
        
        if not user_has_access:
            return JsonResponse({
                'status': 'error',
                'message': _('You do not have permission to download this report')
            })
        
        # Check if file exists
        if not execution.file_path or not os.path.exists(execution.file_path):
            return JsonResponse({
                'status': 'error',
                'message': _('Report file not found or has been deleted')
            })
        
        # Serve the file
        import mimetypes
        from django.http import FileResponse
        
        file_path = execution.file_path
        file_name = os.path.basename(file_path)
        
        # Determine content type
        content_type, _ = mimetypes.guess_type(file_path)
        if not content_type:
            content_type = 'application/octet-stream'
        
        # Create response
        response = FileResponse(
            open(file_path, 'rb'),
            content_type=content_type,
            as_attachment=True,
            filename=file_name
        )
        
        return response
        
    except ScheduledReportExecution.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': _('Report execution not found')
        })
    except Exception as e:
        logger.error(f"Error downloading scheduled report file: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error downloading report file: {}').format(str(e))
        })


@login_required
@user_passes_test(check_risk_report_access)
def get_cabinet_users(request):
    """Get list of cabinet users for email recipients selection"""
    try:
        # Get company filter from request
        company_id = request.GET.get('company_id')
        
        # Get users from the same companies as the current user
        user_companies = get_user_companies(request.user)
        
        if company_id:
            # Filter by specific company if provided
            try:
                selected_company = Company.objects.get(id=company_id)
                if selected_company in user_companies or request.user.is_superuser:
                    cabinet_users = CabinetUser.objects.filter(
                        company=selected_company,
                        user__is_active=True
                    ).select_related('company', 'user', 'department', 'position')
                else:
                    # User doesn't have access to this company
                    cabinet_users = CabinetUser.objects.none()
            except Company.DoesNotExist:
                cabinet_users = CabinetUser.objects.none()
        else:
            # No specific company selected, show users from all accessible companies
            if user_companies:
                cabinet_users = CabinetUser.objects.filter(
                    company__in=user_companies,
                    user__is_active=True
                ).select_related('company', 'user', 'department', 'position')
            else:
                # If no companies, get all active users (admin case)
                cabinet_users = CabinetUser.objects.filter(user__is_active=True).select_related('company', 'user', 'department', 'position')
        
        users_data = []
        for cabinet_user in cabinet_users:
            user = cabinet_user.user
            
            # Get department info
            department_name = ''
            if cabinet_user.department:
                department_name = cabinet_user.department.get_name()
            
            # Get position info  
            position_name = ''
            if cabinet_user.position:
                position_name = cabinet_user.position.get_name()
            
            users_data.append({
                'id': cabinet_user.id,
                'name': user.get_full_name() or user.username,
                'email': user.email,
                'company': cabinet_user.company.name if cabinet_user.company else _('No Company'),
                'company_id': cabinet_user.company.id if cabinet_user.company else None,
                'department': department_name,
                'position': position_name
            })
        
        return JsonResponse({
            'status': 'success',
            'data': users_data
        })
        
    except Exception as e:
        logger.error(f"Error getting cabinet users: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error getting cabinet users: {}').format(str(e))
        })


@login_required
@user_passes_test(check_risk_report_access)
@require_http_methods(["POST"])
def send_profile_link(request):
    """Send Risk Assessment Report link for a profile to a specific email using global email settings."""
    try:
        profile_id = request.POST.get('profile_id')
        recipient_email = request.POST.get('email')
        if not profile_id or not recipient_email:
            return JsonResponse({'status': 'error', 'message': _('Profile ID and email are required')}, status=400)

        # Validate profile
        try:
            profile = ReportProfile.objects.get(id=profile_id)
        except ReportProfile.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': _('Report profile not found')}, status=404)

        if not profile.can_be_used_by(request.user):
            raise PermissionDenied(_("You don't have permission to use this profile"))

        # Build preview HTML (fresh) using existing helpers
        sections_config = profile.get_sections_config()
        company_id = sections_config.get('_company_id', '')

        # Prepare params and generate unified content
        params = {
            'reportType': 'profile',
            'format': 'html',
            'language': profile.default_language,
            'company_id': str(company_id) if company_id else '',
            'selectedSections': sections_config,
            'profile': {
                'name': profile.name,
                'description': profile.description,
                'created_by': profile.created_by.get_full_name() or profile.created_by.username,
                'created_at': profile.created_at,
                'company': _('All Companies'),
            }
        }
        # If company_id exists, resolve name
        if company_id:
            try:
                from app_conf.models import Company
                company = Company.objects.get(id=company_id)
                params['profile']['company'] = company.name
            except Exception:
                pass

        report_data = generate_report_data(request.user, params)
        preview_html = generate_html_preview(report_data, 'profile', profile.default_language, sections_config)

        # Save a transient ScheduledReportExecution-like snapshot for linking
        execution = ScheduledReportExecution.objects.create(
            scheduled_report=ScheduledReport.objects.create(
                name=f"Ad-hoc: {profile.name}",
                description=_('Ad-hoc email from profile'),
                report_profile=profile,
                report_type='full',
                report_format='html',
                report_language=profile.default_language,
                frequency='once',
                start_date=timezone.now().date(),
                execution_time=timezone.now().time(),
                company=None,
                created_by=request.user,
                status='active'
            ),
            status='completed',
            snapshot_html=preview_html,
            snapshot_language=profile.default_language,
            snapshot_created_at=timezone.now(),
            email_sent=False
        )

        # Build absolute URL to snapshot
        def _build_absolute_url(path: str) -> str:
            # Try to get Site Domain from SiteSettings (primary source)
            try:
                from app_conf.models import SiteSettings
                site_settings = SiteSettings.get_settings()
                if site_settings and site_settings.site_domain:
                    # Use get_site_url() method which combines protocol + domain
                    base = site_settings.get_site_url().rstrip('/')
                    logger.info(f"[send_profile_link] Using Site Domain from SiteSettings: {base}")
                    return f"{base}{path}"
            except Exception as e:
                logger.warning(f"[send_profile_link] Could not load SiteSettings: {e}")
            
            # Fallback to PUBLIC_BASE_URL
            base = getattr(__import__('django.conf').conf.settings, 'PUBLIC_BASE_URL', '').rstrip('/')
            if base:
                logger.info(f"[send_profile_link] Using PUBLIC_BASE_URL: {base}")
                return f"{base}{path}"
            
            # Fallback to ALLOWED_HOSTS
            scheme = 'https' if getattr(__import__('django.conf').conf.settings, 'PRODUCTION', False) else 'http'
            host = 'localhost:8000'
            from django.conf import settings as dj_settings
            try:
                if dj_settings.ALLOWED_HOSTS:
                    host = next((h for h in dj_settings.ALLOWED_HOSTS if h not in ('testserver',)), dj_settings.ALLOWED_HOSTS[0])
                    if ':' not in host and host in ('localhost', '127.0.0.1'):
                        host = f"{host}:8000"
            except Exception:
                pass
            fallback_url = f"{scheme}://{host}"
            logger.info(f"[send_profile_link] Using ALLOWED_HOSTS fallback: {fallback_url}")
            return f"{fallback_url}{path}"

        from django.urls import reverse
        snapshot_path = reverse('view_scheduled_report_snapshot', args=[str(execution.id)])
        link_url = _build_absolute_url(snapshot_path)

        # Send email using global settings (direct SMTP if configured)
        cfg = RiskReportEmailConfig.objects.first()
        subject = f"{cfg.default_subject if cfg and cfg.default_subject else 'Risk Assessment Report'}"
        body = (cfg.default_body if cfg and cfg.default_body else _('Please view your report at the link below.')) + f"\n\n{link_url}"

        try:
            sent = False
            if cfg and cfg.mail_account:
                import smtplib, ssl
                from email.mime.multipart import MIMEMultipart
                from email.mime.text import MIMEText
                smtp_host = cfg.mail_account.server.smtp_host
                smtp_port = cfg.mail_account.server.smtp_port
                use_ssl = getattr(cfg.mail_account.server, 'use_ssl', False)
                use_tls = getattr(cfg.mail_account.server, 'use_tls', False)
                username = cfg.mail_account.username
                password = cfg.mail_account.password

                msg = MIMEMultipart()
                msg['From'] = username
                msg['To'] = recipient_email
                msg['Subject'] = subject
                msg.attach(MIMEText(body, 'plain'))

                if use_ssl:
                    context = ssl.create_default_context()
                    context.check_hostname = False
                    context.verify_mode = ssl.CERT_NONE
                    server = smtplib.SMTP_SSL(host=smtp_host, port=smtp_port, context=context)
                else:
                    server = smtplib.SMTP(host=smtp_host, port=smtp_port)
                    if use_tls:
                        server.starttls()
                server.login(username, password)
                server.send_message(msg)
                server.quit()
                sent = True
            else:
                from django.core.mail import EmailMessage
                from_email = cfg.mail_account.username if (cfg and cfg.mail_account) else None
                EmailMessage(subject=subject, body=body, from_email=from_email, to=[recipient_email]).send(fail_silently=False)
                sent = True

            if sent:
                return JsonResponse({'status': 'success', 'message': _('Email sent successfully')})
            return JsonResponse({'status': 'error', 'message': _('Email was not sent')}, status=500)

        except Exception as e:
            logger.error(f"Error sending profile link: {e}", exc_info=True)
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

    except Exception as e:
        logger.error(f"send_profile_link error: {e}", exc_info=True)
        return JsonResponse({'status': 'error', 'message': _('Unexpected error: {}').format(str(e))}, status=500)

@login_required
@user_passes_test(check_risk_report_access)
@require_http_methods(["POST"])
def generate_risk_report_from_scheduled_execution(request):
    """Generate risk report from a scheduled report execution snapshot"""
    
    # Check permissions
    if not check_risk_report_access(request.user, 'view'):
        raise PermissionDenied(_("You don't have permission to generate reports"))
    
    try:
        execution_id = request.POST.get('execution_id')
        logger.info(f"generate_risk_report_from_scheduled_execution: received execution_id={execution_id}, POST data={dict(request.POST)}")
        if not execution_id:
            return JsonResponse({
                'status': 'error',
                'message': _('Execution ID is required')
            })
        
        # Get the execution
        try:
            logger.info(f"generate_risk_report_from_scheduled_execution: Looking for execution with ID: {execution_id}")
            execution = ScheduledReportExecution.objects.get(id=execution_id)
            logger.info(f"generate_risk_report_from_scheduled_execution: Found execution: {execution.id}")
        except ScheduledReportExecution.DoesNotExist:
            logger.error(f"generate_risk_report_from_scheduled_execution: Execution not found with ID: {execution_id}")
            return JsonResponse({
                'status': 'error',
                'message': _('Scheduled report execution not found')
            })
        
        # Check permissions: owner, same company, or email recipient
        report = execution.scheduled_report
        user_companies = get_user_companies(request.user)
        user_has_access = False
        
        # Check if user is the owner
        if report.created_by == request.user:
            user_has_access = True
        # Check if user belongs to the same company
        elif report.company and report.company in user_companies:
            user_has_access = True
        # Check if user is in the email recipients list
        elif report.email_recipients.filter(user=request.user).exists():
            user_has_access = True
        
        if not user_has_access:
            raise PermissionDenied(_("You don't have permission to use this report"))
        
        # Optional overrides from request
        override_format = request.POST.get('format') or 'pdf'
        override_language = request.POST.get('language') or execution.snapshot_language or 'uk'
        
        # Activate the language
        from django.utils.translation import activate
        activate(override_language)
        
        # Set language for the current request
        request.LANGUAGE_CODE = override_language
        
        # Validate format availability
        if not is_format_available(override_format):
            return JsonResponse({
                'status': 'error',
                'message': _('Selected format is not available. Please install required dependencies.')
            })
        
        # Get the snapshot HTML
        snapshot_html = execution.snapshot_html
        if not snapshot_html:
            return JsonResponse({
                'status': 'error',
                'message': _('Report snapshot is not available')
            })
        
        # Generate report file using the same approach as Report Profiles
        logger.info(f"generate_risk_report_from_scheduled_execution: Format requested: {override_format}")
        
        # Get the original scheduled report configuration
        scheduled_report = execution.scheduled_report
        
        # Prepare data from scheduled report (similar to profile approach)
        # Get sections_config from the associated report profile
        sections_config = {}
        company_id = ''
        if scheduled_report.report_profile:
            sections_config = scheduled_report.report_profile.get_sections_config()
            company_id = sections_config.get('_company_id', '')
        else:
            # Fallback: use basic configuration if no profile is associated
            sections_config = {
                'statistics': True,
                'risk_distribution': True,
                'compliance_overview': True,
                'top_risks': True,
                'asset_details': True,
                'vulnerability_details': True,
                'treatment_details': True,
                'recommendations': True
            }
        
        # Prepare parameters in the format expected by generate_report_data
        # Use company from scheduled report if available, otherwise from profile
        report_company_id = ''
        if scheduled_report.company:
            report_company_id = str(scheduled_report.company.id)
        elif company_id:
            report_company_id = str(company_id)
        
        params = {
            'reportType': 'scheduled',  # Use scheduled type
            'format': override_format,
            'language': override_language,
            'company_id': report_company_id,
            'selectedSections': sections_config,
            'scheduled_report': scheduled_report
        }
        
        # Add scheduled report information to params
        company_name = _('All Companies')
        if scheduled_report.company:
            company_name = scheduled_report.company.name
        elif company_id:
            try:
                from app_cabinet.models import Company
                company = Company.objects.get(id=company_id)
                company_name = company.name
            except:
                pass
        
        params['profile'] = {
            'name': scheduled_report.name,
            'description': scheduled_report.description or '',
            'created_by': scheduled_report.created_by.get_full_name(),
            'created_at': scheduled_report.created_at,
            'company': company_name,
        }
        
        # Generate report data using the same method as profiles
        report_data = generate_report_data(request.user, params)
        
        # Debug logging
        logger.info(f"Scheduled report info in params: {params.get('profile', {})}")
        logger.info(f"Scheduled report info in report_data: {report_data.get('profile', {})}")
        
        # Generate report file using the same method as profiles
        response = generate_report_file(report_data, 'scheduled', override_format)
        
        return response
        
    except Exception as e:
        logger.error(f"Error generating report from scheduled execution: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error generating report: {}').format(str(e))
        }, status=500)


@login_required
@user_passes_test(check_risk_report_access)
@require_http_methods(["POST"])
def generate_risk_report_from_profile(request):
    """Generate risk report using a saved profile"""
    
    # Check permissions
    if not check_risk_report_access(request.user, 'view'):
        raise PermissionDenied(_("You don't have permission to generate reports"))
    
    try:
        profile_id = request.POST.get('profile_id')
        logger.info(f"generate_risk_report_from_profile: received profile_id={profile_id}, POST data={dict(request.POST)}")
        if not profile_id:
            return JsonResponse({
                'status': 'error',
                'message': _('Profile ID is required')
            })
        
        # Get the profile
        try:
            profile = ReportProfile.objects.get(id=profile_id)
        except ReportProfile.DoesNotExist:
            return JsonResponse({
                'status': 'error',
                'message': _('Report profile not found')
            })
        
        # Check if user has access to this profile
        if not profile.can_be_used_by(request.user):
            raise PermissionDenied(_("You don't have permission to use this profile"))
        
        # Optional overrides from request
        override_format = request.POST.get('format') or profile.default_format
        override_language = request.POST.get('language') or profile.default_language
        
        # Activate the profile's default language for proper multilingual support
        from django.utils.translation import activate
        activate(override_language)
        
        # Set language for the current request
        request.LANGUAGE_CODE = override_language
        
        # Validate format availability
        if not is_format_available(override_format):
            return JsonResponse({
                'status': 'error',
                'message': _('Selected format is not available. Please install required dependencies.')
            })
        
        # Prepare data from profile
        # Get company_id from sections_config if available, otherwise empty
        sections_config = profile.get_sections_config()
        company_id = sections_config.get('_company_id', '')
        
        # Prepare parameters in the format expected by generate_report_data
        params = {
            'reportType': 'profile',  # Use profile type
            'format': override_format,
            'language': override_language,
            'company_id': str(company_id) if company_id else '',
            'selectedSections': sections_config,
            'profile': profile
        }
        
        # Add profile information to params before generating report data
        # Get company name from sections_config if available
        company_name = _('All Companies')
        if company_id:
            try:
                from app_cabinet.models import Company
                company = Company.objects.get(id=company_id)
                company_name = company.name
            except:
                pass
        
        params['profile'] = {
            'name': profile.name,
            'description': profile.description,
            'created_by': profile.created_by.get_full_name(),
            'created_at': profile.created_at,
            'company': company_name,
        }
        
        # Generate report data
        report_data = generate_report_data(request.user, params)
        
        # Debug logging
        logger.info(f"Profile info in params: {params.get('profile', {})}")
        logger.info(f"Profile info in report_data: {report_data.get('profile', {})}")
        
        # Generate report file
        response = generate_report_file(report_data, 'profile', override_format)
        
        # Update profile usage statistics
        profile.usage_count += 1
        profile.last_used_at = timezone.now()
        profile.save()
        
        return response
        
    except Exception as e:
        logger.error(f"Error generating report from profile: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error generating report: {}').format(str(e))
        })


@login_required
@user_passes_test(check_risk_report_access)
@require_http_methods(["POST"])
def preview_risk_report_from_profile(request):
    """Generate HTML preview for a report profile (always up-to-date)."""
    if not check_risk_report_access(request.user, 'view'):
        raise PermissionDenied(_("You don't have permission to preview reports"))
    try:
        profile_id = request.POST.get('profile_id')
        if not profile_id:
            return JsonResponse({'status': 'error', 'message': _('Profile ID is required')})

        try:
            profile = ReportProfile.objects.get(id=profile_id)
        except ReportProfile.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': _('Report profile not found')})

        if not profile.can_be_used_by(request.user):
            raise PermissionDenied(_("You don't have permission to use this profile"))

        # Activate language
        from django.utils.translation import activate
        activate(profile.default_language)
        request.LANGUAGE_CODE = profile.default_language

        # Prepare sections and company
        sections_config = profile.get_sections_config()
        company_id = sections_config.get('_company_id', '')

        # Build params for data generation
        params = {
            'reportType': 'profile',
            'format': 'html',
            'language': profile.default_language,
            'company_id': str(company_id) if company_id else '',
            'selectedSections': sections_config,
            'profile': profile
        }

        # Enrich profile info
        company_name = _('All Companies')
        if company_id:
            try:
                from app_cabinet.models import Company
                company = Company.objects.get(id=company_id)
                company_name = company.name
            except Exception:
                pass

        params['profile'] = {
            'name': profile.name,
            'description': profile.description,
            'created_by': profile.created_by.get_full_name(),
            'created_at': profile.created_at,
            'company': company_name,
        }

        # Generate fresh data and preview HTML
        report_data = generate_report_data(request.user, params)
        preview_html = generate_html_preview(report_data, 'profile', profile.default_language, sections_config)
        return HttpResponse(preview_html)

    except Exception as e:
        logger.error(f"Error previewing report from profile: {str(e)}", exc_info=True)
        error_html = f'<div class="alert alert-danger"><i class="fas fa-exclamation-triangle me-2"></i>{_("Error generating preview")}: {str(e)}</div>'
        return HttpResponse(error_html)


def generate_audit_logs_data(assets, start_date, end_date):
    """Generate audit logs data for reporting"""
    try:
        # Get audit logs for the specified date range
        audit_logs = RiskAssessmentAuditLog.objects.filter(
            asset__in=assets,
            timestamp__date__range=[start_date, end_date]
        ).select_related('user', 'asset')[:50]  # Limit to 50 most recent
        
        return [{
            'timestamp': log.timestamp.strftime('%Y-%m-%d %H:%M:%S'),
            'user': log.user.get_full_name() or log.user.username if log.user else 'System',
            'action': log.action,
            'details': log.details or '',
            'asset': log.asset.name if log.asset else ''
        } for log in audit_logs]
    except Exception:
        return []


def generate_governance_data(assets, companies):
    """Generate governance data for reporting"""
    try:
        # Get users associated with the companies
        from app_cabinet.models import CabinetUser
        users = CabinetUser.objects.filter(company__in=companies)
        
        governance_data = []
        for user in users[:20]:  # Limit to 20
            governance_data.append({
                'role': user.position or 'User',
                'responsibility': f"Asset management for {user.company.name if user.company else 'N/A'}",
                'level': 'Operational',
                'status': 'Active' if user.is_active else 'Inactive'
            })
        
        return governance_data
    except Exception:
        return []


def generate_trend_data(assets, vulnerabilities, start_date, end_date):
    """Generate trend analysis data for reporting"""
    try:
        from datetime import timedelta
        import calendar
        
        # Generate monthly trend data
        trend_data = []
        current_date = start_date
        
        while current_date <= end_date:
            month_start = current_date.replace(day=1)
            month_end = (month_start + timedelta(days=32)).replace(day=1) - timedelta(days=1)
            
            # Count vulnerabilities for this month
            month_vulns = vulnerabilities.filter(
                modified_at__date__range=[month_start, min(month_end, end_date)]
            ).count()
            
            trend_data.append({
                'period': f"{calendar.month_name[current_date.month]} {current_date.year}",
                'metric': 'Vulnerabilities',
                'value': str(month_vulns),
                'direction': 'stable'  # Could be calculated based on previous month
            })
            
            # Move to next month
            if current_date.month == 12:
                current_date = current_date.replace(year=current_date.year + 1, month=1)
            else:
                current_date = current_date.replace(month=current_date.month + 1)
        
        return trend_data[:12]  # Limit to 12 months
    except Exception:
        return []


def generate_dependency_data(assets, vulnerabilities):
    """Generate dependency analysis data for reporting"""
    try:
        dependency_data = []
        
        # Simple dependency analysis based on asset relationships
        for vuln in vulnerabilities.select_related('asset', 'vulnerability')[:15]:
            if vuln.asset and vuln.vulnerability:
                dependency_data.append({
                    'source': vuln.asset.name,
                    'target': vuln.vulnerability.get_name() or vuln.vulnerability.get_name('en'),
                    'type': 'Vulnerability',
                    'impact': 'High' if vuln.status == 'Yes' else 'Low'
                })
        
        return dependency_data
    except Exception:
        return []


def generate_threat_data(vulnerabilities):
    """Generate threat analysis data for reporting"""
    try:
        from .models import Threat
        
        # Get unique threats from vulnerabilities
        threats = Threat.objects.filter(
            vulnerability__in=vulnerabilities.values_list('vulnerability', flat=True)
        ).distinct()[:15]
        
        threat_data = []
        for threat in threats:
            threat_data.append({
                'name': threat.threat_uk or threat.threat_en or 'Unknown',
                'type': 'Security Threat',
                'probability': 'Medium',  # Could be calculated
                'source': 'Internal/External'  # Could be determined from threat data
            })
        
        return threat_data
    except Exception:
        return []


def generate_pdf_from_html(html_content, execution):
    """Generate PDF from HTML content using ReportLab"""
    try:
        if not REPORTLAB_AVAILABLE:
            return JsonResponse({
                'status': 'error',
                'message': _('PDF generation is not available. Please install reportlab.')
            })
        
        # Log HTML content for debugging
        logger.info(f"generate_pdf_from_html: HTML content length: {len(html_content)}")
        logger.info(f"generate_pdf_from_html: First 500 chars: {html_content[:500]}")
        logger.info(f"generate_pdf_from_html: Contains 'black' or 'rectangle': {'black' in html_content.lower() or 'rectangle' in html_content.lower()}")
        
        # Create PDF using ReportLab
        from io import BytesIO
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.fonts import addMapping
        
        # Create buffer for PDF
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(A4))
        story = []
        
        # Register fonts for Cyrillic support
        try:
            # Try to register a font that supports Cyrillic
            pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
            cyrillic_font = 'DejaVuSans'
        except:
            try:
                # Fallback to Arial Unicode MS if available
                pdfmetrics.registerFont(TTFont('ArialUnicodeMS', 'ARIALUNI.TTF'))
                cyrillic_font = 'ArialUnicodeMS'
            except:
                # Use default font if no Cyrillic fonts available
                cyrillic_font = 'Helvetica'
        
        # Get styles and customize for Cyrillic
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=cyrillic_font,
            fontSize=16,
            spaceAfter=12
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=cyrillic_font,
            fontSize=10,
            spaceAfter=6
        )
        
        # Add title
        title = Paragraph(f"Risk Assessment Report - {execution.scheduled_report.name}", title_style)
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Add metadata
        story.append(Paragraph(f"Generated: {execution.started_at.strftime('%Y-%m-%d %H:%M')}", normal_style))
        story.append(Paragraph(f"Report: {execution.scheduled_report.name}", normal_style))
        story.append(Paragraph(f"Company: {execution.scheduled_report.company.name if execution.scheduled_report.company else 'All Companies'}", normal_style))
        story.append(Spacer(1, 12))
        
        # Extract and process HTML content
        if BEAUTIFULSOUP_AVAILABLE:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Clean up the HTML content
            html_content = str(soup)
            
            # Log processed content
            logger.info(f"generate_pdf_from_html: Processed HTML length: {len(html_content)}")
            
            # Process different HTML elements
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'span', 'table', 'tr', 'td', 'th']):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    # Headers
                    text = element.get_text().strip()
                    if text and not text.startswith('█'):  # Filter out black rectangles
                        if element.name == 'h1':
                            story.append(Paragraph(text, title_style))
                        elif element.name == 'h2':
                            header_style = ParagraphStyle(
                                'CustomHeading2',
                                parent=styles['Heading2'],
                                fontName=cyrillic_font,
                                fontSize=14,
                                spaceAfter=8
                            )
                            story.append(Paragraph(text, header_style))
                        else:
                            header_style = ParagraphStyle(
                                'CustomHeading3',
                                parent=styles['Heading3'],
                                fontName=cyrillic_font,
                                fontSize=12,
                                spaceAfter=6
                            )
                            story.append(Paragraph(text, header_style))
                        story.append(Spacer(1, 6))
                
                elif element.name == 'p':
                    # Paragraphs
                    text = element.get_text().strip()
                    if text and not text.startswith('█'):  # Filter out black rectangles
                        story.append(Paragraph(text, normal_style))
                        story.append(Spacer(1, 6))
                
                elif element.name == 'table':
                    # Tables - create proper table structure for PDF
                    try:
                        from reportlab.platypus import Table, TableStyle
                        from reportlab.lib import colors
                        
                        # Count rows and columns
                        rows = element.find_all('tr')
                        if rows:
                            # Get max columns
                            max_cols = 0
                            for row in rows:
                                cells = row.find_all(['td', 'th'])
                                max_cols = max(max_cols, len(cells))
                            
                            if max_cols > 0:
                                # Create table data
                                table_data = []
                                for row in rows:
                                    cells = row.find_all(['td', 'th'])
                                    row_data = []
                                    for cell in cells:
                                        cell_text = cell.get_text().strip()
                                        if cell_text and not cell_text.startswith('█'):
                                            row_data.append(cell_text)
                                        else:
                                            row_data.append('')
                                    # Pad row if needed
                                    while len(row_data) < max_cols:
                                        row_data.append('')
                                    table_data.append(row_data)
                                
                                if table_data:
                                    # Create PDF table
                                    pdf_table = Table(table_data)
                                    pdf_table.setStyle(TableStyle([
                                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                        ('FONTNAME', (0, 0), (-1, 0), cyrillic_font),
                                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                                    ]))
                                    story.append(pdf_table)
                                    story.append(Spacer(1, 12))
                                else:
                                    # If table creation failed, fall back to text
                                    table_text = element.get_text().strip()
                                    if table_text:
                                        lines = table_text.split('\n')
                                        for line in lines:
                                            if line.strip() and not line.strip().startswith('█'):
                                                story.append(Paragraph(line.strip(), normal_style))
                                        story.append(Spacer(1, 6))
                    except:
                        # Fallback to text if table creation fails
                        table_text = element.get_text().strip()
                        if table_text:
                            lines = table_text.split('\n')
                            for line in lines:
                                if line.strip() and not line.strip().startswith('█'):
                                    story.append(Paragraph(line.strip(), normal_style))
                            story.append(Spacer(1, 6))
                
                elif element.name in ['div', 'span']:
                    # Divs and spans - extract text if it's meaningful
                    text = element.get_text().strip()
                    if text and len(text) > 10 and not text.startswith('█'):  # Only add if it's substantial text and not black rectangles
                        story.append(Paragraph(text, normal_style))
                        story.append(Spacer(1, 6))
            
            # If no structured content found, fall back to simple text extraction
            if not story:
                text_content = soup.get_text()
                paragraphs = text_content.split('\n\n')
                for para in paragraphs:
                    if para.strip() and not para.strip().startswith('█'):  # Filter out black rectangles
                        clean_text = para.strip().replace('\n', ' ')
                        story.append(Paragraph(clean_text, normal_style))
                        story.append(Spacer(1, 6))
        else:
            # Fallback: simple text extraction
            import re
            text_content = re.sub(r'<[^>]+>', '', html_content)
            paragraphs = text_content.split('\n\n')
            for para in paragraphs:
                if para.strip() and not para.strip().startswith('█'):  # Filter out black rectangles
                    clean_text = para.strip().replace('\n', ' ')
                    story.append(Paragraph(clean_text, normal_style))
                    story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        # Create response
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="risk_report_{execution.scheduled_report.name}_{execution.started_at.strftime("%Y%m%d")}.pdf"'
        
        return response
        
    except Exception as e:
        logger.error(f"Error generating PDF from HTML: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error generating PDF: {}').format(str(e))
        })


def generate_word_from_html(html_content, execution):
    """Generate Word document from HTML content"""
    from docx.shared import Inches, RGBColor
    logger.info(f"generate_word_from_html: Function called with execution_id: {execution.id}")
    try:
        if not DOCX_AVAILABLE:
            logger.error("generate_word_from_html: DOCX_AVAILABLE is False")
            return JsonResponse({
                'status': 'error',
                'message': _('Word generation is not available. Please install python-docx.')
            })
        
        logger.info(f"generate_word_from_html: DOCX_AVAILABLE is True, proceeding...")
        
        # Log HTML content for debugging
        logger.info(f"generate_word_from_html: HTML content length: {len(html_content)}")
        logger.info(f"generate_word_from_html: First 500 chars: {html_content[:500]}")
        logger.info(f"generate_word_from_html: Contains 'black' or 'rectangle': {'black' in html_content.lower() or 'rectangle' in html_content.lower()}")
        
        # Create Word document
        doc = Document()
        
        # Set landscape orientation for all sections
        try:
            for section in doc.sections:
                # Set landscape orientation (A4 landscape: 11.69" x 8.27")
                section.page_width = Inches(11.69)
                section.page_height = Inches(8.27)
        except Exception:
            pass  # Continue if orientation setting fails
        
        # Set document language for better Cyrillic support
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            # Set language to Ukrainian
            element = OxmlElement('w:lang')
            element.set(qn('w:val'), 'uk-UA')
            element.set(qn('w:eastAsia'), 'uk-UA')
            element.set(qn('w:bidi'), 'uk-UA')
            
            # Add to document settings
            settings = doc.settings._element
            settings.append(element)
        except:
            pass  # Continue if language setting fails
        
        # Add title
        title = doc.add_heading(f"Risk Assessment Report - {execution.scheduled_report.name}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Generated: {execution.started_at.strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Report: {execution.scheduled_report.name}")
        doc.add_paragraph(f"Company: {execution.scheduled_report.company.name if execution.scheduled_report.company else 'All Companies'}")
        
        # Add content from HTML with better parsing
        if BEAUTIFULSOUP_AVAILABLE:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Clean up the HTML content
            html_content = str(soup)
            
            # Log processed content
            logger.info(f"generate_word_from_html: Processed HTML length: {len(html_content)}")
            
            # Log all elements found
            all_elements = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'span', 'table', 'tr', 'td', 'th'])
            logger.info(f"generate_word_from_html: Found {len(all_elements)} elements to process")
            
            # Count elements by type
            element_counts = {}
            for elem in all_elements:
                element_counts[elem.name] = element_counts.get(elem.name, 0) + 1
            logger.info(f"generate_word_from_html: Element counts: {element_counts}")
            
            # Process different HTML elements
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'span', 'table', 'tr', 'td', 'th']):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    # Headers
                    text = element.get_text().strip()
                    if text and not text.startswith('█'):  # Filter out black rectangles
                        level = int(element.name[1])  # Extract number from h1, h2, etc.
                        heading = doc.add_heading(text, level=min(level, 6))
                        
                        # Add color to headings based on classes
                        classes = element.get('class', [])
                        if isinstance(classes, str):
                            classes = [classes]
                        
                        if 'text-primary' in classes:
                            for run in heading.runs:
                                run.font.color.rgb = RGBColor(0, 123, 255)  # Blue
                        elif 'text-secondary' in classes:
                            for run in heading.runs:
                                run.font.color.rgb = RGBColor(108, 117, 125)  # Gray
                
                elif element.name == 'p':
                    # Paragraphs
                    text = element.get_text().strip()
                    if text and not text.startswith('█'):  # Filter out black rectangles
                        p = doc.add_paragraph(text)
                        
                        # Add styling based on classes
                        classes = element.get('class', [])
                        if isinstance(classes, str):
                            classes = [classes]
                        
                        if 'text-muted' in classes:
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(108, 117, 125)  # Gray
                        elif 'text-center' in classes:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                elif element.name == 'table':
                    # Tables - create proper table structure
                    logger.info(f"generate_word_from_html: Found table element")
                    try:
                        
                        # Count rows and columns
                        rows = element.find_all('tr')
                        logger.info(f"generate_word_from_html: Table has {len(rows)} rows")
                        
                        if rows:
                            # Get max columns
                            max_cols = 0
                            for row in rows:
                                cells = row.find_all(['td', 'th'])
                                max_cols = max(max_cols, len(cells))
                            
                            logger.info(f"generate_word_from_html: Table has {max_cols} columns")
                            
                            if max_cols > 0:
                                # Create table
                                table = doc.add_table(rows=len(rows), cols=max_cols)
                                table.style = 'Table Grid'
                                
                                # Fill table data
                                for i, row in enumerate(rows):
                                    cells = row.find_all(['td', 'th'])
                                    for j, cell in enumerate(cells):
                                        if j < max_cols:
                                            cell_text = cell.get_text().strip()
                                            if cell_text and not cell_text.startswith('█'):
                                                table_cell = table.cell(i, j)
                                                table_cell.text = cell_text
                                                
                                                # Style header row (first row)
                                                if i == 0:
                                                    for paragraph in table_cell.paragraphs:
                                                        for run in paragraph.runs:
                                                            run.bold = True
                                                            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                                
                                                logger.info(f"generate_word_from_html: Added cell [{i},{j}]: {cell_text[:50]}...")
                        
                        # If table creation failed, fall back to text
                        else:
                            logger.info(f"generate_word_from_html: No rows found in table, falling back to text")
                            table_text = element.get_text().strip()
                            if table_text:
                                lines = table_text.split('\n')
                                for line in lines:
                                    if line.strip() and not line.strip().startswith('█'):
                                        doc.add_paragraph(line.strip())
                    except Exception as e:
                        # Fallback to text if table creation fails
                        logger.error(f"generate_word_from_html: Error creating table: {str(e)}")
                        table_text = element.get_text().strip()
                        if table_text:
                            lines = table_text.split('\n')
                            for line in lines:
                                if line.strip() and not line.strip().startswith('█'):
                                    doc.add_paragraph(line.strip())
                
                elif element.name in ['div', 'span']:
                    # Divs and spans - extract text if it's meaningful
                    text = element.get_text().strip()
                    if text and len(text) > 10 and not text.startswith('█'):  # Only add if it's substantial text and not black rectangles
                        # Check if this might be a colored block (look for specific classes or styles)
                        classes = element.get('class', [])
                        if isinstance(classes, str):
                            classes = [classes]
                        
                        # Skip div elements that contain tables or other processed elements
                        if element.name == 'div':
                            if element.find('table'):
                                logger.info(f"generate_word_from_html: Skipping div with table inside")
                                continue
                            # Skip div elements that contain card elements (they will be processed separately)
                            if element.find(class_=lambda x: x and 'card' in x):
                                logger.info(f"generate_word_from_html: Skipping div with card inside")
                                continue
                            # Skip div elements that contain metric elements (they will be processed separately)
                            if element.find(class_=lambda x: x and ('metric' in x or 'block' in x)):
                                logger.info(f"generate_word_from_html: Skipping div with metric/block inside")
                                continue
                        
                        logger.info(f"generate_word_from_html: Processing {element.name} with classes: {classes}, text: {text[:100]}...")
                        
                        # Add colored background for certain elements (simulate colored blocks)
                        if any('card' in cls.lower() for cls in classes):
                            # Only process card elements, skip others to avoid duplication
                            # Only process if text is substantial (not just whitespace or very short)
                            if len(text.strip()) > 5:
                                p = doc.add_paragraph()
                                run = p.add_run(text)
                                run.bold = True
                            
                            # Add color based on card class
                            if 'bg-primary' in classes:
                                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                logger.info(f"generate_word_from_html: Made blue card bold with white text: {text[:50]}...")
                            elif 'bg-warning' in classes:
                                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                logger.info(f"generate_word_from_html: Made yellow card bold with white text: {text[:50]}...")
                            elif 'bg-danger' in classes:
                                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                logger.info(f"generate_word_from_html: Made red card bold with white text: {text[:50]}...")
                            elif 'bg-success' in classes:
                                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                logger.info(f"generate_word_from_html: Made green card bold with white text: {text[:50]}...")
                            else:
                                logger.info(f"generate_word_from_html: Made text bold for colored block: {text[:50]}...")
                        elif any('metric' in cls.lower() or 'block' in cls.lower() for cls in classes):
                            # Process other metric/block elements
                            # Only process if text is substantial
                            if len(text.strip()) > 5:
                                p = doc.add_paragraph()
                                run = p.add_run(text)
                                run.bold = True
                                logger.info(f"generate_word_from_html: Made text bold for metric/block: {text[:50]}...")
                        else:
                            # Only add if it's not a card or metric element
                            doc.add_paragraph(text)
            
            # If no structured content found, fall back to simple text extraction
            if len(doc.paragraphs) <= 4:  # Only title and metadata
                text_content = soup.get_text()
                paragraphs = text_content.split('\n\n')
                for para in paragraphs:
                    if para.strip() and not para.strip().startswith('█'):  # Filter out black rectangles
                        doc.add_paragraph(para.strip())
        else:
            # Fallback: simple text extraction
            import re
            text_content = re.sub(r'<[^>]+>', '', html_content)
            paragraphs = text_content.split('\n\n')
            for para in paragraphs:
                if para.strip() and not para.strip().startswith('█'):  # Filter out black rectangles
                    doc.add_paragraph(para.strip())
        
        # Save to BytesIO
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Create response
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="risk_report_{execution.scheduled_report.name}_{execution.started_at.strftime("%Y%m%d")}.docx"'
        
        return response
        
    except Exception as e:
        logger.error(f"Error generating Word from HTML: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error generating Word document: {}').format(str(e))
        })