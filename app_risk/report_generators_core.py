# SecBoard/app_risk/report_generators_core.py

import json
import logging
import os
import tempfile
from datetime import datetime, timedelta
from decimal import Decimal
import decimal
from io import BytesIO

from django.contrib.auth.decorators import login_required, user_passes_test
from django.core.exceptions import PermissionDenied
from django.db.models import Count, Q, Sum, Avg, Max, Min
from django.http import JsonResponse, HttpResponse, Http404
from django.shortcuts import render, get_object_or_404
from django.template.loader import render_to_string
from django.utils import timezone
from django.utils.translation import gettext as _, get_language
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods

# Third-party imports
try:
    from weasyprint import HTML, CSS  # pyright: ignore[reportMissingImports]
    from weasyprint.text.fonts import FontConfiguration  # pyright: ignore[reportMissingImports]
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import xlsxwriter  # pyright: ignore[reportMissingImports] 
    XLSXWRITER_AVAILABLE = True 
except ImportError:
    XLSXWRITER_AVAILABLE = False

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BEAUTIFULSOUP_AVAILABLE = True
except ImportError:
    BEAUTIFULSOUP_AVAILABLE = False

# Local imports
from .models import (
    RiskTreatment, AssetVulnerability, Vulnerability, Threat, 
    RiskLevel, AccessRisk, RiskAssessmentAuditLog,
    ScheduledReport, ScheduledReportExecution, ScheduledReportAttachment, ReportProfile, RiskReportEmailConfig
)
from app_cabinet.models import CabinetUser
from app_asset.models import InformationAsset
from app_conf.models import Company
from .risk_assessment_views import calculate_risk_level, calculate_value_of_risk
from .access_utils import can_add_risk_report, can_edit_risk_report, can_delete_risk_report
from .report_data import generate_report_data

logger = logging.getLogger(__name__)


def check_risk_assessment_access(user, action='view'):
    """Check if user has access to risk assessment functionality"""
    if user.is_superuser:
        return True
    
    # Check if user has access through groups
    user_groups = user.groups.all()
    access_records = AccessRisk.objects.filter(
        group__in=user_groups,
        has_access_assessment=True
    )
    
    return access_records.exists()

def check_risk_report_access(user, action='view'):
    """Check if user has access to risk report functionality"""
    if user.is_superuser:
        return True
    
    # Check if user has access through groups
    user_groups = user.groups.all()
    access_records = AccessRisk.objects.filter(
        group__in=user_groups,
        has_access_report=True
    )
    
    return access_records.exists()

@login_required
@user_passes_test(check_risk_report_access)
@require_http_methods(["POST"])
def preview_email_content(request):
    """Preview email content with processed tags"""
    try:
        # Get form data
        email_subject = request.POST.get('email_subject', '')
        email_body = request.POST.get('email_body', '')
        schedule_id = request.POST.get('schedule_id')
        schedule_name = request.POST.get('schedule_name', '')
        company_id = request.POST.get('company_id', '')
        
        # Get scheduled report if exists, otherwise create a temporary one for preview
        if schedule_id:
            try:
                schedule = ScheduledReport.objects.get(id=schedule_id)
            except ScheduledReport.DoesNotExist:
                return JsonResponse({
                    'status': 'error',
                    'message': _('Scheduled report not found')
                })
        else:
            # Create a temporary ScheduledReport object for preview
            # This allows preview to work even before saving the schedule
            schedule = ScheduledReport(
                name=schedule_name or _('Preview Report'),
                created_by=request.user
            )
            
            # Set company if provided
            if company_id:
                try:
                    schedule.company = Company.objects.get(id=company_id)
                except Company.DoesNotExist:
                    pass
        
        # Create a mock execution for preview
        # We need to create it without saving to database
        class MockExecution:
            def __init__(self, scheduled_report, started_at):
                self.scheduled_report = scheduled_report
                self.started_at = started_at
                self.status = 'completed'
                self.id = 'preview-mock-id'
            
            def get_snapshot_url(self):
                # Return a placeholder URL path for preview (without domain)
                # The domain will be added by _build_absolute_url in process_email_tags
                from django.urls import reverse
                try:
                    return reverse('view_scheduled_report_snapshot', args=['preview-mock-id'])
                except:
                    return '/app_risk/scheduled-reports/snapshot/preview-mock-id/'
        
        mock_execution = MockExecution(schedule, timezone.now())
        
        # Process tags
        processed_subject = schedule.process_email_tags(email_subject, mock_execution)
        processed_body = schedule.process_email_tags(email_body, mock_execution)
        
        return JsonResponse({
            'status': 'success',
            'data': {
                'subject': processed_subject,
                'body': processed_body
            }
        })
        
    except Exception as e:
        logger.error(f"Error previewing email content: {str(e)}", exc_info=True)
        return JsonResponse({
            'status': 'error',
            'message': _('Error previewing email content: {}').format(str(e))
        })


@login_required
@user_passes_test(check_risk_report_access)
def download_scheduled_report_attachment(request, attachment_id):
    """Download attachment file"""
    from .report_views import get_user_companies
    try:
        from .models import ScheduledReportAttachment
        attachment = get_object_or_404(ScheduledReportAttachment, id=attachment_id)
        
        # Check permissions
        user_companies = get_user_companies(request.user)
        if attachment.scheduled_report.created_by != request.user and \
           (attachment.scheduled_report.company and attachment.scheduled_report.company not in user_companies):
            raise PermissionDenied(_("You do not have permission to download this file"))
        
        # Check if file exists
        if not attachment.file:
            raise Http404(_("File not found"))
        
        # Serve file
        response = HttpResponse(attachment.file, content_type='application/octet-stream')
        response['Content-Disposition'] = f'attachment; filename="{attachment.original_filename}"'
        return response
        
    except Exception as e:
        logger.error(f"Error downloading attachment {attachment_id}: {str(e)}")
        raise Http404(_("File not found"))


@login_required
@user_passes_test(check_risk_report_access)
def view_scheduled_report_snapshot(request, execution_id):
    """Render a frozen snapshot of a scheduled report execution that was saved at send time."""
    from .report_views import get_user_companies
    from django.utils.translation import gettext as _
    try:
        execution = get_object_or_404(ScheduledReportExecution, id=execution_id)
        report = execution.scheduled_report
        
        # Check permissions: owner, same company, or email recipient
        user_companies = get_user_companies(request.user)
        user_has_access = False
        
        # Check if user is the owner
        if report.created_by == request.user:
            user_has_access = True
        # Check if user belongs to the same company
        elif report.company and report.company in user_companies:
            user_has_access = True
        # Check if user is in the email recipients list
        elif report.email_recipients.filter(user=request.user).exists():
            user_has_access = True
        
        if not user_has_access:
            raise PermissionDenied(_("You do not have permission to view this report"))

        html = execution.snapshot_html or ''
        if not html:
            html = '<div class="alert alert-warning"><i class="fas fa-info-circle me-2"></i>' + \
                  _('Snapshot is not available for this execution.') + '</div>'

        return render(request, 'app_risk/risk_report_preview_page.html', {
            'preview_html': html,
            'page_title': _('Risk Assessment Report'),
            'is_profile': True,
            'profile_id': str(execution.id),  # Use execution ID as profile_id for export
            'language': execution.snapshot_language or 'uk',
        })
    except Exception as e:
        logger.error(f"Error rendering scheduled report snapshot {execution_id}: {e}", exc_info=True)
        return render(request, 'app_risk/risk_report_preview_page.html', {
            'preview_html': f'<div class="alert alert-danger">{_("Error generating preview")}: {str(e)}</div>',
            'page_title': _('Risk Assessment Report')
        })
@require_http_methods(["GET", "POST"])
def generate_risk_report(request):
    """Generate comprehensive risk assessment reports with multilingual support"""
    from django.utils.translation import get_language
    from django.utils.translation import gettext as _
    from .report_views import get_user_companies
    
    # Check permissions
    if not check_risk_report_access(request.user, 'view'):
        raise PermissionDenied(_("You don't have permission to generate reports"))
    
    if request.method == 'GET':
        # Check if this is a direct report generation request with parameters
        report_type = request.GET.get('reportType', 'summary')
        report_format = request.GET.get('format', 'word')
        report_language = request.GET.get('language', get_language()[:2])
        
        if report_type and report_format:
            # Direct generation via GET parameters
            try:
                # Activate the selected language
                from django.utils.translation import activate
                activate(report_language)
                
                # Validate format availability
                if not is_format_available(report_format):
                    return JsonResponse({
                        'status': 'error',
                        'message': _('Selected format is not available. Please install required dependencies.')
                    })
                
                # Generate report data
                report_data = generate_report_data(request.user, {
                    'reportType': report_type,
                    'format': report_format,
                    'language': report_language
                })
                
                # Generate report file
                response = generate_report_file(report_data, report_type, report_format)
                
                return response
                
            except Exception as e:
                logger.error(f"Error generating risk report: {str(e)}", exc_info=True)
                return JsonResponse({
                    'status': 'error',
                    'message': _('Error generating report: {}').format(str(e))
                })
        else:
            # Return report configuration page
            return render(request, 'app_risk/reports/report_config.html', {
                'available_formats': get_available_formats(),
                'companies': get_user_companies(request.user),
            })
    
    elif request.method == 'POST':
        try:
            # Parse request data for POST - handle both JSON and form data
            if request.content_type == 'application/json':
                data = json.loads(request.body) if request.body else {}
            else:
                # Handle form data
                data = {
                    'reportType': request.POST.get('reportType', 'full'),
                    'format': request.POST.get('reportFormat', 'pdf'),
                    'language': request.POST.get('reportLanguage', get_language()[:2]),
                    'company_id': request.POST.get('reportCompany', ''),
                    'notes': request.POST.get('reportNotes', ''),
                }
            
            # Parse selected sections if provided
            selected_sections_json = request.POST.get('selectedSections', '{}')
            try:
                selected_sections = json.loads(selected_sections_json)
                data['selectedSections'] = selected_sections
            except json.JSONDecodeError:
                # If JSON parsing fails, use empty dict
                data['selectedSections'] = {}
            
            report_type = data.get('reportType', 'full')
            report_format = data.get('format', data.get('reportFormat', 'pdf'))
            report_language = data.get('language', data.get('reportLanguage', get_language()[:2]))
            
            # Activate the selected language
            from django.utils.translation import activate
            activate(report_language)
            
            # Validate format availability
            if not is_format_available(report_format):
                return JsonResponse({
                    'status': 'error',
                    'message': _('Selected format is not available. Please install required dependencies.')
                })
            
            # Generate report data
            report_data = generate_report_data(request.user, data)
            
            # Generate report file
            response = generate_report_file(report_data, report_type, report_format)
            
            return response
            
        except Exception as e:
            logger.error(f"Error generating risk report: {str(e)}", exc_info=True)
            return JsonResponse({
                'status': 'error',
                'message': _('Error generating report: {}').format(str(e))
            })


# Helper functions
def get_available_formats():
    """Get list of available report formats"""
    formats = []
    
    logger.info(f"Checking available formats: WEASYPRINT={WEASYPRINT_AVAILABLE}, REPORTLAB={REPORTLAB_AVAILABLE}, DOCX={DOCX_AVAILABLE}")
    
    # Only support Word and PDF formats
    if DOCX_AVAILABLE:
        formats.append({'value': 'word', 'name': 'Word'})
    
    if WEASYPRINT_AVAILABLE or REPORTLAB_AVAILABLE:
        formats.append({'value': 'pdf', 'name': 'PDF'})
    
    logger.info(f"Available formats: {formats}")
    return formats


def is_format_available(format_type):
    """Check if specific format is available"""
    if format_type == 'pdf':
        return WEASYPRINT_AVAILABLE or REPORTLAB_AVAILABLE
    elif format_type == 'word':
        return DOCX_AVAILABLE
    return False


def generate_report_file(data, report_type, format_type):
    """Generate report file in specified format"""
    
    # Get selected sections from data
    selected_sections = data.get('selected_sections', {})
    
    if format_type == 'pdf':
        return generate_pdf_report(data, report_type, selected_sections)
    elif format_type == 'word':
        return generate_word_report(data, report_type, selected_sections)
    else:
        raise ValueError(f"Unsupported format: {format_type}")


def generate_pdf_report(data, report_type, selected_sections):
    """Generate PDF report using ReportLab"""
    
    if not (WEASYPRINT_AVAILABLE or REPORTLAB_AVAILABLE):
        raise ImportError("Neither WeasyPrint nor ReportLab is available. Please install one of them to generate PDF reports.")
    
    # Use ReportLab for PDF generation (more Windows-friendly)
    if REPORTLAB_AVAILABLE:
        return generate_pdf_with_reportlab(data, report_type, selected_sections)
    else:
        # Fallback to WeasyPrint if available
        return generate_pdf_with_weasyprint(data, report_type, selected_sections)

def generate_pdf_with_reportlab(data, report_type, selected_sections):
    """Generate PDF using ReportLab with unified content structure"""
    
    # Get current language and translations
    current_language = data.get('language', get_language()[:2])
    translations = get_report_translations(current_language)
    
    # Generate unified content structure
    unified_content = generate_unified_report_content(data, report_type, selected_sections, translations, current_language)
    
    # Create response
    response = HttpResponse(content_type='application/pdf')
    filename = f"{report_type}_report_{data['generated_date'].strftime('%Y%m%d_%H%M%S')}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    # Register Unicode fonts for Cyrillic support
    unicode_font = 'Helvetica'
    unicode_font_bold = 'Helvetica-Bold'
    
    def register_cyrillic_fonts():
        """Register fonts that support Cyrillic characters"""
        import os
        import platform
        
        system = platform.system()
        fonts_registered = {'regular': False, 'bold': False}
        
        # Font priorities - fonts known to work well with Cyrillic
        font_priorities = {
            'Windows': [
                ('C:/Windows/Fonts/arial.ttf', 'C:/Windows/Fonts/arialbd.ttf'),
                ('C:/Windows/Fonts/tahoma.ttf', 'C:/Windows/Fonts/tahomabd.ttf'),
                ('C:/Windows/Fonts/verdana.ttf', 'C:/Windows/Fonts/verdanab.ttf'),
                ('C:/Windows/Fonts/calibri.ttf', 'C:/Windows/Fonts/calibrib.ttf'),
                ('C:/Windows/Fonts/times.ttf', 'C:/Windows/Fonts/timesbd.ttf'),
            ],
            'Darwin': [
                ('/System/Library/Fonts/Arial.ttf', '/System/Library/Fonts/Arial Bold.ttf'),
                ('/System/Library/Fonts/Times.ttc', '/System/Library/Fonts/Times Bold.ttc'),
            ],
            'Linux': [
                ('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'),
                ('/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf', '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf'),
                ('/usr/share/fonts/truetype/ubuntu/Ubuntu-R.ttf', '/usr/share/fonts/truetype/ubuntu/Ubuntu-B.ttf'),
            ]
        }
        
        font_pairs = font_priorities.get(system, [])
        
        for regular_path, bold_path in font_pairs:
            # Try to register regular font
            if not fonts_registered['regular'] and os.path.exists(regular_path):
                try:
                    pdfmetrics.registerFont(TTFont('CyrillicFont', regular_path))
                    fonts_registered['regular'] = True
                    logger.info(f"Successfully registered regular Cyrillic font: {regular_path}")
                except Exception as e:
                    logger.warning(f"Failed to register regular font {regular_path}: {e}")
            
            # Try to register bold font
            if not fonts_registered['bold'] and os.path.exists(bold_path):
                try:
                    pdfmetrics.registerFont(TTFont('CyrillicFontBold', bold_path))
                    fonts_registered['bold'] = True
                    logger.info(f"Successfully registered bold Cyrillic font: {bold_path}")
                except Exception as e:
                    logger.warning(f"Failed to register bold font {bold_path}: {e}")
            
            # If both fonts are registered, we're done
            if fonts_registered['regular'] and fonts_registered['bold']:
                break
        
        return fonts_registered
    
    try:
        fonts_registered = register_cyrillic_fonts()
        
        if fonts_registered['regular']:
            unicode_font = 'CyrillicFont'
        if fonts_registered['bold']:
            unicode_font_bold = 'CyrillicFontBold'
        
        if not (fonts_registered['regular'] or fonts_registered['bold']):
            logger.warning("No Cyrillic fonts found, using default Helvetica - Ukrainian text may display as squares")
        else:
            logger.info(f"Cyrillic font support enabled: Regular={fonts_registered['regular']}, Bold={fonts_registered['bold']}")
            logger.info(f"Using fonts: Regular='{unicode_font}', Bold='{unicode_font_bold}'")
    
    except Exception as e:
        logger.warning(f"Error setting up Cyrillic fonts: {e}")
        logger.info(f"Falling back to default fonts: Regular='{unicode_font}', Bold='{unicode_font_bold}'")
    
    # Helper function to safely encode text for PDF
    def safe_pdf_text(text):
        """Safely encode text for PDF output with Cyrillic support"""
        if text is None:
            return ""
        try:
            # Convert to string and handle encoding
            text_str = str(text)
            
            # If we have Cyrillic fonts available, return text as-is
            if unicode_font in ['CyrillicFont'] or unicode_font_bold in ['CyrillicFontBold']:
                return text_str
            
            # Try to encode to latin-1 first (ReportLab default)
            try:
                text_str.encode('latin-1')
                return text_str
            except UnicodeEncodeError:
                # If latin-1 fails and no Cyrillic font, try transliteration for Ukrainian
                if any(ord(char) > 127 for char in text_str):
                    # Basic Ukrainian to Latin transliteration
                    transliteration_map = {
                        'а': 'a', 'б': 'b', 'в': 'v', 'г': 'h', 'ґ': 'g', 'д': 'd', 'е': 'e', 'є': 'ye',
                        'ж': 'zh', 'з': 'z', 'и': 'y', 'і': 'i', 'ї': 'yi', 'й': 'y', 'к': 'k', 'л': 'l',
                        'м': 'm', 'н': 'n', 'о': 'o', 'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'у': 'u',
                        'ф': 'f', 'х': 'kh', 'ц': 'ts', 'ч': 'ch', 'ш': 'sh', 'щ': 'shch', 'ь': '', 'ю': 'yu',
                        'я': 'ya', 'А': 'A', 'Б': 'B', 'В': 'V', 'Г': 'H', 'Ґ': 'G', 'Д': 'D', 'Е': 'E',
                        'Є': 'Ye', 'Ж': 'Zh', 'З': 'Z', 'И': 'Y', 'І': 'I', 'Ї': 'Yi', 'Й': 'Y', 'К': 'K',
                        'Л': 'L', 'М': 'M', 'Н': 'N', 'О': 'O', 'П': 'P', 'Р': 'R', 'С': 'S', 'Т': 'T',
                        'У': 'U', 'Ф': 'F', 'Х': 'Kh', 'Ц': 'Ts', 'Ч': 'Ch', 'Ш': 'Sh', 'Щ': 'Shch',
                        'Ь': '', 'Ю': 'Yu', 'Я': 'Ya'
                    }
                    
                    transliterated = ""
                    for char in text_str:
                        transliterated += transliteration_map.get(char, char)
                    
                    # Try the transliterated version
                    try:
                        transliterated.encode('latin-1')
                        return transliterated
                    except UnicodeEncodeError:
                        # Final fallback: replace remaining non-latin characters
                        return transliterated.encode('ascii', 'replace').decode('ascii')
                else:
                    return text_str
        except Exception as e:
            logger.warning(f"Error encoding text for PDF: {e}")
            return str(text).encode('ascii', 'replace').decode('ascii')
    
    def wrap_text_by_words(text, font_name, font_size, max_width):
        """Wrap text by words to fit within the specified width"""
        if not text:
            return [""]
        
        words = str(text).split()
        if not words:
            return [""]
        
        lines = []
        current_line = ""
        
        for word in words:
            test_line = current_line + (" " if current_line else "") + word
            test_width = p.stringWidth(test_line, font_name, font_size)
            
            if test_width <= max_width:
                current_line = test_line
            else:
                if current_line:
                    lines.append(current_line)
                    current_line = word
                else:
                    # Single word is too long, truncate it
                    while word and p.stringWidth(word + "...", font_name, font_size) > max_width:
                        word = word[:-1]
                    lines.append(word + "..." if len(word) < len(str(text)) else word)
                    current_line = ""
        
        if current_line:
            lines.append(current_line)
        
        # Final check: ensure no line exceeds max_width
        final_lines = []
        for line in lines:
            line_width = p.stringWidth(line, font_name, font_size)
            if line_width > max_width:
                # Force truncate if still too wide
                while line and p.stringWidth(line + "...", font_name, font_size) > max_width:
                    line = line[:-1]
                final_lines.append(line + "..." if len(line) < len(str(text)) else line)
            else:
                final_lines.append(line)
        
        return final_lines if final_lines else [""] 
    
    # Create PDF document
    from io import BytesIO
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=landscape(A4))
    width, height = landscape(A4)
    
    # Set up margins
    left_margin = 50
    right_margin = width - 50
    top_margin = height - 50
    bottom_margin = 50
    
    # Current position
    y_position = top_margin
    
    # Title
    p.setFont(unicode_font_bold, 16)
    title_text = safe_pdf_text(unified_content['header']['title'])
    p.drawString(left_margin, y_position, title_text)
    y_position -= 30
    
    # Header information
    p.setFont(unicode_font, 10)
    header_info = [
        f"{safe_pdf_text(translations['generated_on'])}: {safe_pdf_text(unified_content['header']['generated_date'])}",
        f"{safe_pdf_text(translations['report_period'])}: {safe_pdf_text(unified_content['header']['report_period'])}",
        f"{safe_pdf_text(translations['generated_by'])}: {safe_pdf_text(unified_content['header']['generated_by'])}"
    ]
    
    # Company info
    if unified_content['header']['company_info']:
        company_info = unified_content['header']['company_info']
        if company_info['type'] == 'single':
            header_info.append(f"{safe_pdf_text(translations['company'])}: {safe_pdf_text(company_info['name'])}")
        else:
            header_info.append(f"{safe_pdf_text(translations['company'])}: {safe_pdf_text(company_info['text'])}")
    
    for info in header_info:
        p.drawString(left_margin, y_position, safe_pdf_text(info))
        y_position -= 15
    
    y_position -= 20
    
    # Process sections
    for section in unified_content['sections']:
        # Check if we need a new page
        if y_position < bottom_margin + 100:
            p.showPage()
            y_position = top_margin
        
        # Section title
        p.setFont(unicode_font_bold, 14)
        p.drawString(left_margin, y_position, safe_pdf_text(section['title']))
        y_position -= 25
        
        # Special handling for AI Conclusion section (it doesn't have subsections)
        if section.get('type') == 'ai_conclusion':
            # Get the full conclusion content
            conclusion_content = section.get("content", "")
            if conclusion_content:
                # Set font for content
                p.setFont(unicode_font, 10)
                p.setFillColor(colors.black)
                
                # Wrap text by words to fit within page width
                content_width = right_margin - left_margin - 40
                wrapped_lines = wrap_text_by_words(conclusion_content, unicode_font, 10, content_width)
                
                # Render each line of the conclusion
                for line in wrapped_lines:
                    # Check if we need a new page
                    if y_position < bottom_margin + 30:
                        p.showPage()
                        y_position = top_margin
                    
                    # Draw the line
                    p.drawString(left_margin + 20, y_position, safe_pdf_text(line))
                    y_position -= 12  # Line spacing
                
                # Add model information if available
                if section.get('model_name'):
                    y_position -= 10  # Add spacing before model info
                    if y_position < bottom_margin + 30:
                        p.showPage()
                        y_position = top_margin
                    
                    provider = section.get("provider", "")
                    model_name = section.get("model_name", "")
                    p.setFont(unicode_font, 8)
                    p.setFillColor(colors.HexColor('#666666'))  # Gray color for model info
                    model_info = f"Generated by: {provider.upper()} - {model_name}"
                    p.drawString(left_margin + 20, y_position, safe_pdf_text(model_info))
                    y_position -= 15
                
                y_position -= 10  # Add spacing after conclusion section
            
            # Skip the subsection loop for AI conclusion
            continue
        
        # Check if section has subsections before processing
        if 'subsections' not in section:
            continue
        
        for subsection in section['subsections']:
            # Check if we need a new page
            if y_position < bottom_margin + 80:
                p.showPage()
                y_position = top_margin
            
            # Subsection title
            p.setFont(unicode_font_bold, 12)
            p.drawString(left_margin, y_position, safe_pdf_text(subsection['title']))
            y_position -= 20
            
            p.setFont(unicode_font, 10)
            
            if subsection['type'] == 'statistics':
                # Add statistics as colored blocks (similar to HTML)
                block_width = 120
                block_height = 60
                # For Mandatory Processes Summary, display all blocks in one row
                if subsection.get('display_inline'):
                    blocks_per_row = len(subsection['data']) if subsection.get('data') else 2
                else:
                    blocks_per_row = 2
                block_spacing = 20
                
                # Calculate starting position to center the blocks
                total_width = blocks_per_row * block_width + (blocks_per_row - 1) * block_spacing
                start_x = left_margin + (right_margin - left_margin - total_width) / 2
                
                for i, item in enumerate(subsection['data']):
                    if y_position < bottom_margin + block_height + 20:
                        p.showPage()
                        y_position = top_margin
                    
                    # Calculate block position
                    row = i // blocks_per_row
                    col = i % blocks_per_row
                    block_x = start_x + col * (block_width + block_spacing)
                    block_y = y_position - row * (block_height + 10)
                    
                    # Define colors based on item type
                    color_map = {
                        'primary': colors.HexColor('#4472C4'),    # Blue
                        'warning': colors.HexColor('#FFC000'),    # Yellow
                        'danger': colors.HexColor('#C00000'),     # Red
                        'success': colors.HexColor('#00B050'),    # Green
                    }
                    
                    block_color = color_map.get(item.get('type', 'primary'), colors.HexColor('#4472C4'))
                    
                    # Draw colored block background
                    p.setFillColor(block_color)
                    p.rect(block_x, block_y - block_height, block_width, block_height, fill=1)
                    
                    # Draw block border
                    p.setStrokeColor(colors.black)
                    p.setFillColor(colors.white)
                    p.rect(block_x, block_y - block_height, block_width, block_height, fill=0, stroke=1)
                    
                    # Add value (large, bold, white)
                    p.setFont(unicode_font_bold, 16)
                    p.setFillColor(colors.white)
                    value_text = safe_pdf_text(item['value'])
                    value_width = p.stringWidth(value_text, unicode_font_bold, 16)
                    value_x = block_x + (block_width - value_width) / 2
                    p.drawString(value_x, block_y - 25, value_text)
                    
                    # Add label (smaller, white)
                    p.setFont(unicode_font, 8)
                    label_text = safe_pdf_text(item['label'])
                    # Wrap long labels
                    if len(label_text) > 20:
                        words = label_text.split()
                        lines = []
                        current_line = ""
                        for word in words:
                            if len(current_line + " " + word) <= 20:
                                current_line += (" " + word) if current_line else word
                            else:
                                if current_line:
                                    lines.append(current_line)
                                current_line = word
                        if current_line:
                            lines.append(current_line)
                    else:
                        lines = [label_text]
                    
                    for j, line in enumerate(lines):
                        line_width = p.stringWidth(line, unicode_font, 8)
                        line_x = block_x + (block_width - line_width) / 2
                        p.drawString(line_x, block_y - 40 - j * 10, line)
                
                # Update y_position for next section
                total_rows = (len(subsection['data']) + blocks_per_row - 1) // blocks_per_row
                y_position -= total_rows * (block_height + 10) + 20
            
            elif 'headers' in subsection and 'data' in subsection:
                # Create styled table with colored headers
                table_y = y_position
                table_width = right_margin - left_margin - 20
                col_width = table_width / len(subsection['headers'])
                row_height = 20  # Increased row height for better vertical alignment
                cell_padding = 10  # Increased padding inside cells to prevent overflow
                

                
                # Draw header background
                p.setFillColor(colors.HexColor('#4472C4'))  # Blue header
                p.rect(left_margin + 20, table_y - row_height, right_margin - left_margin - 20, row_height, fill=1)
                
                # Headers (white text on blue background)
                p.setFont(unicode_font_bold, 9)
                p.setFillColor(colors.white)
                for i, header in enumerate(subsection['headers']):
                    # Center header text in column
                    text_width = p.stringWidth(safe_pdf_text(header), unicode_font_bold, 9)
                    x_pos = left_margin + 20 + (i * col_width) + (col_width - text_width) / 2
                    # Center vertically in header row
                    y_pos = table_y - (row_height / 2) - 3
                    p.drawString(x_pos, y_pos, safe_pdf_text(header))
                
                # Draw header border and column separators
                p.setStrokeColor(colors.black)
                p.rect(left_margin + 20, table_y - row_height, right_margin - left_margin - 20, row_height, fill=0, stroke=1)
                
                # Draw vertical column separators for header
                for i in range(1, len(subsection['headers'])):
                    x_sep = left_margin + 20 + (i * col_width)
                    p.line(x_sep, table_y - row_height, x_sep, table_y)
                
                table_y -= row_height
                
                # Data rows with alternating colors
                for i, row_data in enumerate(subsection['data'][:10]):  # Limit to 10 rows to fit on page
                    if table_y < bottom_margin + 20:
                        p.showPage()
                        table_y = top_margin - 20
                    
                    # Alternate row colors
                    if i % 2 == 0:
                        p.setFillColor(colors.HexColor('#F2F2F2'))  # Light grey
                    else:
                        p.setFillColor(colors.white)
                    
                    # Draw row background
                    p.rect(left_margin + 20, table_y - row_height, right_margin - left_margin - 20, row_height, fill=1)
                    
                    p.setFont(unicode_font, 8)
                    p.setFillColor(colors.black)
                    
                    if subsection['type'] == 'risk_distribution_level':
                        values = [str(row_data['risk_level']), str(row_data['count']), row_data['percentage']]
                    # Note: risk_distribution_criticality removed to avoid duplication with assets_by_criticality
                    elif subsection['type'] == 'top_risks':
                        values = [row_data['asset_name'], row_data['vulnerability_name'], row_data['risk_level']]
                    elif subsection['type'] == 'compliance_gaps':
                        values = [row_data['type'], row_data['requirement'], row_data['description'][:50] + '...' if len(str(row_data['description'])) > 50 else row_data['description']]
                    elif subsection['type'] == 'assets_by_criticality':
                        values = [row_data['criticality'], row_data['count'], f"{row_data['percentage']}%"]
                    elif subsection['type'] == 'assets_by_vulnerability':
                        values = [row_data['criticality'], row_data['count'], f"{row_data['percentage']}%"]
                    elif subsection['type'] == 'asset_table':
                        values = [row_data['asset_id'], row_data['asset_name'], row_data['criticality'], row_data.get('company', '')]
                    elif subsection['type'] == 'top_10_asset':
                        vulnerabilities_display = row_data.get('vulnerabilities', 'Undefined')
                        values = [row_data['asset_id'], row_data['asset_name'], row_data['criticality'], vulnerabilities_display, row_data.get('company', '')]
                    elif subsection['type'] == 'vulnerability_table':
                        values = [row_data['asset_name'], row_data['vulnerability_name'], row_data['status']]
                    elif subsection['type'] == 'top_10_vulnerability':
                        values = [row_data['asset_name'], row_data['vulnerability_name'], row_data['status'], row_data['threat'], row_data['probability_impact'], row_data['impact'], row_data['risk_value'], row_data['risk_level'], row_data['risk_mitigation_controls']]
                    elif subsection['type'] == 'overdue_treatments':
                        values = [row_data['asset_name'], row_data['treatment'], safe_text_assignment(row_data['deadline'])]
                    elif subsection['type'] == 'treatment_details':
                        values = [row_data['asset_name'], row_data['vulnerability_name'], row_data.get('vulnerability_status', ''), row_data.get('risk_level', ''), row_data.get('treatment_type', ''), row_data.get('treatment_status', ''), row_data.get('description_responsible', ''), row_data.get('due_date', '')]
                    elif subsection['type'] == 'framework_company_requirements':
                        values = [
                            row_data.get('name', ''),
                            row_data.get('framework_type', ''),
                            row_data.get('version', ''),
                            row_data.get('company', ''),
                            row_data.get('status', ''),
                            str(row_data.get('controls_total', 0)),
                            str(row_data.get('controls_completed', 0)),
                            row_data.get('completion_percentage', '0%')
                        ]
                    elif subsection['type'] == 'company_requirements':
                        values = [
                            row_data.get('code', ''),
                            row_data.get('name', ''),
                            row_data.get('requirement_type', ''),
                            row_data.get('company', ''),
                            row_data.get('regulator', ''),
                            row_data.get('status', ''),
                            str(row_data.get('controls_total', 0)),
                            str(row_data.get('controls_completed', 0)),
                            row_data.get('completion_percentage', '0%')
                        ]
                    elif subsection['type'] == 'internal_requirements':
                        values = [
                            row_data.get('code', ''),
                            row_data.get('name', ''),
                            row_data.get('requirement_type', ''),
                            row_data.get('company', ''),
                            row_data.get('status', ''),
                            str(row_data.get('controls_total', 0)),
                            str(row_data.get('controls_completed', 0)),
                            row_data.get('completion_percentage', '0%')
                        ]
                    elif subsection['type'] == 'table':
                        # Generic table rendering for acceptable risk tables
                        values = [str(row_data.get(key, '')) for key in row_data.keys()]
                    else:
                        values = ['', '', '']
                    

                    
                    # First pass: calculate maximum row height needed for all cells
                    max_lines = 1
                    cell_data = []
                    for i, value in enumerate(values):
                        display_value = str(value)
                        available_width = col_width - 2 * cell_padding
                        wrapped_lines = wrap_text_by_words(display_value, unicode_font, 8, available_width)
                        cell_data.append(wrapped_lines)
                        max_lines = max(max_lines, len(wrapped_lines))
                    
                    # Calculate new row height based on maximum lines needed
                    line_height = 10
                    new_row_height = max_lines * line_height
                    
                    # Adjust table_y if row height increased
                    if new_row_height > row_height:
                        table_y -= (new_row_height - row_height)
                        row_height = new_row_height
                    
                    # Second pass: render all cells with consistent row height
                    for i, (value, wrapped_lines) in enumerate(zip(values, cell_data)):
                        display_value = str(value)
                        
                        # Draw each line of wrapped text
                        for line_idx, line in enumerate(wrapped_lines):
                            line_width = p.stringWidth(safe_pdf_text(line), unicode_font, 8)
                            
                            # Calculate cell boundaries
                            cell_left = left_margin + 20 + (i * col_width) + cell_padding
                            cell_right = cell_left + col_width - 2 * cell_padding
                            
                            # Center text within cell boundaries
                            x_pos = cell_left + (col_width - 2 * cell_padding - line_width) / 2
                            
                            # Ensure text doesn't go beyond cell boundaries
                            if x_pos < cell_left:
                                x_pos = cell_left
                            elif x_pos + line_width > cell_right:
                                x_pos = cell_right - line_width
                            
                            y_pos = table_y - (row_height / 2) - 3 + (line_idx * line_height)
                            p.drawString(x_pos, y_pos, safe_pdf_text(line))
                    
                    # Draw row border and column separators
                    p.setStrokeColor(colors.HexColor('#CCCCCC'))
                    p.rect(left_margin + 20, table_y - row_height, right_margin - left_margin - 20, row_height, fill=0, stroke=1)
                    
                    # Draw vertical column separators
                    for i in range(1, len(subsection['headers'])):
                        x_sep = left_margin + 20 + (i * col_width)
                        p.line(x_sep, table_y - row_height, x_sep, table_y)
                    
                    table_y -= row_height
                
                y_position = table_y - 10
            
            elif subsection['type'] == 'financial_overview':
                # Add financial overview as colored blocks (similar to statistics)
                block_width = 120
                block_height = 60
                blocks_per_row = 2
                block_spacing = 20
                
                # Calculate starting position to center the blocks
                total_width = blocks_per_row * block_width + (blocks_per_row - 1) * block_spacing
                start_x = left_margin + (right_margin - left_margin - total_width) / 2
                
                for i, item in enumerate(subsection['data']):
                    if y_position < bottom_margin + block_height + 20:
                        p.showPage()
                        y_position = top_margin
                    
                    # Calculate block position
                    row = i // blocks_per_row
                    col = i % blocks_per_row
                    block_x = start_x + col * (block_width + block_spacing)
                    block_y = y_position - row * (block_height + 10)
                    
                    # Use blue color for financial data
                    block_color = colors.HexColor('#17A2B8')  # Info blue
                    
                    # Draw colored block background
                    p.setFillColor(block_color)
                    p.rect(block_x, block_y - block_height, block_width, block_height, fill=1)
                    
                    # Draw block border
                    p.setStrokeColor(colors.black)
                    p.setFillColor(colors.white)
                    p.rect(block_x, block_y - block_height, block_width, block_height, fill=0, stroke=1)
                    
                    # Add value (large, bold, white)
                    p.setFont(unicode_font_bold, 16)
                    p.setFillColor(colors.white)
                    value_text = safe_pdf_text(item['value'])
                    value_width = p.stringWidth(value_text, unicode_font_bold, 16)
                    value_x = block_x + (block_width - value_width) / 2
                    p.drawString(value_x, block_y - 25, value_text)
                    
                    # Add label (smaller, white)
                    p.setFont(unicode_font, 8)
                    label_text = safe_pdf_text(item['label'])
                    # Wrap long labels
                    if len(label_text) > 20:
                        words = label_text.split()
                        lines = []
                        current_line = ""
                        for word in words:
                            if len(current_line + " " + word) <= 20:
                                current_line += (" " + word) if current_line else word
                            else:
                                if current_line:
                                    lines.append(current_line)
                                current_line = word
                        if current_line:
                            lines.append(current_line)
                    else:
                        lines = [label_text]
                    
                    for j, line in enumerate(lines):
                        line_width = p.stringWidth(line, unicode_font, 8)
                        line_x = block_x + (block_width - line_width) / 2
                        p.drawString(line_x, block_y - 40 - j * 10, line)
                
                # Update y_position for next section
                total_rows = (len(subsection['data']) + blocks_per_row - 1) // blocks_per_row
                y_position -= total_rows * (block_height + 10) + 20
            
            elif subsection['type'] == 'residual_overview':
                # Add residual overview as colored blocks (similar to performance metrics)
                block_width = 120
                block_height = 60
                blocks_per_row = 2
                block_spacing = 20
                
                # Calculate starting position to center the blocks
                total_width = blocks_per_row * block_width + (blocks_per_row - 1) * block_spacing
                start_x = left_margin + (right_margin - left_margin - total_width) / 2
                
                for i, item in enumerate(subsection['data']):
                    if y_position < bottom_margin + block_height + 20:
                        p.showPage()
                        y_position = top_margin
                    
                    # Calculate block position
                    row = i // blocks_per_row
                    col = i % blocks_per_row
                    block_x = start_x + col * (block_width + block_spacing)
                    block_y = y_position - row * (block_height + 10)
                    
                    # Use blue color for residual overview
                    block_color = colors.HexColor('#17A2B8')  # Blue
                    
                    # Draw colored block background
                    p.setFillColor(block_color)
                    p.rect(block_x, block_y - block_height, block_width, block_height, fill=1)
                    
                    # Draw block border
                    p.setStrokeColor(colors.black)
                    p.setFillColor(colors.white)
                    p.rect(block_x, block_y - block_height, block_width, block_height, fill=0, stroke=1)
                    
                    # Add value (large, bold, white)
                    p.setFont(unicode_font_bold, 16)
                    p.setFillColor(colors.white)
                    value_text = safe_pdf_text(item['value'])
                    value_width = p.stringWidth(value_text, unicode_font_bold, 16)
                    value_x = block_x + (block_width - value_width) / 2
                    p.drawString(value_x, block_y - 25, value_text)
                    
                    # Add label (smaller, white)
                    p.setFont(unicode_font, 8)
                    label_text = safe_pdf_text(item['label'])
                    # Wrap long labels
                    if len(label_text) > 20:
                        words = label_text.split()
                        lines = []
                        current_line = ""
                        for word in words:
                            if len(current_line + " " + word) <= 20:
                                current_line += (" " + word) if current_line else word
                            else:
                                if current_line:
                                    lines.append(current_line)
                                current_line = word
                        if current_line:
                            lines.append(current_line)
                    else:
                        lines = [label_text]
                    
                    for j, line in enumerate(lines):
                        line_width = p.stringWidth(line, unicode_font, 8)
                        line_x = block_x + (block_width - line_width) / 2
                        p.drawString(line_x, block_y - 40 - j * 10, line)
                
                # Update y_position for next section
                total_rows = (len(subsection['data']) + blocks_per_row - 1) // blocks_per_row
                y_position -= total_rows * (block_height + 10) + 20
            
            elif subsection['type'] == 'governance_overview':
                # Add governance overview as colored blocks (similar to performance metrics)
                block_width = 120
                block_height = 60
                blocks_per_row = 2
                block_spacing = 20
                
                # Calculate starting position to center the blocks
                total_width = blocks_per_row * block_width + (blocks_per_row - 1) * block_spacing
                start_x = left_margin + (right_margin - left_margin - total_width) / 2
                
                for i, item in enumerate(subsection['data']):
                    if y_position < bottom_margin + block_height + 20:
                        p.showPage()
                        y_position = top_margin
                    
                    # Calculate block position
                    row = i // blocks_per_row
                    col = i % blocks_per_row
                    block_x = start_x + col * (block_width + block_spacing)
                    block_y = y_position - row * (block_height + 10)
                    
                    # Use purple color for governance overview
                    block_color = colors.HexColor('#6F42C1')  # Purple
                    
                    # Draw colored block background
                    p.setFillColor(block_color)
                    p.rect(block_x, block_y - block_height, block_width, block_height, fill=1)
                    
                    # Draw block border
                    p.setStrokeColor(colors.black)
                    p.setFillColor(colors.white)
                    p.rect(block_x, block_y - block_height, block_width, block_height, fill=0, stroke=1)
                    
                    # Add value (large, bold, white)
                    p.setFont(unicode_font_bold, 16)
                    p.setFillColor(colors.white)
                    value_text = safe_pdf_text(item['value'])
                    value_width = p.stringWidth(value_text, unicode_font_bold, 16)
                    value_x = block_x + (block_width - value_width) / 2
                    p.drawString(value_x, block_y - 25, value_text)
                    
                    # Add label (smaller, white)
                    p.setFont(unicode_font, 8)
                    label_text = safe_pdf_text(item['label'])
                    # Wrap long labels
                    if len(label_text) > 20:
                        words = label_text.split()
                        lines = []
                        current_line = ""
                        for word in words:
                            if len(current_line + " " + word) <= 20:
                                current_line += (" " + word) if current_line else word
                            else:
                                if current_line:
                                    lines.append(current_line)
                                current_line = word
                        if current_line:
                            lines.append(current_line)
                    else:
                        lines = [label_text]
                    
                    for j, line in enumerate(lines):
                        line_width = p.stringWidth(line, unicode_font, 8)
                        line_x = block_x + (block_width - line_width) / 2
                        p.drawString(line_x, block_y - 40 - j * 10, line)
                
                # Update y_position for next section
                total_rows = (len(subsection['data']) + blocks_per_row - 1) // blocks_per_row
                y_position -= total_rows * (block_height + 10) + 20
            
            elif subsection['type'] == 'pci_dss' or subsection['type'] == 'iso27001':
                # Add compliance data as text
                for item in subsection['data']:
                    if y_position < bottom_margin + 20:
                        p.showPage()
                        y_position = top_margin
                    
                    text = f"{safe_pdf_text(item['label'])}: {safe_pdf_text(item['value'])}"
                    p.drawString(left_margin + 20, y_position, safe_pdf_text(text))
                    y_position -= 15
            
            elif subsection['type'] == 'priority_actions':
                # Add recommendations as bullet points
                for item in subsection['data']:
                    if y_position < bottom_margin + 20:
                        p.showPage()
                        y_position = top_margin
                    
                    text = f"• {safe_pdf_text(item)}"
                    p.drawString(left_margin + 20, y_position, safe_pdf_text(text))
                    y_position -= 15
            
            y_position -= 10
    
    # Remove test text - no longer needed
    
    # Save PDF
    p.save()
    pdf = buffer.getvalue()
    buffer.close()
    response.write(pdf)
    
    return response

def generate_pdf_with_weasyprint(data, report_type, selected_sections):
    """Generate PDF using WeasyPrint (fallback)"""
    
    response = HttpResponse(content_type='application/pdf')
    filename = f"{report_type}_report_{data['generated_date'].strftime('%Y%m%d_%H%M%S')}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    # Simple fallback - create basic PDF content
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from io import BytesIO
    
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=landscape(A4))
    
    # Get translations
    translations = data.get('translations', {})
    
    # Helper function for fallback PDF
    def safe_fallback_text(text):
        """Safely encode text for fallback PDF"""
        if text is None:
            return ""
        try:
            return str(text).encode('ascii', 'replace').decode('ascii')
        except:
            return str(text)
    
    # Title
    p.drawString(100, 750, safe_fallback_text(f"{translations.get('risk_report', 'Risk Report')} - {report_type.title()}"))
    p.drawString(100, 720, safe_fallback_text(f"{translations.get('generated_date', 'Generated Date')}: {data['generated_date'].strftime('%Y-%m-%d %H:%M')}"))
    
    y_position = 680
    
    # Only include selected sections
    if selected_sections.get('statistics') and 'statistics' in data:
        stats = data['statistics']
        p.drawString(100, y_position, safe_fallback_text(f"{translations.get('statistics', 'Statistics')}:"))
        y_position -= 30
        p.drawString(120, y_position, safe_fallback_text(f"{translations.get('total_assets', 'Total Assets')}: {stats.get('total_assets', 0)}"))
        y_position -= 20
        p.drawString(120, y_position, safe_fallback_text(f"{translations.get('total_vulnerabilities', 'Total Vulnerabilities')}: {stats.get('total_vulnerabilities', 0)}"))
        y_position -= 30
    
    if selected_sections.get('recommendations'):
        p.drawString(100, y_position, safe_fallback_text(f"{translations.get('recommendations', 'Recommendations')}:"))
        y_position -= 20
        p.drawString(120, y_position, safe_fallback_text(translations.get('recommendations_text', 'Please implement appropriate security controls.')))
    
    p.showPage()
    p.save()
    
    pdf_data = buffer.getvalue()
    buffer.close()
    response.write(pdf_data)
    
    return response



def safe_text_assignment(value):
    """Safely convert any value to string for Word document text assignment"""
    if value is None:
        return ""
    elif hasattr(value, 'strftime'):
        # It's a date or datetime object
        if hasattr(value, 'hour'):
            # It's a datetime object
            return value.strftime('%Y-%m-%d %H:%M')
        else:
            # It's a date object
            return value.strftime('%Y-%m-%d')
    else:
        # Convert to string
        return str(value)

def generate_word_report(data, report_type, selected_sections):
    """Generate Word report using python-docx with unified content structure"""
    from docx.shared import RGBColor, Pt, Inches
    
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed. Please install it to generate Word reports.")
    
    # Get current language and translations
    current_language = data.get('language', get_language()[:2])
    translations = get_report_translations(current_language)
    
    # Generate unified content structure
    unified_content = generate_unified_report_content(data, report_type, selected_sections, translations, current_language)
    
    # Create Word document
    doc = Document()
    
    # Global styles: margins and default font
    try:
        for section in doc.sections:
            # Set landscape orientation (A4 landscape: 11.69" x 8.27")
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.header_distance = Inches(0.3)
            section.footer_distance = Inches(0.5)
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
    except Exception:
        pass
    
    # Header and footer with page numbers
    try:
        first_section = doc.sections[0]
        header_p = first_section.header.paragraphs[0] if first_section.header.paragraphs else first_section.header.add_paragraph()
        header_p.text = unified_content['header']['title']
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        footer_p = first_section.footer.paragraphs[0] if first_section.footer.paragraphs else first_section.footer.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # "Page X of Y"
        footer_p.add_run('Page ')
        run_page = footer_p.add_run()
        r = run_page._r
        fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin'); r.append(fld_begin)
        instr_text = OxmlElement('w:instrText'); instr_text.set(qn('xml:space'), 'preserve'); instr_text.text = 'PAGE'; r.append(instr_text)
        fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end'); r.append(fld_end)
        footer_p.add_run(' of ')
        run_numpages = footer_p.add_run()
        r2 = run_numpages._r
        fld_begin2 = OxmlElement('w:fldChar'); fld_begin2.set(qn('w:fldCharType'), 'begin'); r2.append(fld_begin2)
        instr_text2 = OxmlElement('w:instrText'); instr_text2.set(qn('xml:space'), 'preserve'); instr_text2.text = 'NUMPAGES'; r2.append(instr_text2)
        fld_end2 = OxmlElement('w:fldChar'); fld_end2.set(qn('w:fldCharType'), 'end'); r2.append(fld_end2)
    except Exception:
        pass
    
    # Add title
    title = doc.add_heading(unified_content['header']['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add header information
    doc.add_paragraph(f"{translations['generated_on']}: {unified_content['header']['generated_date']}")
    doc.add_paragraph(f"{translations['report_period']}: {unified_content['header']['report_period']}")
    doc.add_paragraph(f"{translations['generated_by']}: {unified_content['header']['generated_by']}")
    
    # Company info
    if unified_content['header']['company_info']:
        company_info = unified_content['header']['company_info']
        if company_info['type'] == 'single':
            doc.add_paragraph(f"{translations['company']}: {company_info['name']}")
        else:
            doc.add_paragraph(f"{translations['company']}: {company_info['text']}")
    
    # Optional Table of Contents (Word will require updating fields to render)
    try:
        doc.add_paragraph('Table of Contents')
        toc_p = doc.add_paragraph()
        r_toc = toc_p.add_run()._r
        fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin'); r_toc.append(fld_begin)
        instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'TOC \\o "1-3" \\h \\z \\u'; r_toc.append(instr)
        fld_separate = OxmlElement('w:fldChar'); fld_separate.set(qn('w:fldCharType'), 'separate'); r_toc.append(fld_separate)
        fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end'); r_toc.append(fld_end)
    except Exception:
        pass

    # Executive Summary metrics (match HTML preview style)
    try:
        stats = data.get('statistics') or {}
        if selected_sections.get('statistics') and stats:
            doc.add_heading(translations.get('key_metrics', 'Key Metrics'), level=1)
            metrics = [
                (str(stats.get('total_assets', 0)), translations.get('total_assets', 'Total Assets'), '4472C4'),  # blue
                (str(stats.get('total_vulnerabilities', 0)), translations.get('total_vulnerabilities', 'Total Vulnerabilities'), 'FFC000'),  # amber
                (str(stats.get('high_risk_count', 0)), translations.get('high_risk_assets', 'High Risk Assets'), 'C00000'),  # red
                (f"{stats.get('completion_rate', 0):.1f}%", translations.get('completion_rate', 'Completion Rate (%)'), '00B050'),  # green
            ]
            table = doc.add_table(rows=1, cols=len(metrics))
            try:
                table.autofit = True
            except Exception:
                pass
            row_cells = table.rows[0].cells
            for idx, (value, label, hex_fill) in enumerate(metrics):
                cell = row_cells[idx]
                # Shading background
                try:
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), hex_fill)
                    tcPr.append(shd)
                except Exception:
                    pass
                # Clear existing paragraph
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r1 = p.add_run(value)
                r1.bold = True
                r1.font.size = Pt(18)
                r1.font.color.rgb = RGBColor(255, 255, 255)
                p.add_run('\n')
                r2 = p.add_run(label)
                r2.font.size = Pt(10)
                r2.font.color.rgb = RGBColor(255, 255, 255)
    except Exception:
        pass

    # Process sections with improved styling and pagination
    sections_len = len(unified_content['sections'])
    for idx, section in enumerate(unified_content['sections']):
        doc.add_heading(section['title'], level=1)
        
        for subsection in section['subsections']:
            doc.add_heading(subsection['title'], level=2)
            
            if subsection['type'] == 'statistics':
                # Add statistics as paragraphs
                for item in subsection['data']:
                    paragraph = doc.add_paragraph()
                    key_run = paragraph.add_run(f"{item['label']}: ")
                    key_run.bold = True
                    paragraph.add_run(item['value'])
            
            elif 'headers' in subsection and 'data' in subsection:
                # Create table
                table = doc.add_table(rows=1, cols=len(subsection['headers']))
                try:
                    table.style = 'Light Shading Accent 1'
                except Exception:
                    table.style = 'Table Grid'
                try:
                    table.autofit = True
                except Exception:
                    pass
                
                # Add headers
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(subsection['headers']):
                    hdr_cells[i].text = header
                    # Make headers bold
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                for row_data in subsection['data']:
                    row_cells = table.add_row().cells
                    
                    if subsection['type'] == 'risk_distribution_level':
                        row_cells[0].text = safe_text_assignment(row_data['risk_level'])
                        row_cells[1].text = safe_text_assignment(row_data['count'])
                        row_cells[2].text = safe_text_assignment(row_data['percentage'])
                    # Note: risk_distribution_criticality removed to avoid duplication with assets_by_criticality
                    elif subsection['type'] == 'top_risks':
                        row_cells[0].text = safe_text_assignment(row_data['asset_name'])
                        row_cells[1].text = safe_text_assignment(row_data['vulnerability_name'])
                        row_cells[2].text = safe_text_assignment(row_data['risk_level'])
                    elif subsection['type'] == 'compliance_gaps':
                        row_cells[0].text = safe_text_assignment(row_data['type'])
                        row_cells[1].text = safe_text_assignment(row_data['requirement'])
                        row_cells[2].text = safe_text_assignment(row_data['description'])
                    elif subsection['type'] == 'assets_by_criticality':
                        row_cells[0].text = safe_text_assignment(row_data['criticality'])
                        row_cells[1].text = safe_text_assignment(row_data['count'])
                        row_cells[2].text = safe_text_assignment(f"{row_data['percentage']}%")
                    elif subsection['type'] == 'assets_by_vulnerability':
                        row_cells[0].text = safe_text_assignment(row_data['criticality'])
                        row_cells[1].text = safe_text_assignment(row_data['count'])
                        row_cells[2].text = safe_text_assignment(f"{row_data['percentage']}%")
                    elif subsection['type'] == 'asset_table':
                        row_cells[0].text = safe_text_assignment(row_data['asset_id'])
                        row_cells[1].text = safe_text_assignment(row_data['asset_name'])
                        
                        # Add criticality with color
                        criticality_cell = row_cells[2]
                        criticality_cell.text = safe_text_assignment(row_data['criticality'])
                        
                        # Apply color to criticality text if available
                        if 'criticality_color' in row_data and row_data['criticality_color']:
                            for paragraph in criticality_cell.paragraphs:
                                for run in paragraph.runs:
                                    # Convert hex color to RGB for Word
                                    hex_color = row_data['criticality_color'].lstrip('#')
                                    if len(hex_color) == 6:
                                        rgb_color = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                                        run.font.color.rgb = RGBColor(*rgb_color)
                        
                        row_cells[3].text = safe_text_assignment(row_data.get('company', ''))
                    elif subsection['type'] == 'top_10_asset':
                        row_cells[0].text = safe_text_assignment(row_data['asset_id'])
                        row_cells[1].text = safe_text_assignment(row_data['asset_name'])
                        
                        # Add criticality with color
                        criticality_cell = row_cells[2]
                        criticality_cell.text = safe_text_assignment(row_data['criticality'])
                        
                        # Apply color to criticality text if available
                        if 'criticality_color' in row_data and row_data['criticality_color']:
                            for paragraph in criticality_cell.paragraphs:
                                for run in paragraph.runs:
                                    # Convert hex color to RGB for Word
                                    hex_color = row_data['criticality_color'].lstrip('#')
                                    if len(hex_color) == 6:
                                        rgb_color = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                                        run.font.color.rgb = RGBColor(*rgb_color)
                        
                        row_cells[3].text = safe_text_assignment(row_data.get('company', ''))
                    elif subsection['type'] == 'vulnerability_table':
                        row_cells[0].text = safe_text_assignment(row_data['asset_name'])
                        row_cells[1].text = safe_text_assignment(row_data['vulnerability_name'])
                        row_cells[2].text = safe_text_assignment(row_data['status'])
                    elif subsection['type'] == 'top_10_vulnerability':
                        row_cells[0].text = safe_text_assignment(row_data['asset_name'])
                        row_cells[1].text = safe_text_assignment(row_data['vulnerability_name'])
                        row_cells[2].text = safe_text_assignment(row_data['status'])
                        row_cells[3].text = safe_text_assignment(row_data['threat'])
                        row_cells[4].text = safe_text_assignment(row_data['probability_impact'])
                        row_cells[5].text = safe_text_assignment(row_data['impact'])
                        row_cells[6].text = safe_text_assignment(row_data['risk_value'])
                        row_cells[7].text = safe_text_assignment(row_data['risk_level'])
                        row_cells[8].text = safe_text_assignment(row_data['risk_mitigation_controls'])
                    elif subsection['type'] == 'overdue_treatments':
                        row_cells[0].text = safe_text_assignment(row_data['asset_name'])
                        row_cells[1].text = safe_text_assignment(row_data['treatment'])
                        row_cells[2].text = safe_text_assignment(row_data['deadline'])
                    elif subsection['type'] == 'treatment_details':
                        row_cells[0].text = safe_text_assignment(row_data['asset_name'])
                        row_cells[1].text = safe_text_assignment(row_data['vulnerability_name'])
                        row_cells[2].text = safe_text_assignment(row_data['treatment_type'])
                        row_cells[3].text = safe_text_assignment(row_data.get('status', ''))
                        row_cells[4].text = safe_text_assignment(row_data.get('assigned_to', ''))
                        row_cells[5].text = safe_text_assignment(row_data.get('due_date', ''))
                    elif subsection['type'] == 'framework_company_requirements':
                        row_cells[0].text = safe_text_assignment(row_data.get('name', ''))
                        row_cells[1].text = safe_text_assignment(row_data.get('framework_type', ''))
                        row_cells[2].text = safe_text_assignment(row_data.get('version', ''))
                        row_cells[3].text = safe_text_assignment(row_data.get('company', ''))
                        row_cells[4].text = safe_text_assignment(row_data.get('status', ''))
                        row_cells[5].text = safe_text_assignment(str(row_data.get('controls_total', 0)))
                        row_cells[6].text = safe_text_assignment(str(row_data.get('controls_completed', 0)))
                        row_cells[7].text = safe_text_assignment(row_data.get('completion_percentage', '0%'))
                    elif subsection['type'] == 'company_requirements':
                        row_cells[0].text = safe_text_assignment(row_data.get('code', ''))
                        row_cells[1].text = safe_text_assignment(row_data.get('name', ''))
                        row_cells[2].text = safe_text_assignment(row_data.get('requirement_type', ''))
                        row_cells[3].text = safe_text_assignment(row_data.get('company', ''))
                        row_cells[4].text = safe_text_assignment(row_data.get('regulator', ''))
                        row_cells[5].text = safe_text_assignment(row_data.get('status', ''))
                        row_cells[6].text = safe_text_assignment(str(row_data.get('controls_total', 0)))
                        row_cells[7].text = safe_text_assignment(str(row_data.get('controls_completed', 0)))
                        row_cells[8].text = safe_text_assignment(row_data.get('completion_percentage', '0%'))
                    elif subsection['type'] == 'internal_requirements':
                        row_cells[0].text = safe_text_assignment(row_data.get('code', ''))
                        row_cells[1].text = safe_text_assignment(row_data.get('name', ''))
                        row_cells[2].text = safe_text_assignment(row_data.get('requirement_type', ''))
                        row_cells[3].text = safe_text_assignment(row_data.get('company', ''))
                        row_cells[4].text = safe_text_assignment(row_data.get('status', ''))
                        row_cells[5].text = safe_text_assignment(str(row_data.get('controls_total', 0)))
                        row_cells[6].text = safe_text_assignment(str(row_data.get('controls_completed', 0)))
                        row_cells[7].text = safe_text_assignment(row_data.get('completion_percentage', '0%'))
                    elif subsection['type'] == 'table':
                        # Generic table rendering for acceptable risk tables
                        for i, (key, value) in enumerate(row_data.items()):
                            if i < len(row_cells):
                                row_cells[i].text = safe_text_assignment(str(value))
                    elif subsection['type'] == 'quiz_statistics_table':
                        row_cells[0].text = safe_text_assignment(row_data.get('quiz_title', ''))
                        row_cells[1].text = safe_text_assignment(row_data.get('total_attempts', ''))
                        row_cells[2].text = safe_text_assignment(row_data.get('successful_attempts', ''))
                        row_cells[3].text = safe_text_assignment(row_data.get('failed_attempts', ''))
                        row_cells[4].text = safe_text_assignment(row_data.get('success_rate', ''))
                        row_cells[5].text = safe_text_assignment(row_data.get('average_score', ''))
                        row_cells[6].text = safe_text_assignment(row_data.get('passing_score', ''))
                    elif subsection['type'] == 'users_at_risk_table':
                        row_cells[0].text = safe_text_assignment(row_data.get('user', ''))
                        row_cells[1].text = safe_text_assignment(row_data.get('company', ''))
                        row_cells[2].text = safe_text_assignment(row_data.get('attempts_count', ''))
            
            elif subsection['type'] == 'financial_overview':
                # Add financial overview as paragraphs with bold labels
                for item in subsection['data']:
                    paragraph = doc.add_paragraph()
                    key_run = paragraph.add_run(f"{item['label']}: ")
                    key_run.bold = True
                    paragraph.add_run(item['value'])
            
            elif subsection['type'] == 'residual_overview':
                # Add residual overview as paragraphs with bold labels
                for item in subsection['data']:
                    paragraph = doc.add_paragraph()
                    key_run = paragraph.add_run(f"{item['label']}: ")
                    key_run.bold = True
                    paragraph.add_run(item['value'])
            
            elif subsection['type'] == 'governance_overview':
                # Add governance overview as paragraphs with bold labels
                for item in subsection['data']:
                    paragraph = doc.add_paragraph()
                    key_run = paragraph.add_run(f"{item['label']}: ")
                    key_run.bold = True
                    paragraph.add_run(item['value'])
            
            elif subsection['type'] == 'pci_dss' or subsection['type'] == 'iso27001':
                # Add compliance data as paragraphs
                for item in subsection['data']:
                    paragraph = doc.add_paragraph()
                    key_run = paragraph.add_run(f"{item['label']}: ")
                    key_run.bold = True
                    paragraph.add_run(item['value'])
            
            elif subsection['type'] == 'priority_actions':
                # Add recommendations as bullet points
                for item in subsection['data']:
                    doc.add_paragraph(item, style='List Bullet')
        
        # Page break between major sections (except last)
        if idx < sections_len - 1:
            doc.add_page_break()
    
    # Save to response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"{report_type}_report_{data['generated_date'].strftime('%Y%m%d_%H%M%S')}.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    doc.save(response)
    return response


def get_localized_field(obj, field_prefix, language=None):
    """Get localized field value based on current language"""
    if language is None:
        language = get_language()[:2]
    
    # Try to get field in current language
    field_name = f"{field_prefix}_{language}"
    value = getattr(obj, field_name, None)
    
    # Fallback to Ukrainian, then English, then Russian
    if not value:
        for fallback_lang in ['uk', 'en', 'ru']:
            if fallback_lang != language:
                fallback_field = f"{field_prefix}_{fallback_lang}"
                value = getattr(obj, fallback_field, None)
                if value:
                    break
    
    return value or _("Undefined")


def get_localized_vulnerability_field(vulnerability, field_name, language=None):
    """Get localized vulnerability field"""
    if language is None:
        language = get_language()[:2]
    
    localized_field = f"{field_name}_{language}"
    value = getattr(vulnerability, localized_field, None)
    
    if not value:
        # Fallback to other languages
        for fallback_lang in ['uk', 'en', 'ru']:
            if fallback_lang != language:
                fallback_field = f"{field_name}_{fallback_lang}"
                value = getattr(vulnerability, fallback_field, None)
                if value:
                    break
    
    return value or _("Undefined")


def get_report_translations(language=None):
    """Get all report translations for the given language"""
    if language is None:
        language = get_language()[:2]
    
    translations = {
        'uk': {
            'report_title': 'Звіт з оцінки ризиків',
            'executive_summary': 'Стислий звіт',
            'generated_on': 'Згенеровано',
            'report_period': 'Період звіту',
            'total_assets': 'Загальна кількість активів',
            'total_vulnerabilities': 'Загальна кількість вразливостей',
            'total_treatments': 'Загальна кількість заходів',
            'completion_rate': 'Рівень завершення (%)',
            'risk_levels_distribution': 'Розподіл рівнів ризику',
            'risk_level': 'Рівень ризику',
            'count': 'Кількість',
            'compliance_summary': 'Підсумок відповідності',
            'pci_dss_compliance': 'Відповідність PCI DSS',
            'iso27001_compliance': 'Відповідність ISO 27001',
            'framework_company_requirements': 'Вимоги фреймворків компанії',
            'local_company_requirements': 'Локальні вимоги компанії',
            'internal_company_requirements': 'Внутрішні вимоги компанії',
            'overall_compliance': 'Загальна відповідність (%)',
            'framework_type': 'Тип фреймворку',
            'version': 'Версія',
            'regulator': 'Регулятор',
            'source': 'Джерело',
            'controls_total': 'Всього контролів',
            'controls_completed': 'Завершено',
            'completion_percentage': 'Відсоток завершення',
            'controls_count': 'Кількість контролів',
            'total_requirements': 'Загальна кількість вимог',
            'total_controls': 'Загальна кількість контролів',
            'compliant_vulnerabilities': 'Відповідні вразливості',
            'total_vulnerabilities_comp': 'Загальна кількість вразливостей',
            'assets_by_criticality': 'Активи за критичністю',
        'assets_by_vulnerability': 'Активи за вразливістю',
            'criticality': 'Критичність',
            'high_risk_assets': 'Активи високого ризику',
            'overdue_treatments': 'Прострочені заходи',
            'asset': 'Актив',
            'treatment': 'Захід',
            'deadline': 'Термін',
            'metric': 'Метрика',
            'value': 'Значення',
            'full_report_title': 'Повний звіт з оцінки ризиків',
            'risk_assessment_report': 'Звіт з оцінки ризиків',
            'compliance_report': 'Звіт про відповідність',
            'generated_by': 'Створив',
            'risk_distribution': 'Розподіл ризиків',
            'by_risk_level': 'За рівнем ризику',
            'by_criticality': 'За критичністю',
            'percentage': 'Відсоток',
            'compliance_overview': 'Огляд відповідності',
            'non_compliant_vulnerabilities': 'Невідповідні вразливості',
            'compliance_rate': 'Рівень відповідності',
            'vulnerability': 'Вразливість',
            'due_date': 'Термін виконання',
            'status': 'Статус',
            'total_pages': 'Загальна кількість сторінок',
            'risk_summary': 'Резюме ризиків',
            'key_metrics': 'Ключові метрики',
            'treatments_completed': 'Заходи завершено',
            'top_risks': 'Найбільші ризики',
            'risk_details': 'Деталі ризиків',
            'top_10_treatment': 'ТОП-10 Заходів',
            'top_10_asset': 'ТОП-10 Активів',
            'top_10_vulnerability': 'ТОП-10 Вразливостей',
            'top_10_assets_with_acceptable_risk_config': 'ТОП-10 Активів з конфігурацією прийнятного ризику',
            'top_10_assets_without_acceptable_risk_config': 'ТОП-10 Активів без конфігурації прийнятного ризику',
            'acceptable_risk_analysis': 'Аналіз прийнятного ризику',
            'acceptable_risk_summary': 'Підсумок прийнятного ризику',
            'assets_with_acceptable_risk': 'Активи з прийнятним ризиком',
            'assets_without_acceptable_risk': 'Активи без прийнятного ризику',
            'exceeding_acceptable_risk': 'Перевищуючі прийнятний ризик',
            'current_criticality': 'Поточна критичність',
            'acceptable_risk_level': 'Рівень прийнятного ризику',
            'exceeds_acceptable': 'Перевищує прийнятний',
            'risk_value': 'Значення ризику',
            'compliance': 'Відповідність',
            'compliant': 'Відповідний',
            'non_compliant': 'Невідповідний',
            'compliance_gaps': 'Прогалини у відповідності',
            'gaps': 'Прогалини',
            'recommendations': 'Рекомендації',
            'priority_actions': 'Пріоритетні дії',
            'address_high_risk_vulnerabilities': 'Усунути вразливості високого ризику',
            'improve_compliance_rates': 'Покращити рівень відповідності',
            'implement_regular_assessments': 'Впровадити регулярні оцінки',
            'update_security_policies': 'Оновити політики безпеки',
            'preview_note': 'Це попередній перегляд звіту',
            'high_risk_count': 'Кількість високого ризику',
            'report_generated_by': 'Звіт створено користувачем',
            'company': 'Компанія',
            'multiple_companies': 'Кілька компаній',
            'asset_name': 'Назва активу',
            'asset_id': 'ID активу',
            'asset_type': 'Тип активу',
            'asset_group': 'Група активу',
            'asset_location': 'Розташування активу',
            'asset_description': 'Опис активу',
            'vulnerability_name': 'Назва вразливості',
            'vulnerability_status': 'Статус вразливості',
            'vulnerability_comment': 'Коментар до вразливості',
            'treatment_type': 'Тип обробки',
            'treatment_status': 'Статус обробки',
            'treatment_deadline': 'Термін виконання обробки',
            'treatment_responsible': 'Відповідальний за обробку',
            'treatment_details': 'Деталі обробки',
            'description_responsible': 'Опис та відповідальний',
            'assigned_to': 'Призначено',
            'vulnerability': 'Вразливість',
            'risk_impact': 'Вплив ризику',
            
            # Financial Analysis
            'financial_analysis': 'Фінансовий аналіз',
            'financial_overview': 'Фінансовий огляд',
            'total_cost': 'Загальна вартість',
            'potential_savings': 'Потенційна економія',
            'roi_percentage': 'ROI',
            'budget_allocated': 'Виділений бюджет',
            
            # Audit & Tracking
            'audit_tracking': 'Аудит та відстеження',
            'audit_trail': 'Аудит-трейл',
            'action': 'Дія',
            'timestamp': 'Часова мітка',
            'user': 'Користувач',
            
            # Monitoring & Performance
            'monitoring_performance': 'Моніторинг та продуктивність',
            'performance_metrics': 'Метрики продуктивності',
            'system_performance': 'Продуктивність системи',
            'response_time': 'Час відповіді',
            'availability': 'Доступність',
            'risk_mitigation_efficiency': 'Ефективність зменшення ризиків',
            'treatment_completion_rate': 'Рівень завершення обробки',
            'vulnerability_resolution_time': 'Час вирішення вразливостей',
            
            # Timeline & Trends specific translations
            'high_risk_vulnerabilities': 'Вразливості високого ризику',
            'medium_risk_vulnerabilities': 'Вразливості середнього ризику',
            'low_risk_vulnerabilities': 'Вразливості низького ризику',
            'treatment_in_progress': 'Заходи в процесі',
            'compliance_resolution_rate': 'Рівень вирішення відповідності',
            
            # Management & Governance
            'management_governance': 'Управління та корпоративне управління',
            'governance_overview': 'Огляд корпоративного управління',
            'companies_count': 'Кількість компаній',
            'compliance_status': 'Статус відповідності',
            'total_vulnerabilities': 'Загальна кількість вразливостей',
            'compliant_vulnerabilities': 'Відповідні вразливості',
            'asset_owners_count': 'Власники активів',
            'administrators_count': 'Адміністратори',
            'stakeholders_count': 'Зацікавлені сторони',
            
            # Timeline & Trends
            'timeline_trends': 'Часові рамки та тренди',
            'trend_overview': 'Огляд трендів',
            'vulnerability_trends': 'Тренди вразливостей',
            'compliance_trends': 'Тренди відповідності',
            
            # Risk Dependencies
            'risk_dependencies': 'Залежності ризиків',
            'dependency_overview': 'Огляд залежностей',
            'critical_dependencies': 'Критичні залежності',
            'interdependent_assets': 'Взаємозалежні активи',
            'cascading_risks': 'Каскадні ризики',
            
            # Threat Analysis
            'threat_analysis': 'Аналіз загроз',
            'threat_overview': 'Огляд загроз',
            'external_threats': 'Зовнішні загрози',
            'internal_threats': 'Внутрішні загрози',
            'emerging_threats': 'Нові загрози',
            
            # Residual Risk & Priority
            'residual_risk_priority': 'Залишковий ризик та пріоритет',
            'residual_overview': 'Огляд залишкового ризику',
            'residual_risk_level': 'Рівень залишкового ризику',
            'appetite_threshold': 'Поріг апетиту до ризику',
            'priority_level': 'Рівень пріоритету',
            'no_data': 'Немає даних',
            'risk_probability': 'Ймовірність ризику',
            'risk_level_name': 'Рівень ризику',
            'name': 'Назва',
            'description': 'Опис',
            'location': 'Розташування',
            'type': 'Тип',
            'group': 'Група',
            'responsible_person': 'Відповідальна особа',
            'registration_date': 'Дата реєстрації',
            'modification_date': 'Дата модифікації',
            'notes': 'Примітки',
            'comment': 'Коментар',
            
            # New Financial Analysis translations
            'financial_analysis': 'Фінансовий аналіз',
            'cost_benefit_analysis': 'Аналіз витрат і вигод',
            'roi_assessment': 'Оцінка рентабельності інвестицій',
            'budget_allocation': 'Розподіл бюджету',
            'total_cost': 'Загальна вартість',
            'potential_savings': 'Потенційна економія',
            'roi_percentage': 'ROI (%)',
            'budget_allocated': 'Виділено бюджету',
            
            # New Audit & Tracking translations
            'audit_trail': 'Аудиторський слід',
            'activity_logs': 'Журнали активності',
            'user_sessions': 'Сесії користувачів',
            'modification_history': 'Історія змін',
            'audit_entries': 'Записи аудиту',
            'user_activity': 'Активність користувачів',
            'session_duration': 'Тривалість сесії',
            'last_modified': 'Останні зміни',
            
            # New Monitoring & Performance translations
            'monitoring_dashboard': 'Панель моніторингу',
            'effectiveness_metrics': 'Метрики ефективності',
            'performance_indicators': 'Ключові показники ефективності',
            'treatment_effectiveness': 'Ефективність обробки',
            'system_performance': 'Продуктивність системи',
            'response_time': 'Час відгуку',
            'availability': 'Доступність',
            'success_rate': 'Рівень успішності',
            
            # New Management & Governance translations
            'responsibility_matrix': 'Матриця відповідальності',
            'approval_workflow': 'Робочий процес затвердження',
            'stakeholder_analysis': 'Аналіз зацікавлених сторін',
            'governance_structure': 'Структура управління',
            'responsible_party': 'Відповідальна сторона',
            'approval_status': 'Статус затвердження',
            'stakeholder_role': 'Роль зацікавленої сторони',
            'governance_level': 'Рівень управління',
            
            # New Timeline & Trends translations
            'timeline_analysis': 'Аналіз часових рамок',
            'trend_analysis': 'Аналіз тенденцій',
            'historical_data': 'Історичні дані',
            'deadline_tracking': 'Відстеження термінів',
            'time_period': 'Часовий період',
            'trend_direction': 'Напрямок тенденції',
            'historical_value': 'Історичне значення',
            'deadline_status': 'Статус терміну',
            
            # New Risk Dependencies translations
            'dependency_analysis': 'Аналіз залежностей',
            'impact_assessment': 'Оцінка впливу',
            'cascading_risks': 'Каскадні ризики',
            'interdependency_matrix': 'Матриця взаємозалежностей',
            'dependency_type': 'Тип залежності',
            'impact_level': 'Рівень впливу',
            'cascade_effect': 'Каскадний ефект',
            'interdependency': 'Взаємозалежність',
            
            # New Extended Asset Details translations
            'asset_criticality': 'Критичність активів',
            'asset_location': 'Розташування активів',
            'asset_owners': 'Власники активів',
            'asset_lifecycle': 'Життєвий цикл активів',
            'criticality_level': 'Рівень критичності',
            'physical_location': 'Фізичне розташування',
            'asset_owner': 'Власник активу',
            'lifecycle_stage': 'Етап життєвого циклу',
            
            # New Threat Analysis translations
            'threat_analysis': 'Аналіз загроз',
            'threat_scenarios': 'Сценарії загроз',
            'probability_analysis': 'Аналіз ймовірності',
            'threat_landscape': 'Ландшафт загроз',
            'threat_type': 'Тип загрози',
            'threat_scenario': 'Сценарій загрози',
            'probability_level': 'Рівень ймовірності',
            'threat_source': 'Джерело загрози',
            
            # New Residual Risk & Priority translations
            'residual_risk_analysis': 'Аналіз залишкового ризику',
            'risk_appetite': 'Ризик-апетит',
            'priority_matrix': 'Матриця пріоритетів',
            'resource_allocation': 'Розподіл ресурсів',
            'residual_risk_level': 'Рівень залишкового ризику',
            'appetite_threshold': 'Поріг апетиту',
            'priority_level': 'Рівень пріоритету',
            'allocated_resources': 'Виділені ресурси',
            
            # Section titles
            'asset_details': 'Деталі активів',
            'vulnerability_details': 'Деталі вразливостей',
            'vulnerability_tables': 'Таблиці вразливостей',
    
            'asset_tables': 'Таблиці активів',
            'unknown': 'Невідомо',
            'no_description': 'Без опису',
            'treatment_details': 'Деталі обробки',
            
            # New vulnerability table columns
            'threat': 'Загроза',
            'probability_impact': 'Ймовірність/Вплив',
            'impact': 'Вплив',
            'risk_mitigation_controls': 'Контрольні заходи зменшення ризику',
            'no_threats': 'Немає загроз',
            'not_applicable': 'Не застосовується',
            
            # Access Risk Summary translations
            'access_risk_summary': 'Зведення ризиків доступів',
            'total_active_accesses': 'Кількість активних доступів',
            'overdue_reviews': 'Прострочені access reviews',
            'overdue_access_reviews': 'Прострочені access reviews',
            'third_party_access': 'Доступи третіх сторін',
            'third_party_access_higher_risk': 'Доступи третіх сторін (higher risk)',
            'access_without_last_review': 'Доступи без останнього review',
            'privileged_access': 'Привілейовані доступи',
            'privileged_access_high_privileges': 'Привілейовані доступи - Доступи з високими правами',
            'privileged_access_without_regular_reviews': 'Привілейовані доступи без регулярних reviews',
            'access_summary_statistics': 'Статистика доступів',
            'access_description': 'Опис доступу',
            'last_review': 'Останній review',
            'reviewed_by': 'Перевірено',
            'days_overdue': 'Днів прострочено',
            'start_date': 'Дата початку',
            'days_since_start': 'Днів з початку',
            'access_right': 'Право доступу',
            'roles': 'Ролі',
            'never': 'Ніколи',
            
            # SIEM translations
            'siem': 'SIEM',
            'total_active_agents_wazuh': 'Кількість активних агентів Wazuh',
            'agents_with_problems': 'Агенти з проблемами',
            'agents_with_problems_inactive_outdated': 'Агенти з проблемами (inactive, outdated)',
            'coverage_rate': 'Coverage rate',
            'coverage_rate_monitoring': 'Coverage rate (скільки активів під моніторингом)',
            'monitored_assets': 'Активи під моніторингом',
            'wazuh_agents_summary': 'Зведення агентів Wazuh',
            'agent_name': 'Назва агента',
            'agent_id': 'ID агента',
            'agent_ip': 'IP агента',
            'agent_version': 'Версія',
            'last_seen': 'Останній раз бачено',
            'problems': 'Проблеми',
            'days_since_seen': 'Днів з моменту останнього з\'єднання',
            'threat_detection_trends': 'Тенденції виявлення загроз',
            'alert_count': 'Кількість алертів',
            'not_specified': 'Не вказано'
        },
        'en': {
            'report_title': 'Risk Assessment Report',
            'executive_summary': 'Executive Summary',
            'generated_on': 'Generated on',
            'report_period': 'Report period',
            'total_assets': 'Total Assets',
            'total_vulnerabilities': 'Total Vulnerabilities',
            'total_treatments': 'Total Treatments',
            'completion_rate': 'Completion Rate (%)',
            'risk_levels_distribution': 'Risk Levels Distribution',
            'risk_level': 'Risk Level',
            'count': 'Count',
            'compliance_summary': 'Compliance Summary',
            'pci_dss_compliance': 'PCI DSS Compliance',
            'iso27001_compliance': 'ISO 27001 Compliance',
            'framework_company_requirements': 'Framework Company Requirements',
            'local_company_requirements': 'Local Company Requirements',
            'internal_company_requirements': 'Internal Company Requirements',
            'overall_compliance': 'Overall Compliance (%)',
            'framework_type': 'Framework Type',
            'version': 'Version',
            'regulator': 'Regulator',
            'source': 'Source',
            'controls_total': 'Total Controls',
            'controls_completed': 'Completed',
            'completion_percentage': 'Completion %',
            'controls_count': 'Controls Count',
            'total_requirements': 'Total Requirements',
            'total_controls': 'Total Controls',
            'compliant_vulnerabilities': 'Compliant Vulnerabilities',
            'total_vulnerabilities_comp': 'Total Vulnerabilities',
            'assets_by_criticality': 'Assets by Criticality',
        'assets_by_vulnerability': 'Assets by Vulnerability',
            'criticality': 'Criticality',
            'high_risk_assets': 'High Risk Assets',
            'overdue_treatments': 'Overdue Treatments',
            'asset': 'Asset',
            'treatment': 'Treatment',
            'deadline': 'Deadline',
            'metric': 'Metric',
            'value': 'Value',
            'full_report_title': 'Full Risk Assessment Report',
            'risk_assessment_report': 'Risk Assessment Report',
            'compliance_report': 'Compliance Report',
            'generated_by': 'Generated by',
            'risk_distribution': 'Risk Distribution',
            'by_risk_level': 'By Risk Level',
            'by_criticality': 'By Criticality',
            'percentage': 'Percentage',
            'compliance_overview': 'Compliance Overview',
            'non_compliant_vulnerabilities': 'Non-Compliant Vulnerabilities',
            'compliance_rate': 'Compliance Rate',
            'vulnerability': 'Vulnerability',
            'vulnerability_status': 'Vulnerability Status',
            'treatment_type': 'Treatment Type',
            'treatment_status': 'Treatment Status',
            'treatment_details': 'Treatment Details',
            'description_responsible': 'Description & Responsible',
            'assigned_to': 'Assigned To',
            'due_date': 'Due Date',
            
            # Financial Analysis
            'financial_analysis': 'Financial Analysis',
            'financial_overview': 'Financial Overview',
            'total_cost': 'Total Cost',
            'potential_savings': 'Potential Savings',
            'roi_percentage': 'ROI',
            'budget_allocated': 'Budget Allocated',
            
            # Audit & Tracking
            'audit_tracking': 'Audit & Tracking',
            'audit_trail': 'Audit Trail',
            'action': 'Action',
            'timestamp': 'Timestamp',
            'user': 'User',
            
            # Monitoring & Performance
            'monitoring_performance': 'Monitoring & Performance',
            'performance_metrics': 'Performance Metrics',
            'system_performance': 'System Performance',
            'response_time': 'Response Time',
            'availability': 'Availability',
            'risk_mitigation_efficiency': 'Risk Mitigation Efficiency',
            'treatment_completion_rate': 'Treatment Completion Rate',
            'vulnerability_resolution_time': 'Vulnerability Resolution Time',
            
            # Timeline & Trends specific translations
            'high_risk_vulnerabilities': 'High Risk Vulnerabilities',
            'medium_risk_vulnerabilities': 'Medium Risk Vulnerabilities',
            'low_risk_vulnerabilities': 'Low Risk Vulnerabilities',
            'treatment_in_progress': 'Treatments In Progress',
            'compliance_resolution_rate': 'Compliance Resolution Rate',
            
            # Management & Governance
            'management_governance': 'Management & Governance',
            'governance_overview': 'Governance Overview',
            'companies_count': 'Companies Count',
            'compliance_status': 'Compliance Status',
            'total_vulnerabilities': 'Total Vulnerabilities',
            'compliant_vulnerabilities': 'Compliant Vulnerabilities',
            'asset_owners_count': 'Asset Owners',
            'administrators_count': 'Administrators',
            'stakeholders_count': 'Stakeholders',
            
            # Timeline & Trends
            'timeline_trends': 'Timeline & Trends',
            'trend_overview': 'Trend Overview',
            'vulnerability_trends': 'Vulnerability Trends',
            'compliance_trends': 'Compliance Trends',
            
            # Risk Dependencies
            'risk_dependencies': 'Risk Dependencies',
            'dependency_overview': 'Dependency Overview',
            'critical_dependencies': 'Critical Dependencies',
            'interdependent_assets': 'Interdependent Assets',
            'cascading_risks': 'Cascading Risks',
            
            # Threat Analysis
            'threat_analysis': 'Threat Analysis',
            'threat_overview': 'Threat Overview',
            'external_threats': 'External Threats',
            'internal_threats': 'Internal Threats',
            'emerging_threats': 'Emerging Threats',
            
            # Residual Risk & Priority
            'residual_risk_priority': 'Residual Risk & Priority',
            'residual_overview': 'Residual Risk Overview',
            'residual_risk_level': 'Residual Risk Level',
            'appetite_threshold': 'Risk Appetite Threshold',
            'priority_level': 'Priority Level',
            'no_data': 'No data',
            'status': 'Status',
            'total_pages': 'Total Pages',
            'risk_summary': 'Risk Summary',
            'key_metrics': 'Key Metrics',
            'treatments_completed': 'Treatments Completed',
            'top_risks': 'TOP-10 Risks',
            'risk_details': 'Risk Details',
            'top_10_treatment': 'TOP-10 Treatment',
            'top_10_asset': 'TOP-10 Asset',
            'top_10_vulnerability': 'TOP-10 Vulnerability',
            'top_10_assets_with_acceptable_risk_config': 'TOP-10 Assets with Acceptable Risk Configuration',
            'top_10_assets_without_acceptable_risk_config': 'TOP-10 Assets without Acceptable Risk Configuration',
            'acceptable_risk_analysis': 'Acceptable Risk Analysis',
            'acceptable_risk_summary': 'Acceptable Risk Summary',
            'assets_with_acceptable_risk': 'Assets with Acceptable Risk',
            'assets_without_acceptable_risk': 'Assets without Acceptable Risk',
            'exceeding_acceptable_risk': 'Exceeding Acceptable Risk',
            'current_criticality': 'Current Criticality',
            'acceptable_risk_level': 'Acceptable Risk Level',
            'exceeds_acceptable': 'Exceeds Acceptable',
            'risk_value': 'Risk Value',
            'compliance': 'Compliance',
            'compliant': 'Compliant',
            'non_compliant': 'Non-Compliant',
            'compliance_gaps': 'Compliance Gaps',
            'gaps': 'Gaps',
            'recommendations': 'Recommendations',
            'priority_actions': 'Priority Actions',
            'address_high_risk_vulnerabilities': 'Address high-risk vulnerabilities',
            'improve_compliance_rates': 'Improve compliance rates',
            'implement_regular_assessments': 'Implement regular assessments',
            'update_security_policies': 'Update security policies',
            'preview_note': 'This is a report preview',
            'high_risk_count': 'High Risk Count',
            'report_generated_by': 'Report generated by',
            'company': 'Company',
            'multiple_companies': 'Multiple Companies',
            'asset_name': 'Asset Name',
            'asset_id': 'Asset ID',
            'asset_type': 'Asset Type',
            'asset_group': 'Asset Group',
            'asset_location': 'Asset Location',
            'asset_description': 'Asset Description',
            'vulnerability_name': 'Vulnerability Name',
            'vulnerability_status': 'Vulnerability Status',
            'vulnerability_comment': 'Vulnerability Comment',
            'treatment_type': 'Treatment Type',
            'treatment_status': 'Treatment Status',
            'treatment_deadline': 'Treatment Deadline',
            'treatment_responsible': 'Treatment Responsible',
            'assigned_to': 'Assigned To',
            'vulnerability': 'Vulnerability',
            'risk_impact': 'Risk Impact',
            'risk_probability': 'Risk Probability',
            'risk_level_name': 'Risk Level',
            'name': 'Name',
            'description': 'Description',
            'location': 'Location',
            'type': 'Type',
            'group': 'Group',
            'responsible_person': 'Responsible Person',
            'registration_date': 'Registration Date',
            'modification_date': 'Modification Date',
            'notes': 'Notes',
            'comment': 'Comment',
            
            # New Financial Analysis translations
            'financial_analysis': 'Financial Analysis',
            'cost_benefit_analysis': 'Cost-Benefit Analysis',
            'roi_assessment': 'ROI Assessment',
            'budget_allocation': 'Budget Allocation',
            'total_cost': 'Total Cost',
            'potential_savings': 'Potential Savings',
            'roi_percentage': 'ROI (%)',
            'budget_allocated': 'Budget Allocated',
            
            # New Audit & Tracking translations
            'audit_trail': 'Audit Trail',
            'activity_logs': 'Activity Logs',
            'user_sessions': 'User Sessions',
            'modification_history': 'Modification History',
            'audit_entries': 'Audit Entries',
            'user_activity': 'User Activity',
            'session_duration': 'Session Duration',
            'last_modified': 'Last Modified',
            
            # New Monitoring & Performance translations
            'monitoring_dashboard': 'Monitoring Dashboard',
            'effectiveness_metrics': 'Effectiveness Metrics',
            'performance_indicators': 'Key Performance Indicators',
            'treatment_effectiveness': 'Treatment Effectiveness',
            'system_performance': 'System Performance',
            'response_time': 'Response Time',
            'availability': 'Availability',
            'success_rate': 'Success Rate',
            
            # New Management & Governance translations
            'responsibility_matrix': 'Responsibility Matrix',
            'approval_workflow': 'Approval Workflow',
            'stakeholder_analysis': 'Stakeholder Analysis',
            'governance_structure': 'Governance Structure',
            'responsible_party': 'Responsible Party',
            'approval_status': 'Approval Status',
            'stakeholder_role': 'Stakeholder Role',
            'governance_level': 'Governance Level',
            
            # New Timeline & Trends translations
            'timeline_analysis': 'Timeline Analysis',
            'trend_analysis': 'Trend Analysis',
            'historical_data': 'Historical Data',
            'deadline_tracking': 'Deadline Tracking',
            'time_period': 'Time Period',
            'trend_direction': 'Trend Direction',
            'historical_value': 'Historical Value',
            'deadline_status': 'Deadline Status',
            
            # New Risk Dependencies translations
            'dependency_analysis': 'Dependency Analysis',
            'impact_assessment': 'Impact Assessment',
            'cascading_risks': 'Cascading Risks',
            'interdependency_matrix': 'Interdependency Matrix',
            'dependency_type': 'Dependency Type',
            'impact_level': 'Impact Level',
            'cascade_effect': 'Cascade Effect',
            'interdependency': 'Interdependency',
            
            # New Extended Asset Details translations
            'asset_criticality': 'Asset Criticality',
            'asset_location': 'Asset Location',
            'asset_owners': 'Asset Owners',
            'asset_lifecycle': 'Asset Lifecycle',
            'criticality_level': 'Criticality Level',
            'physical_location': 'Physical Location',
            'asset_owner': 'Asset Owner',
            'lifecycle_stage': 'Lifecycle Stage',
            
            # New Threat Analysis translations
            'threat_analysis': 'Threat Analysis',
            'threat_scenarios': 'Threat Scenarios',
            'probability_analysis': 'Probability Analysis',
            'threat_landscape': 'Threat Landscape',
            'threat_type': 'Threat Type',
            'threat_scenario': 'Threat Scenario',
            'probability_level': 'Probability Level',
            'threat_source': 'Threat Source',
            
            # New Residual Risk & Priority translations
            'residual_risk_analysis': 'Residual Risk Analysis',
            'risk_appetite': 'Risk Appetite',
            'priority_matrix': 'Priority Matrix',
            'resource_allocation': 'Resource Allocation',
            'residual_risk_level': 'Residual Risk Level',
            'appetite_threshold': 'Appetite Threshold',
            'priority_level': 'Priority Level',
            'allocated_resources': 'Allocated Resources',
            
            # Section titles
            'asset_details': 'Asset Details',
            'vulnerability_details': 'Vulnerability Details',
            'vulnerability_tables': 'Vulnerability Tables',
    
            'asset_tables': 'Asset Tables',
            'unknown': 'Unknown',
            'no_description': 'No description',
            'treatment_details': 'Treatment Details',
            
            # New vulnerability table columns
            'threat': 'Threat',
            'probability_impact': 'Probability/Impact',
            'impact': 'Impact',
            'risk_mitigation_controls': 'Risk Mitigation Controls',
            'no_threats': 'No threats',
            'not_applicable': 'Not applicable',
            
            # Access Risk Summary translations
            'access_risk_summary': 'Access Risk Summary',
            'total_active_accesses': 'Total Active Accesses',
            'overdue_reviews': 'Overdue Access Reviews',
            'overdue_access_reviews': 'Overdue Access Reviews',
            'third_party_access': 'Third-Party Access',
            'third_party_access_higher_risk': 'Third-Party Access (Higher Risk)',
            'access_without_last_review': 'Access Without Last Review',
            'privileged_access': 'Privileged Access',
            'privileged_access_high_privileges': 'Privileged Access - High Privileges',
            'privileged_access_without_regular_reviews': 'Privileged Access Without Regular Reviews',
            'access_summary_statistics': 'Access Summary Statistics',
            'access_description': 'Access Description',
            'last_review': 'Last Review',
            'reviewed_by': 'Reviewed By',
            'days_overdue': 'Days Overdue',
            'start_date': 'Start Date',
            'days_since_start': 'Days Since Start',
            'access_right': 'Access Right',
            'roles': 'Roles',
            'never': 'Never',
            
            # SIEM translations
            'siem': 'SIEM',
            'total_active_agents_wazuh': 'Total Active Wazuh Agents',
            'agents_with_problems': 'Agents with Problems',
            'agents_with_problems_inactive_outdated': 'Agents with Problems (inactive, outdated)',
            'coverage_rate': 'Coverage Rate',
            'coverage_rate_monitoring': 'Coverage Rate (Assets Under Monitoring)',
            'monitored_assets': 'Monitored Assets',
            'wazuh_agents_summary': 'Wazuh Agents Summary',
            'agent_name': 'Agent Name',
            'agent_id': 'Agent ID',
            'agent_ip': 'Agent IP',
            'agent_version': 'Version',
            'last_seen': 'Last Seen',
            'problems': 'Problems',
            'days_since_seen': 'Days Since Seen',
            'threat_detection_trends': 'Threat Detection Trends',
            'alert_count': 'Alert Count',
            'not_specified': 'Not specified'
        },
        'ru': {
            'report_title': 'Отчет по оценке рисков',
            'executive_summary': 'Исполнительное резюме',
            'generated_on': 'Сгенерировано',
            'report_period': 'Период отчета',
            'total_assets': 'Общее количество активов',
            'total_vulnerabilities': 'Общее количество уязвимостей',
            'total_treatments': 'Общее количество мер',
            'completion_rate': 'Уровень завершения (%)',
            'risk_levels_distribution': 'Распределение уровней риска',
            'risk_level': 'Уровень риска',
            'count': 'Количество',
            'compliance_summary': 'Сводка соответствия',
            'pci_dss_compliance': 'Соответствие PCI DSS',
            'iso27001_compliance': 'Соответствие ISO 27001',
            'framework_company_requirements': 'Требования фреймворков компании',
            'local_company_requirements': 'Локальные требования компании',
            'internal_company_requirements': 'Внутренние требования компании',
            'overall_compliance': 'Общее соответствие (%)',
            'framework_type': 'Тип фреймворка',
            'version': 'Версия',
            'regulator': 'Регулятор',
            'source': 'Источник',
            'controls_total': 'Всего контролей',
            'controls_completed': 'Завершено',
            'completion_percentage': 'Процент завершения',
            'controls_count': 'Количество контролей',
            'total_requirements': 'Общее количество требований',
            'total_controls': 'Общее количество контролей',
            'compliant_vulnerabilities': 'Соответствующие уязвимости',
            'total_vulnerabilities_comp': 'Общее количество уязвимостей',
            'assets_by_criticality': 'Активы по критичности',
        'assets_by_vulnerability': 'Активы по уязвимости',
            'criticality': 'Критичность',
            'high_risk_assets': 'Активы высокого риска',
            'overdue_treatments': 'Просроченные меры',
            'asset': 'Актив',
            'treatment': 'Мера',
            'deadline': 'Срок',
            'metric': 'Метрика',
            'value': 'Значение',
            'full_report_title': 'Полный отчет по оценке рисков',
            'risk_assessment_report': 'Отчет по оценке рисков',
            'compliance_report': 'Отчет о соответствии',
            'generated_by': 'Создал',
            'risk_distribution': 'Распределение рисков',
            'by_risk_level': 'По уровню риска',
            'by_criticality': 'По критичности',
            'percentage': 'Процент',
            'compliance_overview': 'Обзор соответствия',
            'non_compliant_vulnerabilities': 'Несоответствующие уязвимости',
            'compliance_rate': 'Уровень соответствия',
            'vulnerability': 'Уязвимость',
            'vulnerability_status': 'Статус уязвимости',
            'treatment_type': 'Тип меры',
            'treatment_status': 'Статус меры',
            'treatment_details': 'Детали лечения',
            'description_responsible': 'Описание и ответственный',
            'assigned_to': 'Назначено',
            'due_date': 'Срок выполнения',
            
            # Financial Analysis
            'financial_analysis': 'Финансовый анализ',
            'financial_overview': 'Финансовый обзор',
            'total_cost': 'Общая стоимость',
            'potential_savings': 'Потенциальная экономия',
            'roi_percentage': 'ROI',
            'budget_allocated': 'Выделенный бюджет',
            
            # Audit & Tracking
            'audit_tracking': 'Аудит и отслеживание',
            'audit_trail': 'Аудит-трейл',
            'action': 'Действие',
            'timestamp': 'Временная метка',
            'user': 'Пользователь',
            
            # Monitoring & Performance
            'monitoring_performance': 'Мониторинг и производительность',
            'performance_metrics': 'Метрики производительности',
            'system_performance': 'Производительность системы',
            'response_time': 'Время отклика',
            'availability': 'Доступность',
            'risk_mitigation_efficiency': 'Эффективность снижения рисков',
            'treatment_completion_rate': 'Уровень завершения лечения',
            'vulnerability_resolution_time': 'Время разрешения уязвимостей',
            
            # Timeline & Trends specific translations
            'high_risk_vulnerabilities': 'Уязвимости высокого риска',
            'medium_risk_vulnerabilities': 'Уязвимости среднего риска',
            'low_risk_vulnerabilities': 'Уязвимости низкого риска',
            'treatment_in_progress': 'Меры в процессе',
            'compliance_resolution_rate': 'Уровень разрешения соответствия',
            
            # Management & Governance
            'management_governance': 'Управление и корпоративное управление',
            'governance_overview': 'Обзор корпоративного управления',
            'companies_count': 'Количество компаний',
            'compliance_status': 'Статус соответствия',
            'total_vulnerabilities': 'Общее количество уязвимостей',
            'compliant_vulnerabilities': 'Соответствующие уязвимости',
            'asset_owners_count': 'Владельцы активов',
            'administrators_count': 'Администраторы',
            'stakeholders_count': 'Заинтересованные стороны',
            
            # Timeline & Trends
            'timeline_trends': 'Временные рамки и тренды',
            'trend_overview': 'Обзор трендов',
            'vulnerability_trends': 'Тренды уязвимостей',
            'compliance_trends': 'Тренды соответствия',
            
            # Risk Dependencies
            'risk_dependencies': 'Зависимости рисков',
            'dependency_overview': 'Обзор зависимостей',
            'critical_dependencies': 'Критические зависимости',
            'interdependent_assets': 'Взаимозависимые активы',
            'cascading_risks': 'Каскадные риски',
            
            # Threat Analysis
            'threat_analysis': 'Анализ угроз',
            'threat_overview': 'Обзор угроз',
            'external_threats': 'Внешние угрозы',
            'internal_threats': 'Внутренние угрозы',
            'emerging_threats': 'Новые угрозы',
            
            # Residual Risk & Priority
            'residual_risk_priority': 'Остаточный риск и приоритет',
            'residual_overview': 'Обзор остаточного риска',
            'residual_risk_level': 'Уровень остаточного риска',
            'appetite_threshold': 'Порог аппетита к риску',
            'priority_level': 'Уровень приоритета',
            'no_data': 'Нет данных',
            'status': 'Статус',
            'total_pages': 'Общее количество страниц',
            'risk_summary': 'Резюме рисков',
            'key_metrics': 'Ключевые метрики',
            'treatments_completed': 'Меры завершены',
            'top_risks': 'Основные риски',
            'risk_details': 'Детали рисков',
            'top_10_treatment': 'ТОП-10 Мер',
            'top_10_asset': 'ТОП-10 Активов',
            'top_10_vulnerability': 'ТОП-10 Уязвимостей',
            'top_10_assets_with_acceptable_risk_config': 'ТОП-10 Активов с конфигурацией приемлемого риска',
            'top_10_assets_without_acceptable_risk_config': 'ТОП-10 Активов без конфигурации приемлемого риска',
            'acceptable_risk_analysis': 'Анализ приемлемого риска',
            'acceptable_risk_summary': 'Сводка приемлемого риска',
            'assets_with_acceptable_risk': 'Активы с приемлемым риском',
            'assets_without_acceptable_risk': 'Активы без приемлемого риска',
            'exceeding_acceptable_risk': 'Превышающие приемлемый риск',
            'current_criticality': 'Текущая критичность',
            'acceptable_risk_level': 'Уровень приемлемого риска',
            'exceeds_acceptable': 'Превышает приемлемый',
            'risk_value': 'Значение риска',
            'compliance': 'Соответствие',
            'compliant': 'Соответствующий',
            'non_compliant': 'Несоответствующий',
            'compliance_gaps': 'Пробелы в соответствии',
            'gaps': 'Пробелы',
            'recommendations': 'Рекомендации',
            'priority_actions': 'Приоритетные действия',
            'address_high_risk_vulnerabilities': 'Устранить уязвимости высокого риска',
            'improve_compliance_rates': 'Улучшить уровень соответствия',
            'implement_regular_assessments': 'Внедрить регулярные оценки',
            'update_security_policies': 'Обновить политики безопасности',
            'preview_note': 'Это предварительный просмотр отчета',
            'high_risk_count': 'Количество высокого риска',
            'report_generated_by': 'Отчет создан пользователем',
            'company': 'Компания',
            'multiple_companies': 'Несколько компаний',
            'asset_name': 'Название актива',
            'asset_id': 'ID актива',
            'asset_type': 'Тип актива',
            'asset_group': 'Группа актива',
            'asset_location': 'Расположение актива',
            'asset_description': 'Описание актива',
            'vulnerability_name': 'Название уязвимости',
            'vulnerability_status': 'Статус уязвимости',
            'vulnerability_comment': 'Комментарий к уязвимости',
            'treatment_type': 'Тип меры',
            'treatment_status': 'Статус меры',
            'treatment_deadline': 'Срок выполнения меры',
            'treatment_responsible': 'Ответственный за меру',
            'assigned_to': 'Назначено',
            'vulnerability': 'Уязвимость',
            'risk_impact': 'Воздействие риска',
            'risk_probability': 'Вероятность риска',
            'risk_level_name': 'Уровень риска',
            'name': 'Название',
            'description': 'Описание',
            'location': 'Расположение',
            'type': 'Тип',
            'group': 'Группа',
            'responsible_person': 'Ответственное лицо',
            'registration_date': 'Дата регистрации',
            'modification_date': 'Дата модификации',
            'notes': 'Примечания',
            'comment': 'Комментарий',
            
            # New Financial Analysis translations
            'financial_analysis': 'Финансовый анализ',
            'cost_benefit_analysis': 'Анализ затрат и выгод',
            'roi_assessment': 'Оценка рентабельности инвестиций',
            'budget_allocation': 'Распределение бюджета',
            'total_cost': 'Общая стоимость',
            'potential_savings': 'Потенциальная экономия',
            'roi_percentage': 'ROI (%)',
            'budget_allocated': 'Выделено бюджета',
            
            # New Audit & Tracking translations
            'audit_trail': 'Аудиторский след',
            'activity_logs': 'Журналы активности',
            'user_sessions': 'Сессии пользователей',
            'modification_history': 'История изменений',
            'audit_entries': 'Записи аудита',
            'user_activity': 'Активность пользователей',
            'session_duration': 'Длительность сессии',
            'last_modified': 'Последние изменения',
            
            # New Monitoring & Performance translations
            'monitoring_dashboard': 'Панель мониторинга',
            'effectiveness_metrics': 'Метрики эффективности',
            'performance_indicators': 'Ключевые показатели эффективности',
            'treatment_effectiveness': 'Эффективность лечения',
            'system_performance': 'Производительность системы',
            'response_time': 'Время отклика',
            'availability': 'Доступность',
            'success_rate': 'Уровень успешности',
            
            # New Management & Governance translations
            'responsibility_matrix': 'Матрица ответственности',
            'approval_workflow': 'Рабочий процесс утверждения',
            'stakeholder_analysis': 'Анализ заинтересованных сторон',
            'governance_structure': 'Структура управления',
            'responsible_party': 'Ответственная сторона',
            'approval_status': 'Статус утверждения',
            'stakeholder_role': 'Роль заинтересованной стороны',
            'governance_level': 'Уровень управления',
            
            # New Timeline & Trends translations
            'timeline_analysis': 'Анализ временных рамок',
            'trend_analysis': 'Анализ тенденций',
            'historical_data': 'Исторические данные',
            'deadline_tracking': 'Отслеживание сроков',
            'time_period': 'Временной период',
            'trend_direction': 'Направление тенденции',
            'historical_value': 'Историческое значение',
            'deadline_status': 'Статус срока',
            
            # New Risk Dependencies translations
            'dependency_analysis': 'Анализ зависимостей',
            'impact_assessment': 'Оценка воздействия',
            'cascading_risks': 'Каскадные риски',
            'interdependency_matrix': 'Матрица взаимозависимостей',
            'dependency_type': 'Тип зависимости',
            'impact_level': 'Уровень воздействия',
            'cascade_effect': 'Каскадный эффект',
            'interdependency': 'Взаимозависимость',
            
            # New Extended Asset Details translations
            'asset_criticality': 'Критичность активов',
            'asset_location': 'Расположение активов',
            'asset_owners': 'Владельцы активов',
            'asset_lifecycle': 'Жизненный цикл активов',
            'criticality_level': 'Уровень критичности',
            'physical_location': 'Физическое расположение',
            'asset_owner': 'Владелец актива',
            'lifecycle_stage': 'Этап жизненного цикла',
            
            # New Threat Analysis translations
            'threat_analysis': 'Анализ угроз',
            'threat_scenarios': 'Сценарии угроз',
            'probability_analysis': 'Анализ вероятности',
            'threat_landscape': 'Ландшафт угроз',
            'threat_type': 'Тип угрозы',
            'threat_scenario': 'Сценарий угрозы',
            'probability_level': 'Уровень вероятности',
            'threat_source': 'Источник угрозы',
            
            # New Residual Risk & Priority translations
            'residual_risk_analysis': 'Анализ остаточного риска',
            'risk_appetite': 'Риск-аппетит',
            'priority_matrix': 'Матрица приоритетов',
            'resource_allocation': 'Распределение ресурсов',
            'residual_risk_level': 'Уровень остаточного риска',
            'appetite_threshold': 'Порог аппетита',
            'priority_level': 'Уровень приоритета',
            'allocated_resources': 'Выделенные ресурсы',
            
            # Section titles
            'asset_details': 'Детали активов',
            'vulnerability_details': 'Детали уязвимостей',
            'vulnerability_tables': 'Таблицы уязвимостей',
    
            'asset_tables': 'Таблицы активов',
            'unknown': 'Неизвестно',
            'no_description': 'Без описания',
            'treatment_details': 'Детали лечения',
            
            # New vulnerability table columns
            'threat': 'Угроза',
            'probability_impact': 'Вероятность/Воздействие',
            'impact': 'Воздействие',
            'risk_mitigation_controls': 'Контрольные меры снижения риска',
            'no_threats': 'Нет угроз',
            'not_applicable': 'Не применимо',
            
            # Access Risk Summary translations
            'access_risk_summary': 'Сводка рисков доступов',
            'total_active_accesses': 'Количество активных доступов',
            'overdue_reviews': 'Просроченные access reviews',
            'overdue_access_reviews': 'Просроченные access reviews',
            'third_party_access': 'Доступы третьих сторон',
            'third_party_access_higher_risk': 'Доступы третьих сторон (higher risk)',
            'access_without_last_review': 'Доступы без последнего review',
            'privileged_access': 'Привилегированные доступы',
            'privileged_access_high_privileges': 'Привилегированные доступы - Доступы с высокими правами',
            'privileged_access_without_regular_reviews': 'Привилегированные доступы без регулярных reviews',
            'access_summary_statistics': 'Статистика доступов',
            'access_description': 'Описание доступа',
            'last_review': 'Последний review',
            'reviewed_by': 'Проверено',
            'days_overdue': 'Дней просрочено',
            'start_date': 'Дата начала',
            'days_since_start': 'Дней с начала',
            'access_right': 'Право доступа',
            'roles': 'Роли',
            'never': 'Никогда',
            
            # SIEM translations
            'siem': 'SIEM',
            'total_active_agents_wazuh': 'Количество активных агентов Wazuh',
            'agents_with_problems': 'Агенты с проблемами',
            'agents_with_problems_inactive_outdated': 'Агенты с проблемами (inactive, outdated)',
            'coverage_rate': 'Coverage rate',
            'coverage_rate_monitoring': 'Coverage rate (сколько активов под мониторингом)',
            'monitored_assets': 'Активы под мониторингом',
            'wazuh_agents_summary': 'Сводка агентов Wazuh',
            'agent_name': 'Название агента',
            'agent_id': 'ID агента',
            'agent_ip': 'IP агента',
            'agent_version': 'Версия',
            'last_seen': 'Последний раз видели',
            'problems': 'Проблемы',
            'days_since_seen': 'Дней с момента последнего соединения',
            'threat_detection_trends': 'Тенденции выявления угроз',
            'alert_count': 'Количество алертов',
            'not_specified': 'Не указано'
        }
    }
    
    return translations.get(language, translations['uk'])


def get_localized_report_type(report_type, language=None):
    """Get localized report type name"""
    if language is None:
        language = get_language()[:2]
    
    report_types = {
        'uk': {
            'full': 'повний_звіт',
            'summary': 'резюме',
            'compliance': 'відповідність'
        },
        'en': {
            'full': 'full_report',
            'summary': 'summary',
            'compliance': 'compliance'
        },
        'ru': {
            'full': 'полный_отчет',
            'summary': 'резюме',
            'compliance': 'соответствие'
        }
    }
    
    return report_types.get(language, report_types['en']).get(report_type, report_type)


def format_localized_date(date_obj, language=None):
    """Format date according to the selected language"""
    if language is None:
        language = get_language()[:2]
    
    if language == 'uk':
        return date_obj.strftime('%d.%m.%Y %H:%M')
    elif language == 'ru':
        return date_obj.strftime('%d.%m.%Y %H:%M')
    else:  # English
        return date_obj.strftime('%m/%d/%Y %I:%M %p')


@login_required
@user_passes_test(check_risk_report_access)
@require_http_methods(["POST"])
def preview_risk_report(request):
    """Generate HTML preview of risk assessment report for modal display"""
    
    # Check permissions
    if not check_risk_report_access(request.user, 'view'):
        raise PermissionDenied(_("You don't have permission to preview reports"))
    
    try:
        # Parse request data
        report_type = request.POST.get('reportType', 'summary')
        report_format = request.POST.get('reportFormat', 'word')
        report_language = request.POST.get('reportLanguage', get_language()[:2])
        
        # Parse selected sections from the configuration matrix
        selected_sections_json = request.POST.get('selectedSections', '{}')
        try:
            selected_sections = json.loads(selected_sections_json) if selected_sections_json else {}
        except json.JSONDecodeError:
            logger.warning(f"Invalid JSON in selectedSections: {selected_sections_json}")
            selected_sections = {}
        
        # Activate the selected language
        from django.utils.translation import activate
        activate(report_language)
        
        # Generate report data with selected sections
        report_data = generate_report_data(request.user, {
            'reportType': report_type,
            'format': report_format,
            'language': report_language,
            'company_id': request.POST.get('reportCompany', ''),
            'selectedSections': selected_sections,  # Pass selected sections to report data generation
        })
        
        # Generate HTML preview with selected sections
        preview_html = generate_html_preview(report_data, report_type, report_language, selected_sections)
        
        return HttpResponse(preview_html)
        
    except Exception as e:
        logger.error(f"Error generating report preview: {str(e)}", exc_info=True)
        error_html = f'''
        <div class="alert alert-danger">
            <i class="fas fa-exclamation-triangle me-2"></i>
            {_("Error generating preview")}: {str(e)}
        </div>
        '''
        return HttpResponse(error_html)


@login_required
@user_passes_test(check_risk_report_access)
@require_http_methods(["GET"])
def preview_risk_report_page(request):
    """Render preview as a standalone page using query params."""
    try:
        report_type = request.GET.get('reportType', 'summary')
        report_format = request.GET.get('format', 'word')
        report_language = request.GET.get('language', get_language()[:2])
        selected_sections_json = request.GET.get('selectedSections', '{}')
        try:
            selected_sections = json.loads(selected_sections_json) if selected_sections_json else {}
        except json.JSONDecodeError:
            selected_sections = {}

        # Activate language
        from django.utils.translation import activate
        activate(report_language)

        report_data = generate_report_data(request.user, {
            'reportType': report_type,
            'format': report_format,
            'language': report_language,
            'company_id': request.GET.get('reportCompany', ''),
            'selectedSections': selected_sections,
        })

        html = generate_html_preview(report_data, report_type, report_language, selected_sections)
        return render(request, 'app_risk/risk_report_preview_page.html', {
            'preview_html': html,
            'page_title': _('Report Preview'),
            'is_profile': False,
            'report_params': {
                'reportType': report_type,
                'language': report_language,
                'format': report_format,
                'company_id': request.GET.get('reportCompany', ''),
                'selectedSections': selected_sections_json,
            }
        })
    except Exception as e:
        logger.error(f"Error generating standalone preview: {str(e)}", exc_info=True)
        return render(request, 'app_risk/risk_report_preview_page.html', {
            'preview_html': f'<div class="alert alert-danger">{_("Error generating preview")}: {str(e)}</div>',
            'page_title': _('Report Preview')
        })


@login_required
@user_passes_test(check_risk_report_access)
@require_http_methods(["GET"])
def preview_risk_report_from_profile_page(request):
    """Render profile-based preview as a standalone page."""
    try:
        profile_id = request.GET.get('profile_id')
        if not profile_id:
            raise ValueError(_('Profile ID is required'))

        profile = ReportProfile.objects.get(id=profile_id)
        if not profile.can_be_used_by(request.user):
            raise PermissionDenied(_("You don't have permission to use this profile"))

        override_language = request.GET.get('language') or profile.default_language
        from django.utils.translation import activate
        activate(override_language)

        sections_config = profile.get_sections_config()
        company_id = sections_config.get('_company_id', '')

        params = {
            'reportType': 'profile',
            'format': 'html',
            'language': override_language,
            'company_id': str(company_id) if company_id else '',
            'selectedSections': sections_config,
            'profile': profile
        }

        # Enrich profile header
        company_name = _('All Companies')
        if company_id:
            try:
                from app_cabinet.models import Company
                company = Company.objects.get(id=company_id)
                company_name = company.name
            except Exception:
                pass
        params['profile'] = {
            'name': profile.name,
            'description': profile.description,
            'created_by': profile.created_by.get_full_name(),
            'created_at': profile.created_at,
            'company': company_name,
        }

        report_data = generate_report_data(request.user, params)
        html = generate_html_preview(report_data, 'profile', override_language, sections_config)

        return render(request, 'app_risk/risk_report_preview_page.html', {
            'preview_html': html,
            'page_title': _('Report Preview'),
            'is_profile': True,
            'profile_id': str(profile.id),
            'language': override_language,
        })
    except Exception as e:
        logger.error(f"Error generating standalone profile preview: {str(e)}", exc_info=True)
        return render(request, 'app_risk/risk_report_preview_page.html', {
            'preview_html': f'<div class="alert alert-danger">{_("Error generating preview")}: {str(e)}</div>',
            'page_title': _('Report Preview')
        })

def generate_ai_conclusion(report_data, selected_sections, language='uk'):
    """
    Generate AI Conclusion based on report data using selected AI model.
    
    Args:
        report_data: Dictionary containing all report data
        selected_sections: Dictionary with selected sections configuration
        language: Language code for the conclusion
    
    Returns:
        str: Generated AI conclusion text or None if generation failed
    """
    try:
        # Check if conclusions section is enabled (handle both boolean and string values)
        conclusions_enabled = selected_sections.get('conclusions', False)
        if not (conclusions_enabled == True or (isinstance(conclusions_enabled, str) and conclusions_enabled.lower() in ('true', '1'))):
            logger.debug("AI Conclusion section is not enabled")
            return None
        
        # Get AI model ID from selected sections (handle both string and int)
        ai_model_id = selected_sections.get('_ai_conclusion_model')
        if not ai_model_id:
            logger.warning("AI Conclusion is enabled but no AI model is selected")
            return None
        
        # Convert to int if it's a string
        if isinstance(ai_model_id, str):
            try:
                ai_model_id = int(ai_model_id)
            except ValueError:
                logger.error(f"Invalid AI model ID format: {ai_model_id}")
                return None
        
        # Get the AI model
        from app_ai.models import ModelChoice
        try:
            ai_model = ModelChoice.objects.get(id=ai_model_id, is_active=True)
        except ModelChoice.DoesNotExist:
            logger.error(f"AI model with ID {ai_model_id} not found or inactive")
            return None
        
        # Prepare report summary for AI analysis
        report_summary = prepare_report_summary_for_ai(report_data, language)
        
        # Create prompt for AI
        if language == 'uk':
            system_prompt = "Ти експерт з аналізу ризиків та безпеки. Проаналізуй дані звіту про оцінку ризиків та надай професійний висновок українською мовою."
            user_prompt = f"""Проаналізуй наступні дані звіту про оцінку ризиків та надай детальний професійний висновок:

{report_summary}

Висновок має містити:
1. Загальну оцінку поточного стану безпеки
2. Основні виявлені ризики та їх критичність
3. Рекомендації щодо пріоритетних дій
4. Оцінку ефективності запроваджених заходів безпеки
5. Загальні висновки та рекомендації для керівництва

Висновок має бути професійним, структурованим та зрозумілим."""
        elif language == 'ru':
            system_prompt = "Ты эксперт по анализу рисков и безопасности. Проанализируй данные отчета об оценке рисков и предоставь профессиональный вывод на русском языке."
            user_prompt = f"""Проанализируй следующие данные отчета об оценке рисков и предоставь детальный профессиональный вывод:

{report_summary}

Вывод должен содержать:
1. Общую оценку текущего состояния безопасности
2. Основные выявленные риски и их критичность
3. Рекомендации по приоритетным действиям
4. Оценку эффективности внедренных мер безопасности
5. Общие выводы и рекомендации для руководства

Вывод должен быть профессиональным, структурированным и понятным."""
        else:  # English
            system_prompt = "You are an expert in risk analysis and security. Analyze the risk assessment report data and provide a professional conclusion in English."
            user_prompt = f"""Analyze the following risk assessment report data and provide a detailed professional conclusion:

{report_summary}

The conclusion should include:
1. Overall assessment of current security status
2. Main identified risks and their criticality
3. Recommendations for priority actions
4. Assessment of effectiveness of implemented security measures
5. General conclusions and recommendations for management

The conclusion should be professional, structured, and clear."""
        
        # Get AI response based on provider
        from app_ai import ai_utils
        
        conversation_history = []
        ai_response = None
        
        # Normalize provider to lowercase for comparison
        provider = ai_model.provider.lower() if ai_model.provider else ''
        
        if provider == 'ollama':
            ai_response, _ = ai_utils.get_ollama_response(user_prompt, conversation_history, system_prompt)
        elif provider == 'google':
            ai_response, _ = ai_utils.get_google_response(user_prompt, conversation_history, system_prompt)
        elif provider == 'claude':
            ai_response, _ = ai_utils.get_claude_response(user_prompt, conversation_history, system_prompt)
        elif provider == 'groq':
            ai_response, _ = ai_utils.get_groq_response(user_prompt, conversation_history, system_prompt)
        elif provider == 'deepseek':
            ai_response, _ = ai_utils.get_deepseek_response(user_prompt, conversation_history, system_prompt)
        else:
            logger.error(f"Unsupported AI provider: {ai_model.provider}")
            return None
        
        if ai_response and not ai_response.startswith("Error"):
            return ai_response
        else:
            logger.error(f"AI response error: {ai_response}")
            return None
            
    except Exception as e:
        logger.error(f"Error generating AI conclusion: {str(e)}", exc_info=True)
        return None


def prepare_report_summary_for_ai(report_data, language='uk'):
    """
    Prepare a text summary of report data for AI analysis.
    
    Args:
        report_data: Dictionary containing all report data
        language: Language code for the summary
    
    Returns:
        str: Formatted text summary
    """
    summary_parts = []
    
    # Statistics
    if 'statistics' in report_data:
        stats = report_data['statistics']
        if language == 'uk':
            summary_parts.append(f"Загальна статистика:")
            summary_parts.append(f"- Всього активів: {stats.get('total_assets', 0)}")
            summary_parts.append(f"- Всього вразливостей: {stats.get('total_vulnerabilities', 0)}")
            summary_parts.append(f"- Активів з високим ризиком: {stats.get('high_risk_count', 0)}")
            summary_parts.append(f"- Рівень завершеності: {stats.get('completion_rate', 0):.1f}%")
        elif language == 'ru':
            summary_parts.append(f"Общая статистика:")
            summary_parts.append(f"- Всего активов: {stats.get('total_assets', 0)}")
            summary_parts.append(f"- Всего уязвимостей: {stats.get('total_vulnerabilities', 0)}")
            summary_parts.append(f"- Активов с высоким риском: {stats.get('high_risk_count', 0)}")
            summary_parts.append(f"- Уровень завершенности: {stats.get('completion_rate', 0):.1f}%")
        else:
            summary_parts.append(f"Overall Statistics:")
            summary_parts.append(f"- Total Assets: {stats.get('total_assets', 0)}")
            summary_parts.append(f"- Total Vulnerabilities: {stats.get('total_vulnerabilities', 0)}")
            summary_parts.append(f"- High Risk Assets: {stats.get('high_risk_count', 0)}")
            summary_parts.append(f"- Completion Rate: {stats.get('completion_rate', 0):.1f}%")
    
    # Risk Distribution
    if 'risk_distribution' in report_data and 'by_level' in report_data['risk_distribution']:
        risk_levels = report_data['risk_distribution']['by_level']
        if language == 'uk':
            summary_parts.append(f"\nРозподіл ризиків за рівнями:")
            for level, count in risk_levels.items():
                summary_parts.append(f"- {level}: {count}")
        elif language == 'ru':
            summary_parts.append(f"\nРаспределение рисков по уровням:")
            for level, count in risk_levels.items():
                summary_parts.append(f"- {level}: {count}")
        else:
            summary_parts.append(f"\nRisk Distribution by Level:")
            for level, count in risk_levels.items():
                summary_parts.append(f"- {level}: {count}")
    
    # Top Risks
    if 'statistics' in report_data and 'top_risks' in report_data['statistics']:
        top_risks = report_data['statistics']['top_risks']
        if top_risks:
            if language == 'uk':
                summary_parts.append(f"\nТоп-10 найбільш критичних ризиків:")
            elif language == 'ru':
                summary_parts.append(f"\nТоп-10 наиболее критичных рисков:")
            else:
                summary_parts.append(f"\nTop-10 Most Critical Risks:")
            
            for i, risk in enumerate(top_risks[:10], 1):
                asset_name = risk.get('asset_name', 'Unknown')
                vuln_name = risk.get('vulnerability_name', 'Unknown')
                risk_level = risk.get('risk_level', 'Unknown')
                if language == 'uk':
                    summary_parts.append(f"{i}. {asset_name} - {vuln_name} (Рівень ризику: {risk_level})")
                elif language == 'ru':
                    summary_parts.append(f"{i}. {asset_name} - {vuln_name} (Уровень риска: {risk_level})")
                else:
                    summary_parts.append(f"{i}. {asset_name} - {vuln_name} (Risk Level: {risk_level})")
    
    # Compliance
    if 'compliance' in report_data:
        compliance = report_data['compliance']
        if language == 'uk':
            summary_parts.append(f"\nСтатус відповідності:")
            if 'iso27001' in compliance:
                iso_rate = compliance['iso27001'].get('compliance_rate', 0)
                summary_parts.append(f"- ISO 27001: {iso_rate:.1f}% відповідності")
            if 'pcidss' in compliance:
                pci_rate = compliance['pcidss'].get('compliance_rate', 0)
                summary_parts.append(f"- PCI DSS: {pci_rate:.1f}% відповідності")
        elif language == 'ru':
            summary_parts.append(f"\nСтатус соответствия:")
            if 'iso27001' in compliance:
                iso_rate = compliance['iso27001'].get('compliance_rate', 0)
                summary_parts.append(f"- ISO 27001: {iso_rate:.1f}% соответствия")
            if 'pcidss' in compliance:
                pci_rate = compliance['pcidss'].get('compliance_rate', 0)
                summary_parts.append(f"- PCI DSS: {pci_rate:.1f}% соответствия")
        else:
            summary_parts.append(f"\nCompliance Status:")
            if 'iso27001' in compliance:
                iso_rate = compliance['iso27001'].get('compliance_rate', 0)
                summary_parts.append(f"- ISO 27001: {iso_rate:.1f}% compliance")
            if 'pcidss' in compliance:
                pci_rate = compliance['pcidss'].get('compliance_rate', 0)
                summary_parts.append(f"- PCI DSS: {pci_rate:.1f}% compliance")
    
    # Treatment Details
    if 'statistics' in report_data and 'overdue_treatments' in report_data['statistics']:
        overdue = report_data['statistics']['overdue_treatments']
        if overdue:
            overdue_count = len(overdue)
            if language == 'uk':
                summary_parts.append(f"\nПрострочені заходи з ліквідації ризиків: {overdue_count}")
            elif language == 'ru':
                summary_parts.append(f"\nПросроченные меры по устранению рисков: {overdue_count}")
            else:
                summary_parts.append(f"\nOverdue Risk Treatments: {overdue_count}")
    
    return "\n".join(summary_parts)


def generate_unified_report_content(data, report_type, selected_sections, translations, language=None):
    """
    Generate unified report content structure that can be used by HTML preview, Word, and PDF generators.
    This ensures all formats have identical data and structure.
    """
    
    # Use provided language or fallback to data language or default
    if language is None:
        language = data.get('language', 'uk')
    # Activate the language for proper multilingual support
    from django.utils.translation import activate
    activate(language)
    
    content = {
        'header': {},
        'sections': []
    }
    
    # Header information (always included)
    # Determine title based on profile or report type
    if data.get('profile') and data['profile'].get('name'):
        # Use profile name for title
        title = f"{translations['report_title']} - {data['profile']['name']}"
    else:
        # Fallback to report type (for backward compatibility)
        title = f"{translations['report_title']} - {report_type.title()}"
    
    content['header'] = {
        'title': title,
        'generated_date': format_localized_date(data['generated_date'], data.get('language', 'uk')),
        'report_period': f"{data['date_range']['start']} - {data['date_range']['end']}",
        'generated_by': data.get('generated_by', 'System'),
        'company_info': None
    }
    
    # Company information
    if data.get('companies'):
        if data['companies'].count() == 1:
            content['header']['company_info'] = {
                'type': 'single',
                'name': data['companies'].first().name
            }
        elif data['companies'].count() > 1:
            content['header']['company_info'] = {
                'type': 'multiple',
                'count': data['companies'].count(),
                'text': f"{translations['multiple_companies']} ({data['companies'].count()})"
            }
    
    # Executive Summary Section
    if selected_sections.get('statistics') or selected_sections.get('risk_distribution') or selected_sections.get('compliance_overview') or selected_sections.get('top_risks'):
        exec_summary = {
            'type': 'executive_summary',
            'title': translations['executive_summary'],
            'subsections': []
        }
        
        # Key Statistics
        if selected_sections.get('statistics') and 'statistics' in data:
            stats = data['statistics']
            exec_summary['subsections'].append({
                'type': 'statistics',
                'title': translations.get('key_metrics', 'Key Metrics'),
                'data': [
                    {
                        'label': translations['total_assets'],
                        'value': str(stats.get('total_assets', 0)),
                        'type': 'primary'
                    },
                    {
                        'label': translations['total_vulnerabilities'],
                        'value': str(stats.get('total_vulnerabilities', 0)),
                        'type': 'warning'
                    },
                    {
                        'label': translations['high_risk_assets'],
                        'value': str(stats.get('high_risk_count', 0)),
                        'type': 'danger'
                    },
                    {
                        'label': translations['completion_rate'],
                        'value': f"{stats.get('completion_rate', 0):.1f}%",
                        'type': 'success'
                    }
                ]
            })
        
        # Risk Details (Combined By Risk Level and TOP-10 Risks)
        risk_details_section = None
        
        # Check if either risk_distribution or top_risks is selected
        if (selected_sections.get('risk_distribution') and 'risk_distribution' in data) or (selected_sections.get('top_risks') and 'statistics' in data):
            risk_details_section = {
                'type': 'risk_details',
                'title': translations.get('risk_details', 'Risk Details'),
                'subsections': []
            }
            
            # By Risk Level subsection
            if selected_sections.get('risk_distribution') and 'risk_distribution' in data:
                risk_dist = data['risk_distribution']
                if risk_dist.get('by_level'):
                    total_risks = sum(risk_dist['by_level'].values())
                    risk_level_data = []
                    
                    # Get risk levels with colors from statistics if available
                    risk_levels_with_colors = data.get('statistics', {}).get('risk_levels_with_colors', {})
                    
                    # Create list with all info for sorting
                    for level, count in risk_dist['by_level'].items():
                        percentage = (count / total_risks * 100) if total_risks > 0 else 0
                        color_info = risk_levels_with_colors.get(level, {})
                        risk_level_data.append({
                            'risk_level': level,
                            'count': count,
                            'percentage': f"{percentage:.1f}%",
                            'color': color_info.get('color', '#000000'),
                            'min_value': color_info.get('min_value', 0),
                            'max_value': color_info.get('max_value', 0)
                        })
                    
                    # Sort by max_value in descending order (highest risk first)
                    risk_level_data.sort(key=lambda x: x['max_value'], reverse=True)
                    
                    risk_details_section['subsections'].append({
                        'type': 'risk_distribution_level',
                        'title': translations.get('by_risk_level', 'By Risk Level'),
                        'headers': [translations['risk_level'], translations['count'], translations.get('percentage', 'Percentage')],
                        'data': risk_level_data
                    })
            
            # TOP-10 Risks subsection
            if selected_sections.get('top_risks') and 'statistics' in data:
                stats = data['statistics']
                if stats.get('high_risk_assets'):
                    top_risks_data = []
                    for asset in stats['high_risk_assets'][:10]:  # Limit to 10
                        top_risks_data.append({
                            'asset_name': asset.get('asset_name', translations.get('unknown', 'Unknown')),
                            'vulnerability_name': asset.get('vulnerability_name', translations.get('unknown', 'Unknown')),
                            'risk_level': asset.get('risk_level', translations.get('unknown', 'Unknown'))
                        })
                    
                    risk_details_section['subsections'].append({
                        'type': 'top_risks',
                        'title': translations.get('top_risks', 'TOP-10 Risks'),
                        'headers': [translations.get('asset_name', 'Asset Name'), translations['vulnerability'], translations['risk_level']],
                        'data': top_risks_data
                    })
        
        content['sections'].append(exec_summary)
        
        # Add the combined section to content
        if risk_details_section and risk_details_section.get('subsections'):
            content['sections'].append(risk_details_section)
    
    # Compliance Section
    if (selected_sections.get('compliance_overview') or selected_sections.get('pci_dss') or 
        selected_sections.get('iso27001') or selected_sections.get('compliance_gaps') or
        selected_sections.get('framework_company_requirements') or 
        selected_sections.get('company_requirements') or 
        selected_sections.get('internal_requirements')):
        compliance_section = {
            'type': 'compliance',
            'title': translations['compliance_summary'],
            'subsections': []
        }
        
        # PCI DSS Compliance
        if selected_sections.get('pci_dss') and 'pcidss_compliance' in data:
            pcidss = data['pcidss_compliance']
            compliance_section['subsections'].append({
                'type': 'pci_dss',
                'title': translations['pci_dss_compliance'],
                'data': [
                    {
                        'label': translations['overall_compliance'],
                        'value': f"{pcidss.get('overall_compliance', 0):.1f}%"
                    },
                    {
                        'label': translations['total_requirements'],
                        'value': str(pcidss.get('total_requirements', 0))
                    },
                    {
                        'label': translations['compliant_vulnerabilities'],
                        'value': str(pcidss.get('compliant_vulnerabilities', 0))
                    }
                ]
            })
        
        # ISO 27001 Compliance
        if selected_sections.get('iso27001') and 'iso27001_compliance' in data:
            iso27001 = data['iso27001_compliance']
            compliance_section['subsections'].append({
                'type': 'iso27001',
                'title': translations['iso27001_compliance'],
                'data': [
                    {
                        'label': translations['overall_compliance'],
                        'value': f"{iso27001.get('overall_compliance', 0):.1f}%"
                    },
                    {
                        'label': translations['total_controls'],
                        'value': str(iso27001.get('total_controls', 0))
                    },
                    {
                        'label': translations['compliant_vulnerabilities'],
                        'value': str(iso27001.get('compliant_vulnerabilities', 0))
                    }
                ]
            })
        
        # Compliance Gaps
        if selected_sections.get('compliance_gaps'):
            gaps_data = []
            
            # PCI DSS Gaps
            if 'pcidss_gaps' in data and data['pcidss_gaps']:
                for gap in data['pcidss_gaps'][:10]:  # Limit to 10
                    gaps_data.append({
                        'type': 'PCI DSS',
                        'requirement': gap.get('requirement', translations.get('unknown', 'Unknown')),
                        'description': gap.get('description', translations.get('no_description', 'No description'))
                    })
            
            # ISO 27001 Gaps
            if 'iso27001_gaps' in data and data['iso27001_gaps']:
                for gap in data['iso27001_gaps'][:10]:  # Limit to 10
                    gaps_data.append({
                        'type': 'ISO 27001',
                        'requirement': gap.get('control', translations.get('unknown', 'Unknown')),
                        'description': gap.get('description', translations.get('no_description', 'No description'))
                    })
            
            if gaps_data:
                compliance_section['subsections'].append({
                    'type': 'compliance_gaps',
                    'title': translations.get('compliance_gaps', 'Compliance Gaps'),
                    'headers': [translations.get('type', 'Type'), translations.get('requirement', 'Requirement'), translations['description']],
                    'data': gaps_data
                })
        
        # Framework Company Requirements
        if selected_sections.get('framework_company_requirements') and 'compliance' in data:
            compliance_data = data.get('compliance', {})
            framework_data = compliance_data.get('framework_company_requirements')
            # Check if data exists (even if empty list, we should show the section if enabled)
            if framework_data is not None:
                frameworks_list = []
                if framework_data.get('frameworks') and len(framework_data['frameworks']) > 0:
                    for framework in framework_data['frameworks']:
                        frameworks_list.append({
                            'name': framework.get('name', ''),
                            'framework_type': framework.get('framework_type', ''),
                            'version': framework.get('version', ''),
                            'company': framework.get('company', ''),
                            'status': framework.get('status', ''),
                            'controls_total': framework.get('controls_total', 0),
                            'controls_completed': framework.get('controls_completed', 0),
                            'completion_percentage': f"{framework.get('completion_percentage', 0):.1f}%"
                        })
                
                compliance_section['subsections'].append({
                    'type': 'framework_company_requirements',
                    'title': translations.get('framework_company_requirements', 'Framework Company Requirements'),
                    'headers': [
                        translations.get('name', 'Name'),
                        translations.get('framework_type', 'Framework Type'),
                        translations.get('version', 'Version'),
                        translations.get('company', 'Company'),
                        translations.get('status', 'Status'),
                        translations.get('controls_total', 'Total Controls'),
                        translations.get('controls_completed', 'Completed'),
                        translations.get('completion_percentage', 'Completion %')
                    ],
                    'data': frameworks_list if frameworks_list else [{'name': translations.get('no_data', 'No data available'), 'framework_type': '', 'version': '', 'company': '', 'status': '', 'controls_total': 0, 'controls_completed': 0, 'completion_percentage': '0%'}]
                })
        
        # Local Company Requirements
        if selected_sections.get('company_requirements') and 'compliance' in data:
            compliance_data = data.get('compliance', {})
            local_data = compliance_data.get('company_requirements')
            # Check if data exists (even if empty list, we should show the section if enabled)
            if local_data is not None:
                requirements_list = []
                if local_data.get('requirements') and len(local_data['requirements']) > 0:
                    for req in local_data['requirements']:
                        requirements_list.append({
                            'code': req.get('code', ''),
                            'name': req.get('name', ''),
                            'requirement_type': req.get('requirement_type', ''),
                            'company': req.get('company', ''),
                            'regulator': req.get('regulator', ''),
                            'status': req.get('status', ''),
                            'controls_total': req.get('controls_total', 0),
                            'controls_completed': req.get('controls_completed', 0),
                            'completion_percentage': f"{req.get('completion_percentage', 0):.1f}%"
                        })
                
                compliance_section['subsections'].append({
                    'type': 'company_requirements',
                    'title': translations.get('local_company_requirements', 'Local Company Requirements'),
                    'headers': [
                        translations.get('code', 'Code'),
                        translations.get('name', 'Name'),
                        translations.get('requirement_type', 'Type'),
                        translations.get('company', 'Company'),
                        translations.get('regulator', 'Regulator'),
                        translations.get('status', 'Status'),
                        translations.get('controls_total', 'Total Controls'),
                        translations.get('controls_completed', 'Completed'),
                        translations.get('completion_percentage', 'Completion %')
                    ],
                    'data': requirements_list if requirements_list else [{'code': '', 'name': translations.get('no_data', 'No data available'), 'requirement_type': '', 'company': '', 'regulator': '', 'status': '', 'controls_total': 0, 'controls_completed': 0, 'completion_percentage': '0%'}]
                })
        
        # Internal Company Requirements
        if selected_sections.get('internal_requirements') and 'compliance' in data:
            compliance_data = data.get('compliance', {})
            internal_data = compliance_data.get('internal_requirements')
            # Check if data exists (even if empty list, we should show the section if enabled)
            if internal_data is not None:
                requirements_list = []
                if internal_data.get('requirements') and len(internal_data['requirements']) > 0:
                    for req in internal_data['requirements']:
                        requirements_list.append({
                            'code': req.get('code', ''),
                            'name': req.get('name', ''),
                            'requirement_type': req.get('requirement_type', ''),
                            'company': req.get('company', ''),
                            'status': req.get('status', ''),
                            'controls_total': req.get('controls_total', 0),
                            'controls_completed': req.get('controls_completed', 0),
                            'completion_percentage': f"{req.get('completion_percentage', 0):.1f}%"
                        })
                
                compliance_section['subsections'].append({
                    'type': 'internal_requirements',
                    'title': translations.get('internal_company_requirements', 'Internal Company Requirements'),
                    'headers': [
                        translations.get('code', 'Code'),
                        translations.get('name', 'Name'),
                        translations.get('requirement_type', 'Type'),
                        translations.get('company', 'Company'),
                        translations.get('status', 'Status'),
                        translations.get('controls_total', 'Total Controls'),
                        translations.get('controls_completed', 'Completed'),
                        translations.get('completion_percentage', 'Completion %')
                    ],
                    'data': requirements_list if requirements_list else [{'code': '', 'name': translations.get('no_data', 'No data available'), 'requirement_type': '', 'company': '', 'status': '', 'controls_total': 0, 'controls_completed': 0, 'completion_percentage': '0%'}]
                })
        
        if compliance_section['subsections']:
            content['sections'].append(compliance_section)
    
    # Incident Register Section
    if selected_sections.get('incident') and 'incident_data' in data:
        incident_data = data.get('incident_data')
        # Check if data exists (even if empty list, we should show the section if enabled)
        if incident_data is not None:
            incident_section = {
                'type': 'incident',
                'title': translations.get('incident_register', 'Incident Register'),
                'subsections': []
            }
            
            # Summary statistics
            if incident_data.get('total_incidents', 0) > 0:
                stats_data = [
                    {
                        'label': translations.get('total_incidents', 'Total Incidents'),
                        'value': str(incident_data.get('total_incidents', 0)),
                        'type': 'primary'
                    }
                ]
                
                # Add classification statistics
                if incident_data.get('statistics', {}).get('by_classification'):
                    classification_stats = incident_data['statistics']['by_classification']
                    for classification, count in sorted(classification_stats.items(), key=lambda x: x[1], reverse=True):
                        stats_data.append({
                            'label': f"{translations.get('by_classification', 'By Classification')}: {classification}",
                            'value': str(count),
                            'type': 'info'
                        })
                
                # Add state statistics
                if incident_data.get('statistics', {}).get('by_state'):
                    state_stats = incident_data['statistics']['by_state']
                    for state, count in sorted(state_stats.items(), key=lambda x: x[1], reverse=True):
                        stats_data.append({
                            'label': f"{translations.get('by_state', 'By State')}: {state}",
                            'value': str(count),
                            'type': 'warning'
                        })
                
                incident_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('incident_summary', 'Incident Summary'),
                    'data': stats_data
                })
            
            # Incident table
            incidents_list = []
            if incident_data.get('incidents') and len(incident_data['incidents']) > 0:
                for incident in incident_data['incidents']:
                    incidents_list.append({
                        'id': incident.get('id', ''),
                        'company': incident.get('company', ''),
                        'occurrence_datetime': incident.get('occurrence_datetime', ''),
                        'place': incident.get('place', ''),
                        'description': incident.get('description', '')[:200] + '...' if len(incident.get('description', '')) > 200 else incident.get('description', ''),
                        'classification': incident.get('classification', ''),
                        'incident_type': incident.get('incident_type', ''),
                        'responsible': incident.get('responsible', ''),
                        'reported_by': incident.get('reported_by', ''),
                        'current_state': incident.get('current_state', ''),
                        'impact': incident.get('impact', '')[:200] + '...' if len(incident.get('impact', '')) > 200 else incident.get('impact', ''),
                    })
            
            incident_section['subsections'].append({
                'type': 'incident_table',
                'title': translations.get('incidents', 'Incidents'),
                'headers': [
                    translations.get('id', 'ID'),
                    translations.get('company', 'Company'),
                    translations.get('occurrence_datetime', 'Occurrence Date'),
                    translations.get('place', 'Place'),
                    translations.get('incident_type', 'Type'),
                    translations.get('classification', 'Classification'),
                    translations.get('current_state', 'State'),
                    translations.get('responsible', 'Responsible'),
                    translations.get('description', 'Description')
                ],
                'data': incidents_list if incidents_list else [{'id': '', 'company': translations.get('no_data', 'No data available'), 'occurrence_datetime': '', 'place': '', 'incident_type': '', 'classification': '', 'current_state': '', 'responsible': '', 'description': ''}]
            })
            
            if incident_section['subsections']:
                content['sections'].append(incident_section)
    
    # Mandatory Processes Section
    if selected_sections.get('mandatory_processes') and 'mandatory_processes_data' in data:
        mandatory_processes_data = data.get('mandatory_processes_data')
        # Check if data exists (even if empty list, we should show the section if enabled)
        if mandatory_processes_data is not None:
            mandatory_processes_section = {
                'type': 'mandatory_processes',
                'title': translations.get('mandatory_processes', 'Mandatory Processes'),
                'subsections': []
            }
            
            # Summary statistics
            if mandatory_processes_data.get('total_processes', 0) > 0:
                stats_data = [
                    {
                        'label': translations.get('total_processes', 'Total Processes'),
                        'value': str(mandatory_processes_data.get('total_processes', 0)),
                        'type': 'primary'
                    }
                ]
                
                # Add status statistics
                if mandatory_processes_data.get('statistics', {}).get('by_status'):
                    status_stats = mandatory_processes_data['statistics']['by_status']
                    if status_stats.get('overdue', 0) > 0:
                        stats_data.append({
                            'label': translations.get('overdue', 'Overdue'),
                            'value': str(status_stats.get('overdue', 0)),
                            'type': 'danger'
                        })
                    if status_stats.get('upcoming', 0) > 0:
                        stats_data.append({
                            'label': translations.get('upcoming', 'Upcoming'),
                            'value': str(status_stats.get('upcoming', 0)),
                            'type': 'warning'
                        })
                    if status_stats.get('completed', 0) > 0:
                        stats_data.append({
                            'label': translations.get('completed', 'Completed'),
                            'value': str(status_stats.get('completed', 0)),
                            'type': 'success'
                        })
                
                # Add priority statistics
                if mandatory_processes_data.get('statistics', {}).get('by_priority'):
                    priority_stats = mandatory_processes_data['statistics']['by_priority']
                    for priority, count in sorted(priority_stats.items(), key=lambda x: x[1], reverse=True):
                        priority_display = priority.capitalize()
                        stats_data.append({
                            'label': f"{translations.get('by_priority', 'By Priority')}: {priority_display}",
                            'value': str(count),
                            'type': 'info'
                        })
                
                mandatory_processes_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('mandatory_processes_summary', 'Mandatory Processes Summary'),
                    'data': stats_data,
                    'display_inline': True  # Flag to display all blocks in one row
                })
            
            # Processes table - show only last 5 processes sorted by Next Due Date
            processes_list = []
            if mandatory_processes_data.get('processes') and len(mandatory_processes_data['processes']) > 0:
                # Sort processes by next_due_date (None values go to the end)
                sorted_processes = sorted(
                    mandatory_processes_data['processes'],
                    key=lambda x: (
                        x.get('next_due_date') if x.get('next_due_date') else '9999-12-31',
                        x.get('process_name', '')
                    )
                )
                # Take only first 5 processes (with nearest due dates)
                for process in sorted_processes[:5]:
                    processes_list.append({
                        'process_name': process.get('process_name', ''),
                        'company': process.get('company', ''),
                        'description': process.get('description', '')[:200] + '...' if len(process.get('description', '')) > 200 else process.get('description', ''),
                        'priority': process.get('priority', ''),
                        'frequency': process.get('frequency', ''),
                        'next_due_date': process.get('next_due_date', ''),
                        'last_completed_date': process.get('last_completed_date', ''),
                        'responsible_person': process.get('responsible_person', ''),
                        'status': process.get('status', ''),
                        'source_document': process.get('source_document', ''),
                    })
            
            mandatory_processes_section['subsections'].append({
                'type': 'mandatory_processes_table',
                'title': translations.get('last_5_mandatory_processes', 'Last 5 Mandatory Processes'),
                'headers': [
                    translations.get('process_name', 'Process Name'),
                    translations.get('company', 'Company'),
                    translations.get('priority', 'Priority'),
                    translations.get('frequency', 'Frequency'),
                    translations.get('next_due_date', 'Next Due Date'),
                    translations.get('last_completed_date', 'Last Completed Date'),
                    translations.get('responsible_person', 'Responsible Person'),
                    translations.get('status', 'Status'),
                    translations.get('description', 'Description')
                ],
                'data': processes_list if processes_list else [{'process_name': translations.get('no_data', 'No data available'), 'company': '', 'priority': '', 'frequency': '', 'next_due_date': '', 'last_completed_date': '', 'responsible_person': '', 'status': '', 'description': ''}]
            })
            
            if mandatory_processes_section['subsections']:
                content['sections'].append(mandatory_processes_section)
    
    # Certificate & Key Management Section
    if selected_sections.get('certificate_key_management') and 'certificate_key_management_data' in data:
        certificate_key_management_data = data.get('certificate_key_management_data')
        # Check if data exists (even if empty list, we should show the section if enabled)
        if certificate_key_management_data is not None:
            certificate_key_management_section = {
                'type': 'certificate_key_management',
                'title': translations.get('certificate_key_management', 'Certificate & Key Management'),
                'subsections': []
            }
            
            # Summary statistics
            if certificate_key_management_data.get('total_certificates', 0) > 0:
                stats_data = [
                    {
                        'label': translations.get('total_certificates', 'Total Certificates'),
                        'value': str(certificate_key_management_data.get('total_certificates', 0)),
                        'type': 'primary'
                    }
                ]
                
                # Add expiry status statistics
                if certificate_key_management_data.get('statistics', {}).get('by_expiry_status'):
                    expiry_stats = certificate_key_management_data['statistics']['by_expiry_status']
                    if expiry_stats.get('expired', 0) > 0:
                        stats_data.append({
                            'label': translations.get('expired', 'Expired'),
                            'value': str(expiry_stats.get('expired', 0)),
                            'type': 'danger'
                        })
                    if expiry_stats.get('expiring_soon', 0) > 0:
                        stats_data.append({
                            'label': translations.get('expiring_soon', 'Expiring Soon (≤30 days)'),
                            'value': str(expiry_stats.get('expiring_soon', 0)),
                            'type': 'warning'
                        })
                    if expiry_stats.get('valid', 0) > 0:
                        stats_data.append({
                            'label': translations.get('valid', 'Valid'),
                            'value': str(expiry_stats.get('valid', 0)),
                            'type': 'success'
                        })
                
                # Add type statistics
                if certificate_key_management_data.get('statistics', {}).get('by_type'):
                    type_stats = certificate_key_management_data['statistics']['by_type']
                    for cert_type, count in sorted(type_stats.items(), key=lambda x: x[1], reverse=True)[:5]:
                        stats_data.append({
                            'label': f"{translations.get('by_type', 'By Type')}: {cert_type}",
                            'value': str(count),
                            'type': 'info'
                        })
                
                certificate_key_management_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('certificate_key_management_summary', 'Certificate & Key Management Summary'),
                    'data': stats_data,
                    'display_inline': True  # Flag to display all blocks in one row
                })
            
            # Certificates table - show only last 5 certificates sorted by Expiry Date
            certificates_list = []
            if certificate_key_management_data.get('certificates') and len(certificate_key_management_data['certificates']) > 0:
                # Sort certificates by expiry_date (expiring soon first, then expired, then valid)
                sorted_certificates = sorted(
                    certificate_key_management_data['certificates'],
                    key=lambda x: (
                        x.get('expiry_status') == 'expired' and 0 or (x.get('expiry_status') == 'expiring_soon' and 1 or 2),
                        x.get('expiry_date', '9999-12-31'),
                        x.get('key_cert_num', '')
                    )
                )
                # Take only first 5 certificates (expiring soon/expired first)
                for cert in sorted_certificates[:5]:
                    certificates_list.append({
                        'key_cert_num': cert.get('key_cert_num', ''),
                        'company': cert.get('company', ''),
                        'type': cert.get('type', ''),
                        'purpose': cert.get('purpose', '')[:200] + '...' if len(cert.get('purpose', '')) > 200 else cert.get('purpose', ''),
                        'expiry_date': cert.get('expiry_date', ''),
                        'expiry_status': cert.get('expiry_status', ''),
                        'days_until_expiry': cert.get('days_until_expiry'),
                        'revocation_status': cert.get('revocation_status', ''),
                        'owner_name': cert.get('owner_name', ''),
                        'location': cert.get('location', ''),
                    })
            
            certificate_key_management_section['subsections'].append({
                'type': 'certificate_key_management_table',
                'title': translations.get('last_5_certificates', 'Last 5 Certificates (by Expiry Date)'),
                'headers': [
                    translations.get('key_cert_num', 'Key/Cert Number'),
                    translations.get('company', 'Company'),
                    translations.get('type', 'Type'),
                    translations.get('expiry_date', 'Expiry Date'),
                    translations.get('expiry_status', 'Status'),
                    translations.get('days_until_expiry', 'Days Until Expiry'),
                    translations.get('revocation_status', 'Revocation Status'),
                    translations.get('owner', 'Owner'),
                    translations.get('location', 'Location'),
                    translations.get('purpose', 'Purpose')
                ],
                'data': certificates_list
            })
            
            content['sections'].append(certificate_key_management_section)
    
    # Quiz Results Section
    if selected_sections.get('quiz_results') and 'quiz_results_data' in data:
        quiz_results_data = data.get('quiz_results_data')
        # Check if data exists (even if empty, we should show the section if enabled)
        if quiz_results_data is not None:
            quiz_results_section = {
                'type': 'quiz_results',
                'title': translations.get('quiz_results', 'Security Training & Quiz Results'),
                'subsections': []
            }
            
            # Summary statistics
            if quiz_results_data.get('total_attempts', 0) > 0:
                stats_data = [
                    {
                        'label': translations.get('total_attempts', 'Total Attempts'),
                        'value': str(quiz_results_data.get('total_attempts', 0)),
                        'type': 'primary'
                    },
                    {
                        'label': translations.get('success_rate', 'Success Rate'),
                        'value': f"{quiz_results_data.get('success_rate', 0):.1f}%",
                        'type': 'success' if quiz_results_data.get('success_rate', 0) >= 80 else 'warning'
                    },
                    {
                        'label': translations.get('average_score', 'Average Score'),
                        'value': f"{quiz_results_data.get('average_score', 0):.1f}",
                        'type': 'info'
                    },
                    {
                        'label': translations.get('failed_attempts', 'Failed Attempts'),
                        'value': str(quiz_results_data.get('failed_attempts', 0)),
                        'type': 'danger' if quiz_results_data.get('failed_attempts', 0) > 0 else 'success'
                    }
                ]
                
                quiz_results_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('training_summary', 'Training Summary'),
                    'data': stats_data,
                    'display_inline': True  # Display all blocks in one row
                })
            
            # Quiz Statistics Table
            if quiz_results_data.get('quiz_statistics') and len(quiz_results_data['quiz_statistics']) > 0:
                quiz_stats_list = []
                for quiz_stat in quiz_results_data['quiz_statistics']:
                    quiz_stats_list.append({
                        'quiz_title': quiz_stat.get('quiz_title', ''),
                        'total_attempts': str(quiz_stat.get('total_attempts', 0)),
                        'successful_attempts': str(quiz_stat.get('successful_attempts', 0)),
                        'failed_attempts': str(quiz_stat.get('failed_attempts', 0)),
                        'success_rate': f"{quiz_stat.get('success_rate', 0):.1f}%",
                        'average_score': f"{quiz_stat.get('average_score', 0):.1f}",
                        'passing_score': str(quiz_stat.get('passing_score', 0))
                    })
                
                quiz_results_section['subsections'].append({
                    'type': 'quiz_statistics_table',
                    'title': translations.get('quiz_statistics', 'Quiz Statistics'),
                    'headers': [
                        translations.get('quiz_title', 'Quiz Title'),
                        translations.get('total_attempts', 'Total Attempts'),
                        translations.get('successful_attempts', 'Successful'),
                        translations.get('failed_attempts', 'Failed'),
                        translations.get('success_rate', 'Success Rate'),
                        translations.get('average_score', 'Average Score'),
                        translations.get('passing_score', 'Passing Score')
                    ],
                    'data': quiz_stats_list if quiz_stats_list else [{'quiz_title': translations.get('no_data', 'No data available'), 'total_attempts': '', 'successful_attempts': '', 'failed_attempts': '', 'success_rate': '', 'average_score': '', 'passing_score': ''}]
                })
            
            # Users at Risk Table
            if quiz_results_data.get('users_at_risk') and len(quiz_results_data['users_at_risk']) > 0:
                users_at_risk_list = []
                for user_risk in quiz_results_data['users_at_risk']:
                    users_at_risk_list.append({
                        'user': user_risk.get('user', ''),
                        'company': user_risk.get('company', ''),
                        'attempts_count': str(user_risk.get('attempts_count', 0))
                    })
                
                quiz_results_section['subsections'].append({
                    'type': 'users_at_risk_table',
                    'title': translations.get('users_requiring_training', 'Users Requiring Additional Training'),
                    'headers': [
                        translations.get('user', 'User'),
                        translations.get('company', 'Company'),
                        translations.get('attempts_count', 'Attempts Count')
                    ],
                    'data': users_at_risk_list
                })
            
            if quiz_results_section['subsections']:
                content['sections'].append(quiz_results_section)
    
    # Third Party Risk Section
    if selected_sections.get('third_party_risk') and 'third_party_risk_data' in data:
        third_party_risk_data = data.get('third_party_risk_data')
        # Check if data exists (even if empty, we should show the section if enabled)
        if third_party_risk_data is not None:
            third_party_risk_section = {
                'type': 'third_party_risk',
                'title': translations.get('third_party_risk', 'Third Party Risk'),
                'subsections': []
            }
            
            # Vendor Risk Summary - Statistics
            if third_party_risk_data.get('total_vendors', 0) >= 0:
                stats_data = [
                    {
                        'label': translations.get('total_vendors', 'Total Vendors'),
                        'value': str(third_party_risk_data.get('total_vendors', 0)),
                        'type': 'primary'
                    },
                    {
                        'label': translations.get('high_risk_vendors', 'High Risk Vendors'),
                        'value': str(third_party_risk_data.get('high_risk_vendors_count', 0)),
                        'type': 'danger' if third_party_risk_data.get('high_risk_vendors_count', 0) > 0 else 'success'
                    }
                ]
                
                # Add risk distribution
                risk_dist = third_party_risk_data.get('risk_distribution', {})
                if risk_dist:
                    for risk_level in ['critical', 'high', 'medium', 'low']:
                        count = risk_dist.get(risk_level, 0)
                        if count > 0:
                            risk_level_display = risk_level.capitalize()
                            stats_data.append({
                                'label': f"{translations.get('risk_level', 'Risk Level')}: {risk_level_display}",
                                'value': str(count),
                                'type': 'warning' if risk_level in ['high', 'critical'] else 'info'
                            })
                
                third_party_risk_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('vendor_risk_summary', 'Vendor Risk Summary'),
                    'data': stats_data,
                    'display_inline': True
                })
            
            # Risk Distribution Table
            risk_dist = third_party_risk_data.get('risk_distribution', {})
            if risk_dist and sum(risk_dist.values()) > 0:
                risk_dist_list = []
                total = sum(risk_dist.values())
                for risk_level in ['low', 'medium', 'high', 'critical']:
                    count = risk_dist.get(risk_level, 0)
                    if count > 0 or total == 0:
                        percentage = (count / total * 100) if total > 0 else 0
                        risk_dist_list.append({
                            'risk_level': risk_level.capitalize(),
                            'count': str(count),
                            'percentage': f"{percentage:.1f}%"
                        })
                
                third_party_risk_section['subsections'].append({
                    'type': 'risk_distribution_table',
                    'title': translations.get('risk_distribution_by_level', 'Risk Distribution by Level'),
                    'headers': [
                        translations.get('risk_level', 'Risk Level'),
                        translations.get('count', 'Count'),
                        translations.get('percentage', 'Percentage')
                    ],
                    'data': risk_dist_list if risk_dist_list else [{'risk_level': translations.get('no_data', 'No data available'), 'count': '', 'percentage': ''}]
                })
            
            # Assessment Statuses Table
            assessment_statuses = third_party_risk_data.get('assessment_statuses', {})
            if assessment_statuses:
                status_list = []
                for status_key in ['pending', 'in_progress', 'completed', 'overdue']:
                    count = assessment_statuses.get(status_key, 0)
                    status_display = {
                        'pending': translations.get('pending', 'Pending'),
                        'in_progress': translations.get('in_progress', 'In Progress'),
                        'completed': translations.get('completed', 'Completed'),
                        'overdue': translations.get('overdue', 'Overdue')
                    }.get(status_key, status_key.capitalize())
                    
                    status_list.append({
                        'status': status_display,
                        'count': str(count)
                    })
                
                third_party_risk_section['subsections'].append({
                    'type': 'assessment_status_table',
                    'title': translations.get('assessment_status', 'Assessment Status'),
                    'headers': [
                        translations.get('status', 'Status'),
                        translations.get('count', 'Count')
                    ],
                    'data': status_list
                })
            
            # Overdue Assessments Table
            overdue_assessments = third_party_risk_data.get('overdue_assessments', [])
            if overdue_assessments:
                overdue_list = []
                for assessment in overdue_assessments:
                    overdue_list.append({
                        'vendor_name': assessment.get('vendor_name', ''),
                        'assessment_date': assessment.get('assessment_date', ''),
                        'next_review_date': assessment.get('next_review_date', ''),
                        'status': assessment.get('status', ''),
                        'days_overdue': str(assessment.get('days_overdue', 0))
                    })
                
                third_party_risk_section['subsections'].append({
                    'type': 'overdue_assessments_table',
                    'title': translations.get('overdue_assessments', 'Overdue Assessments'),
                    'headers': [
                        translations.get('vendor_name', 'Vendor Name'),
                        translations.get('assessment_date', 'Assessment Date'),
                        translations.get('next_review_date', 'Next Review Date'),
                        translations.get('status', 'Status'),
                        translations.get('days_overdue', 'Days Overdue')
                    ],
                    'data': overdue_list
                })
            
            # Incomplete Questionnaires Table
            incomplete_questionnaires = third_party_risk_data.get('incomplete_questionnaires', [])
            if incomplete_questionnaires:
                incomplete_list = []
                for questionnaire in incomplete_questionnaires:
                    incomplete_list.append({
                        'vendor_name': questionnaire.get('vendor_name', ''),
                        'template_name': questionnaire.get('template_name', ''),
                        'status': questionnaire.get('status', ''),
                        'started_date': questionnaire.get('started_date', ''),
                        'percentage_score': f"{questionnaire.get('percentage_score', 0):.1f}%"
                    })
                
                third_party_risk_section['subsections'].append({
                    'type': 'incomplete_questionnaires_table',
                    'title': translations.get('incomplete_questionnaires', 'Incomplete Questionnaires'),
                    'headers': [
                        translations.get('vendor_name', 'Vendor Name'),
                        translations.get('template_name', 'Template Name'),
                        translations.get('status', 'Status'),
                        translations.get('started_date', 'Started Date'),
                        translations.get('percentage_score', 'Score')
                    ],
                    'data': incomplete_list
                })
            
            # Risk Level Trends
            risk_level_trends = third_party_risk_data.get('risk_level_trends', [])
            if risk_level_trends:
                trends_list = []
                for trend in risk_level_trends:
                    trends_list.append({
                        'risk_level': trend.get('risk_level', '').capitalize(),
                        'count': str(trend.get('count', 0)),
                        'trend': trend.get('trend', 'stable').capitalize()
                    })
                
                third_party_risk_section['subsections'].append({
                    'type': 'risk_level_trends_table',
                    'title': translations.get('risk_level_trends', 'Risk Level Trends'),
                    'headers': [
                        translations.get('risk_level', 'Risk Level'),
                        translations.get('count', 'Count'),
                        translations.get('trend', 'Trend')
                    ],
                    'data': trends_list
                })
            
            if third_party_risk_section['subsections']:
                content['sections'].append(third_party_risk_section)
    
    # GDPR Compliance Section
    if selected_sections.get('gdpr_compliance') and 'gdpr_compliance_data' in data:
        gdpr_compliance_data = data.get('gdpr_compliance_data')
        # Check if data exists (even if empty, we should show the section if enabled)
        if gdpr_compliance_data is not None:
            gdpr_compliance_section = {
                'type': 'gdpr_compliance',
                'title': translations.get('gdpr_compliance', 'GDPR Compliance'),
                'subsections': []
            }
            
            # ===== DATA BREACH INCIDENTS =====
            # Breach Summary Statistics
            if gdpr_compliance_data.get('total_breaches', 0) >= 0:
                breach_stats_data = [
                    {
                        'label': translations.get('total_breaches', 'Total Breaches'),
                        'value': str(gdpr_compliance_data.get('total_breaches', 0)),
                        'type': 'primary'
                    },
                    {
                        'label': translations.get('affected_subjects', 'Affected Subjects'),
                        'value': str(gdpr_compliance_data.get('total_affected_subjects', 0)),
                        'type': 'danger' if gdpr_compliance_data.get('total_affected_subjects', 0) > 0 else 'success'
                    }
                ]
                
                # Add severity distribution
                severity_dist = gdpr_compliance_data.get('severity_distribution', {})
                if severity_dist:
                    for severity in ['critical', 'high', 'medium', 'low']:
                        count = severity_dist.get(severity, 0)
                        if count > 0:
                            severity_display = severity.capitalize()
                            breach_stats_data.append({
                                'label': f"{translations.get('severity', 'Severity')}: {severity_display}",
                                'value': str(count),
                                'type': 'danger' if severity in ['high', 'critical'] else 'warning' if severity == 'medium' else 'info'
                            })
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('data_breach_incidents', 'Data Breach Incidents'),
                    'data': breach_stats_data,
                    'display_inline': True
                })
            
            # Severity Distribution Table
            severity_dist = gdpr_compliance_data.get('severity_distribution', {})
            if severity_dist and sum(severity_dist.values()) > 0:
                severity_dist_list = []
                total = sum(severity_dist.values())
                for severity in ['low', 'medium', 'high', 'critical']:
                    count = severity_dist.get(severity, 0)
                    if count > 0 or total == 0:
                        percentage = (count / total * 100) if total > 0 else 0
                        severity_dist_list.append({
                            'severity': severity.capitalize(),
                            'count': str(count),
                            'percentage': f"{percentage:.1f}%"
                        })
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'breach_severity_table',
                    'title': translations.get('breach_severity_distribution', 'Breach Severity Distribution'),
                    'headers': [
                        translations.get('severity', 'Severity'),
                        translations.get('count', 'Count'),
                        translations.get('percentage', 'Percentage')
                    ],
                    'data': severity_dist_list if severity_dist_list else [{'severity': translations.get('no_data', 'No data available'), 'count': '', 'percentage': ''}]
                })
            
            # Notification Statuses Table
            notification_statuses = gdpr_compliance_data.get('notification_statuses', {})
            if notification_statuses:
                notification_list = []
                for status_key in ['reported_on_time', 'reported_late', 'not_reported_overdue', 'not_reported_within_deadline']:
                    count = notification_statuses.get(status_key, 0)
                    status_display = {
                        'reported_on_time': translations.get('reported_on_time', 'Reported On Time'),
                        'reported_late': translations.get('reported_late', 'Reported Late'),
                        'not_reported_overdue': translations.get('not_reported_overdue', 'Not Reported (Overdue)'),
                        'not_reported_within_deadline': translations.get('not_reported_within_deadline', 'Not Reported (Within Deadline)')
                    }.get(status_key, status_key.replace('_', ' ').title())
                    
                    notification_list.append({
                        'status': status_display,
                        'count': str(count)
                    })
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'notification_status_table',
                    'title': translations.get('notification_statuses', 'Notification Statuses (GDPR 72-hour deadlines)'),
                    'headers': [
                        translations.get('status', 'Status'),
                        translations.get('count', 'Count')
                    ],
                    'data': notification_list
                })
            
            # Breach Notification Details (overdue/late)
            breach_notification_details = gdpr_compliance_data.get('breach_notification_details', [])
            if breach_notification_details:
                notification_details_list = []
                for breach in breach_notification_details:
                    notification_details_list.append({
                        'incident_number': breach.get('incident_number', ''),
                        'title': breach.get('title', ''),
                        'incident_date': breach.get('incident_date', ''),
                        'discovery_date': breach.get('discovery_date', ''),
                        'notification_deadline': breach.get('notification_deadline', ''),
                        'reported': translations.get('yes', 'Yes') if breach.get('reported_to_authority') else translations.get('no', 'No'),
                        'authority_report_date': breach.get('authority_report_date', ''),
                        'days_overdue': str(breach.get('days_overdue', 0))
                    })
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'breach_notification_details_table',
                    'title': translations.get('breach_notification_details', 'Breach Notification Details (Overdue/Late)'),
                    'headers': [
                        translations.get('incident_number', 'Incident Number'),
                        translations.get('title', 'Title'),
                        translations.get('incident_date', 'Incident Date'),
                        translations.get('discovery_date', 'Discovery Date'),
                        translations.get('notification_deadline', 'Notification Deadline'),
                        translations.get('reported', 'Reported'),
                        translations.get('authority_report_date', 'Authority Report Date'),
                        translations.get('days_overdue', 'Days Overdue')
                    ],
                    'data': notification_details_list
                })
            
            # ===== DATA PROCESSING ACTIVITIES =====
            # Activities Summary Statistics
            if gdpr_compliance_data.get('total_active_activities', 0) >= 0:
                activities_stats_data = [
                    {
                        'label': translations.get('active_activities', 'Active Activities'),
                        'value': str(gdpr_compliance_data.get('total_active_activities', 0)),
                        'type': 'primary'
                    },
                    {
                        'label': translations.get('international_transfers', 'International Transfers'),
                        'value': str(gdpr_compliance_data.get('international_transfers_count', 0)),
                        'type': 'warning' if gdpr_compliance_data.get('international_transfers_count', 0) > 0 else 'success'
                    },
                    {
                        'label': translations.get('outdated_retention', 'Outdated Retention'),
                        'value': str(len(gdpr_compliance_data.get('outdated_retention_activities', []))),
                        'type': 'danger' if len(gdpr_compliance_data.get('outdated_retention_activities', [])) > 0 else 'success'
                    },
                    {
                        'label': translations.get('non_compliant_retention', 'Non-Compliant Retention'),
                        'value': str(len(gdpr_compliance_data.get('non_compliant_retention_activities', []))),
                        'type': 'danger' if len(gdpr_compliance_data.get('non_compliant_retention_activities', [])) > 0 else 'success'
                    }
                ]
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('data_processing_activities', 'Data Processing Activities'),
                    'data': activities_stats_data,
                    'display_inline': True
                })
            
            # Data Categories Table
            data_categories = gdpr_compliance_data.get('data_categories', [])
            if data_categories:
                categories_list = []
                for category in data_categories:
                    categories_list.append({
                        'category': category.get('category', ''),
                        'count': str(category.get('count', 0))
                    })
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'data_categories_table',
                    'title': translations.get('data_categories', 'Data Categories'),
                    'headers': [
                        translations.get('category', 'Category'),
                        translations.get('count', 'Count')
                    ],
                    'data': categories_list
                })
            
            # International Transfers Table
            international_transfers_list = gdpr_compliance_data.get('international_transfers_list', [])
            if international_transfers_list:
                transfers_list = []
                for transfer in international_transfers_list:
                    transfers_list.append({
                        'activity_name': transfer.get('activity_name', ''),
                        'company': transfer.get('company', ''),
                        'transfer_safeguards': transfer.get('transfer_safeguards', '')
                    })
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'international_transfers_table',
                    'title': translations.get('international_data_transfers', 'International Data Transfers (Risk)'),
                    'headers': [
                        translations.get('activity_name', 'Activity Name'),
                        translations.get('company', 'Company'),
                        translations.get('transfer_safeguards', 'Transfer Safeguards')
                    ],
                    'data': transfers_list
                })
            
            # Outdated Retention Periods Table
            outdated_retention = gdpr_compliance_data.get('outdated_retention_activities', [])
            if outdated_retention:
                outdated_list = []
                for activity in outdated_retention:
                    outdated_list.append({
                        'activity_name': activity.get('activity_name', ''),
                        'company': activity.get('company', ''),
                        'retention_period_days': str(activity.get('retention_period_days', 0)),
                        'created_date': activity.get('created_date', ''),
                        'updated_date': activity.get('updated_date', ''),
                        'days_overdue': str(activity.get('days_overdue', 0))
                    })
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'outdated_retention_table',
                    'title': translations.get('outdated_retention_periods', 'Outdated Retention Periods'),
                    'headers': [
                        translations.get('activity_name', 'Activity Name'),
                        translations.get('company', 'Company'),
                        translations.get('retention_period_days', 'Retention Period (days)'),
                        translations.get('created_date', 'Created Date'),
                        translations.get('updated_date', 'Updated Date'),
                        translations.get('days_overdue', 'Days Overdue')
                    ],
                    'data': outdated_list
                })
            
            # Non-Compliant Retention Periods Table
            non_compliant_retention = gdpr_compliance_data.get('non_compliant_retention_activities', [])
            if non_compliant_retention:
                non_compliant_list = []
                for activity in non_compliant_retention:
                    issue = activity.get('issue', '')
                    issue_display = {
                        'excessive_retention': translations.get('excessive_retention', 'Excessive Retention (>10 years)'),
                        'insufficient_retention': translations.get('insufficient_retention', 'Insufficient Retention (<30 days)')
                    }.get(issue, issue.replace('_', ' ').title())
                    
                    non_compliant_list.append({
                        'activity_name': activity.get('activity_name', ''),
                        'company': activity.get('company', ''),
                        'retention_period_days': str(activity.get('retention_period_days', 0)),
                        'issue': issue_display
                    })
                
                gdpr_compliance_section['subsections'].append({
                    'type': 'non_compliant_retention_table',
                    'title': translations.get('non_compliant_retention_periods', 'Non-Compliant Retention Periods'),
                    'headers': [
                        translations.get('activity_name', 'Activity Name'),
                        translations.get('company', 'Company'),
                        translations.get('retention_period_days', 'Retention Period (days)'),
                        translations.get('issue', 'Issue')
                    ],
                    'data': non_compliant_list
                })
            
            if gdpr_compliance_section['subsections']:
                content['sections'].append(gdpr_compliance_section)
    
    # Access Risk Summary Section
    if selected_sections.get('access_risk_summary') and 'access_risk_summary_data' in data:
        access_risk_summary_data = data.get('access_risk_summary_data')
        # Check if data exists (even if empty, we should show the section if enabled)
        if access_risk_summary_data is not None:
            access_risk_section = {
                'type': 'access_risk_summary',
                'title': translations.get('access_risk_summary', 'Access Risk Summary'),
                'subsections': []
            }
            
            # Summary Statistics
            if access_risk_summary_data.get('total_active_accesses', 0) >= 0:
                stats_data = [
                    {
                        'label': translations.get('total_active_accesses', 'Total Active Accesses'),
                        'value': str(access_risk_summary_data.get('total_active_accesses', 0)),
                        'type': 'primary'
                    },
                    {
                        'label': translations.get('overdue_reviews', 'Overdue Access Reviews'),
                        'value': str(access_risk_summary_data.get('overdue_reviews_count', 0)),
                        'type': 'danger' if access_risk_summary_data.get('overdue_reviews_count', 0) > 0 else 'success'
                    },
                    {
                        'label': translations.get('third_party_access', 'Third-Party Access'),
                        'value': str(access_risk_summary_data.get('third_party_access_count', 0)),
                        'type': 'warning' if access_risk_summary_data.get('third_party_access_count', 0) > 0 else 'info'
                    },
                    {
                        'label': translations.get('privileged_access', 'Privileged Access'),
                        'value': str(access_risk_summary_data.get('privileged_access_count', 0)),
                        'type': 'danger' if access_risk_summary_data.get('privileged_access_count', 0) > 0 else 'info'
                    }
                ]
                
                access_risk_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('access_summary_statistics', 'Access Summary Statistics'),
                    'data': stats_data,
                    'display_inline': True
                })
            
            # Overdue Access Reviews Table
            overdue_reviews = access_risk_summary_data.get('overdue_reviews', [])
            if overdue_reviews:
                overdue_list = []
                for review in overdue_reviews:
                    overdue_list.append({
                        'asset_name': review.get('asset_name', ''),
                        'access_description': review.get('access_description', ''),
                        'last_review': review.get('last_review', '') or translations.get('never', 'Never'),
                        'reviewed_by': review.get('reviewed_by', ''),
                        'days_overdue': str(review.get('days_overdue', 0))
                    })
                
                access_risk_section['subsections'].append({
                    'type': 'overdue_reviews_table',
                    'title': translations.get('overdue_access_reviews', 'Overdue Access Reviews'),
                    'headers': [
                        translations.get('asset_name', 'Asset Name'),
                        translations.get('access_description', 'Access Description'),
                        translations.get('last_review', 'Last Review'),
                        translations.get('reviewed_by', 'Reviewed By'),
                        translations.get('days_overdue', 'Days Overdue')
                    ],
                    'data': overdue_list
                })
            
            # Third-Party Access Table
            third_party_access = access_risk_summary_data.get('third_party_access', [])
            if third_party_access:
                third_party_list = []
                for access in third_party_access:
                    third_party_list.append({
                        'asset_name': access.get('asset_name', ''),
                        'access_description': access.get('access_description', ''),
                        'last_review': access.get('last_review', '') or translations.get('never', 'Never'),
                        'start_date': access.get('start_date', '')
                    })
                
                access_risk_section['subsections'].append({
                    'type': 'third_party_access_table',
                    'title': translations.get('third_party_access_higher_risk', 'Third-Party Access (Higher Risk)'),
                    'headers': [
                        translations.get('asset_name', 'Asset Name'),
                        translations.get('access_description', 'Access Description'),
                        translations.get('last_review', 'Last Review'),
                        translations.get('start_date', 'Start Date')
                    ],
                    'data': third_party_list
                })
            
            # Access Without Last Review Table
            access_without_review = access_risk_summary_data.get('access_without_review', [])
            if access_without_review:
                without_review_list = []
                for access in access_without_review:
                    without_review_list.append({
                        'asset_name': access.get('asset_name', ''),
                        'access_description': access.get('access_description', ''),
                        'start_date': access.get('start_date', ''),
                        'days_since_start': str(access.get('days_since_start', 0))
                    })
                
                access_risk_section['subsections'].append({
                    'type': 'access_without_review_table',
                    'title': translations.get('access_without_last_review', 'Access Without Last Review'),
                    'headers': [
                        translations.get('asset_name', 'Asset Name'),
                        translations.get('access_description', 'Access Description'),
                        translations.get('start_date', 'Start Date'),
                        translations.get('days_since_start', 'Days Since Start')
                    ],
                    'data': without_review_list
                })
            
            # Privileged Access Section
            privileged_access = access_risk_summary_data.get('privileged_access', [])
            if privileged_access:
                privileged_list = []
                for access in privileged_access:
                    privileged_list.append({
                        'asset_name': access.get('asset_name', ''),
                        'access_description': access.get('access_description', ''),
                        'access_right': access.get('access_right', ''),
                        'roles': access.get('roles', ''),
                        'last_review': access.get('last_review', '') or translations.get('never', 'Never')
                    })
                
                access_risk_section['subsections'].append({
                    'type': 'privileged_access_table',
                    'title': translations.get('privileged_access_high_privileges', 'Privileged Access - High Privileges'),
                    'headers': [
                        translations.get('asset_name', 'Asset Name'),
                        translations.get('access_description', 'Access Description'),
                        translations.get('access_right', 'Access Right'),
                        translations.get('roles', 'Roles'),
                        translations.get('last_review', 'Last Review')
                    ],
                    'data': privileged_list
                })
            
            # Privileged Access Without Regular Reviews Table
            privileged_without_review = access_risk_summary_data.get('privileged_without_review', [])
            if privileged_without_review:
                privileged_no_review_list = []
                for access in privileged_without_review:
                    privileged_no_review_list.append({
                        'asset_name': access.get('asset_name', ''),
                        'access_description': access.get('access_description', ''),
                        'access_right': access.get('access_right', ''),
                        'roles': access.get('roles', ''),
                        'last_review': access.get('last_review', '') or translations.get('never', 'Never'),
                        'days_overdue': str(access.get('days_overdue', 0))
                    })
                
                access_risk_section['subsections'].append({
                    'type': 'privileged_without_review_table',
                    'title': translations.get('privileged_access_without_regular_reviews', 'Privileged Access Without Regular Reviews'),
                    'headers': [
                        translations.get('asset_name', 'Asset Name'),
                        translations.get('access_description', 'Access Description'),
                        translations.get('access_right', 'Access Right'),
                        translations.get('roles', 'Roles'),
                        translations.get('last_review', 'Last Review'),
                        translations.get('days_overdue', 'Days Overdue')
                    ],
                    'data': privileged_no_review_list
                })
            
            if access_risk_section['subsections']:
                content['sections'].append(access_risk_section)
    
    # SIEM Section
    if selected_sections.get('siem') and 'siem_data' in data:
        siem_data = data.get('siem_data')
        # Check if data exists (even if empty, we should show the section if enabled)
        if siem_data is not None:
            siem_section = {
                'type': 'siem',
                'title': translations.get('siem', 'SIEM'),
                'subsections': []
            }
            
            # Summary Statistics
            if siem_data.get('total_active_agents', 0) >= 0:
                stats_data = [
                    {
                        'label': translations.get('total_active_agents_wazuh', 'Total Active Wazuh Agents'),
                        'value': str(siem_data.get('total_active_agents', 0)),
                        'type': 'primary'
                    },
                    {
                        'label': translations.get('agents_with_problems', 'Agents with Problems'),
                        'value': str(siem_data.get('agents_with_problems_count', 0)),
                        'type': 'danger' if siem_data.get('agents_with_problems_count', 0) > 0 else 'success'
                    },
                    {
                        'label': translations.get('coverage_rate', 'Coverage Rate'),
                        'value': f"{siem_data.get('coverage_rate', 0):.1f}%",
                        'type': 'info'
                    },
                    {
                        'label': translations.get('monitored_assets', 'Monitored Assets'),
                        'value': f"{siem_data.get('monitored_assets', 0)}/{siem_data.get('total_assets', 0)}",
                        'type': 'info'
                    }
                ]
                
                siem_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('wazuh_agents_summary', 'Wazuh Agents Summary'),
                    'data': stats_data,
                    'display_inline': True
                })
            
            # Agents with Problems Table
            agents_with_problems = siem_data.get('agents_with_problems', [])
            if agents_with_problems:
                problems_list = []
                for agent in agents_with_problems:
                    problems_list.append({
                        'agent_name': agent.get('agent_name', ''),
                        'agent_id': agent.get('agent_id', ''),
                        'agent_ip': agent.get('agent_ip', ''),
                        'status': agent.get('status', ''),
                        'agent_version': agent.get('agent_version', '') or translations.get('not_specified', 'Not specified'),
                        'last_seen': agent.get('last_seen', '') or translations.get('never', 'Never'),
                        'problems': agent.get('problems', ''),
                        'days_since_seen': str(agent.get('days_since_seen', 0))
                    })
                
                siem_section['subsections'].append({
                    'type': 'agents_with_problems_table',
                    'title': translations.get('agents_with_problems_inactive_outdated', 'Agents with Problems (inactive, outdated)'),
                    'headers': [
                        translations.get('agent_name', 'Agent Name'),
                        translations.get('agent_id', 'Agent ID'),
                        translations.get('agent_ip', 'Agent IP'),
                        translations.get('status', 'Status'),
                        translations.get('agent_version', 'Version'),
                        translations.get('last_seen', 'Last Seen'),
                        translations.get('problems', 'Problems'),
                        translations.get('days_since_seen', 'Days Since Seen')
                    ],
                    'data': problems_list
                })
            
            # Coverage Rate Information
            coverage_rate = siem_data.get('coverage_rate', 0)
            monitored_assets = siem_data.get('monitored_assets', 0)
            total_assets = siem_data.get('total_assets', 0)
            
            if total_assets > 0:
                coverage_info = [
                    {
                        'label': translations.get('total_assets', 'Total Assets'),
                        'value': str(total_assets),
                        'type': 'info'
                    },
                    {
                        'label': translations.get('monitored_assets', 'Monitored Assets'),
                        'value': str(monitored_assets),
                        'type': 'success' if monitored_assets > 0 else 'warning'
                    },
                    {
                        'label': translations.get('coverage_rate', 'Coverage Rate'),
                        'value': f"{coverage_rate:.1f}%",
                        'type': 'primary' if coverage_rate >= 80 else 'warning' if coverage_rate >= 50 else 'danger'
                    }
                ]
                
                siem_section['subsections'].append({
                    'type': 'statistics',
                    'title': translations.get('coverage_rate_monitoring', 'Coverage Rate (Assets Under Monitoring)'),
                    'data': coverage_info,
                    'display_inline': True
                })
            
            # Threat Detection Trends
            threat_detection_trends = siem_data.get('threat_detection_trends', [])
            if threat_detection_trends:
                trends_list = []
                for trend in threat_detection_trends:
                    trends_list.append({
                        'date': trend.get('date', ''),
                        'alert_count': str(trend.get('alert_count', 0)),
                        'trend': trend.get('trend', 'stable').capitalize()
                    })
                
                siem_section['subsections'].append({
                    'type': 'threat_detection_trends_table',
                    'title': translations.get('threat_detection_trends', 'Threat Detection Trends'),
                    'headers': [
                        translations.get('date', 'Date'),
                        translations.get('alert_count', 'Alert Count'),
                        translations.get('trend', 'Trend')
                    ],
                    'data': trends_list
                })
            
            if siem_section['subsections']:
                content['sections'].append(siem_section)
    
    # Asset Details Section - Combined Assets by Criticality and TOP-10 Asset
    if selected_sections.get('asset_details') or selected_sections.get('asset_tables'):
        asset_details_section = {
            'type': 'asset_details_combined',
            'title': translations.get('asset_details', 'Asset Details'),
            'subsections': []
        }
        
        # Assets by Criticality
        if selected_sections.get('asset_details') and 'statistics' in data:
            stats = data['statistics']
            if stats.get('assets_by_criticality'):
                criticality_data = []
                for item in stats['assets_by_criticality']:
                    criticality_name = item.get('criticality__name_uk', translations.get('unknown', 'Unknown'))
                    criticality_data.append({
                        'criticality': criticality_name,
                        'count': str(item.get('count', 0)),
                        'percentage': str(item.get('percentage', 0)),
                        'color': item.get('color', '#000000'),
                        'cost': item.get('cost', 0)
                    })
                
                # Sort by cost in descending order (highest criticality first)
                criticality_data.sort(key=lambda x: x['cost'], reverse=True)
                
                asset_details_section['subsections'].append({
                    'type': 'assets_by_criticality',
                    'title': translations.get('assets_by_criticality', 'Assets by Criticality'),
                    'headers': [translations['criticality'], translations['count'], translations.get('percentage', 'Percentage')],
                    'data': criticality_data
                })
        
        # TOP-10 Asset
        if selected_sections.get('asset_tables') and 'assets' in data:
            asset_table_data = []
            # Get vulnerability data for assets
            assets_vuln_data = {}
            if 'statistics' in data and 'assets_vulnerability_data' in data['statistics']:
                for vuln_info in data['statistics']['assets_vulnerability_data']:
                    assets_vuln_data[vuln_info['asset_id']] = vuln_info
            
            for asset in data['assets']:
                # Handle both dictionary and model object cases
                if hasattr(asset, 'get'):  # Dictionary
                    criticality_cost = asset.get('criticality_cost', 0)
                    asset_id = asset.get('id', asset.get('asset_id', ''))
                    # Get vulnerability count info
                    vuln_info = assets_vuln_data.get(asset_id, {})
                    vulnerabilities_display = 'Undefined' if vuln_info.get('is_undefined', True) else str(vuln_info.get('vulnerabilities_count', 0))
                    
                    asset_table_data.append({
                        'asset_id': asset.get('asset_id', ''),
                        'asset_name': asset.get('name', ''),
                        'criticality': asset.get('criticality_name', ''),
                        'criticality_cost': criticality_cost,
                        'vulnerabilities': vulnerabilities_display
                    })
                else:  # Model object
                    # Get criticality using the proper method
                    criticality_info = asset.get_criticality()
                    criticality_name = criticality_info.get('name', '') if criticality_info else ''
                    criticality_color = criticality_info.get('color', '#000000') if criticality_info else '#000000'
                    criticality_cost = criticality_info.get('cost', 0) if criticality_info else 0
                    
                    # Get vulnerability count info
                    asset_id = getattr(asset, 'id', None)
                    vuln_info = assets_vuln_data.get(asset_id, {}) if asset_id else {}
                    vulnerabilities_display = 'Undefined' if vuln_info.get('is_undefined', True) else str(vuln_info.get('vulnerabilities_count', 0))
                    
                    asset_table_data.append({
                        'asset_id': getattr(asset, 'asset_id', ''),
                        'asset_name': getattr(asset, 'name', ''),
                        'criticality': criticality_name,
                        'criticality_color': criticality_color,
                        'criticality_cost': criticality_cost,
                        'vulnerabilities': vulnerabilities_display
                    })
            
            # Sort by criticality cost (highest first) and limit to 10
            asset_table_data.sort(key=lambda x: x['criticality_cost'], reverse=True)
            asset_table_data = asset_table_data[:10]
            
            asset_details_section['subsections'].append({
                'type': 'top_10_asset',
                'title': translations.get('top_10_asset', 'TOP-10 Asset'),
                'headers': [
                    translations.get('asset_id', 'Asset ID'),
                    translations.get('asset_name', 'Asset Name'),
                    translations['criticality'],
                    translations.get('vulnerabilities', 'Vulnerabilities'),
                    translations.get('company', 'Company')
                ],
                'data': asset_table_data
            })
        
        # Add the combined section to content
        if asset_details_section['subsections']:
            content['sections'].append(asset_details_section)
    
    # Vulnerability Details Section
    if selected_sections.get('vulnerability_details') or selected_sections.get('vulnerability_tables'):
        vuln_section = {
            'type': 'vulnerability_details_combined',
            'title': translations.get('vulnerability_details', 'Vulnerability Details'),
            'subsections': []
        }
        
        # Initialize vuln_data outside the condition to avoid "referenced before assignment" error
        vuln_data = []
        
        if selected_sections.get('vulnerability_details') and 'vulnerabilities' in data:
            
            # Process vulnerabilities and calculate additional data
            processed_vulns = []
            for vuln in data['vulnerabilities']:
                # Handle both dictionary and model object cases
                if hasattr(vuln, 'get'):  # Dictionary
                    asset_name = vuln.get('asset_name', '')
                    vulnerability_name = vuln.get('vulnerability_name', '')
                    status = vuln.get('status', '')
                    # For dictionary case, we need to get the actual objects
                    asset_obj = None
                    vuln_obj = None
                else:  # Model object
                    asset_obj = vuln.asset if hasattr(vuln, 'asset') else None
                    vuln_obj = vuln.vulnerability if hasattr(vuln, 'vulnerability') else None
                    asset_name = getattr(vuln, 'asset_name', '') if hasattr(vuln, 'asset_name') else (getattr(asset_obj, 'name', '') if asset_obj else '')
                    vulnerability_name = getattr(vuln, 'vulnerability_name', '') if hasattr(vuln, 'vulnerability_name') else (get_localized_vulnerability_field(vuln_obj, 'vulnerability', language) if vuln_obj else '')
                    status = getattr(vuln, 'status', '')
                
                # Calculate additional data
                threat_info = ''
                probability_impact = ''
                impact_value = ''
                risk_value = ''
                risk_level = ''
                risk_mitigation_controls = ''
                asset_criticality_cost = 0
                
                if asset_obj and vuln_obj:
                    # Get asset criticality for sorting
                    try:
                        criticality = asset_obj.get_criticality()
                        asset_criticality_cost = criticality.get('cost', 0)
                    except:
                        asset_criticality_cost = 0
                    
                    # Get risk mitigation controls and limit to 400 characters
                    risk_mitigation_controls = get_localized_vulnerability_field(vuln_obj, 'risk_mitigation_controls', language)
                    if risk_mitigation_controls and len(risk_mitigation_controls) > 400:
                        # Truncate to 400 characters and ensure it ends with a complete sentence
                        truncated = risk_mitigation_controls[:400]
                        # Find the last sentence ending (period, exclamation, question mark)
                        last_sentence_end = max(
                            truncated.rfind('.'), 
                            truncated.rfind('!'), 
                            truncated.rfind('?')
                        )
                        if last_sentence_end > 350:  # If we have a sentence ending in the last 50 chars
                            risk_mitigation_controls = truncated[:last_sentence_end + 1]
                        else:
                            # If no sentence ending found, just truncate and add ellipsis
                            risk_mitigation_controls = truncated + '...'
                    
                    # Calculate threat and risk information for all vulnerabilities (regardless of status)
                    try:
                        risk_value_calc = calculate_value_of_risk(asset_obj, vuln_obj)
                        risk_value = f"{risk_value_calc:.2f}"
                        
                        # Debug logging for risk calculation
                        logger.debug(f"Risk calculation for {vulnerability_name}: value={risk_value_calc}, asset={asset_name}, status={status}")
                        
                        risk_level_obj = calculate_risk_level(risk_value_calc)
                        if risk_level_obj:
                            risk_level = risk_level_obj.get_name_by_language(language) or risk_level_obj.get_name()
                            logger.debug(f"Found risk level: {risk_level} for value {risk_value_calc}")
                        else:
                            # If no risk level found, try to find the highest level for high risk values
                            if risk_value_calc > 50:  # High risk threshold
                                try:
                                    # First try to find "Extraordinary" level specifically
                                    extraordinary_level = RiskLevel.objects.filter(
                                        Q(name__icontains='Extraordinary') | 
                                        Q(translations__name_local__icontains='Extraordinary')
                                    ).first()
                                    
                                    if extraordinary_level:
                                        risk_level = extraordinary_level.get_name_by_language(language) or extraordinary_level.get_name()
                                        logger.debug(f"Using Extraordinary risk level: {risk_level} for high risk value {risk_value_calc}")
                                    else:
                                        # Fallback to highest level
                                        highest_level = RiskLevel.objects.order_by('-max_value').first()
                                        if highest_level:
                                            risk_level = highest_level.get_name_by_language(language) or highest_level.get_name()
                                            logger.debug(f"Using highest risk level: {risk_level} for high risk value {risk_value_calc}")
                                        else:
                                            risk_level = translations.get('unknown', 'Unknown')
                                            logger.warning(f"No risk levels found in database for high risk value {risk_value_calc}")
                                except Exception as e:
                                    risk_level = translations.get('unknown', 'Unknown')
                                    logger.error(f"Error finding highest risk level: {str(e)}")
                            else:
                                risk_level = translations.get('unknown', 'Unknown')
                                logger.debug(f"No risk level found for low risk value {risk_value_calc}")
                    except Exception as e:
                        logger.warning(f"Error calculating risk for vulnerability {vulnerability_name}: {str(e)}")
                        risk_value = '0.00'
                        risk_level = translations.get('unknown', 'Unknown')
                    
                    # Get threat information for all vulnerabilities
                    threats = vuln_obj.threats.all()
                    if threats.exists():
                        threat_names = []
                        for threat in threats:
                            threat_name = threat.get_name() or getattr(threat, f'name_{language}', '')
                            threat_names.append(threat_name)
                            
                            # Calculate probability/impact for this threat
                            try:
                                probability = threat.calculate_probability()
                                impact = threat.calculate_overall_impact() if hasattr(threat, 'calculate_overall_impact') else threat.impact / 100
                                prob_impact = probability * impact * 100
                                probability_impact = f"{prob_impact:.2f}%"
                                impact_value = f"{impact * 100:.2f}%"
                            except:
                                probability_impact = '0.00%'
                                impact_value = '0.00%'
                        
                        threat_info = ', '.join(threat_names)
                    else:
                        threat_info = translations.get('no_threats', 'No threats')
                        probability_impact = '0.00%'
                        impact_value = '0.00%'
                
                # Calculate risk level priority for sorting (higher number = higher priority)
                risk_level_priority = 0
                if risk_level:
                    if 'Надзвичайний' in risk_level or 'Extraordinary' in risk_level:
                        risk_level_priority = 6
                    elif 'Критичний' in risk_level or 'Critical' in risk_level:
                        risk_level_priority = 5
                    elif 'Високий' in risk_level or 'High' in risk_level:
                        risk_level_priority = 4
                    elif 'Середній' in risk_level or 'Medium' in risk_level:
                        risk_level_priority = 3
                    elif 'Низький' in risk_level or 'Low' in risk_level:
                        risk_level_priority = 2
                    elif 'Невизначено' in risk_level or 'Undefined' in risk_level:
                        risk_level_priority = 1
                
                processed_vulns.append({
                    'asset_name': asset_name,
                        'vulnerability_name': vulnerability_name,
                    'status': status,
                    'threat': threat_info,
                    'probability_impact': probability_impact,
                    'impact': impact_value,
                    'risk_value': risk_value,
                    'risk_level': risk_level,
                    'risk_mitigation_controls': risk_mitigation_controls,
                    'asset_criticality_cost': asset_criticality_cost,
                    'status_priority': 1 if status == 'Yes' else (2 if status == 'Undefined' else 3),
                    'risk_level_priority': risk_level_priority
                })
            
            # Sort vulnerabilities: first by status (Yes first), then by asset criticality (highest first), then by risk level (highest first)
            processed_vulns.sort(key=lambda x: (x['status_priority'], -x['asset_criticality_cost'], -x['risk_level_priority']))
            
            # Take top 10
            for vuln in processed_vulns[:10]:
                vuln_data.append({
                    'asset_name': vuln['asset_name'],
                    'vulnerability_name': vuln['vulnerability_name'],
                    'status': vuln['status'],
                    'threat': vuln['threat'],
                    'probability_impact': vuln['probability_impact'],
                    'impact': vuln['impact'],
                    'risk_value': vuln['risk_value'],
                    'risk_level': vuln['risk_level'],
                    'risk_mitigation_controls': vuln['risk_mitigation_controls']
                    })
            
        # Add Assets by Vulnerability subsection first (left side)
        if selected_sections.get('vulnerability_details') and 'statistics' in data:
            stats = data['statistics']
            if stats.get('assets_by_vulnerability'):
                vuln_by_asset_data = []
                for item in stats['assets_by_vulnerability']:
                    vuln_by_asset_data.append({
                        'criticality': item.get('criticality', ''),
                        'count': item.get('count', 0),
                        'percentage': item.get('percentage', 0),
                        'color': item.get('color', '#000000')
                    })
                
                vuln_section['subsections'].append({
                    'type': 'assets_by_vulnerability',
                    'title': translations.get('assets_by_vulnerability', 'Assets by Vulnerability'),
                    'headers': [
                        translations.get('criticality', 'Criticality'),
                        translations.get('count', 'Count'),
                        translations.get('percentage', 'Percentage')
                    ],
                    'data': vuln_by_asset_data
                })
        
        # Add TOP-10 Vulnerability subsection second (right side)
        vuln_section['subsections'].append({
            'type': 'top_10_vulnerability',
                'title': translations.get('top_10_vulnerability', 'TOP-10 Vulnerability'),
                'headers': [
                    translations.get('asset', 'Asset'),
                    translations['vulnerability'],
                    translations['status'],
                    translations['threat'],
                    translations['probability_impact'],
                    translations['impact'],
                    translations['risk_value'],
                    translations['risk_level'],
                    translations['risk_mitigation_controls']
                ],
                'data': vuln_data
            })
        
        if vuln_section['subsections']:
            content['sections'].append(vuln_section)
    
    # Treatment Details Section
    if selected_sections.get('treatment_details') or selected_sections.get('treatment_tables'):
        treatment_section = {
            'type': 'treatment_details',
            'title': translations.get('treatment_details', 'Treatment Details'),
            'subsections': []
        }
        
        if selected_sections.get('treatment_details') and 'statistics' in data:
            stats = data['statistics']
            if stats.get('overdue_treatments'):
                treatment_data = []
                for treatment in stats['overdue_treatments'][:10]:  # Limit to 10
                    treatment_data.append({
                        'asset_name': treatment.get('asset_name', translations.get('unknown', 'Unknown')),
                        'treatment': treatment.get('vulnerability_name', translations.get('unknown', 'Unknown')),
                        'deadline': treatment.get('due_date', translations.get('unknown', 'Unknown'))
                    })
                
                treatment_section['subsections'].append({
                    'type': 'overdue_treatments',
                    'title': translations.get('overdue_treatments', 'Overdue Treatments'),
                    'headers': [translations['asset'], translations.get('treatment', 'Treatment'), translations.get('deadline', 'Deadline')],
                    'data': treatment_data
                })
        
        # Add Treatment Details subsection if requested
        if selected_sections.get('treatment_details') and 'risk_treatments' in data:
            treatment_details_data = []
            for treatment in data['risk_treatments']:  # Process all treatments for sorting
                # Get vulnerability status and risk level if available
                vulnerability_status = translations.get('unknown', 'Unknown')
                risk_level = translations.get('unknown', 'Unknown')
                
                # Try to find the corresponding AssetVulnerability for this treatment
                if 'asset_vulnerabilities' in data:
                    for vuln in data['asset_vulnerabilities']:
                        # Handle both dictionary and model object cases
                        if hasattr(vuln, 'get'):  # Dictionary
                            vuln_asset_name = vuln.get('asset_name', '')
                            vuln_name = vuln.get('vulnerability_name', '')
                        else:  # Model object
                            vuln_asset_name = getattr(vuln.asset, 'name', '') if hasattr(vuln, 'asset') and vuln.asset else ''
                            vuln_name = get_localized_vulnerability_field(vuln.vulnerability, 'vulnerability', language) if hasattr(vuln, 'vulnerability') and vuln.vulnerability else ''
                        
                        # Match by asset name and vulnerability name
                        if (vuln_asset_name == treatment.get('asset_name', '') and 
                            vuln_name == treatment.get('vulnerability_name', '')):
                            vulnerability_status = getattr(vuln, 'status', translations.get('unknown', 'Unknown'))
                            
                            # Calculate risk level for this vulnerability
                            try:
                                if hasattr(vuln, 'asset') and hasattr(vuln, 'vulnerability'):
                                    # For model objects
                                    risk_value = calculate_value_of_risk(vuln.asset, vuln.vulnerability)
                                    risk_level_obj = calculate_risk_level(risk_value)
                                    if risk_level_obj:
                                        risk_level = risk_level_obj.get_name_by_language(language) or risk_level_obj.get_name()
                                    else:
                                        risk_level = translations.get('unknown', 'Unknown')
                                elif hasattr(vuln, 'get'):
                                    # For dictionary objects, try to get risk level if available
                                    risk_level = vuln.get('risk_level', translations.get('unknown', 'Unknown'))
                                else:
                                    risk_level = translations.get('unknown', 'Unknown')
                            except Exception as e:
                                logger.warning(f"Error calculating risk level: {e}")
                                risk_level = translations.get('unknown', 'Unknown')
                            break
                
                # Calculate priority for sorting
                status_priority = 0
                if vulnerability_status == 'Yes':
                    status_priority = 1
                
                risk_level_priority = 0
                if risk_level:
                    if 'Надзвичайний' in risk_level or 'Extraordinary' in risk_level:
                        risk_level_priority = 6
                    elif 'Критичний' in risk_level or 'Critical' in risk_level:
                        risk_level_priority = 5
                    elif 'Високий' in risk_level or 'High' in risk_level:
                        risk_level_priority = 4
                    elif 'Середній' in risk_level or 'Medium' in risk_level:
                        risk_level_priority = 3
                    elif 'Низький' in risk_level or 'Low' in risk_level:
                        risk_level_priority = 2
                    elif 'Невизначено' in risk_level or 'Undefined' in risk_level:
                        risk_level_priority = 1
                
                treatment_details_data.append({
                    'asset_name': treatment.get('asset_name', translations.get('unknown', 'Unknown')),
                    'vulnerability_name': treatment.get('vulnerability_name', translations.get('unknown', 'Unknown')),
                    'vulnerability_status': vulnerability_status,
                    'risk_level': risk_level,
                    'treatment_type': treatment.get('treatment_type', translations.get('unknown', 'Unknown')),
                    'treatment_status': treatment.get('status', translations.get('unknown', 'Unknown')),
                    'description_responsible': treatment.get('description', '') + ' | ' + treatment.get('assigned_to', translations.get('unknown', 'Unknown')),
                    'due_date': treatment.get('due_date', translations.get('unknown', 'Unknown')),
                    'status_priority': status_priority,
                    'risk_level_priority': risk_level_priority
                })
            
            # Sort by status priority (Yes first) and then by risk level priority (highest first)
            treatment_details_data.sort(key=lambda x: (x['status_priority'], x['risk_level_priority']), reverse=True)
            
            # Limit to top 10
            treatment_details_data = treatment_details_data[:10]
            
            # Remove priority fields from final data
            for item in treatment_details_data:
                item.pop('status_priority', None)
                item.pop('risk_level_priority', None)
            
            treatment_section['subsections'].append({
                'type': 'treatment_details',
                'title': translations.get('top_10_treatment', 'TOP-10 Treatment'),
                'headers': [translations['asset'], translations.get('vulnerability', 'Vulnerability'), translations.get('vulnerability_status', 'Vulnerability Status'), translations.get('risk_level', 'Risk Level'), translations.get('treatment_type', 'Treatment Type'), translations.get('treatment_status', 'Treatment Status'), translations.get('description_responsible', 'Description & Responsible'), translations.get('due_date', 'Due Date')],
                'data': treatment_details_data
            })
        
        if treatment_section['subsections']:
            content['sections'].append(treatment_section)
    
    # Recommendations Section
    if selected_sections.get('recommendations'):
        recommendations_section = {
            'type': 'recommendations',
            'title': translations['recommendations'],
            'subsections': [{
                'type': 'priority_actions',
                'title': translations.get('priority_actions', 'Priority Actions'),
                'data': [
                    translations.get('address_high_risk_vulnerabilities', 'Address high-risk vulnerabilities'),
                    translations.get('improve_compliance_rates', 'Improve compliance rates'),
                    translations.get('implement_regular_assessments', 'Implement regular assessments'),
                    translations.get('update_security_policies', 'Update security policies')
                ]
            }]
        }
        content['sections'].append(recommendations_section)
    
    # Financial Analysis Section
    if selected_sections.get('financial_analysis') or selected_sections.get('cost_benefit_analysis') or selected_sections.get('roi_assessment') or selected_sections.get('budget_allocation'):
        financial_section = {
            'type': 'financial_analysis',
            'title': translations.get('financial_analysis', 'Financial Analysis'),
            'subsections': []
        }
        
        if selected_sections.get('financial_analysis') and 'statistics' in data:
            stats = data['statistics']
            
            # Ensure stats is a dictionary
            if isinstance(stats, dict):
                financial_section['subsections'].append({
                    'type': 'financial_overview',
                    'title': translations.get('financial_overview', 'Financial Overview'),
                    'data': [
                        {
                            'label': translations.get('total_cost', 'Total Cost'),
                            'value': f"${stats.get('total_cost', 0):,.2f}"
                        },
                        {
                            'label': translations.get('potential_savings', 'Potential Savings'),
                            'value': f"${stats.get('potential_savings', 0):,.2f}"
                        },
                        {
                            'label': translations.get('roi_percentage', 'ROI'),
                            'value': f"{stats.get('roi_percentage', 0):.1f}%"
                        },
                        {
                            'label': translations.get('budget_allocated', 'Budget Allocated'),
                            'value': f"${stats.get('budget_allocated', 0):,.2f}"
                        }
                    ]
                })
            else:
                # Fallback if stats is not a dictionary
                financial_section['subsections'].append({
                    'type': 'financial_overview',
                    'title': translations.get('financial_overview', 'Financial Overview'),
                    'data': [
                        {
                            'label': translations.get('total_cost', 'Total Cost'),
                            'value': '$0.00'
                        },
                        {
                            'label': translations.get('potential_savings', 'Potential Savings'),
                            'value': '$0.00'
                        },
                        {
                            'label': translations.get('roi_percentage', 'ROI'),
                            'value': '0.0%'
                        },
                        {
                            'label': translations.get('budget_allocated', 'Budget Allocated'),
                            'value': '$0.00'
                        }
                    ]
                })
        
        content['sections'].append(financial_section)
    
    # Residual Risk & Priority Section
    if selected_sections.get('residual_risk_analysis') or selected_sections.get('risk_appetite') or selected_sections.get('priority_matrix') or selected_sections.get('resource_allocation'):
        residual_section = {
            'type': 'residual_risk_priority',
            'title': translations.get('residual_risk_priority', 'Residual Risk & Priority'),
            'subsections': []
        }
        
        if selected_sections.get('residual_risk_analysis') and 'statistics' in data:
            stats = data['statistics']
            
            # Ensure stats is a dictionary
            if isinstance(stats, dict):
                residual_section['subsections'].append({
                    'type': 'residual_overview',
                    'title': translations.get('residual_overview', 'Residual Risk Overview'),
                    'data': [
                        {
                            'label': translations.get('residual_risk_level', 'Residual Risk Level'),
                            'value': stats.get('residual_risk_level', translations.get('unknown', 'Unknown'))
                        },
                        {
                            'label': translations.get('appetite_threshold', 'Risk Appetite Threshold'),
                            'value': f"{stats.get('appetite_threshold', 0):.1f}%"
                        },
                        {
                            'label': translations.get('priority_level', 'Priority Level'),
                            'value': stats.get('priority_level', translations.get('unknown', 'Unknown'))
                        }
                    ]
                })
            else:
                # Fallback if stats is not a dictionary
                residual_section['subsections'].append({
                    'type': 'residual_overview',
                    'title': translations.get('residual_overview', 'Residual Risk Overview'),
                    'data': [
                        {
                            'label': translations.get('residual_risk_level', 'Residual Risk Level'),
                            'value': translations.get('unknown', 'Unknown')
                        },
                        {
                            'label': translations.get('appetite_threshold', 'Risk Appetite Threshold'),
                            'value': '0.0%'
                        },
                        {
                            'label': translations.get('priority_level', 'Priority Level'),
                            'value': translations.get('unknown', 'Unknown')
                        }
                    ]
                })
        
        content['sections'].append(residual_section)
    
    # Acceptable Risk Section
    if selected_sections.get('acceptable_risk_analysis') and 'acceptable_risk_data' in data:
        acceptable_risk_section = {
            'type': 'acceptable_risk_analysis',
            'title': translations.get('acceptable_risk_analysis', 'Acceptable Risk Analysis'),
            'subsections': []
        }
        
        # Summary statistics
        if data['acceptable_risk_data']['summary']:
            summary = data['acceptable_risk_data']['summary']
            acceptable_risk_section['subsections'].append({
                'type': 'statistics',
                'title': translations.get('acceptable_risk_summary', 'Acceptable Risk Summary'),
                'data': [
                    {
                        'label': translations.get('total_assets', 'Total Assets'),
                        'value': str(summary['total_assets']),
                        'type': 'primary'
                    },
                    {
                        'label': translations.get('assets_with_acceptable_risk', 'Assets with Acceptable Risk'),
                        'value': str(summary['assets_with_acceptable_risk']),
                        'type': 'success'
                    },
                    {
                        'label': translations.get('assets_without_acceptable_risk', 'Assets without Acceptable Risk'),
                        'value': str(summary['assets_without_acceptable_risk']),
                        'type': 'warning'
                    },
                    {
                        'label': translations.get('exceeding_acceptable_risk', 'Exceeding Acceptable Risk'),
                        'value': str(summary['exceeding_acceptable_risk']),
                        'type': 'danger'
                    }
                ]
            })
        
        # Assets with acceptable risk configuration
        if data['acceptable_risk_data']['assets_with_acceptable_risk']:
            acceptable_risk_section['subsections'].append({
                'type': 'table',
                'title': translations.get('top_10_assets_with_acceptable_risk_config', 'TOP-10 Assets with Acceptable Risk Configuration'),
                'headers': [
                    translations.get('asset_id', 'Asset ID'),
                    translations.get('asset_name', 'Asset Name'),
                    translations.get('current_criticality', 'Current Criticality'),
                    translations.get('acceptable_risk_level', 'Acceptable Risk Level'),
                    translations.get('exceeds_acceptable', 'Exceeds Acceptable')
                ],
                'data': [
                    {
                        'asset_id': asset['asset_id'],
                        'asset_name': asset['asset_name'],
                        'current_criticality': asset['current_criticality'],
                        'acceptable_risk_level': asset['acceptable_risk_level'],
                        'exceeds_acceptable': 'Yes' if asset['exceeds_acceptable'] else 'No'
                    }
                    for asset in data['acceptable_risk_data']['assets_with_acceptable_risk']
                ]
            })
        
        # Assets without acceptable risk configuration
        if data['acceptable_risk_data']['assets_without_acceptable_risk']:
            acceptable_risk_section['subsections'].append({
                'type': 'table',
                'title': translations.get('top_10_assets_without_acceptable_risk_config', 'TOP-10 Assets without Acceptable Risk Configuration'),
                'headers': [
                    translations.get('asset_id', 'Asset ID'),
                    translations.get('asset_name', 'Asset Name'),
                    translations.get('current_criticality', 'Current Criticality')
                ],
                'data': [
                    {
                        'asset_id': asset['asset_id'],
                        'asset_name': asset['asset_name'],
                        'current_criticality': asset['current_criticality']
                    }
                    for asset in data['acceptable_risk_data']['assets_without_acceptable_risk']
                ]
            })
        
        content['sections'].append(acceptable_risk_section)
    
    # AI Conclusion Section
    # Check if conclusions is enabled (handle both boolean and string values)
    conclusions_enabled = selected_sections.get('conclusions', False)
    logger.info(f"AI Conclusion check: conclusions={conclusions_enabled}, type={type(conclusions_enabled)}, model_id={selected_sections.get('_ai_conclusion_model')}")
    if conclusions_enabled == True or (isinstance(conclusions_enabled, str) and conclusions_enabled.lower() in ('true', '1')):
        logger.info(f"Generating AI Conclusion. Model ID: {selected_sections.get('_ai_conclusion_model')}")
        ai_conclusion_text = generate_ai_conclusion(data, selected_sections, language)
        logger.info(f"AI Conclusion generated: {ai_conclusion_text[:100] if ai_conclusion_text else 'None'}...")
        if ai_conclusion_text:
            conclusion_section = {
                'type': 'ai_conclusion',
                'title': translations.get('conclusions', 'Conclusions'),
                'content': ai_conclusion_text,
                'model_name': None
            }
            
            # Get AI model name if available
            ai_model_id = selected_sections.get('_ai_conclusion_model')
            if ai_model_id:
                # Convert to int if it's a string
                if isinstance(ai_model_id, str):
                    try:
                        ai_model_id = int(ai_model_id)
                    except ValueError:
                        logger.error(f"Invalid AI model ID format: {ai_model_id}")
                        ai_model_id = None
                
                if ai_model_id:
                    try:
                        from app_ai.models import ModelChoice
                        ai_model = ModelChoice.objects.get(id=ai_model_id, is_active=True)
                        conclusion_section['model_name'] = ai_model.model_name
                        conclusion_section['provider'] = ai_model.provider
                    except ModelChoice.DoesNotExist:
                        logger.error(f"AI model with ID {ai_model_id} not found")
                        pass
            
            content['sections'].append(conclusion_section)
    
    # Add translations to content for use in HTML rendering
    content['translations'] = translations
    
    return content


def markdown_to_html(text):
    """
    Convert basic markdown formatting to HTML.
    Supports: **bold**, *italic*, lists, paragraphs.
    """
    if not text:
        return ''
    
    import re
    from html import escape
    from django.utils.safestring import mark_safe
    
    # Split into lines for processing
    lines = text.split('\n')
    html_parts = []
    in_list = False
    list_type = None  # 'ul' or 'ol'
    list_items = []
    
    def flush_list():
        nonlocal in_list, list_type, list_items, html_parts
        if list_items:
            tag = 'ol' if list_type == 'ol' else 'ul'
            list_html = ''.join(list_items)
            html_parts.append(f'<{tag}>{list_html}</{tag}>')
            list_items = []
            in_list = False
            list_type = None
    
    def process_inline_formatting(text_content):
        """Process bold and italic formatting in text content"""
        # Process bold text **text** or __text__
        text_content = re.sub(r'\*\*([^*]+?)\*\*', r'<strong>\1</strong>', text_content)
        text_content = re.sub(r'__([^_]+?)__', r'<strong>\1</strong>', text_content)
        # Process italic text *text* (but not **text**)
        text_content = re.sub(r'(?<!\*)\*([^*\n]+?)\*(?!\*)', r'<em>\1</em>', text_content)
        text_content = re.sub(r'(?<!_)_([^_\n\s][^_\n]*?[^_\n\s])_(?!_)', r'<em>\1</em>', text_content)
        return text_content
    
    for line in lines:
        trimmed = line.strip()
        
        # Check for unordered list items (* or -)
        if re.match(r'^[\*\-]\s+', trimmed):
            content = re.sub(r'^[\*\-]\s+', '', trimmed)
            # Process markdown formatting
            content = process_inline_formatting(content)
            # Escape HTML but preserve formatting tags
            content = escape(content)
            content = content.replace('&lt;strong&gt;', '<strong>').replace('&lt;/strong&gt;', '</strong>')
            content = content.replace('&lt;em&gt;', '<em>').replace('&lt;/em&gt;', '</em>')
            
            if not in_list or list_type != 'ul':
                flush_list()
                in_list = True
                list_type = 'ul'
            list_items.append(f'<li>{content}</li>')
            continue
        
        # Check for ordered list items (1. 2. etc.)
        if re.match(r'^\d+\.\s+', trimmed):
            content = re.sub(r'^\d+\.\s+', '', trimmed)
            # Process markdown formatting
            content = process_inline_formatting(content)
            # Escape HTML but preserve formatting tags
            content = escape(content)
            content = content.replace('&lt;strong&gt;', '<strong>').replace('&lt;/strong&gt;', '</strong>')
            content = content.replace('&lt;em&gt;', '<em>').replace('&lt;/em&gt;', '</em>')
            
            if not in_list or list_type != 'ol':
                flush_list()
                in_list = True
                list_type = 'ol'
            list_items.append(f'<li>{content}</li>')
            continue
        
        # If not a list item, flush current list
        if in_list:
            flush_list()
        
        # Store regular line (will be processed as paragraph later)
        if trimmed:
            html_parts.append(('text', trimmed))
        else:
            html_parts.append(('blank', ''))
    
    # Flush any remaining list
    flush_list()
    
    # Now process paragraphs from stored text lines
    result_parts = []
    current_para_lines = []
    
    def finalize_paragraph():
        """Finalize current paragraph and add to result"""
        nonlocal current_para_lines, result_parts
        if current_para_lines:
            para_text = ' '.join(current_para_lines)
            para_text = process_inline_formatting(para_text)
            # Escape HTML but preserve formatting tags
            para_text = escape(para_text)
            para_text = para_text.replace('&lt;strong&gt;', '<strong>').replace('&lt;/strong&gt;', '</strong>')
            para_text = para_text.replace('&lt;em&gt;', '<em>').replace('&lt;/em&gt;', '</em>')
            result_parts.append(f'<p>{para_text}</p>')
            current_para_lines = []
    
    for item in html_parts:
        if isinstance(item, tuple):
            item_type, content = item
            if item_type == 'blank':
                # Blank line - end of paragraph
                finalize_paragraph()
            elif item_type == 'text':
                current_para_lines.append(content)
        else:
            # This is an HTML element (list) - finalize paragraph and add list
            finalize_paragraph()
            result_parts.append(item)
    
    # Handle remaining paragraph
    finalize_paragraph()
    
    return mark_safe(''.join(result_parts))


def render_unified_content_as_html(unified_content):
    """Convert unified content structure to HTML for preview"""
    
    # Get translations for confidentiality classification
    confidentiality_classification = _("Confidentiality Classification")
    confidential_information = _("Confidential Information")
    
    html = f'''
    <div class="report-preview">
        <div class="report-header text-center mb-4">
            <h2 class="text-primary">{unified_content['header']['title']}</h2>
    '''
    
    # Profile info (if available)
    if unified_content.get('profile_info'):
        profile_info = unified_content['profile_info']
        html += f'''
        <div class="profile-info mb-3">
            <h4 class="text-info">Profile Information</h4>
            <p><strong>Profile Name:</strong> {profile_info.get('profile_name', 'N/A')}</p>
            <p><strong>Report Type:</strong> {profile_info.get('report_type', 'N/A')}</p>
            <p><strong>Language:</strong> {profile_info.get('language', 'N/A')}</p>
            <p><strong>Date Range:</strong> {profile_info.get('date_range', 'N/A')}</p>
        </div>
    '''
    
    # Company info
    if unified_content['header']['company_info']:
        company_info = unified_content['header']['company_info']
        if company_info['type'] == 'single':
            html += f'<p class="text-muted mb-1"><i class="fas fa-building me-2"></i><strong>Company:</strong> {company_info["name"]}</p>'
        else:
            html += f'<p class="text-muted mb-1"><i class="fas fa-building me-2"></i><strong>Company:</strong> {company_info["text"]}</p>'
    
    # Add confidentiality classification
    html += f'''
            <p class="text-muted">{unified_content['header']['generated_date']}</p>
            <div class="confidentiality-classification mt-3 mb-2">
                <p class="text-danger fw-bold"><i class="fas fa-shield-alt me-2"></i><strong>{confidentiality_classification}:</strong> {confidential_information}</p>
            </div>
            <hr>
        </div>
    '''
    
    # Render sections
    logger.info(f"Rendering {len(unified_content['sections'])} sections")
    for section in unified_content['sections']:
        logger.debug(f"Rendering section: type={section.get('type')}, title={section.get('title')}")
        html += f'<div class="mb-4"><h3 class="text-secondary"><i class="fas fa-chart-line me-2"></i>{section["title"]}</h3>'
        
        # Special handling for Risk Details section - display subsections side by side
        if section['type'] == 'risk_details':
            html += '<div class="row">'
            
            for subsection in section['subsections']:
                if subsection['type'] == 'risk_distribution_level':
                    html += '<div class="col-md-2">'  # 15% width (closest to 15% is col-md-2 which is ~16.67%)
                else:
                    html += '<div class="col-md-10">'  # 85% width for TOP-10 Risks
                html += f'<h5>{subsection["title"]}</h5>'
                
                # Render subsection content
                if 'headers' in subsection and 'data' in subsection:
                    html += '<div class="table-responsive"><table class="table table-striped"><thead><tr>'
                    for header in subsection['headers']:
                        # Use translations for common field names
                        translated_header = header
                        if header.lower() in ['type', 'title', 'name', 'description', 'location', 'group', 'status', 'comment', 'notes']:
                            # Get translation from unified_content translations if available
                            if 'translations' in unified_content:
                                translation_key = header.lower()
                                translated_header = unified_content['translations'].get(translation_key, header)
                        html += f'<th>{translated_header}</th>'
                    html += '</tr></thead><tbody>'
                    
                    for row in subsection['data']:
                        html += '<tr>'
                        if subsection['type'] == 'risk_distribution_level':
                            color = row.get('color', '#000000')
                            html += f'<td><span class="badge" style="background-color: {color}; color: white;">{row["risk_level"]}</span></td>'
                            html += f'<td>{row["count"]}</td>'
                            html += f'<td>{row["percentage"]}</td>'
                        elif subsection['type'] == 'top_risks':
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td>{row["vulnerability_name"]}</td>'
                            html += f'<td>{row["risk_level"]}</td>'
                        html += '</tr>'
                    
                    html += '</tbody></table></div>'
                
                html += '</div>'
            
            html += '</div>'
        
        # Special handling for Asset Details section - display subsections side by side
        elif section['type'] == 'asset_details_combined':
            html += '<div class="row">'
            
            for subsection in section['subsections']:
                if subsection['type'] == 'assets_by_criticality':
                    html += '<div class="col-md-2">'  # 15% width (closest to 15% is col-md-2 which is ~16.67%)
                else:
                    html += '<div class="col-md-10">'  # 85% width for TOP-10 Asset
                html += f'<h5>{subsection["title"]}</h5>'
                
                # Render subsection content
                if 'headers' in subsection and 'data' in subsection:
                        html += '<div class="table-responsive"><table class="table table-striped"><thead><tr>'
                        for header in subsection['headers']:
                        # Use translations for common field names
                            translated_header = header
                            if header.lower() in ['type', 'title', 'name', 'description', 'location', 'group', 'status', 'comment', 'notes']:
                            # Get translation from unified_content translations if available
                                if 'translations' in unified_content:
                                    translation_key = header.lower()
                                    translated_header = unified_content['translations'].get(translation_key, header)
                            html += f'<th>{translated_header}</th>'
                        html += '</tr></thead><tbody>'
                        
                        for row in subsection['data']:
                            html += '<tr>'
                        if subsection['type'] == 'assets_by_criticality':
                            # Use color from data if available, otherwise use default
                            color = row.get('color', '#000000')
                            html += f'<td><span class="badge" style="background-color: {color}; color: white;">{row["criticality"]}</span></td>'
                            html += f'<td>{row["count"]}</td>'
                            html += f'<td>{row["percentage"]}%</td>'
                        elif subsection['type'] == 'top_10_asset':
                            html += f'<td>{row["asset_id"]}</td>'
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td><span style="color:{row.get("criticality_color", "#000000")}">{row["criticality"]}</span></td>'
                            html += '</tr>'
                    
                        html += '</tbody></table></div>'
                
                html += '</div>'
            
            html += '</div>'
        
        # Special handling for Vulnerability Details section - display subsections side by side with 15%/85% width distribution
        elif section['type'] == 'vulnerability_details_combined':
            html += '<div class="row">'
            
            for subsection in section['subsections']:
                # Set width based on subsection type: 15% for Assets by Vulnerability, 85% for TOP-10 Vulnerability
                if subsection['type'] == 'assets_by_vulnerability':
                    html += '<div class="col-md-2">'  # 15% width (closest to 15% is col-md-2 which is ~16.67%)
                else:
                    html += '<div class="col-md-10">'  # 85% width for TOP-10 Vulnerability
                
                html += f'<h5>{subsection["title"]}</h5>'
                
                # Render subsection content
                if 'headers' in subsection and 'data' in subsection:
                    html += '<div class="table-responsive"><table class="table table-striped"><thead><tr>'
                    for header in subsection['headers']:
                        # Use translations for common field names
                        translated_header = header
                        if header.lower() in ['type', 'title', 'name', 'description', 'location', 'group', 'status', 'comment', 'notes']:
                            # Get translation from unified_content translations if available
                            if 'translations' in unified_content:
                                translation_key = header.lower()
                                translated_header = unified_content['translations'].get(translation_key, header)
                        html += f'<th>{translated_header}</th>'
                    html += '</tr></thead><tbody>'
                    
                    for row in subsection['data']:
                        html += '<tr>'
                        if subsection['type'] == 'top_10_vulnerability':
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td>{row["vulnerability_name"]}</td>'
                            html += f'<td>{row["status"]}</td>'
                            html += f'<td>{row["threat"]}</td>'
                            html += f'<td>{row["probability_impact"]}</td>'
                            html += f'<td>{row["impact"]}</td>'
                            html += f'<td>{row["risk_value"]}</td>'
                            html += f'<td>{row["risk_level"]}</td>'
                            html += f'<td>{row["risk_mitigation_controls"]}</td>'
                        elif subsection['type'] == 'assets_by_vulnerability':
                            # Use color from data if available, otherwise use default
                            color = row.get('color', '#000000')
                            html += f'<td><span class="badge" style="background-color: {color}; color: white;">{row["criticality"]}</span></td>'
                            html += f'<td>{row["count"]}</td>'
                            html += f'<td>{row["percentage"]}%</td>'
                        html += '</tr>'
                    
                    html += '</tbody></table></div>'
                
                html += '</div>'
            
            html += '</div>'
        
        # Special handling for AI Conclusion section
        elif section['type'] == 'ai_conclusion':
            # Convert markdown formatting to HTML
            conclusion_content = section.get("content", "")
            formatted_content = markdown_to_html(conclusion_content)
            
            # Escape title but keep formatted_content as SafeString
            from django.utils.html import escape
            from django.utils.safestring import mark_safe
            
            escaped_title = escape(section["title"])
            # Use string concatenation to preserve SafeString
            html_part = '<div class="ai-conclusion-section mb-4"><div class="card border-info"><div class="card-header bg-info text-white"><h4 class="mb-0"><i class="fas fa-lightbulb me-2"></i>' + escaped_title + '</h4></div><div class="card-body"><div class="ai-conclusion-content" style="line-height: 1.6;">'
            html += mark_safe(html_part) + formatted_content + mark_safe('</div>')
            if section.get('model_name'):
                provider = section.get("provider", "")
                model_name = section.get("model_name", "")
                html += f'''
                        <hr>
                        <small class="text-muted">
                            <i class="fas fa-robot me-1"></i>
                            {_("Generated by")}: {escape(provider.upper())} - {escape(model_name)}
                        </small>
                '''
            html += '''
                    </div>
                </div>
            </div>
            '''
            html += '</div>'  # Close section div
        
        else:
            # Normal rendering for other sections
            for subsection in section['subsections']:
                html += f'<h5>{subsection["title"]}</h5>'
                
                if subsection['type'] == 'statistics':
                    # Render statistics as cards
                    # Check if display_inline flag is set to display all blocks in one row
                    if subsection.get('display_inline'):
                        html += '<div class="row flex-nowrap">'
                        for item in subsection['data']:
                            color_class = f"bg-{item['type']}"
                            html += f'''
                            <div class="col">
                                <div class="card {color_class} text-white">
                                    <div class="card-body text-center">
                                        <h4>{item['value']}</h4>
                                        <p class="mb-0">{item['label']}</p>
                                    </div>
                                </div>
                            </div>
                            '''
                        html += '</div>'
                    else:
                        # Default rendering for other statistics
                        html += '<div class="row">'
                        for item in subsection['data']:
                            color_class = f"bg-{item['type']}"
                            html += f'''
                            <div class="col-md-3">
                                <div class="card {color_class} text-white">
                                    <div class="card-body text-center">
                                        <h4>{item['value']}</h4>
                                        <p class="mb-0">{item['label']}</p>
                                    </div>
                                </div>
                            </div>
                            '''
                        html += '</div>'
                
                elif 'headers' in subsection and 'data' in subsection:
                    # Normal table rendering for all subsection types
                    html += '<div class="table-responsive"><table class="table table-striped"><thead><tr>'
                    for header in subsection['headers']:
                        # Use translations for common field names
                        translated_header = header
                        if header.lower() in ['type', 'title', 'name', 'description', 'location', 'group', 'status', 'comment', 'notes']:
                            # Get translation from unified_content translations if available
                            if 'translations' in unified_content:
                                translation_key = header.lower()
                                translated_header = unified_content['translations'].get(translation_key, header)
                        html += f'<th>{translated_header}</th>'
                    html += '</tr></thead><tbody>'
                    
                    for row in subsection['data']:
                        html += '<tr>'
                        if subsection['type'] == 'risk_distribution_level':
                            color = row.get('color', '#000000')
                            html += f'<td><span class="badge" style="background-color: {color}; color: white;">{row["risk_level"]}</span></td>'
                            html += f'<td>{row["count"]}</td>'
                            html += f'<td>{row["percentage"]}</td>'
                        elif subsection['type'] == 'top_risks':
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td>{row["vulnerability_name"]}</td>'
                            html += f'<td>{row["risk_level"]}</td>'
                        elif subsection['type'] == 'compliance_gaps':
                            html += f'<td>{row["type"]}</td>'
                            html += f'<td>{row["requirement"]}</td>'
                            html += f'<td>{row["description"]}</td>'
                        elif subsection['type'] == 'framework_company_requirements':
                            html += f'<td>{row.get("name", "")}</td>'
                            html += f'<td>{row.get("framework_type", "")}</td>'
                            html += f'<td>{row.get("version", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("status", "")}</td>'
                            html += f'<td>{row.get("controls_total", 0)}</td>'
                            html += f'<td>{row.get("controls_completed", 0)}</td>'
                            html += f'<td>{row.get("completion_percentage", "0%")}</td>'
                        elif subsection['type'] == 'company_requirements':
                            html += f'<td>{row.get("code", "")}</td>'
                            html += f'<td>{row.get("name", "")}</td>'
                            html += f'<td>{row.get("requirement_type", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("regulator", "")}</td>'
                            html += f'<td>{row.get("status", "")}</td>'
                            html += f'<td>{row.get("controls_total", 0)}</td>'
                            html += f'<td>{row.get("controls_completed", 0)}</td>'
                            html += f'<td>{row.get("completion_percentage", "0%")}</td>'
                        elif subsection['type'] == 'internal_requirements':
                            html += f'<td>{row.get("code", "")}</td>'
                            html += f'<td>{row.get("name", "")}</td>'
                            html += f'<td>{row.get("requirement_type", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("status", "")}</td>'
                            html += f'<td>{row.get("controls_total", 0)}</td>'
                            html += f'<td>{row.get("controls_completed", 0)}</td>'
                            html += f'<td>{row.get("completion_percentage", "0%")}</td>'
                        elif subsection['type'] == 'incident_table':
                            html += f'<td>{row.get("id", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("occurrence_datetime", "")}</td>'
                            html += f'<td>{row.get("place", "")}</td>'
                            html += f'<td>{row.get("incident_type", "")}</td>'
                            html += f'<td>{row.get("classification", "")}</td>'
                            html += f'<td>{row.get("current_state", "")}</td>'
                            html += f'<td>{row.get("responsible", "")}</td>'
                            html += f'<td>{row.get("description", "")}</td>'
                        elif subsection['type'] == 'mandatory_processes_table':
                            html += f'<td>{row.get("process_name", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("priority", "")}</td>'
                            html += f'<td>{row.get("frequency", "")}</td>'
                            html += f'<td>{row.get("next_due_date", "")}</td>'
                            html += f'<td>{row.get("last_completed_date", "")}</td>'
                            html += f'<td>{row.get("responsible_person", "")}</td>'
                            # Add status badge with color
                            status = row.get("status", "")
                            status_color = "secondary"
                            if status == "overdue":
                                status_color = "danger"
                            elif status == "upcoming":
                                status_color = "warning"
                            elif status == "completed":
                                status_color = "success"
                            html += f'<td><span class="badge bg-{status_color}">{status.capitalize() if status else ""}</span></td>'
                            html += f'<td>{row.get("description", "")}</td>'
                        elif subsection['type'] == 'certificate_key_management_table':
                            html += f'<td>{row.get("key_cert_num", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("type", "")}</td>'
                            html += f'<td>{row.get("expiry_date", "")}</td>'
                            # Add expiry status badge with color
                            expiry_status = row.get("expiry_status", "")
                            status_color = "secondary"
                            if expiry_status == "expired":
                                status_color = "danger"
                            elif expiry_status == "expiring_soon":
                                status_color = "warning"
                            elif expiry_status == "valid":
                                status_color = "success"
                            html += f'<td><span class="badge bg-{status_color}">{expiry_status.replace("_", " ").title() if expiry_status else ""}</span></td>'
                            days_until_expiry = row.get("days_until_expiry")
                            if days_until_expiry is not None:
                                html += f'<td>{days_until_expiry}</td>'
                            else:
                                html += '<td>-</td>'
                            html += f'<td>{row.get("revocation_status", "")}</td>'
                            html += f'<td>{row.get("owner_name", "")}</td>'
                            html += f'<td>{row.get("location", "")}</td>'
                            html += f'<td>{row.get("purpose", "")}</td>'
                        elif subsection['type'] == 'quiz_statistics_table':
                            html += f'<td>{row.get("quiz_title", "")}</td>'
                            html += f'<td>{row.get("total_attempts", "")}</td>'
                            html += f'<td>{row.get("successful_attempts", "")}</td>'
                            html += f'<td>{row.get("failed_attempts", "")}</td>'
                            html += f'<td>{row.get("success_rate", "")}</td>'
                            html += f'<td>{row.get("average_score", "")}</td>'
                            html += f'<td>{row.get("passing_score", "")}</td>'
                        elif subsection['type'] == 'users_at_risk_table':
                            html += f'<td>{row.get("user", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("attempts_count", "")}</td>'
                            if status == "overdue":
                                status_color = "danger"
                            elif status == "upcoming":
                                status_color = "warning"
                            elif status == "completed":
                                status_color = "success"
                            html += f'<td><span class="badge bg-{status_color}">{status.capitalize() if status else ""}</span></td>'
                            html += f'<td>{row.get("description", "")}</td>'
                        elif subsection['type'] == 'risk_distribution_table':
                            risk_level = row.get("risk_level", "")
                            risk_level_lower = risk_level.lower()
                            color = "#28a745"  # green for low
                            if risk_level_lower == "medium":
                                color = "#ffc107"  # yellow
                            elif risk_level_lower == "high":
                                color = "#fd7e14"  # orange
                            elif risk_level_lower == "critical":
                                color = "#dc3545"  # red
                            html += f'<td><span class="badge" style="background-color: {color}; color: white;">{risk_level}</span></td>'
                            html += f'<td>{row.get("count", "")}</td>'
                            html += f'<td>{row.get("percentage", "")}</td>'
                        elif subsection['type'] == 'assessment_status_table':
                            html += f'<td>{row.get("status", "")}</td>'
                            html += f'<td>{row.get("count", "")}</td>'
                        elif subsection['type'] == 'overdue_assessments_table':
                            html += f'<td>{row.get("vendor_name", "")}</td>'
                            html += f'<td>{row.get("assessment_date", "")}</td>'
                            html += f'<td>{row.get("next_review_date", "")}</td>'
                            html += f'<td>{row.get("status", "")}</td>'
                            days_overdue = row.get("days_overdue", "0")
                            html += f'<td><span class="badge bg-danger">{days_overdue}</span></td>'
                        elif subsection['type'] == 'incomplete_questionnaires_table':
                            html += f'<td>{row.get("vendor_name", "")}</td>'
                            html += f'<td>{row.get("template_name", "")}</td>'
                            html += f'<td>{row.get("status", "")}</td>'
                            html += f'<td>{row.get("started_date", "")}</td>'
                            html += f'<td>{row.get("percentage_score", "")}</td>'
                        elif subsection['type'] == 'risk_level_trends_table':
                            html += f'<td>{row.get("risk_level", "")}</td>'
                            html += f'<td>{row.get("count", "")}</td>'
                            html += f'<td>{row.get("trend", "")}</td>'
                        elif subsection['type'] == 'breach_severity_table':
                            severity = row.get("severity", "")
                            severity_lower = severity.lower()
                            color = "#28a745"  # green for low
                            if severity_lower == "medium":
                                color = "#ffc107"  # yellow
                            elif severity_lower == "high":
                                color = "#fd7e14"  # orange
                            elif severity_lower == "critical":
                                color = "#dc3545"  # red
                            html += f'<td><span class="badge" style="background-color: {color}; color: white;">{severity}</span></td>'
                            html += f'<td>{row.get("count", "")}</td>'
                            html += f'<td>{row.get("percentage", "")}</td>'
                        elif subsection['type'] == 'notification_status_table':
                            html += f'<td>{row.get("status", "")}</td>'
                            html += f'<td>{row.get("count", "")}</td>'
                        elif subsection['type'] == 'breach_notification_details_table':
                            html += f'<td>{row.get("incident_number", "")}</td>'
                            html += f'<td>{row.get("title", "")}</td>'
                            html += f'<td>{row.get("incident_date", "")}</td>'
                            html += f'<td>{row.get("discovery_date", "")}</td>'
                            html += f'<td>{row.get("notification_deadline", "")}</td>'
                            reported = row.get("reported", "")
                            reported_color = "success" if reported.lower() == "yes" or reported.lower() == "так" else "danger"
                            html += f'<td><span class="badge bg-{reported_color}">{reported}</span></td>'
                            html += f'<td>{row.get("authority_report_date", "")}</td>'
                            days_overdue = row.get("days_overdue", "0")
                            try:
                                if int(days_overdue) > 0:
                                    html += f'<td><span class="badge bg-danger">{days_overdue}</span></td>'
                                else:
                                    html += f'<td>{days_overdue}</td>'
                            except (ValueError, TypeError):
                                html += f'<td>{days_overdue}</td>'
                        elif subsection['type'] == 'data_categories_table':
                            html += f'<td>{row.get("category", "")}</td>'
                            html += f'<td>{row.get("count", "")}</td>'
                        elif subsection['type'] == 'international_transfers_table':
                            html += f'<td>{row.get("activity_name", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("transfer_safeguards", "")}</td>'
                        elif subsection['type'] == 'outdated_retention_table':
                            html += f'<td>{row.get("activity_name", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("retention_period_days", "")}</td>'
                            html += f'<td>{row.get("created_date", "")}</td>'
                            html += f'<td>{row.get("updated_date", "")}</td>'
                            days_overdue = row.get("days_overdue", "0")
                            try:
                                if int(days_overdue) > 0:
                                    html += f'<td><span class="badge bg-danger">{days_overdue}</span></td>'
                                else:
                                    html += f'<td>{days_overdue}</td>'
                            except (ValueError, TypeError):
                                html += f'<td>{days_overdue}</td>'
                        elif subsection['type'] == 'non_compliant_retention_table':
                            html += f'<td>{row.get("activity_name", "")}</td>'
                            html += f'<td>{row.get("company", "")}</td>'
                            html += f'<td>{row.get("retention_period_days", "")}</td>'
                            issue = row.get("issue", "")
                            issue_lower = issue.lower()
                            issue_color = "danger"
                            if "excessive" in issue_lower or "надмірний" in issue_lower:
                                issue_color = "warning"
                            elif "insufficient" in issue_lower or "недостатній" in issue_lower:
                                issue_color = "info"
                            html += f'<td><span class="badge bg-{issue_color}">{issue}</span></td>'
                        elif subsection['type'] == 'assets_by_criticality':
                            # Use color from data if available, otherwise use default
                            color = row.get('color', '#000000')
                            html += f'<td><span class="badge" style="background-color: {color}; color: white;">{row["criticality"]}</span></td>'
                            html += f'<td>{row["count"]}</td>'
                            html += f'<td>{row["percentage"]}%</td>'
                        elif subsection['type'] == 'assets_by_vulnerability':
                            # Use color from data if available, otherwise use default
                            color = row.get('color', '#000000')
                            html += f'<td><span class="badge" style="background-color: {color}; color: white;">{row["criticality"]}</span></td>'
                            html += f'<td>{row["count"]}</td>'
                            html += f'<td>{row["percentage"]}%</td>'
                        elif subsection['type'] == 'asset_table':
                            html += f'<td>{row["asset_id"]}</td>'
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td><span style="color:{row["criticality_color"]}">{row["criticality"]}</span></td>'
                            html += f'<td>{row.get("company", "")}</td>'
                        elif subsection['type'] == 'top_10_asset':
                            html += f'<td>{row["asset_id"]}</td>'
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td><span style="color:{row.get("criticality_color", "#000000")}">{row["criticality"]}</span></td>'
                            # Display vulnerabilities with proper logic: "Undefined" if no assessment, "0" if assessed but no vulnerabilities, or count
                            vulnerabilities_display = row.get("vulnerabilities", "Undefined")
                            if vulnerabilities_display == "Undefined":
                                html += f'<td><span class="badge bg-secondary">Undefined</span></td>'
                            elif vulnerabilities_display == "0":
                                html += f'<td><span class="badge bg-secondary">0</span></td>'
                            else:
                                html += f'<td><span class="badge bg-danger">{vulnerabilities_display}</span></td>'
                            html += f'<td>{row.get("company", "")}</td>'
                        elif subsection['type'] == 'vulnerability_table':
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td>{row["vulnerability_name"]}</td>'
                            html += f'<td>{row["status"]}</td>'
                            html += f'<td>{row["threat"]}</td>'
                            html += f'<td>{row["probability_impact"]}</td>'
                            html += f'<td>{row["impact"]}</td>'
                            html += f'<td>{row["risk_value"]}</td>'
                            html += f'<td>{row["risk_level"]}</td>'
                            html += f'<td>{row["risk_mitigation_controls"]}</td>'
                        elif subsection['type'] == 'overdue_treatments':
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td>{row["treatment"]}</td>'
                            html += f'<td>{row["deadline"]}</td>'
                        elif subsection['type'] == 'treatment_details':
                            html += f'<td>{row["asset_name"]}</td>'
                            html += f'<td>{row["vulnerability_name"]}</td>'
                            html += f'<td>{row["vulnerability_status"]}</td>'
                            html += f'<td>{row["risk_level"]}</td>'
                            html += f'<td>{row["treatment_type"]}</td>'
                            html += f'<td>{row["treatment_status"]}</td>'
                            html += f'<td>{row["description_responsible"]}</td>'
                            html += f'<td>{row["due_date"]}</td>'
                        elif subsection['type'] == 'table':
                            # Generic table rendering for acceptable risk tables
                            for key, value in row.items():
                                html += f'<td>{value}</td>'
                    html += '</tr>'
                
                    html += '</tbody></table></div>'
                
                elif subsection['type'] == 'pci_dss' or subsection['type'] == 'iso27001':
                    # Render compliance data as list
                    for item in subsection['data']:
                        html += f'<p><strong>{item["label"]}:</strong> {item["value"]}</p>'
                
                elif subsection['type'] == 'financial_overview':
                    # Render financial overview as cards
                    html += '<div class="row">'
                    for item in subsection['data']:
                        html += f'''
                        <div class="col-md-3">
                            <div class="card bg-info text-white">
                                <div class="card-body text-center">
                                    <h4>{item["value"]}</h4>
                                    <p class="mb-0">{item["label"]}</p>
                                </div>
                            </div>
                        </div>
                        '''
                    html += '</div>'
                
                elif subsection['type'] == 'residual_overview':
                    # Render residual overview as cards
                    html += '<div class="row">'
                    for item in subsection['data']:
                        html += f'''
                        <div class="col-md-4">
                            <div class="card bg-info text-white">
                                <div class="card-body text-center">
                                    <h4>{item["value"]}</h4>
                                    <p class="mb-0">{item["label"]}</p>
                                </div>
                            </div>
                        </div>
                        '''
                    html += '</div>'
                
                elif subsection['type'] == 'governance_overview':
                    # Render governance overview as cards
                    html += '<div class="row">'
                    for item in subsection['data']:
                        html += f'''
                        <div class="col-md-3">
                            <div class="card bg-primary text-white">
                                <div class="card-body text-center">
                                    <h4>{item["value"]}</h4>
                                    <p class="mb-0">{item["label"]}</p>
                                </div>
                            </div>
                        </div>
                        '''
                    html += '</div>'
                
                elif subsection['type'] == 'priority_actions':
                    # Render recommendations as list
                    html += '<ul>'
                    for item in subsection['data']:
                        html += f'<li>{item}</li>'
                    html += '</ul>'
        
            html += '</div>'
    
    html += '</div>'
    # Mark the entire HTML as safe since it contains SafeString parts
    from django.utils.safestring import mark_safe
    return mark_safe(html)


def generate_html_preview(data, report_type, language, selected_sections):
    """Generate HTML preview of the report content using unified content structure"""
    
    translations = get_report_translations(language)
    
    # Generate unified content structure
    unified_content = generate_unified_report_content(data, report_type, selected_sections, translations, language)
    
    # Convert unified content to HTML
    return render_unified_content_as_html(unified_content)


def generate_full_report_preview(data, translations, selected_sections):
    """Generate full report preview HTML"""
    
    # Get company information
    company_name = None
    if data.get('companies'):
        if data['companies'].count() == 1:
            company_name = data['companies'].first().name
        elif data['companies'].count() > 1:
            company_name = f"{translations.get('multiple_companies', 'Multiple Companies')} ({data['companies'].count()})"
    
    # Build company display text
    company_display = ""
    if company_name:
        company_display = f'<p class="text-muted mb-1"><i class="fas fa-building me-2"></i><strong>{translations.get("company", "Company")}:</strong> {company_name}</p>'
    
    # Get translations for confidentiality classification
    confidentiality_classification = _("Confidentiality Classification")
    confidential_information = _("Confidential Information")
    
    html = f'''
    <div class="report-preview">
        <div class="report-header text-center mb-4">
            <h2 class="text-primary">{translations['full_report_title']}</h2>
            {company_display}
            <p class="text-muted">{translations['generated_on']}: {format_localized_date(data['generation_date'], data['language'])}</p>
            <div class="confidentiality-classification mt-3 mb-2">
                <p class="text-danger fw-bold"><i class="fas fa-shield-alt me-2"></i><strong>{confidentiality_classification}:</strong> {confidential_information}</p>
            </div>
            <hr>
        </div>
    '''
    
    # Executive Summary Section
    if selected_sections.get('statistics') or selected_sections.get('top_risks'):
        html += f'''
        <div class="executive-summary mb-4">
            <h3 class="text-secondary"><i class="fas fa-chart-line me-2"></i>{translations['executive_summary']}</h3>
        '''
        
        # Key Statistics
        if selected_sections.get('statistics'):
            html += f'''
            <div class="row">
                <div class="col-md-3">
                    <div class="card bg-primary text-white">
                        <div class="card-body text-center">
                            <h4>{data['statistics']['total_assets']}</h4>
                            <p class="mb-0">{translations['total_assets']}</p>
                        </div>
                    </div>
                </div>
                <div class="col-md-3">
                    <div class="card bg-warning text-white">
                        <div class="card-body text-center">
                            <h4>{data['statistics']['total_vulnerabilities']}</h4>
                            <p class="mb-0">{translations['total_vulnerabilities']}</p>
                        </div>
                    </div>
                </div>
                <div class="col-md-3">
                    <div class="card bg-danger text-white">
                        <div class="card-body text-center">
                            <h4>{data['statistics']['high_risk_count']}</h4>
                            <p class="mb-0">{translations['high_risk_assets']}</p>
                        </div>
                    </div>
                </div>
                <div class="col-md-3">
                    <div class="card bg-success text-white">
                        <div class="card-body text-center">
                            <h4>{data['statistics']['completion_rate']:.1f}%</h4>
                            <p class="mb-0">{translations['completion_rate']}</p>
                        </div>
                    </div>
                </div>
            </div>
            '''
        
        html += '</div>'
    
    # Risk Details Section (Combined By Risk Level and TOP-10 Risks)
    if selected_sections.get('risk_distribution') or selected_sections.get('top_risks'):
        html += f'''
        <div class="risk-details mb-4">
            <h3 class="text-secondary"><i class="fas fa-chart-pie me-2"></i>{translations.get('risk_details', 'Risk Details')}</h3>
            <div class="row">
        '''
        
        # By Risk Level (left column)
        if selected_sections.get('risk_distribution'):
            html += f'''
                <div class="col-md-6">
                    <h5>{translations['by_risk_level']}</h5>
                    <div class="table-responsive">
                        <table class="table table-striped">
                            <thead>
                                <tr>
                                    <th>{translations['risk_level']}</th>
                                    <th>{translations['count']}</th>
                                    <th>{translations['percentage']}</th>
                                </tr>
                            </thead>
                            <tbody>
    '''
        
        # Add risk level distribution with sorting and colors
        total_risks = sum(data['risk_distribution']['by_level'].values())
        risk_levels_with_colors = data.get('statistics', {}).get('risk_levels_with_colors', {})
        
        # Create list for sorting
        risk_level_items = []
        for level, count in data['risk_distribution']['by_level'].items():
            percentage = (count / total_risks * 100) if total_risks > 0 else 0
            color_info = risk_levels_with_colors.get(level, {})
            risk_level_items.append({
                'level': level,
                'count': count,
                'percentage': percentage,
                'color': color_info.get('color', '#000000'),
                'max_value': color_info.get('max_value', 0)
            })
        
        # Sort by max_value in descending order (highest risk first)
        risk_level_items.sort(key=lambda x: x['max_value'], reverse=True)
        
        for item in risk_level_items:
            html += f'''
                                <tr>
                                    <td><span class="badge" style="background-color: {item['color']}; color: white;">{item['level']}</span></td>
                                    <td>{item['count']}</td>
                                    <td>{item['percentage']:.1f}%</td>
                                </tr>
            '''
        
            html += '''
                            </tbody>
                        </table>
                    </div>
                </div>
            '''
        
        # TOP-10 Risks (right column)
        if selected_sections.get('top_risks'):
            html += f'''
                <div class="col-md-6">
                    <h5>{translations.get('top_risks', 'TOP-10 Risks')}</h5>
                    <div class="table-responsive">
                        <table class="table table-striped">
                            <thead>
                                <tr>
                                    <th>{translations.get('asset_name', 'Asset Name')}</th>
                                    <th>{translations['vulnerability']}</th>
                                    <th>{translations['risk_level']}</th>
                                </tr>
                            </thead>
                            <tbody>
    '''
        
            # Add TOP-10 risks data
            if 'statistics' in data and data['statistics'].get('high_risk_assets'):
                for asset in data['statistics']['high_risk_assets'][:10]:  # Limit to 10
                    html += f'''
                                <tr>
                                        <td>{asset.get('asset_name', translations.get('unknown', 'Unknown'))}</td>
                                        <td>{asset.get('vulnerability_name', translations.get('unknown', 'Unknown'))}</td>
                                        <td>{asset.get('risk_level', translations.get('unknown', 'Unknown'))}</td>
                                </tr>
                    '''
        
        html += '''
                            </tbody>
                        </table>
                    </div>
                </div>
            '''
        
        html += '''
            </div>
        </div>
        '''
    
    # Compliance Overview Section
    if selected_sections.get('compliance_overview') or selected_sections.get('pci_dss') or selected_sections.get('iso27001'):
        html += f'''
        <div class="compliance-overview mb-4">
            <h3 class="text-secondary"><i class="fas fa-shield-alt me-2"></i>{translations['compliance_overview']}</h3>
            <div class="row">
        '''
        
        if selected_sections.get('pci_dss'):
            html += f'''
                <div class="col-md-6">
                    <div class="card">
                        <div class="card-header">
                            <h5 class="mb-0">PCI DSS</h5>
                        </div>
                        <div class="card-body">
                            <p><strong>{translations['compliant_vulnerabilities']}:</strong> {data['compliance']['pcidss']['compliant_count']}</p>
                            <p><strong>{translations['non_compliant_vulnerabilities']}:</strong> {data['compliance']['pcidss']['non_compliant_count']}</p>
                            <p><strong>{translations['compliance_rate']}:</strong> {data['compliance']['pcidss']['compliance_rate']:.1f}%</p>
                        </div>
                    </div>
                </div>
            '''
        
        if selected_sections.get('iso27001'):
            html += f'''
                <div class="col-md-6">
                    <div class="card">
                        <div class="card-header">
                            <h5 class="mb-0">ISO 27001</h5>
                        </div>
                        <div class="card-body">
                            <p><strong>{translations['compliant_vulnerabilities']}:</strong> {data['compliance']['iso27001']['compliant_count']}</p>
                            <p><strong>{translations['non_compliant_vulnerabilities']}:</strong> {data['compliance']['iso27001']['non_compliant_count']}</p>
                            <p><strong>{translations['compliance_rate']}:</strong> {data['compliance']['iso27001']['compliance_rate']:.1f}%</p>
                        </div>
                    </div>
                </div>
            '''
        
        html += '''
            </div>
        </div>
        '''
    
    # Asset Details Section - Combined Assets by Criticality and TOP-10 Asset
    if selected_sections.get('asset_details') or selected_sections.get('asset_tables'):
        html += f'''
        <div class="asset-details mb-4">
            <h3 class="text-secondary"><i class="fas fa-server me-2"></i>{translations.get('asset_details', 'Asset Details')}</h3>
            <div class="row">
        '''
        
        # Assets by Criticality
        if selected_sections.get('asset_details') and 'statistics' in data:
            stats = data['statistics']
            if stats.get('assets_by_criticality'):
                html += '''
                <div class="col-md-6">
                    <h5>Assets by Criticality</h5>
            <div class="table-responsive">
                <table class="table table-striped">
                    <thead>
                        <tr>
                                    <th>Criticality</th>
                                    <th>Count</th>
                        </tr>
                    </thead>
                    <tbody>
    '''
        
                for item in stats['assets_by_criticality']:
                    criticality_name = item.get('criticality__name_uk', translations.get('unknown', 'Unknown'))
                    color = item.get('color', '#000000')
                    count = item.get('count', 0)
                    
                    html += f'''
                                <tr>
                                    <td><span class="badge" style="background-color: {color}; color: white;">{criticality_name}</span></td>
                                    <td>{count}</td>
                                </tr>
                    '''
                
                html += '''
                            </tbody>
                        </table>
                    </div>
                </div>
                '''
        
        # TOP-10 Asset
        if selected_sections.get('asset_tables') and 'assets' in data:
            html += '''
            <div class="col-md-6">
                <h5>{translations.get('top_10_asset', 'TOP-10 Asset')}</h5>
                <div class="table-responsive">
                    <table class="table table-striped">
                        <thead>
                            <tr>
                                <th>Asset ID</th>
                                <th>Asset Name</th>
                                <th>Criticality</th>
                            </tr>
                        </thead>
                        <tbody>
            '''
            
            # Get top 10 assets by criticality cost
            asset_table_data = []
            for asset in data['assets']:
                if hasattr(asset, 'get_criticality'):
                    criticality_info = asset.get_criticality()
                    criticality_name = criticality_info.get('name', '') if criticality_info else ''
                    criticality_color = criticality_info.get('color', '#000000') if criticality_info else '#000000'
                    criticality_cost = criticality_info.get('cost', 0) if criticality_info else 0
                    
                    asset_table_data.append({
                        'asset_id': getattr(asset, 'asset_id', ''),
                        'asset_name': getattr(asset, 'name', ''),
                        'criticality': criticality_name,
                        'criticality_color': criticality_color,
                        'criticality_cost': criticality_cost
                    })
            
            # Sort by criticality cost (highest first) and limit to 10
            asset_table_data.sort(key=lambda x: x['criticality_cost'], reverse=True)
            asset_table_data = asset_table_data[:10]
            
            for asset_data in asset_table_data:
                html += f'''
                        <tr>
                                <td>{asset_data["asset_id"]}</td>
                                <td>{asset_data["asset_name"]}</td>
                                <td><span style="color:{asset_data["criticality_color"]}">{asset_data["criticality"]}</span></td>
                        </tr>
                '''
        
        html += '''
                    </tbody>
                </table>
            </div>
        </div>
        '''
    
        html += '''
            </div>
        </div>
        '''
    
    # Treatment Details Section
    if selected_sections.get('treatment_details') or selected_sections.get('treatment_tables'):
        html += f'''
        <div class="treatment-details mb-4">
            <h3 class="text-secondary"><i class="fas fa-tasks me-2"></i>{translations.get('treatment_details', 'Treatment Details')}</h3>
            <div class="table-responsive">
                <table class="table table-striped">
                    <thead>
                        <tr>
                            <th>{translations['asset']}</th>
                            <th>{translations['vulnerability']}</th>
                            <th>{translations.get('due_date', 'Due Date')}</th>
                            <th>{translations['status']}</th>
                        </tr>
                    </thead>
                    <tbody>
    '''
        
        # Add treatment details (first 10 for preview)
        for treatment in data.get('overdue_treatments', [])[:10]:
            html += f'''
                        <tr>
                            <td>{treatment['asset_name']}</td>
                            <td>{treatment['vulnerability_name']}</td>
                            <td>{format_localized_date(treatment['due_date'], data['language'])}</td>
                            <td><span class="badge bg-danger">{treatment['status']}</span></td>
                        </tr>
            '''
        
        html += '''
                    </tbody>
                </table>
            </div>
        </div>
        '''
        
        # Add Treatment Details subsection if requested
        if selected_sections.get('treatment_details') and 'risk_treatments' in data:
            html += f'''
        <div class="treatment-details mb-4">
            <h4 class="text-info"><i class="fas fa-clipboard-list me-2"></i>{translations.get('top_10_treatment', 'TOP-10 Treatment')}</h4>
            <div class="table-responsive">
                <table class="table table-striped">
                    <thead>
                        <tr>
                            <th>{translations['asset']}</th>
                            <th>{translations.get('vulnerability', 'Vulnerability')}</th>
                            <th>{translations.get('vulnerability_status', 'Vulnerability Status')}</th>
                            <th>{translations.get('risk_level', 'Risk Level')}</th>
                            <th>{translations.get('treatment_type', 'Treatment Type')}</th>
                            <th>{translations.get('treatment_status', 'Treatment Status')}</th>
                            <th>{translations.get('description_responsible', 'Description & Responsible')}</th>
                            <th>{translations.get('due_date', 'Due Date')}</th>
                        </tr>
                    </thead>
                    <tbody>
            '''
            
            # Add treatment details (first 10 for preview with sorting)
            treatments_for_preview = []
            for treatment in data.get('risk_treatments', []):
                # Get vulnerability status and risk level for this treatment
                vulnerability_status = translations.get('unknown', 'Unknown')
                risk_level = translations.get('unknown', 'Unknown')
                if 'vulnerabilities' in data:
                    for vuln in data['vulnerabilities']:
                        # Handle both dictionary and model object cases
                        if hasattr(vuln, 'get'):  # Dictionary
                            vuln_asset_name = vuln.get('asset_name', '')
                            vuln_name = vuln.get('vulnerability_name', '')
                        else:  # Model object
                            vuln_asset_name = getattr(vuln, 'asset_name', '') if hasattr(vuln, 'asset_name') else (getattr(vuln.asset, 'name', '') if hasattr(vuln, 'asset') and vuln.asset else '')
                            vuln_name = getattr(vuln, 'vulnerability_name', '') if hasattr(vuln, 'vulnerability_name') else (get_localized_vulnerability_field(vuln.vulnerability, 'vulnerability', data.get('language', 'uk')) if hasattr(vuln, 'vulnerability') and vuln.vulnerability else '')
                        
                        # Match by asset name and vulnerability name
                        if (vuln_asset_name == treatment.get('asset_name', '') and 
                            vuln_name == treatment.get('vulnerability_name', '')):
                            vulnerability_status = getattr(vuln, 'status', translations.get('unknown', 'Unknown'))
                            
                            # Calculate risk level for this vulnerability
                            try:
                                if hasattr(vuln, 'asset') and hasattr(vuln, 'vulnerability'):
                                    # For model objects
                                    risk_value = calculate_value_of_risk(vuln.asset, vuln.vulnerability)
                                    risk_level_obj = calculate_risk_level(risk_value)
                                    if risk_level_obj:
                                        risk_level = risk_level_obj.get_name_by_language(data.get("language", "uk")) or risk_level_obj.get_name()
                                    else:
                                        risk_level = translations.get('unknown', 'Unknown')
                                elif hasattr(vuln, 'get'):
                                    # For dictionary objects, try to get risk level if available
                                    risk_level = vuln.get('risk_level', translations.get('unknown', 'Unknown'))
                                else:
                                    risk_level = translations.get('unknown', 'Unknown')
                            except Exception as e:
                                logger.warning(f"Error calculating risk level: {e}")
                                risk_level = translations.get('unknown', 'Unknown')
                            break
                
                # Calculate priority for sorting
                status_priority = 0
                if vulnerability_status == 'Yes':
                    status_priority = 1
                
                risk_level_priority = 0
                if risk_level:
                    if 'Надзвичайний' in risk_level or 'Extraordinary' in risk_level:
                        risk_level_priority = 6
                    elif 'Критичний' in risk_level or 'Critical' in risk_level:
                        risk_level_priority = 5
                    elif 'Високий' in risk_level or 'High' in risk_level:
                        risk_level_priority = 4
                    elif 'Середній' in risk_level or 'Medium' in risk_level:
                        risk_level_priority = 3
                    elif 'Низький' in risk_level or 'Low' in risk_level:
                        risk_level_priority = 2
                    elif 'Невизначено' in risk_level or 'Undefined' in risk_level:
                        risk_level_priority = 1
                
                treatments_for_preview.append({
                    'asset_name': treatment.get('asset_name', translations.get('unknown', 'Unknown')),
                    'vulnerability_name': treatment.get('vulnerability_name', translations.get('unknown', 'Unknown')),
                    'vulnerability_status': vulnerability_status,
                    'risk_level': risk_level,
                    'treatment_type': treatment.get('treatment_type', translations.get('unknown', 'Unknown')),
                    'treatment_status': treatment.get('status', translations.get('unknown', 'Unknown')),
                    'description_responsible': treatment.get('description', '') + ' | ' + treatment.get('assigned_to', translations.get('unknown', 'Unknown')),
                    'due_date': format_localized_date(treatment.get('due_date'), data['language']) if treatment.get('due_date') else translations.get('unknown', 'Unknown'),
                    'status_priority': status_priority,
                    'risk_level_priority': risk_level_priority
                })
            
            # Sort by status priority (Yes first) and then by risk level priority (highest first)
            treatments_for_preview.sort(key=lambda x: (x['status_priority'], x['risk_level_priority']), reverse=True)
            
            # Limit to top 10
            treatments_for_preview = treatments_for_preview[:10]
            
            # Generate HTML for sorted treatments
            for treatment in treatments_for_preview:
                html += f'''
                        <tr>
                            <td>{treatment["asset_name"]}</td>
                            <td>{treatment["vulnerability_name"]}</td>
                            <td><span class="badge bg-info">{treatment["vulnerability_status"]}</span></td>
                            <td><span class="badge bg-warning">{treatment["risk_level"]}</span></td>
                            <td>{treatment["treatment_type"]}</td>
                            <td><span class="badge bg-primary">{treatment["treatment_status"]}</span></td>
                            <td>{treatment["description_responsible"]}</td>
                            <td>{treatment["due_date"]}</td>
                        </tr>
                '''
            
            html += '''
                    </tbody>
                </table>
            </div>
        </div>
        '''
    
    # Charts & Graphs Section
    if selected_sections.get('charts') or selected_sections.get('graphs') or selected_sections.get('metrics'):
        html += f'''
        <div class="charts-section mb-4">
            <h3 class="text-secondary"><i class="fas fa-chart-bar me-2"></i>{translations.get('charts_graphs', 'Charts & Graphs')}</h3>
            <div class="alert alert-info">
                <i class="fas fa-info-circle me-2"></i>
                {translations.get('charts_preview_note', 'Charts and graphs will be included in the final report.')}
            </div>
        </div>
        '''
    
    # Recommendations Section
    if selected_sections.get('recommendations'):
        html += f'''
        <div class="recommendations mb-4">
            <h3 class="text-secondary"><i class="fas fa-lightbulb me-2"></i>{translations.get('recommendations', 'Recommendations')}</h3>
            <div class="alert alert-warning">
                <h6>{translations.get('priority_recommendations', 'Priority Recommendations')}:</h6>
                <ul>
                    <li>{translations.get('recommendation_1', 'Address high-risk vulnerabilities immediately')}</li>
                    <li>{translations.get('recommendation_2', 'Implement regular security assessments')}</li>
                    <li>{translations.get('recommendation_3', 'Update compliance documentation')}</li>
                </ul>
            </div>
        </div>
        '''
    
    # Report Footer
    html += f'''
        <div class="report-footer text-center mt-4">
            <p class="text-muted">{translations['report_generated_by']}: {data['generated_by']}</p>
            <p class="text-muted">{translations.get('preview_note', 'This is a preview. The final report will contain complete data.')}</p>
        </div>
    </div>
    '''
    
    return html


def generate_summary_report_preview(data, translations, selected_sections):
    """Generate summary report preview HTML"""
    
    # Get company information
    company_name = None
    if data.get('companies'):
        if data['companies'].count() == 1:
            company_name = data['companies'].first().name
        elif data['companies'].count() > 1:
            company_name = f"{translations.get('multiple_companies', 'Multiple Companies')} ({data['companies'].count()})"
    
    # Build company display text
    company_display = ""
    if company_name:
        company_display = f'<p class="text-muted mb-1"><i class="fas fa-building me-2"></i><strong>{translations.get("company", "Company")}:</strong> {company_name}</p>'
    
    # Get translations for confidentiality classification
    confidentiality_classification = _("Confidentiality Classification")
    confidential_information = _("Confidential Information")
    
    html = f'''
    <div class="report-preview">
        <div class="report-header text-center mb-4">
            <h2 class="text-primary">{translations['executive_summary']}</h2>
            {company_display}
            <p class="text-muted">{translations['generated_on']}: {format_localized_date(data['generation_date'], data['language'])}</p>
            <div class="confidentiality-classification mt-3 mb-2">
                <p class="text-danger fw-bold"><i class="fas fa-shield-alt me-2"></i><strong>{confidentiality_classification}:</strong> {confidential_information}</p>
            </div>
            <hr>
        </div>
    '''
    
    # Key Statistics Section
    if selected_sections.get('statistics'):
        html += f'''
        <div class="key-metrics mb-4">
            <h3 class="text-secondary"><i class="fas fa-tachometer-alt me-2"></i>{translations['key_metrics']}</h3>
            <div class="row">
                <div class="col-md-4">
                    <div class="card bg-primary text-white">
                        <div class="card-body text-center">
                            <h3>{data['statistics']['total_assets']}</h3>
                            <p class="mb-0">{translations['total_assets']}</p>
                        </div>
                    </div>
                </div>
                <div class="col-md-4">
                    <div class="card bg-warning text-white">
                        <div class="card-body text-center">
                            <h3>{data['statistics']['total_vulnerabilities']}</h3>
                            <p class="mb-0">{translations['total_vulnerabilities']}</p>
                        </div>
                    </div>
                </div>
                <div class="col-md-4">
                    <div class="card bg-danger text-white">
                        <div class="card-body text-center">
                            <h3>{data['statistics']['high_risk_count']}</h3>
                            <p class="mb-0">{translations['high_risk_assets']}</p>
                        </div>
                    </div>
                </div>
            </div>
        </div>
        '''
    
    # Note: TOP-10 Risks is now combined with Risk Details section above
    
    # Compliance Overview Section
    if selected_sections.get('compliance_overview'):
        html += f'''
        <div class="compliance-summary mb-4">
            <h3 class="text-secondary"><i class="fas fa-shield-alt me-2"></i>{translations['compliance_overview']}</h3>
            <div class="row">
                <div class="col-md-6">
                    <div class="card">
                        <div class="card-header">
                            <h6 class="mb-0">PCI DSS</h6>
                        </div>
                        <div class="card-body">
                            <div class="d-flex justify-content-between">
                                <span>{translations.get('compliance_rate', 'Compliance Rate')}:</span>
                                <strong>{data['compliance']['pcidss']['compliance_rate']:.1f}%</strong>
                            </div>
                        </div>
                    </div>
                </div>
                <div class="col-md-6">
                    <div class="card">
                        <div class="card-header">
                            <h6 class="mb-0">ISO 27001</h6>
                        </div>
                        <div class="card-body">
                            <div class="d-flex justify-content-between">
                                <span>{translations.get('compliance_rate', 'Compliance Rate')}:</span>
                                <strong>{data['compliance']['iso27001']['compliance_rate']:.1f}%</strong>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
        </div>
        '''
    
    # Recommendations Section
    if selected_sections.get('recommendations'):
        html += f'''
        <div class="recommendations mb-4">
            <h3 class="text-secondary"><i class="fas fa-lightbulb me-2"></i>{translations.get('recommendations', 'Recommendations')}</h3>
            <div class="alert alert-info">
                <h6>{translations.get('key_recommendations', 'Key Recommendations')}:</h6>
                <ol>
                    <li>{translations.get('recommendation_1', 'Address high-risk vulnerabilities immediately')}</li>
                    <li>{translations.get('recommendation_2', 'Implement regular security assessments')}</li>
                    <li>{translations.get('recommendation_3', 'Update compliance documentation')}</li>
                </ol>
            </div>
        </div>
        '''
    
    # Charts & Metrics Section
    if selected_sections.get('charts') or selected_sections.get('metrics'):
        html += f'''
        <div class="charts-section mb-4">
            <h3 class="text-secondary"><i class="fas fa-chart-line me-2"></i>{translations.get('performance_metrics', 'Performance Metrics')}</h3>
            <div class="alert alert-info">
                <i class="fas fa-info-circle me-2"></i>
                {translations.get('metrics_preview_note', 'Performance metrics and charts will be included in the final report.')}
            </div>
        </div>
        '''
    
    # Report Footer
    html += f'''
        <div class="report-footer text-center mt-4">
            <p class="text-muted">{translations['report_generated_by']}: {data['generated_by']}</p>
            <p class="text-muted">{translations.get('preview_note', 'This is a preview. The final report will contain complete data.')}</p>
        </div>
    </div>
    '''
    
    return html


def generate_compliance_report_preview(data, translations, selected_sections):
    """Generate compliance report preview HTML"""
    
    # Get company information
    company_name = None
    if data.get('companies'):
        if data['companies'].count() == 1:
            company_name = data['companies'].first().name
        elif data['companies'].count() > 1:
            company_name = f"{translations.get('multiple_companies', 'Multiple Companies')} ({data['companies'].count()})"
    
    # Build company display text
    company_display = ""
    if company_name:
        company_display = f'<p class="text-muted mb-1"><i class="fas fa-building me-2"></i><strong>{translations.get("company", "Company")}:</strong> {company_name}</p>'
    
    # Get translations for confidentiality classification
    confidentiality_classification = _("Confidentiality Classification")
    confidential_information = _("Confidential Information")
    
    html = f'''
    <div class="report-preview">
        <div class="report-header text-center mb-4">
            <h2 class="text-primary">{translations.get('compliance_report', 'Compliance Report')}</h2>
            {company_display}
            <p class="text-muted">{translations['generated_on']}: {format_localized_date(data['generation_date'], data['language'])}</p>
            <div class="confidentiality-classification mt-3 mb-2">
                <p class="text-danger fw-bold"><i class="fas fa-shield-alt me-2"></i><strong>{confidentiality_classification}:</strong> {confidential_information}</p>
            </div>
            <hr>
        </div>
    '''
    
    # Compliance Overview Section
    if selected_sections.get('compliance_overview') or selected_sections.get('pci_dss') or selected_sections.get('iso27001'):
        html += f'''
        <div class="compliance-overview mb-4">
            <h3 class="text-secondary"><i class="fas fa-shield-alt me-2"></i>{translations['compliance_overview']}</h3>
            <div class="row">
        '''
        
        if selected_sections.get('pci_dss'):
            html += f'''
                <div class="col-md-6">
                    <div class="card">
                        <div class="card-header bg-primary text-white">
                            <h5 class="mb-0">PCI DSS {translations.get('compliance', 'Compliance')}</h5>
                        </div>
                        <div class="card-body">
                            <div class="row">
                                <div class="col-6">
                                    <p><strong>{translations.get('compliant', 'Compliant')}:</strong></p>
                                    <h4 class="text-success">{data['compliance']['pcidss']['compliant_count']}</h4>
                                </div>
                                <div class="col-6">
                                    <p><strong>{translations.get('non_compliant', 'Non-Compliant')}:</strong></p>
                                    <h4 class="text-danger">{data['compliance']['pcidss']['non_compliant_count']}</h4>
                                </div>
                            </div>
                            <div class="progress mt-3">
                                <div class="progress-bar bg-success" style="width: {data['compliance']['pcidss']['compliance_rate']:.1f}%">
                                    {data['compliance']['pcidss']['compliance_rate']:.1f}%
                                </div>
                            </div>
                        </div>
                    </div>
                </div>
            '''
        
        if selected_sections.get('iso27001'):
            html += f'''
                <div class="col-md-6">
                    <div class="card">
                        <div class="card-header bg-info text-white">
                            <h5 class="mb-0">ISO 27001 {translations.get('compliance', 'Compliance')}</h5>
                        </div>
                        <div class="card-body">
                            <div class="row">
                                <div class="col-6">
                                    <p><strong>{translations.get('compliant', 'Compliant')}:</strong></p>
                                    <h4 class="text-success">{data['compliance']['iso27001']['compliant_count']}</h4>
                                </div>
                                <div class="col-6">
                                    <p><strong>{translations.get('non_compliant', 'Non-Compliant')}:</strong></p>
                                    <h4 class="text-danger">{data['compliance']['iso27001']['non_compliant_count']}</h4>
                                </div>
                            </div>
                            <div class="progress mt-3">
                                <div class="progress-bar bg-success" style="width: {data['compliance']['iso27001']['compliance_rate']:.1f}%">
                                    {data['compliance']['iso27001']['compliance_rate']:.1f}%
                                </div>
                            </div>
                        </div>
                    </div>
                </div>
            '''
        
        html += '''
            </div>
        </div>
        '''
    
    # Compliance Gaps Section
    if selected_sections.get('compliance_gaps'):
        html += f'''
        <div class="compliance-gaps mb-4">
            <h3 class="text-secondary"><i class="fas fa-exclamation-circle me-2"></i>{translations.get('compliance_gaps', 'Compliance Gaps')}</h3>
            <div class="row">
        '''
        
        if selected_sections.get('pci_dss'):
            html += f'''
                <div class="col-md-6">
                    <h5>PCI DSS {translations.get('gaps', 'Gaps')}</h5>
                    <div class="table-responsive">
                        <table class="table table-striped">
                            <thead>
                                <tr>
                                    <th>{translations.get('vulnerability', 'Vulnerability')}</th>
                                    <th>{translations.get('status', 'Status')}</th>
                                </tr>
                            </thead>
                            <tbody>
            '''
            
            # Add PCI DSS gaps
            for gap in data['compliance']['pcidss'].get('gaps', [])[:5]:  # Show first 5
                html += f'''
                                <tr>
                                    <td>{gap.get('vulnerability_name', 'N/A')}</td>
                                    <td><span class="badge bg-danger">{translations.get('non_compliant', 'Non-Compliant')}</span></td>
                                </tr>
                '''
            
            html += '''
                            </tbody>
                        </table>
                    </div>
                </div>
            '''
        
        if selected_sections.get('iso27001'):
            html += f'''
                <div class="col-md-6">
                    <h5>ISO 27001 {translations.get('gaps', 'Gaps')}</h5>
                    <div class="table-responsive">
                        <table class="table table-striped">
                            <thead>
                                <tr>
                                    <th>{translations.get('vulnerability', 'Vulnerability')}</th>
                                    <th>{translations.get('status', 'Status')}</th>
                                </tr>
                            </thead>
                            <tbody>
            '''
            
            # Add ISO 27001 gaps
            for gap in data['compliance']['iso27001'].get('gaps', [])[:5]:  # Show first 5
                html += f'''
                                <tr>
                                    <td>{gap.get('vulnerability_name', 'N/A')}</td>
                                    <td><span class="badge bg-danger">{translations.get('non_compliant', 'Non-Compliant')}</span></td>
                                </tr>
                '''
            
            html += '''
                            </tbody>
                        </table>
                    </div>
                </div>
            '''
        
        html += '''
            </div>
        </div>
        '''
    
    # Recommendations Section
    if selected_sections.get('recommendations'):
        html += f'''
        <div class="recommendations mb-4">
            <h3 class="text-secondary"><i class="fas fa-lightbulb me-2"></i>{translations.get('recommendations', 'Recommendations')}</h3>
            <div class="alert alert-warning">
                <h6>{translations.get('compliance_recommendations', 'Compliance Recommendations')}:</h6>
                <ol>
                    <li>{translations.get('recommendation_compliance_1', 'Review and address all non-compliant vulnerabilities')}</li>
                    <li>{translations.get('recommendation_compliance_2', 'Implement required security controls for PCI DSS and ISO 27001')}</li>
                    <li>{translations.get('recommendation_compliance_3', 'Establish regular compliance monitoring and reporting')}</li>
                    <li>{translations.get('recommendation_compliance_4', 'Update security policies and procedures')}</li>
                </ol>
            </div>
        </div>
        '''
    
    # Charts Section
    if selected_sections.get('charts'):
        html += f'''
        <div class="charts-section mb-4">
            <h3 class="text-secondary"><i class="fas fa-chart-pie me-2"></i>{translations.get('compliance_charts', 'Compliance Charts')}</h3>
            <div class="alert alert-info">
                <i class="fas fa-info-circle me-2"></i>
                {translations.get('compliance_charts_note', 'Detailed compliance charts and visualizations will be included in the final report.')}
            </div>
        </div>
        '''
    
    # Report Footer
    html += f'''
        <div class="report-footer text-center mt-4">
            <p class="text-muted">{translations['report_generated_by']}: {data['generated_by']}</p>
            <p class="text-muted">{translations.get('preview_note', 'This is a preview. The final report will contain complete data.')}</p>
        </div>
    </div>
    '''
    
    return html


def get_risk_level_color(risk_level):
    """Get Bootstrap color class for risk level"""
    risk_colors = {
        'Low': 'success',
        'Medium': 'warning',
        'High': 'danger',
        'Critical': 'dark',
        'Низький': 'success',
        'Середній': 'warning',
        'Високий': 'danger',
        'Критичний': 'dark'
    }
    return risk_colors.get(risk_level, 'secondary')
