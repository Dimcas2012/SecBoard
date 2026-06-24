# SecBoard/app_risk/services/report_generator_factory.py

import logging
from typing import Dict, Any, Optional, Type, List
from abc import ABC, abstractmethod
from django.utils.translation import gettext as _

logger = logging.getLogger(__name__)


class BaseReportGenerator(ABC):
    """Abstract base class for report generators"""
    
    def __init__(self, config, data_service):
        self.config = config
        self.data_service = data_service
        self.report_data = None
    
    @abstractmethod
    def generate(self) -> Dict[str, Any]:
        """Generate the report and return result"""
        pass
    
    @abstractmethod
    def get_file_extension(self) -> str:
        """Get file extension for this generator"""
        pass
    
    @abstractmethod
    def get_content_type(self) -> str:
        """Get MIME content type for this generator"""
        pass
    
    def prepare_data(self):
        """Prepare data for report generation"""
        if not self.report_data:
            self.report_data = self.data_service.get_comprehensive_report_data()
        return self.report_data


class PDFReportGenerator(BaseReportGenerator):
    """PDF report generator using ReportLab"""
    
    def generate(self) -> Dict[str, Any]:
        """Generate PDF report"""
        try:
            # Prepare data
            data = self.prepare_data()
            
            # Try ReportLab first
            if self._is_reportlab_available():
                return self._generate_with_reportlab(data)
            else:
                raise ImportError(_("ReportLab library is not available"))
        
        except Exception as e:
            logger.error(f"Error generating PDF report: {e}")
            raise
    
    def get_file_extension(self) -> str:
        return '.pdf'
    
    def get_content_type(self) -> str:
        return 'application/pdf'
    
    def _generate_with_reportlab(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate PDF using ReportLab"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            from io import BytesIO
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title = Paragraph(
                f"Risk Report - {data['generation_date'].strftime('%Y-%m-%d')}",
                styles['Title']
            )
            story.append(title)
            story.append(Spacer(1, 12))
            
            # Statistics section
            stats = data['statistics']
            stats_data = [
                ['Total Assets', str(stats.get('total_assets', 0))],
                ['Total Vulnerabilities', str(stats.get('vulnerability_statistics', {}).get('total_vulnerabilities', 0))],
                ['Completion Rate', f"{stats.get('completion_rate', 0):.1f}%"],
                ['High Risk Count', str(stats.get('high_risk_count', 0))],
            ]
            
            stats_table = Table(stats_data)
            stats_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(stats_table)
            story.append(Spacer(1, 12))
            
            # Risk levels section
            if stats.get('risk_levels'):
                risk_heading = Paragraph('Risk Distribution', styles['Heading2'])
                story.append(risk_heading)
                
                risk_data = [['Risk Level', 'Count']]
                for level, count in stats['risk_levels'].items():
                    risk_data.append([str(level), str(count)])
                
                risk_table = Table(risk_data)
                risk_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(risk_table)
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            
            return {
                'content': buffer.getvalue(),
                'filename': f"risk_report_{data['generation_date'].strftime('%Y%m%d')}.pdf",
                'content_type': self.get_content_type()
            }
        
        except ImportError as e:
            logger.error(f"ReportLab import error: {e}")
            raise ImportError(_("ReportLab library is not available"))
    
    def _is_reportlab_available(self) -> bool:
        try:
            import reportlab
            return True
        except ImportError:
            return False


class WordReportGenerator(BaseReportGenerator):
    """Word report generator using python-docx"""
    
    def generate(self) -> Dict[str, Any]:
        """Generate Word report with improved styling"""
        try:
            import docx
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from io import BytesIO
            
            # Prepare data
            data = self.prepare_data()
            
            # Create document
            doc = docx.Document()
            
            # Margins and defaults
            try:
                for section in doc.sections:
                    section.left_margin = Inches(1)
                    section.right_margin = Inches(1)
                    section.top_margin = Inches(1)
                    section.bottom_margin = Inches(1)
                    section.header_distance = Inches(0.3)
                    section.footer_distance = Inches(0.5)
                normal_style = doc.styles['Normal']
                normal_style.font.name = 'Calibri'
                normal_style.font.size = Pt(11)
            except Exception:
                pass
            
            # Header/footer with page numbers
            try:
                first_section = doc.sections[0]
                header_p = first_section.header.paragraphs[0] if first_section.header.paragraphs else first_section.header.add_paragraph()
                header_p.text = f"Risk Report"
                header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                footer_p = first_section.footer.paragraphs[0] if first_section.footer.paragraphs else first_section.footer.add_paragraph()
                footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                footer_p.add_run('Page ')
                run_page = footer_p.add_run(); r = run_page._r
                fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin'); r.append(fld_begin)
                instr_text = OxmlElement('w:instrText'); instr_text.set(qn('xml:space'), 'preserve'); instr_text.text = 'PAGE'; r.append(instr_text)
                fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end'); r.append(fld_end)
                footer_p.add_run(' of ')
                run_numpages = footer_p.add_run(); r2 = run_numpages._r
                fld_begin2 = OxmlElement('w:fldChar'); fld_begin2.set(qn('w:fldCharType'), 'begin'); r2.append(fld_begin2)
                instr_text2 = OxmlElement('w:instrText'); instr_text2.set(qn('xml:space'), 'preserve'); instr_text2.text = 'NUMPAGES'; r2.append(instr_text2)
                fld_end2 = OxmlElement('w:fldChar'); fld_end2.set(qn('w:fldCharType'), 'end'); r2.append(fld_end2)
            except Exception:
                pass
            
            # Title
            title = doc.add_heading(
                f"Risk Report - {data['generation_date'].strftime('%Y-%m-%d')}",
                level=1
            )
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            
            stats = data['statistics']
            p = doc.add_paragraph()
            p.add_run('Total Assets: ').bold = True
            p.add_run(str(stats.get('total_assets', 0)))
            p = doc.add_paragraph()
            p.add_run('Total Vulnerabilities: ').bold = True
            p.add_run(str(stats.get('vulnerability_statistics', {}).get('total_vulnerabilities', 0)))
            p = doc.add_paragraph()
            p.add_run('Completion Rate: ').bold = True
            p.add_run(f"{stats.get('completion_rate', 0):.1f}%")
            p = doc.add_paragraph()
            p.add_run('High Risk Count: ').bold = True
            p.add_run(str(stats.get('high_risk_count', 0)))
            
            # Risk Distribution
            if stats.get('risk_levels'):
                doc.add_heading('Risk Distribution', level=1)
                
                # Create table
                table = doc.add_table(rows=1, cols=2)
                try:
                    table.style = 'Light Shading Accent 1'
                except Exception:
                    table.style = 'Table Grid'
                
                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Risk Level'
                hdr_cells[1].text = 'Count'
                for c in hdr_cells:
                    for par in c.paragraphs:
                        for run in par.runs:
                            run.bold = True
                
                # Data rows
                for level, count in stats['risk_levels'].items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(level)
                    row_cells[1].text = str(count)
            
            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return {
                'content': buffer.getvalue(),
                'filename': f"risk_report_{data['generation_date'].strftime('%Y%m%d')}.docx",
                'content_type': self.get_content_type()
            }
        
        except ImportError:
            raise ImportError(_("python-docx library is not available"))
        except Exception as e:
            logger.error(f"Error generating Word report: {e}")
            raise
    
    def get_file_extension(self) -> str:
        return '.docx'
    
    def get_content_type(self) -> str:
        return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'


class ExcelReportGenerator(BaseReportGenerator):
    """Excel report generator using xlsxwriter"""
    
    def generate(self) -> Dict[str, Any]:
        """Generate Excel report"""
        try:
            import xlsxwriter
            from io import BytesIO
            
            # Prepare data
            data = self.prepare_data()
            
            # Create workbook
            buffer = BytesIO()
            workbook = xlsxwriter.Workbook(buffer)
            
            # Define formats
            header_format = workbook.add_format({
                'bold': True,
                'font_size': 14,
                'bg_color': '#4472C4',
                'font_color': 'white',
                'align': 'center'
            })
            
            cell_format = workbook.add_format({
                'font_size': 11,
                'align': 'left'
            })
            
            number_format = workbook.add_format({
                'font_size': 11,
                'align': 'right',
                'num_format': '#,##0'
            })
            
            percent_format = workbook.add_format({
                'font_size': 11,
                'align': 'right',
                'num_format': '0.0%'
            })
            
            # Summary sheet
            summary_sheet = workbook.add_worksheet('Summary')
            summary_sheet.write('A1', 'Risk Report Summary', header_format)
            summary_sheet.write('A2', f"Generated: {data['generation_date'].strftime('%Y-%m-%d %H:%M')}", cell_format)
            
            # Statistics
            stats = data['statistics']
            row = 4
            summary_sheet.write(row, 0, 'Total Assets', cell_format)
            summary_sheet.write(row, 1, stats.get('total_assets', 0), number_format)
            
            row += 1
            summary_sheet.write(row, 0, 'Total Vulnerabilities', cell_format)
            summary_sheet.write(row, 1, stats.get('vulnerability_statistics', {}).get('total_vulnerabilities', 0), number_format)
            
            row += 1
            summary_sheet.write(row, 0, 'Completion Rate', cell_format)
            summary_sheet.write(row, 1, stats.get('completion_rate', 0) / 100, percent_format)
            
            row += 1
            summary_sheet.write(row, 0, 'High Risk Count', cell_format)
            summary_sheet.write(row, 1, stats.get('high_risk_count', 0), number_format)
            
            # Risk Levels sheet
            if stats.get('risk_levels'):
                risk_sheet = workbook.add_worksheet('Risk Levels')
                risk_sheet.write('A1', 'Risk Level', header_format)
                risk_sheet.write('B1', 'Count', header_format)
                
                row = 1
                for level, count in stats['risk_levels'].items():
                    risk_sheet.write(row, 0, str(level), cell_format)
                    risk_sheet.write(row, 1, count, number_format)
                    row += 1
            
            # Close workbook
            workbook.close()
            buffer.seek(0)
            
            return {
                'content': buffer.getvalue(),
                'filename': f"risk_report_{data['generation_date'].strftime('%Y%m%d')}.xlsx",
                'content_type': self.get_content_type()
            }
        
        except ImportError:
            raise ImportError(_("xlsxwriter library is not available"))
        except Exception as e:
            logger.error(f"Error generating Excel report: {e}")
            raise
    
    def get_file_extension(self) -> str:
        return '.xlsx'
    
    def get_content_type(self) -> str:
        return 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'


class ReportGeneratorFactory:
    """Factory for creating report generators"""
    
    _generators: Dict[str, Type[BaseReportGenerator]] = {
        'pdf': PDFReportGenerator,
        'word': WordReportGenerator,
    }
    
    @classmethod
    def create_generator(cls, config, data_service) -> BaseReportGenerator:
        """Create appropriate report generator based on format"""
        format_type = config.format.lower()
        
        if format_type not in cls._generators:
            raise ValueError(_("Unsupported report format: {}").format(format_type))
        
        generator_class = cls._generators[format_type]
        return generator_class(config, data_service)
    
    @classmethod
    def get_supported_formats(cls) -> List[str]:
        """Get list of supported formats"""
        return list(cls._generators.keys())
    
    @classmethod
    def register_generator(cls, format_type: str, generator_class: Type[BaseReportGenerator]):
        """Register a new generator for a format"""
        cls._generators[format_type] = generator_class
    
    @classmethod
    def is_format_available(cls, format_type: str) -> bool:
        """Check if format is available (dependencies installed)"""
        if format_type not in cls._generators:
            return False
        
        try:
            # Test format-specific dependencies
            if format_type == 'pdf':
                try:
                    import reportlab
                    return True
                except ImportError:
                    return False
            elif format_type == 'word':
                try:
                    import docx
                    return True
                except ImportError:
                    return False
            elif format_type == 'excel':
                try:
                    import xlsxwriter
                    return True
                except ImportError:
                    return False
            
            return True
        
        except Exception:
            return False