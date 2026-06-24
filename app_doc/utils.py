# SecBoard/app_doc/utils.py

import os
import io
import json
import logging
import tempfile
from typing import Dict, List, Any, Tuple, Optional

import google.generativeai as genai
from django.conf import settings
from django.core.files.storage import default_storage
from django.utils.translation import gettext as _

# Document parsing libraries
try:
    import PyPDF2  # type: ignore
except ImportError:
    PyPDF2 = None  # type: ignore

try:
    import docx  # type: ignore
except ImportError:
    docx = None  # type: ignore
from app_ai.models import APISettingsGoogle, APISettingsClaude, APISettingsGroq, APISettingsDeepSeek
from .models import RegisterDocs

# Try to import document structure models (may not exist)
try:
    from .models import DocumentSection, DocumentContent, DocumentMetadata
except ImportError:
    DocumentSection = None
    DocumentContent = None
    DocumentMetadata = None

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str, page_by_page: bool = False) -> Dict[str, Any]:
    """
    Extract text from a PDF file
    
    Args:
        file_path: Path to the PDF file
        page_by_page: If True, return text for each page separately
        
    Returns:
        Dict with 'full_text' and optional 'pages' keys
    """
    if PyPDF2 is None:
        logger.error("PyPDF2 is not installed. Cannot extract text from PDF.")
        return {'full_text': "", 'total_pages': 0, 'pages': []}
    
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            result = {
                'full_text': "",
                'total_pages': len(reader.pages)
            }
            
            if page_by_page:
                result['pages'] = []
                
            for page_num in range(len(reader.pages)):
                page_text = reader.pages[page_num].extract_text()
                result['full_text'] += page_text + "\n\n"
                
                if page_by_page:
                    result['pages'].append({
                        'page_num': page_num + 1,
                        'text': page_text
                    })
            
            return result
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        return {'full_text': "", 'total_pages': 0, 'pages': []}

def extract_text_from_docx(file_path: str, page_by_page: bool = False) -> Dict[str, Any]:
    """
    Extract text from a DOCX file
    
    Args:
        file_path: Path to the DOCX file
        page_by_page: If True, attempt to separate content by paragraphs (not true pages)
        
    Returns:
        Dict with 'full_text' and optional 'pages' keys
    """
    if docx is None:
        logger.error("python-docx is not installed. Cannot extract text from DOCX.")
        return {'full_text': "", 'total_pages': 0, 'pages': []}
    
    try:
        doc = docx.Document(file_path)
        result = {
            'full_text': "",
            'total_pages': 1  # DOCX doesn't have a direct way to get page count
        }
        
        # Extract paragraphs
        paragraphs_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs_text.append(para.text)
                result['full_text'] += para.text + "\n"
        
        # Extract tables
        tables_text = []
        for table in doc.tables:
            table_text = ""
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " | "
                table_text += "\n"
            
            if table_text.strip():
                tables_text.append(table_text)
                result['full_text'] += table_text + "\n"
        
        # If page_by_page is True, create artificial "pages" based on paragraphs
        if page_by_page:
            result['pages'] = []
            
            # Estimate ~15 paragraphs per page as a rough approximation
            paragraphs_per_page = 15
            total_paragraphs = len(paragraphs_text)
            estimated_pages = max(1, (total_paragraphs + paragraphs_per_page - 1) // paragraphs_per_page)
            
            result['total_pages'] = estimated_pages
            
            for page_num in range(estimated_pages):
                start_idx = page_num * paragraphs_per_page
                end_idx = min(start_idx + paragraphs_per_page, total_paragraphs)
                
                page_paragraphs = paragraphs_text[start_idx:end_idx]
                page_text = "\n".join(page_paragraphs)
                
                result['pages'].append({
                    'page_num': page_num + 1,
                    'text': page_text
                })
        
        return result
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        return {'full_text': "", 'total_pages': 0, 'pages': []}

def extract_document_text(file_path: str, page_by_page: bool = False) -> Dict[str, Any]:
    """
    Extract text from a document based on its extension
    
    Args:
        file_path: Path to the document file
        page_by_page: If True, return text for each page separately
        
    Returns:
        Dict with document text information
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path, page_by_page)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path, page_by_page)
    else:
        logger.warning(f"Unsupported file extension: {ext}")
        return {'full_text': "", 'total_pages': 0, 'pages': []}

def parse_document_with_ai(file_path: str, ai_query: str = "", ai_provider: str = "claude", language: str = "en") -> Dict[str, Any]:
    """
    Parse document content using the specified AI provider
    
    Args:
        file_path: Path to the document file
        ai_query: Custom query for AI analysis
        ai_provider: AI provider to use (claude, google, groq, deepseek)
        language: Language code for AI response (en, uk, ru)
        
    Returns:
        Dict with AI analysis results
    """
    try:
        # Extract text from document
        document_text_data = extract_document_text(file_path, page_by_page=False)
        
        if not document_text_data['full_text']:
            logger.warning(f"Could not extract text from document at {file_path}")
            return {'success': False, 'error': 'Could not extract text from document'}
        
        document_text = document_text_data['full_text']
        
        # Prepare the prompt for AI analysis based on language
        if language == "uk":
            system_prompt = "Ви експерт з аналізу документів, що спеціалізується на виявленні обов'язкових процесів та вимог дотримання."
            process_instruction = "Будь ласка, проаналізуйте цей документ та виявіть потенційні обов'язкові процеси, які слід відстежувати та контролювати."
            json_format = """
            Будь ласка, поверніть вашу відповідь у наступному JSON форматі:
            {
                "processes": [
                    {
                        "name": "Назва процесу",
                        "description": "Детальний опис процесу",
                        "frequency": "Частота (наприклад, щодня, щотижня, щомісяця, щокварталу, щороку)",
                        "section": "Розділ вихідного документа або посилання"
                    }
                ],
                "summary": "Короткий підсумок аналізу документа"
            }
            """
        elif language == "ru":
            system_prompt = "Вы эксперт по анализу документов, специализирующийся на выявлении обязательных процессов и требований соответствия."
            process_instruction = "Пожалуйста, проанализируйте этот документ и выявите потенциальные обязательные процессы, которые следует отслеживать и контролировать."
            json_format = """
            Пожалуйста, верните ваш ответ в следующем JSON формате:
            {
                "processes": [
                    {
                        "name": "Название процесса",
                        "description": "Подробное описание процесса",
                        "frequency": "Частота (например, ежедневно, еженедельно, ежемесячно, ежеквартально, ежегодно)",
                        "section": "Раздел исходного документа или ссылка"
                    }
                ],
                "summary": "Краткое резюме анализа документа"
            }
            """
        else:  # English (default)
            system_prompt = "You are an expert document analyzer specializing in identifying mandatory processes and compliance requirements."
            process_instruction = "Please analyze the following document and identify potential mandatory processes that should be tracked and monitored."
            json_format = """
            Please return your response in the following JSON format:
            {
                "processes": [
                    {
                        "name": "Process Name",
                        "description": "Detailed description of the process",
                        "frequency": "Frequency (e.g., daily, weekly, monthly, quarterly, annually)",
                        "section": "Source document section or reference"
                    }
                ],
                "summary": "Brief summary of the document analysis"
            }
            """
        
        if ai_query:
            prompt = f"""
            {ai_query}
            
            Document content:
            {document_text[:8000]}
            
            {json_format}
            """
        else:
            prompt = f"""
            {system_prompt}
            
            {process_instruction}
            
            Document content:
            {document_text[:8000]}
            
            {json_format}
            """
        
        # Get AI response based on provider
        ai_response = get_ai_response_for_parsing(prompt, ai_provider, language)
        
        if not ai_response:
            return {'success': False, 'error': f'Failed to get response from {ai_provider} AI'}
        
        # Try to parse JSON response
        try:
            # Clean the response text
            response_text = ai_response.strip()
            if response_text.startswith('```json'):
                response_text = response_text.replace('```json', '', 1)
            if response_text.endswith('```'):
                response_text = response_text[:-3]
            
            # Parse JSON
            parsed_data = json.loads(response_text.strip())
            
            return {
                'success': True,
                'processes': parsed_data.get('processes', []),
                'summary': parsed_data.get('summary', 'Document analysis completed'),
                'ai_provider': ai_provider
            }
            
        except json.JSONDecodeError as e:
            logger.error(f"Error parsing AI response JSON: {str(e)}")
            logger.error(f"Response text: {ai_response[:500]}")
            
            # Return a fallback response
            return {
                'success': True,
                'processes': [{
                    'name': f"Process from {ai_provider} analysis",
                    'description': ai_response[:500] + "..." if len(ai_response) > 500 else ai_response,
                    'frequency': 'monthly',
                    'section': 'General document content'
                }],
                'summary': 'AI analysis completed with fallback parsing',
                'ai_provider': ai_provider
            }
            
    except Exception as e:
        logger.error(f"Error parsing document with {ai_provider} AI: {str(e)}")
        return {'success': False, 'error': f'Error parsing document: {str(e)}'}


def get_ai_response_for_parsing(prompt: str, ai_provider: str, language: str = "en") -> str:
    """
    Get AI response from the specified provider for document parsing
    
    Args:
        prompt: The prompt to send to AI
        ai_provider: AI provider to use
        language: Language code for AI response
        
    Returns:
        AI response text
    """
    try:
        if ai_provider == 'claude':
            return get_claude_response_for_parsing(prompt, language)
        elif ai_provider == 'google':
            return get_google_response_for_parsing(prompt, language)
        elif ai_provider == 'groq':
            return get_groq_response_for_parsing(prompt, language)
        elif ai_provider == 'deepseek':
            return get_deepseek_response_for_parsing(prompt, language)
        else:
            logger.warning(f"Unknown AI provider: {ai_provider}, defaulting to Claude")
            return get_claude_response_for_parsing(prompt, language)
            
    except Exception as e:
        logger.error(f"Error getting AI response from {ai_provider}: {str(e)}")
        return ""


def get_claude_response_for_parsing(prompt: str, language: str = "en") -> str:
    """Get response from Claude AI for document parsing"""
    try:
        claude_settings = APISettingsClaude.objects.first()
        if not claude_settings:
            logger.error("Claude API settings not found")
            return ""
        
        import anthropic
        client = anthropic.Anthropic(api_key=claude_settings.api_key)
        
        # Add language instruction to the prompt
        language_instruction = ""
        if language == "uk":
            language_instruction = "\n\nВажливо: Відповідайте українською мовою."
        elif language == "ru":
            language_instruction = "\n\nВажно: Отвечайте на русском языке."
        else:
            language_instruction = "\n\nImportant: Respond in English."
        
        enhanced_prompt = prompt + language_instruction
        
        response = client.messages.create(
            model=claude_settings.model_name.model_id,
            max_tokens=claude_settings.max_tokens,
            temperature=claude_settings.temperature,
            messages=[{"role": "user", "content": enhanced_prompt}]
        )
        
        return response.content[0].text
        
    except Exception as e:
        logger.error(f"Error getting Claude response: {str(e)}")
        return ""


def get_google_response_for_parsing(prompt: str, language: str = "en") -> str:
    """Get response from Google AI for document parsing"""
    try:
        google_settings = APISettingsGoogle.objects.first()
        if not google_settings:
            logger.error("Google AI settings not found")
            return ""
        
        genai.configure(api_key=google_settings.api_key)
        model = genai.GenerativeModel(google_settings.model_name.model_id)
        
        # Add language instruction to the prompt
        language_instruction = ""
        if language == "uk":
            language_instruction = "\n\nВажливо: Відповідайте українською мовою."
        elif language == "ru":
            language_instruction = "\n\nВажно: Отвечайте на русском языке."
        else:
            language_instruction = "\n\nImportant: Respond in English."
        
        enhanced_prompt = prompt + language_instruction
        
        response = model.generate_content(enhanced_prompt)
        return response.text
        
    except Exception as e:
        logger.error(f"Error getting Google AI response: {str(e)}")
        return ""


def get_groq_response_for_parsing(prompt: str, language: str = "en") -> str:
    """Get response from Groq AI for document parsing"""
    try:
        groq_settings = APISettingsGroq.objects.first()
        if not groq_settings:
            logger.error("Groq API settings not found")
            return ""
        
        from groq import Groq
        client = Groq(api_key=groq_settings.api_key)
        
        # Add language instruction to the prompt
        language_instruction = ""
        if language == "uk":
            language_instruction = "\n\nВажливо: Відповідайте українською мовою."
        elif language == "ru":
            language_instruction = "\n\nВажно: Отвечайте на русском языке."
        else:
            language_instruction = "\n\nImportant: Respond in English."
        
        enhanced_prompt = prompt + language_instruction
        
        response = client.chat.completions.create(
            messages=[{"role": "user", "content": enhanced_prompt}],
            model=groq_settings.model_name.model_id,
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        logger.error(f"Error getting Groq response: {str(e)}")
        return ""


def get_deepseek_response_for_parsing(prompt: str, language: str = "en") -> str:
    """Get response from DeepSeek AI for document parsing"""
    try:
        deepseek_settings = APISettingsDeepSeek.objects.first()
        if not deepseek_settings:
            logger.error("DeepSeek API settings not found")
            return ""
        
        from openai import OpenAI
        client = OpenAI(api_key=deepseek_settings.api_key, base_url="https://api.deepseek.com/v1")
        
        # Add language instruction to the prompt
        language_instruction = ""
        if language == "uk":
            language_instruction = "\n\nВажливо: Відповідайте українською мовою."
        elif language == "ru":
            language_instruction = "\n\nВажно: Отвечайте на русском языке."
        else:
            language_instruction = "\n\nImportant: Respond in English."
        
        enhanced_prompt = prompt + language_instruction
        
        response = client.chat.completions.create(
            model=deepseek_settings.model_name.model_id,
            messages=[{"role": "user", "content": enhanced_prompt}],
            max_tokens=deepseek_settings.max_tokens,
            temperature=deepseek_settings.temperature
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        logger.error(f"Error getting DeepSeek response: {str(e)}")
        return ""


def parse_document_with_google_ai(document: RegisterDocs, page_by_page: bool = False) -> bool:
    """
    Parse document content using Google AI (maintained for backward compatibility)
    
    Args:
        document: RegisterDocs instance
        page_by_page: If True, parse document page by page
        
    Returns:
        True if successful, False otherwise
    """
    if not document.file_doc:
        logger.warning(f"No file attached to document {document.id}")
        return False
    
    try:
        # Get Google AI settings
        google_settings = APISettingsGoogle.objects.first()
        if not google_settings:
            logger.error("Google AI settings not found")
            return False
        
        # Configure Google AI
        genai.configure(api_key=google_settings.api_key)
        model = genai.GenerativeModel(google_settings.model_name.model_id)
        
        # Get file path
        file_path = document.file_doc.path
        
        # Extract text from document
        document_text_data = extract_document_text(file_path, page_by_page)
        
        if not document_text_data['full_text']:
            logger.warning(f"Could not extract text from document {document.id}")
            return False
        
        if page_by_page and document_text_data.get('pages'):
            # Parse document page by page
            return parse_document_pages(document, document_text_data, model)
        else:
            # Parse entire document at once
            return parse_document_full(document, document_text_data, model)
    
    except Exception as e:
        logger.error(f"Error parsing document with Google AI: {str(e)}")
        return False

def parse_document_full(document: RegisterDocs, document_text_data: Dict[str, Any], model) -> bool:
    """Parse entire document at once"""
    try:
        document_text = document_text_data['full_text']
        
        # Prepare prompt for Google AI
        prompt = f"""
        You are an expert document analyzer. I need you to analyze the following document and extract its structure.
        
        Please identify:
        1. All sections and subsections with their hierarchical structure
        2. The content of each section (text, tables, lists)
        3. Any key metadata (terms, definitions, references)
        
        Format your response as a JSON object with the following structure:
        {{
            "sections": [
                {{
                    "title": "Section title",
                    "level": 1,
                    "parent_id": null,
                    "contents": [
                        {{
                            "type": "text",
                            "text": "Content text",
                            "metadata": [
                                {{"key": "key1", "value": "value1"}},
                                {{"key": "key2", "value": "value2"}}
                            ]
                        }},
                        {{
                            "type": "table",
                            "text": "Table content in JSON format",
                            "metadata": []
                        }}
                    ]
                }},
                {{
                    "title": "Subsection title",
                    "level": 2,
                    "parent_id": 0,
                    "contents": [...]
                }}
            ]
        }}
        
        Here is the document text:
        {document_text[:10000]}  # Limit text to avoid token limits
        """
        
        # Generate response from Google AI
        response = model.generate_content(prompt)
        response_text = response.text
        
        # Extract JSON from response
        try:
            # Try to find JSON in the response
            json_match = response_text.strip()
            if json_match.startswith('```json'):
                json_match = json_match.replace('```json', '', 1)
            if json_match.endswith('```'):
                json_match = json_match[:-3]
            
            # Parse JSON
            parsed_data = json.loads(json_match.strip())
            
            # Save parsed data to database
            save_parsed_document_structure(document, parsed_data)
            
            return True
        except Exception as e:
            logger.error(f"Error parsing Google AI response: {str(e)}")
            logger.error(f"Response text: {response_text}")
            return False
    
    except Exception as e:
        logger.error(f"Error parsing full document: {str(e)}")
        return False

def parse_document_pages(document: RegisterDocs, document_text_data: Dict[str, Any], model) -> bool:
    """Parse document page by page"""
    try:
        pages = document_text_data.get('pages', [])
        if not pages:
            logger.warning(f"No pages found in document {document.id}")
            return False
        
        all_sections = []
        section_id_counter = 0
        
        # Process each page
        for page_data in pages:
            page_num = page_data['page_num']
            page_text = page_data['text']
            
            if not page_text.strip():
                continue
            
            # Prepare prompt for Google AI
            prompt = f"""
            You are an expert document analyzer. I need you to analyze page {page_num} of a document and extract its structure.
            
            Please identify:
            1. All sections and subsections with their hierarchical structure
            2. The content of each section (text, tables, lists)
            3. Any key metadata (terms, definitions, references)
            
            Format your response as a JSON object with the following structure:
            {{
                "sections": [
                    {{
                        "title": "Section title",
                        "level": 1,
                        "parent_id": null,
                        "contents": [
                            {{
                                "type": "text",
                                "text": "Content text",
                                "metadata": [
                                    {{"key": "key1", "value": "value1"}},
                                    {{"key": "key2", "value": "value2"}}
                                ]
                            }},
                            {{
                                "type": "table",
                                "text": "Table content in JSON format",
                                "metadata": []
                            }}
                        ]
                    }},
                    {{
                        "title": "Subsection title",
                        "level": 2,
                        "parent_id": 0,
                        "contents": [...]
                    }}
                ]
            }}
            
            Here is the text from page {page_num}:
            {page_text}
            """
            
            # Generate response from Google AI
            response = model.generate_content(prompt)
            response_text = response.text
            
            # Extract JSON from response
            try:
                # Try to find JSON in the response
                json_match = response_text.strip()
                if json_match.startswith('```json'):
                    json_match = json_match.replace('```json', '', 1)
                if json_match.endswith('```'):
                    json_match = json_match[:-3]
                
                # Parse JSON
                parsed_data = json.loads(json_match.strip())
                
                # Add page number to section titles
                sections = parsed_data.get('sections', [])
                for section in sections:
                    # Update section title to include page number
                    if not section['title'].startswith(f"Page {page_num}:"):
                        section['title'] = f"Page {page_num}: {section['title']}"
                    
                    # Update section IDs to avoid conflicts
                    section['original_id'] = section.get('id', None)
                    section['id'] = section_id_counter
                    section_id_counter += 1
                    
                    # Update parent_id references
                    if section.get('parent_id') is not None:
                        # Find the corresponding section in the current page
                        parent_original_id = section['parent_id']
                        for potential_parent in sections:
                            if potential_parent.get('original_id') == parent_original_id:
                                section['parent_id'] = potential_parent['id']
                                break
                        else:
                            # No parent found in current page, set to None
                            section['parent_id'] = None
                
                all_sections.extend(sections)
            
            except Exception as e:
                logger.error(f"Error parsing Google AI response for page {page_num}: {str(e)}")
                logger.error(f"Response text: {response_text}")
                # Continue with other pages even if one fails
        
        if all_sections:
            # Save all parsed sections
            save_parsed_document_structure(document, {"sections": all_sections})
            return True
        else:
            logger.warning(f"No sections extracted from document {document.id}")
            return False
    
    except Exception as e:
        logger.error(f"Error parsing document pages: {str(e)}")
        return False

def save_parsed_document_structure(document: RegisterDocs, parsed_data: Dict[str, Any]) -> None:
    """Save parsed document structure to database"""
    try:
        # Check if required models are available
        if DocumentSection is None or DocumentContent is None or DocumentMetadata is None:
            logger.warning("DocumentSection, DocumentContent, or DocumentMetadata models not found. Skipping document structure save.")
            return
        
        # Create a mapping of section IDs to database objects
        section_mapping = {}
        
        # Process sections
        for i, section_data in enumerate(parsed_data.get('sections', [])):
            # Create section
            section = DocumentSection(
                document=document,
                section_title=section_data.get('title', f"Section {i+1}"),
                section_level=section_data.get('level', 1),
                order=i
            )
            
            # Handle parent section
            parent_id = section_data.get('parent_id')
            if parent_id is not None and parent_id in section_mapping:
                section.parent_section = section_mapping[parent_id]
            
            section.save()
            section_mapping[i] = section
            
            # Process contents
            for j, content_data in enumerate(section_data.get('contents', [])):
                content_type = content_data.get('type', 'text')
                content_text = content_data.get('text', '')
                
                # Create content
                content = DocumentContent(
                    section=section,
                    content_type=content_type,
                    content_text=content_text,
                    order=j
                )
                content.save()
                
                # Process metadata
                for metadata_data in content_data.get('metadata', []):
                    metadata = DocumentMetadata(
                        content=content,
                        key=metadata_data.get('key', ''),
                        value=metadata_data.get('value', '')
                    )
                    metadata.save()
    
    except Exception as e:
        logger.error(f"Error saving parsed document structure: {str(e)}")
        raise 